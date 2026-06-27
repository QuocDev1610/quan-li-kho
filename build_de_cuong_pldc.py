from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"D:\ProjectWeb\Project - Sao chép\De-cuong-PLDC-19-cau-hoan-chinh.docx"

CONTENT = r"""
# Câu 1. Trình bày khái niệm, bản chất của Nhà nước và ý nghĩa phương pháp luận

## 1. Khái niệm Nhà nước
Nhà nước là một tổ chức đặc biệt của quyền lực chính trị, có bộ máy chuyên làm nhiệm vụ quản lý xã hội và cưỡng chế, quản lý dân cư theo lãnh thổ, ban hành pháp luật, tổ chức thực hiện pháp luật, thiết lập thuế và đại diện chính thức cho quốc gia trong quan hệ đối nội, đối ngoại.

Khái niệm trên cho thấy Nhà nước không phải là mọi tổ chức quyền lực trong xã hội. Nhà nước là tổ chức quyền lực công cộng đặc biệt, tách khỏi dân cư nhưng nhân danh toàn xã hội để quản lý xã hội bằng pháp luật và các công cụ quyền lực khác.

## 2. Bản chất của Nhà nước
### a) Tính giai cấp
Theo quan điểm của chủ nghĩa Mác - Lênin, Nhà nước xuất hiện khi xã hội phân chia thành các giai cấp có lợi ích cơ bản đối lập và mâu thuẫn giai cấp không thể tự điều hòa. Nhà nước trước hết là công cụ quyền lực chính trị của giai cấp hoặc liên minh giai cấp giữ địa vị thống trị về kinh tế.

Tính giai cấp thể hiện ở việc giai cấp thống trị sử dụng Nhà nước để bảo vệ chế độ sở hữu và trật tự xã hội có lợi cho mình; tổ chức, thực hiện quyền lực chính trị; xây dựng hệ tư tưởng và các giá trị chi phối đời sống xã hội. Quyền lực nhà nước được bảo đảm bằng pháp luật, bộ máy quản lý, lực lượng vũ trang, tòa án, cơ sở vật chất và các biện pháp cưỡng chế.

Cần sửa một ý sai trong bản phác thảo: Nhà nước tuy là đại diện chính thức của toàn xã hội nhưng không phải “trước hết và trên hết bảo vệ lợi ích của toàn xã hội” trong mọi kiểu nhà nước. Xét về bản chất giai cấp, Nhà nước trước hết bảo vệ những quan hệ xã hội cơ bản phù hợp với lợi ích của lực lượng cầm quyền; đồng thời phải giải quyết các công việc chung để duy trì xã hội.

### b) Tính xã hội
Nhà nước chỉ có thể tồn tại bền vững khi thực hiện những công việc chung mà xã hội cần: bảo đảm an ninh, trật tự; tổ chức kinh tế; xây dựng kết cấu hạ tầng; bảo vệ môi trường; phát triển giáo dục, y tế, văn hóa; phòng, chống thiên tai, dịch bệnh; bảo vệ quyền và lợi ích hợp pháp của con người.

Tính xã hội và tính giai cấp không tách rời nhau. Mức độ thể hiện của từng mặt phụ thuộc vào kiểu nhà nước, chế độ chính trị, tương quan lực lượng xã hội và điều kiện lịch sử cụ thể.

## 3. Ý nghĩa phương pháp luận
- Khi nghiên cứu Nhà nước phải xem xét đồng thời tính giai cấp và tính xã hội, tránh tuyệt đối hóa một mặt.
- Phải đặt Nhà nước trong điều kiện kinh tế, chính trị, văn hóa và lịch sử cụ thể; không xem Nhà nước là hiện tượng bất biến hoặc do ý chí chủ quan của một cá nhân tạo ra.
- Việc đánh giá một nhà nước không chỉ dựa vào tuyên bố chính trị mà phải căn cứ vào cơ sở kinh tế, lực lượng nắm quyền, tổ chức bộ máy, pháp luật và hoạt động thực tế.
- Nhận thức đúng bản chất Nhà nước giúp xây dựng, hoàn thiện Nhà nước pháp quyền, kiểm soát quyền lực và bảo đảm quyền con người, quyền công dân.

## Kết luận
Nhà nước là hiện tượng lịch sử - xã hội đặc biệt, vừa mang bản chất giai cấp vừa thực hiện vai trò xã hội. Cách tiếp cận toàn diện này là cơ sở khoa học để lý giải sự ra đời, tồn tại, phát triển và hoạt động của Nhà nước.

# Câu 2. Khái niệm, đặc điểm của văn bản quy phạm pháp luật và các nguyên tắc ban hành

## 1. Khái niệm
Văn bản quy phạm pháp luật là văn bản có chứa quy phạm pháp luật, được ban hành đúng thẩm quyền, hình thức, trình tự và thủ tục theo luật định. Quy phạm trong văn bản là quy tắc xử sự chung, có hiệu lực bắt buộc chung, được áp dụng lặp lại đối với cơ quan, tổ chức, cá nhân trong phạm vi điều chỉnh.

Không nên định nghĩa đơn giản rằng đây chỉ là “hình thức pháp luật tiến bộ nhất”. Đó là nhận xét về ưu thế của văn bản, không phải dấu hiệu pháp lý đầy đủ để nhận diện văn bản quy phạm pháp luật.

## 2. Đặc điểm
- Do chủ thể có thẩm quyền ban hành hoặc phối hợp ban hành. Mỗi chủ thể chỉ được ban hành loại văn bản và quy định những nội dung thuộc thẩm quyền của mình.
- Được ban hành dưới tên gọi, hình thức pháp lý do luật quy định; không phải văn bản nào của cơ quan nhà nước cũng là văn bản quy phạm pháp luật.
- Chứa quy tắc xử sự chung, không chỉ giải quyết một vụ việc cá biệt.
- Được áp dụng nhiều lần, đối với nhiều đối tượng hoặc một nhóm đối tượng xác định, khi xuất hiện điều kiện đã được dự liệu.
- Có hiệu lực trong phạm vi thời gian, không gian và đối tượng nhất định; có vị trí xác định trong hệ thống pháp luật.
- Được Nhà nước tổ chức thực hiện và bảo đảm thực hiện, kể cả bằng biện pháp cưỡng chế khi cần thiết.

## 3. Các nguyên tắc xây dựng, ban hành
### a) Bảo đảm tính hợp hiến, hợp pháp và thống nhất
Mọi văn bản phải phù hợp với Hiến pháp; văn bản của cơ quan cấp dưới không được trái văn bản có hiệu lực pháp lý cao hơn; các quy định trong cùng hệ thống phải đồng bộ, không mâu thuẫn, chồng chéo.

### b) Đúng thẩm quyền, hình thức, nội dung, trình tự và thủ tục
Chủ thể ban hành phải đúng thẩm quyền về hình thức và nội dung. Quy trình thông thường gồm đề xuất chính sách hoặc lập chương trình khi pháp luật yêu cầu; soạn thảo; lấy ý kiến; thẩm định, thẩm tra; tiếp thu, chỉnh lý; xem xét, thông qua hoặc ký ban hành; công bố, đăng tải và tổ chức thi hành. Không nên áp dụng máy móc một quy trình của luật do Quốc hội ban hành cho tất cả các loại văn bản.

### c) Bảo đảm dân chủ, công khai và minh bạch
Dự thảo phải được lấy ý kiến phù hợp, nhất là ý kiến của đối tượng chịu tác động; nội dung tiếp thu, giải trình cần rõ ràng; văn bản sau khi ban hành phải được công bố và dễ tiếp cận.

### d) Bảo đảm tính khả thi, kịp thời, ổn định và hiệu quả
Quy định phải rõ ràng, có nguồn lực thực hiện, phù hợp điều kiện kinh tế - xã hội, không tạo thủ tục hoặc chi phí bất hợp lý. Văn bản phải xử lý được vấn đề thực tiễn nhưng cũng cần đủ ổn định để người dân, doanh nghiệp dự liệu hành vi.

### e) Tôn trọng quyền con người, quyền công dân và bảo đảm bình đẳng
Chính sách pháp luật phải bảo vệ quyền, lợi ích hợp pháp; việc hạn chế quyền chỉ được thực hiện theo Hiến pháp và luật, trong trường hợp thật sự cần thiết vì các lý do hiến định.

### f) Bảo đảm yêu cầu về quốc phòng, an ninh, bảo vệ môi trường và điều ước quốc tế
Việc ban hành văn bản phải phù hợp lợi ích quốc gia, dân tộc và không cản trở việc thực hiện điều ước quốc tế mà Việt Nam là thành viên. Khi điều ước quốc tế và văn bản trong nước có quy định khác nhau, việc áp dụng được thực hiện theo quy định của pháp luật về điều ước quốc tế, trừ Hiến pháp.

## Kết luận
Giá trị của văn bản quy phạm pháp luật không chỉ nằm ở việc được ký ban hành mà còn phụ thuộc vào tính hợp hiến, đúng thẩm quyền, minh bạch, khả thi và hiệu quả thực tế.

# Câu 3. Bản chất và đặc trưng của Nhà nước Cộng hòa xã hội chủ nghĩa Việt Nam

## 1. Cơ sở hiến định
Điều 2 Hiến pháp năm 2013 xác định Nhà nước Cộng hòa xã hội chủ nghĩa Việt Nam là Nhà nước pháp quyền xã hội chủ nghĩa của Nhân dân, do Nhân dân, vì Nhân dân; tất cả quyền lực nhà nước thuộc về Nhân dân, nền tảng là liên minh giữa giai cấp công nhân với giai cấp nông dân và đội ngũ trí thức.

## 2. Bản chất
### a) Nhà nước của Nhân dân
Nhân dân là chủ thể tối cao của quyền lực nhà nước. Nhân dân thực hiện quyền lực bằng dân chủ trực tiếp và dân chủ đại diện thông qua Quốc hội, Hội đồng nhân dân và các cơ quan khác của Nhà nước.

### b) Nhà nước do Nhân dân
Bộ máy nhà nước được hình thành và hoạt động trên cơ sở ý chí, sự ủy quyền, tham gia và giám sát của Nhân dân. Cán bộ, công chức phải tôn trọng Nhân dân, tận tụy phục vụ Nhân dân, chịu sự giám sát của Nhân dân.

### c) Nhà nước vì Nhân dân
Mục tiêu của tổ chức và hoạt động nhà nước là bảo vệ độc lập, chủ quyền, lợi ích quốc gia; công nhận, tôn trọng, bảo vệ, bảo đảm quyền con người, quyền công dân; nâng cao đời sống vật chất và tinh thần của Nhân dân.

### d) Tính giai cấp và tính xã hội
Nhà nước mang bản chất của giai cấp công nhân, dựa trên liên minh công nhân - nông dân - trí thức và đặt dưới sự lãnh đạo của Đảng Cộng sản Việt Nam. Đồng thời, Nhà nước có cơ sở xã hội rộng lớn, đại diện cho lợi ích của Nhân dân và của dân tộc Việt Nam.

## 3. Những đặc trưng cơ bản của Nhà nước Việt Nam
### a) Là Nhà nước pháp quyền xã hội chủ nghĩa
Nhà nước được tổ chức và hoạt động trên cơ sở Hiến pháp và pháp luật; Hiến pháp có hiệu lực pháp lý cao nhất; quyền lực phải được kiểm soát; quyền con người, quyền công dân được công nhận, tôn trọng, bảo vệ và bảo đảm.

### b) Quyền lực nhà nước thống nhất, có phân công, phối hợp và kiểm soát
Quyền lực nhà nước thống nhất thuộc về Nhân dân nhưng trong thực hiện có sự phân công, phối hợp và kiểm soát giữa các cơ quan thực hiện quyền lập pháp, hành pháp và tư pháp. Đây không phải mô hình “tam quyền phân lập” tuyệt đối.

### c) Đặt dưới sự lãnh đạo của Đảng Cộng sản Việt Nam
Đảng là lực lượng lãnh đạo Nhà nước và xã hội, gắn bó với Nhân dân, phục vụ Nhân dân, chịu sự giám sát của Nhân dân và chịu trách nhiệm trước Nhân dân về các quyết định của mình; tổ chức đảng và đảng viên hoạt động trong khuôn khổ Hiến pháp và pháp luật.

### d) Tổ chức theo nguyên tắc tập trung dân chủ
Bộ máy nhà nước bảo đảm sự quản lý thống nhất từ trung ương đến địa phương, đồng thời phân định nhiệm vụ, quyền hạn, phát huy tính chủ động và trách nhiệm của từng cơ quan, cấp chính quyền.

### e) Là Nhà nước thống nhất của các dân tộc
Việt Nam là quốc gia thống nhất; các dân tộc bình đẳng, đoàn kết, tôn trọng, giúp nhau cùng phát triển; nghiêm cấm kỳ thị, chia rẽ dân tộc.

### f) Thực hiện đường lối đối ngoại hòa bình, hợp tác
Nhà nước tôn trọng Hiến chương Liên hợp quốc và điều ước quốc tế mà Việt Nam là thành viên; là bạn, đối tác tin cậy và thành viên có trách nhiệm của cộng đồng quốc tế.

## Lưu ý sửa sai
Các ý “dân giàu, nước mạnh, dân chủ, công bằng, văn minh; có nền văn hóa tiên tiến...” trong bản gốc là đặc trưng của xã hội xã hội chủ nghĩa mà Nhân dân ta xây dựng, không phải toàn bộ đặc trưng pháp lý riêng của Nhà nước.

# Câu 4. Khái niệm Luật hình sự và các nguyên tắc cơ bản

## 1. Khái niệm
Luật hình sự là ngành luật độc lập trong hệ thống pháp luật Việt Nam, gồm các quy phạm xác định hành vi nào là tội phạm, quy định hình phạt và các biện pháp hình sự khác, căn cứ và điều kiện chịu trách nhiệm hình sự, đồng thời quy định các nguyên tắc xử lý người và pháp nhân thương mại phạm tội.

## 2. Các nguyên tắc cơ bản
### a) Nguyên tắc pháp chế
Chỉ Bộ luật Hình sự mới có thể quy định tội phạm và hình phạt. Việc khởi tố, điều tra, truy tố, xét xử, quyết định và thi hành hình phạt phải đúng căn cứ, thẩm quyền và thủ tục pháp luật; không được áp dụng tương tự pháp luật để buộc tội.

### b) Nguyên tắc bình đẳng trước pháp luật hình sự
Mọi người phạm tội đều bình đẳng trước pháp luật, không phân biệt giới tính, dân tộc, tín ngưỡng, tôn giáo, thành phần, địa vị xã hội hoặc tình trạng tài sản. Pháp nhân thương mại phạm tội cũng bị xử lý theo điều kiện luật định.

### c) Nguyên tắc có lỗi
Một chủ thể chỉ phải chịu trách nhiệm hình sự khi thực hiện hành vi nguy hiểm cho xã hội với lỗi cố ý hoặc vô ý. Không được quy kết trách nhiệm chỉ vì đã có thiệt hại nếu chủ thể không có lỗi.

### d) Nguyên tắc trách nhiệm theo hành vi và chủ thể
Cá nhân chịu trách nhiệm về hành vi phạm tội do mình thực hiện hoặc tham gia với tư cách đồng phạm. Không được buộc người thân, người đứng đầu hoặc thành viên tổ chức chịu thay nếu họ không tham gia và không có căn cứ pháp luật.

Ý trong đề cương cũ cho rằng hoàn toàn loại trừ trách nhiệm hình sự của tổ chức không còn chính xác. Bộ luật Hình sự hiện hành quy định trách nhiệm hình sự của pháp nhân thương mại đối với những tội danh và điều kiện nhất định; việc pháp nhân chịu trách nhiệm không loại trừ trách nhiệm của cá nhân có liên quan.

### e) Nguyên tắc công bằng, phân hóa và cá thể hóa
Hình phạt phải tương xứng với tính chất, mức độ nguy hiểm của hành vi, lỗi, nhân thân, vai trò đồng phạm, tình tiết tăng nặng và giảm nhẹ. Người phạm tội khác nhau không bị xử lý máy móc như nhau.

### f) Nguyên tắc nhân đạo
Hình phạt không nhằm trả thù, hành hạ hoặc hạ thấp nhân phẩm mà nhằm trừng trị cần thiết, giáo dục, cải tạo và phòng ngừa. Pháp luật có chính sách khoan hồng đối với người tự thú, thành khẩn, lập công, ăn năn; đồng thời có chính sách riêng đối với người dưới 18 tuổi, phụ nữ có thai, người già yếu và các trường hợp đặc biệt.

### g) Kết hợp trừng trị với phòng ngừa
Đấu tranh chống tội phạm phải gắn với phòng ngừa, phát hiện nguyên nhân và điều kiện phạm tội, giáo dục ý thức pháp luật, bảo vệ người bị hại và tái hòa nhập cộng đồng.

## Kết luận
Các nguyên tắc hình sự vừa bảo đảm xử lý nghiêm minh tội phạm, vừa ngăn ngừa oan sai và bảo vệ quyền con người.

# Câu 5. Khái niệm, đặc điểm và các yếu tố cấu thành quan hệ pháp luật

## 1. Khái niệm
Quan hệ pháp luật là quan hệ xã hội được quy phạm pháp luật điều chỉnh, trong đó các chủ thể có quyền và nghĩa vụ pháp lý xác định, được Nhà nước công nhận và bảo đảm thực hiện.

## 2. Đặc điểm
- Là quan hệ xã hội có tính ý chí: vừa thể hiện ý chí của Nhà nước thông qua pháp luật, vừa thể hiện ý chí của các chủ thể khi tham gia trong nhiều trường hợp.
- Phát sinh, thay đổi hoặc chấm dứt trên cơ sở quy phạm pháp luật và sự kiện pháp lý.
- Các bên có quyền chủ thể và nghĩa vụ pháp lý tương ứng, được xác định tương đối cụ thể.
- Việc thực hiện được Nhà nước bảo đảm bằng giáo dục, tổ chức, khuyến khích và khi cần bằng cưỡng chế.
- Có tính xác định về chủ thể, nội dung và khách thể; khác với quan hệ xã hội thuần túy đạo đức hoặc tình cảm.

## 3. Các yếu tố cấu thành
### a) Chủ thể
Chủ thể là cá nhân, pháp nhân, tổ chức hoặc Nhà nước có năng lực chủ thể và tham gia quan hệ pháp luật.

Năng lực pháp luật là khả năng có quyền và nghĩa vụ pháp lý do pháp luật quy định. Đối với cá nhân, năng lực pháp luật dân sự có từ khi sinh ra và chấm dứt khi chết.

Năng lực hành vi là khả năng bằng hành vi của mình xác lập, thực hiện quyền và nghĩa vụ. Năng lực này phụ thuộc vào độ tuổi, tình trạng nhận thức, khả năng làm chủ hành vi và quy định của từng ngành luật.

Năng lực chủ thể của tổ chức phát sinh theo điều kiện thành lập, cơ cấu, chức năng và phạm vi hoạt động. Không nên dùng Điều 84 Bộ luật Dân sự cũ; pháp nhân hiện được xác định theo Bộ luật Dân sự năm 2015 với các điều kiện: được thành lập hợp pháp; có cơ cấu tổ chức; có tài sản độc lập và tự chịu trách nhiệm; nhân danh mình tham gia quan hệ pháp luật độc lập.

### b) Nội dung
Nội dung gồm quyền chủ thể và nghĩa vụ pháp lý.

Quyền chủ thể là khả năng xử sự được pháp luật cho phép: tự mình thực hiện hành vi; yêu cầu bên có nghĩa vụ thực hiện hoặc không thực hiện hành vi; yêu cầu cơ quan có thẩm quyền bảo vệ quyền, lợi ích hợp pháp.

Nghĩa vụ pháp lý là cách xử sự bắt buộc: phải thực hiện một hành vi; phải kiềm chế không thực hiện hành vi; phải chịu hậu quả pháp lý khi không thực hiện đúng nghĩa vụ.

### c) Khách thể
Khách thể là lợi ích vật chất, tinh thần hoặc lợi ích xã hội mà các chủ thể hướng tới khi tham gia quan hệ pháp luật, như tài sản, kết quả công việc, danh dự, sức khỏe, quyền nhân thân, trật tự quản lý.

## 4. Sự kiện pháp lý
Sự kiện pháp lý không phải là một yếu tố nằm trong cơ cấu nội tại của quan hệ pháp luật nhưng là căn cứ trực tiếp làm phát sinh, thay đổi hoặc chấm dứt quan hệ. Sự kiện có thể là hành vi của con người hoặc biến cố khách quan.

# Câu 6. Khái niệm, đặc điểm và các loại hình thức pháp luật

## 1. Khái niệm
Hình thức pháp luật là phương thức tồn tại và biểu hiện của pháp luật. Theo nghĩa bên ngoài, hình thức pháp luật còn được gọi là nguồn pháp luật, tức nơi chứa đựng hoặc căn cứ để nhận biết, áp dụng quy tắc pháp lý.

## 2. Đặc điểm
- Phản ánh điều kiện kinh tế, chính trị, văn hóa, truyền thống và kỹ thuật pháp lý của từng quốc gia, từng thời kỳ.
- Có hình thức xác định, giúp chủ thể nhận biết quyền, nghĩa vụ và hậu quả pháp lý.
- Được Nhà nước thừa nhận và bảo đảm thực hiện.
- Các hình thức có giá trị và phạm vi áp dụng khác nhau; việc lựa chọn nguồn phải tuân theo thứ bậc hiệu lực và nguyên tắc áp dụng pháp luật.

## 3. Các hình thức cơ bản
### a) Tập quán pháp
Là tập quán được Nhà nước thừa nhận hoặc cho phép áp dụng như quy tắc pháp lý. Tập quán thường được áp dụng khi pháp luật không có quy định và các bên không có thỏa thuận, với điều kiện không trái các nguyên tắc cơ bản của pháp luật và đạo đức xã hội.

### b) Tiền lệ pháp, án lệ
Tiền lệ pháp là quyết định giải quyết vụ việc cụ thể được thừa nhận làm khuôn mẫu cho vụ việc tương tự về sau. Ở Việt Nam, án lệ do Hội đồng Thẩm phán Tòa án nhân dân tối cao lựa chọn và công bố là nguồn tham khảo, áp dụng trong xét xử theo quy định.

Không nên kết luận tuyệt đối rằng tiền lệ pháp luôn tạo ra sự tùy tiện. Nếu có cơ chế lựa chọn, công bố, viện dẫn, thay thế và giám sát chặt chẽ, án lệ góp phần thống nhất áp dụng pháp luật và lấp khoảng trống pháp lý.

### c) Văn bản quy phạm pháp luật
Là hình thức chủ yếu trong hệ thống pháp luật Việt Nam. Ưu điểm là nội dung rõ ràng, phạm vi áp dụng rộng, dễ công bố, sửa đổi và hệ thống hóa; hạn chế là có thể chậm hơn thực tiễn hoặc phát sinh chồng chéo nếu chất lượng xây dựng pháp luật không tốt.

### d) Điều ước quốc tế
Điều ước quốc tế mà Việt Nam là thành viên là căn cứ pháp lý quan trọng trong quan hệ có yếu tố quốc tế. Việc ký kết, thực hiện và áp dụng phải tuân theo Hiến pháp và pháp luật về điều ước quốc tế.

### e) Các nguồn bổ trợ khác
Trong một số lĩnh vực, pháp luật có thể cho phép áp dụng nguyên tắc cơ bản của pháp luật, lẽ công bằng, tương tự pháp luật hoặc thỏa thuận hợp pháp của các bên. Đây là nguồn bổ trợ, không được sử dụng tùy tiện để thay thế quy định bắt buộc.

# Câu 7. Khái niệm Luật dân sự, đối tượng và phương pháp điều chỉnh

## 1. Khái niệm
Luật dân sự là ngành luật độc lập gồm các quy phạm điều chỉnh quan hệ tài sản và quan hệ nhân thân được hình thành trên cơ sở bình đẳng, tự do ý chí, độc lập về tài sản và tự chịu trách nhiệm của các chủ thể.

## 2. Đối tượng điều chỉnh
### a) Quan hệ tài sản
Là quan hệ giữa các chủ thể thông qua tài sản, thường có thể xác định bằng tiền và gắn với sự dịch chuyển lợi ích vật chất. Tài sản theo Bộ luật Dân sự năm 2015 gồm vật, tiền, giấy tờ có giá và quyền tài sản; có thể là bất động sản hoặc động sản, tài sản hiện có hoặc tài sản hình thành trong tương lai.

Các nhóm quan hệ tài sản tiêu biểu gồm sở hữu; nghĩa vụ và hợp đồng; bồi thường thiệt hại; thừa kế; quyền đối với tài sản; quan hệ có yếu tố nước ngoài.

### b) Quan hệ nhân thân
Là quan hệ gắn với giá trị nhân thân của cá nhân hoặc pháp nhân, như tên gọi, hình ảnh, danh dự, nhân phẩm, uy tín, đời sống riêng tư, quyền tác giả.

Quan hệ nhân thân không gắn với tài sản tồn tại độc lập với lợi ích tài sản. Quan hệ nhân thân gắn với tài sản có thể làm phát sinh quyền tài sản, ví dụ quyền nhân thân của tác giả gắn với quyền tài sản trong sở hữu trí tuệ.

Lưu ý: quyền kết hôn, ly hôn chủ yếu thuộc đối tượng điều chỉnh của Luật Hôn nhân và gia đình, không nên đưa làm ví dụ điển hình của Luật dân sự.

## 3. Phương pháp điều chỉnh
### a) Bình đẳng về địa vị pháp lý
Không bên nào có quyền dùng quyền lực công để đơn phương áp đặt bên kia trong quan hệ dân sự, kể cả khi một bên là cơ quan nhà nước nhưng tham gia với tư cách chủ thể dân sự.

### b) Tự do, tự nguyện cam kết và thỏa thuận
Chủ thể được tự do xác lập, thay đổi, chấm dứt quyền và nghĩa vụ, miễn là không vi phạm điều cấm của luật, không trái đạo đức xã hội và không xâm phạm quyền của người khác.

### c) Độc lập về tài sản và tự chịu trách nhiệm
Mỗi chủ thể tự định đoạt tài sản hợp pháp và chịu trách nhiệm bằng tài sản của mình trong phạm vi pháp luật quy định.

### d) Tự bảo vệ và yêu cầu cơ quan có thẩm quyền bảo vệ
Khi quyền dân sự bị xâm phạm, chủ thể có thể tự bảo vệ bằng biện pháp hợp pháp hoặc yêu cầu tòa án, trọng tài và cơ quan có thẩm quyền áp dụng biện pháp bảo vệ.

## Kết luận
Đặc trưng nổi bật của phương pháp dân sự là bình đẳng, thỏa thuận, tự định đoạt và tự chịu trách nhiệm, khác với phương pháp quyền lực - phục tùng của Luật hành chính.

# Câu 8. Khái niệm, đặc điểm và cấu trúc của quy phạm pháp luật

## 1. Khái niệm
Quy phạm pháp luật là quy tắc xử sự chung do Nhà nước ban hành hoặc thừa nhận và bảo đảm thực hiện, nhằm điều chỉnh quan hệ xã hội theo định hướng nhất định.

Quy phạm pháp luật không đồng nhất với “điều luật”. Một điều luật có thể chứa một hoặc nhiều quy phạm; một quy phạm cũng có thể được thể hiện trong nhiều điều, khoản khác nhau.

## 2. Đặc điểm
- Là quy tắc xử sự chung, có tính khuôn mẫu và bắt buộc đối với các chủ thể thuộc phạm vi điều chỉnh.
- Do Nhà nước đặt ra hoặc thừa nhận, thể hiện ý chí nhà nước và nhu cầu quản lý xã hội.
- Được áp dụng nhiều lần cho đến khi hết hiệu lực, bị sửa đổi, thay thế hoặc bãi bỏ.
- Xác định quyền, nghĩa vụ, điều kiện xử sự và hậu quả pháp lý.
- Có hình thức chặt chẽ và được Nhà nước bảo đảm thực hiện.

## 3. Cấu trúc
### a) Giả định
Nêu chủ thể, hoàn cảnh, thời gian, địa điểm hoặc điều kiện để quy phạm được áp dụng. Giả định trả lời: ai, khi nào, trong điều kiện nào?

### b) Quy định
Nêu cách xử sự mà chủ thể được làm, phải làm hoặc không được làm. Đây là bộ phận trung tâm vì trực tiếp thể hiện mệnh lệnh hoặc sự cho phép của Nhà nước. Quy định có thể mang tính bắt buộc, cấm đoán, cho phép hoặc tùy nghi.

### c) Chế tài
Nêu hậu quả pháp lý bất lợi hoặc biện pháp tác động khi chủ thể không thực hiện đúng phần quy định. Chế tài có thể là hình sự, hành chính, dân sự, kỷ luật; có thể nhằm trừng phạt, khôi phục tình trạng ban đầu, bồi thường hoặc phủ nhận giá trị pháp lý của hành vi.

## 4. Quan hệ giữa ba bộ phận
Không phải quy phạm nào cũng trình bày đủ ba bộ phận trong một điều luật. Có bộ phận được thể hiện ngầm hoặc nằm ở điều khoản khác. Khi phân tích phải dựa vào nội dung đầy đủ của quy phạm, không chỉ nhìn hình thức câu chữ.

# Câu 9. Sở hữu, quyền sở hữu; căn cứ xác lập, chấm dứt và nội dung quyền sở hữu

## 1. Khái niệm sở hữu
Sở hữu là quan hệ xã hội giữa người với người về việc chiếm hữu, sử dụng và định đoạt của cải. Đối tượng của quan hệ sở hữu là tài sản, nhưng bản chất của sở hữu là quan hệ xã hội chứ không chỉ là quan hệ giữa người và vật.

## 2. Khái niệm quyền sở hữu
Theo Bộ luật Dân sự năm 2015, quyền sở hữu bao gồm quyền chiếm hữu, quyền sử dụng và quyền định đoạt tài sản của chủ sở hữu theo quy định của luật. Có thể hiểu theo nghĩa khách quan là tổng hợp quy phạm điều chỉnh quan hệ sở hữu; theo nghĩa chủ quan là các quyền năng cụ thể của chủ sở hữu.

## 3. Căn cứ xác lập quyền sở hữu
Quyền sở hữu được xác lập trong các trường hợp chủ yếu:
- Do lao động, hoạt động sản xuất, kinh doanh hợp pháp hoặc hoạt động sáng tạo.
- Được chuyển quyền theo thỏa thuận như mua bán, trao đổi, tặng cho hoặc theo bản án, quyết định của cơ quan có thẩm quyền.
- Thu hoa lợi, lợi tức.
- Tạo thành tài sản mới do sáp nhập, trộn lẫn, chế biến.
- Được thừa kế.
- Chiếm hữu tài sản vô chủ, tài sản không xác định được chủ sở hữu; vật bị chôn, giấu, vùi lấp, chìm đắm; vật bị đánh rơi, bỏ quên; gia súc, gia cầm thất lạc; vật nuôi dưới nước di chuyển tự nhiên theo điều kiện luật định.
- Chiếm hữu, được lợi về tài sản không có căn cứ pháp luật nhưng ngay tình, liên tục, công khai trong thời hạn luật định.
- Trường hợp khác do luật quy định.

## 4. Căn cứ chấm dứt quyền sở hữu
Quyền sở hữu chấm dứt khi:
- Chủ sở hữu chuyển quyền cho người khác.
- Chủ sở hữu từ bỏ quyền sở hữu.
- Tài sản đã được tiêu dùng hoặc bị tiêu hủy.
- Tài sản bị xử lý để thực hiện nghĩa vụ của chủ sở hữu.
- Tài sản bị trưng mua theo luật.
- Tài sản bị tịch thu.
- Người khác được xác lập quyền sở hữu đối với tài sản theo quy định về vật thất lạc, tài sản vô chủ hoặc thời hiệu.
- Trường hợp khác do luật quy định.

## 5. Nội dung quyền sở hữu
### a) Quyền chiếm hữu
Là quyền nắm giữ, chi phối tài sản trực tiếp hoặc gián tiếp. Chủ sở hữu có thể tự chiếm hữu hoặc giao cho người khác thông qua thuê, mượn, gửi giữ, ủy quyền. Cần phân biệt chiếm hữu có căn cứ pháp luật với chiếm hữu không có căn cứ pháp luật; chiếm hữu ngay tình với không ngay tình.

### b) Quyền sử dụng
Là quyền khai thác công dụng, hưởng hoa lợi và lợi tức từ tài sản. Việc sử dụng không được gây thiệt hại, ảnh hưởng đến lợi ích quốc gia, dân tộc, lợi ích công cộng hoặc quyền, lợi ích hợp pháp của người khác.

### c) Quyền định đoạt
Là quyền chuyển giao quyền sở hữu, từ bỏ, tiêu dùng hoặc tiêu hủy tài sản. Hình thức phổ biến gồm bán, trao đổi, tặng cho, để thừa kế, góp vốn. Việc định đoạt phải do người có quyền và năng lực phù hợp thực hiện, tuân thủ điều kiện, hình thức và giới hạn do luật quy định.

## Lưu ý cập nhật
Đề bài dẫn Bộ luật Dân sự năm 2005 đã hết hiệu lực. Khi ôn thi hiện nay nên sử dụng Bộ luật Dân sự năm 2015, đặc biệt các quy định về nội dung quyền sở hữu, căn cứ xác lập và chấm dứt quyền sở hữu.

# Câu 10. Khái niệm, dấu hiệu và các loại vi phạm pháp luật

## 1. Khái niệm
Vi phạm pháp luật là hành vi trái pháp luật, có lỗi, do chủ thể có năng lực trách nhiệm pháp lý thực hiện, xâm hại hoặc đe dọa xâm hại các quan hệ xã hội được pháp luật bảo vệ.

Không phải mọi thiệt hại đều là vi phạm pháp luật; cũng không phải mọi vi phạm đều phải có thiệt hại vật chất thực tế. Nhiều hành vi chỉ cần tạo nguy cơ hoặc xâm phạm trật tự pháp luật đã đủ cấu thành.

## 2. Các dấu hiệu
### a) Là hành vi xác định của con người
Hành vi có thể là hành động hoặc không hành động. Suy nghĩ, ý định chưa biểu hiện thành hành vi thì thông thường không bị coi là vi phạm.

### b) Có tính trái pháp luật
Hành vi không thực hiện điều pháp luật bắt buộc, thực hiện điều pháp luật cấm, vượt quá giới hạn được phép hoặc sử dụng quyền trái mục đích.

### c) Có lỗi
Lỗi là thái độ tâm lý của chủ thể đối với hành vi và hậu quả, dưới hình thức cố ý hoặc vô ý. Trường hợp sự kiện bất khả kháng, tình thế cấp thiết, phòng vệ chính đáng hoặc chủ thể không có khả năng nhận thức có thể loại trừ tính vi phạm tùy quy định.

### d) Do chủ thể có năng lực trách nhiệm pháp lý thực hiện
Chủ thể phải đạt độ tuổi, có khả năng nhận thức, điều khiển hành vi và đáp ứng điều kiện của ngành luật tương ứng. Tổ chức, pháp nhân cũng có thể chịu trách nhiệm trong các trường hợp luật quy định.

### e) Xâm hại quan hệ xã hội được pháp luật bảo vệ
Hành vi gây thiệt hại hoặc đe dọa gây thiệt hại cho Nhà nước, tổ chức, cá nhân hoặc trật tự pháp luật.

## 3. Các loại
### a) Vi phạm hình sự
Là tội phạm được quy định trong Bộ luật Hình sự, có mức độ nguy hiểm cao nhất và bị áp dụng hình phạt hoặc biện pháp hình sự.

### b) Vi phạm hành chính
Là hành vi có lỗi, vi phạm quy định quản lý nhà nước, không phải tội phạm và theo luật phải bị xử phạt hành chính.

### c) Vi phạm dân sự
Là hành vi xâm phạm quan hệ tài sản hoặc nhân thân, vi phạm nghĩa vụ dân sự, gây thiệt hại ngoài hợp đồng hoặc xâm phạm quyền dân sự. Hậu quả thường là buộc thực hiện nghĩa vụ, bồi thường, hoàn trả, chấm dứt hành vi.

### d) Vi phạm kỷ luật
Là hành vi vi phạm kỷ luật lao động, công vụ, học tập hoặc quy tắc nội bộ hợp pháp; chủ thể bị áp dụng trách nhiệm kỷ luật theo quan hệ quản lý tương ứng.

# Câu 11. Khái niệm thừa kế và nội dung thừa kế theo di chúc

## 1. Khái niệm
Thừa kế là việc chuyển dịch tài sản của người chết cho người còn sống hoặc tổ chức theo di chúc hoặc theo pháp luật. Di sản gồm tài sản riêng của người chết và phần tài sản của người chết trong tài sản chung.

## 2. Di chúc và thừa kế theo di chúc
Di chúc là sự thể hiện ý chí của cá nhân nhằm chuyển tài sản của mình cho người khác sau khi chết. Thừa kế theo di chúc là việc dịch chuyển di sản theo sự định đoạt hợp pháp của người để lại di sản.

## 3. Người lập di chúc và quyền của họ
Người thành niên minh mẫn, sáng suốt có quyền lập di chúc. Người từ đủ 15 đến dưới 18 tuổi được lập di chúc bằng văn bản nếu cha, mẹ hoặc người giám hộ đồng ý về việc lập.

Người lập di chúc có quyền chỉ định người thừa kế; truất quyền; phân định phần di sản; dành tài sản để di tặng, thờ cúng; giao nghĩa vụ; chỉ định người giữ di chúc, quản lý và phân chia di sản; sửa đổi, bổ sung, thay thế hoặc hủy bỏ di chúc.

## 4. Điều kiện di chúc hợp pháp
- Người lập minh mẫn, sáng suốt; không bị lừa dối, đe dọa hoặc cưỡng ép.
- Nội dung không vi phạm điều cấm của luật, không trái đạo đức xã hội.
- Hình thức phù hợp quy định pháp luật.
- Di chúc của người bị hạn chế về thể chất hoặc không biết chữ phải được người làm chứng lập thành văn bản và có công chứng hoặc chứng thực.

Ý trong bản cũ nói cha, mẹ hoặc người giám hộ phải “ký” vào di chúc của người từ đủ 15 đến dưới 18 tuổi là chưa chuẩn; pháp luật yêu cầu họ đồng ý về việc lập di chúc.

## 5. Hình thức
Di chúc bằng văn bản có thể không có người làm chứng, có người làm chứng, có công chứng hoặc chứng thực. Di chúc miệng chỉ được lập khi tính mạng bị cái chết đe dọa và không thể lập bằng văn bản; phải thể hiện ý chí cuối cùng trước ít nhất hai người làm chứng, được ghi chép và xác nhận theo thời hạn luật định. Sau ba tháng mà người lập còn sống, minh mẫn, sáng suốt thì di chúc miệng bị hủy bỏ.

## 6. Hiệu lực
Di chúc có hiệu lực từ thời điểm mở thừa kế, tức thời điểm người có tài sản chết. Di chúc có thể không có hiệu lực toàn bộ hoặc một phần nếu người thừa kế chết trước hoặc cùng thời điểm, cơ quan/tổ chức không còn tồn tại, di sản không còn, hoặc nội dung trái luật.

## 7. Người thừa kế không phụ thuộc nội dung di chúc
Con chưa thành niên; cha, mẹ, vợ, chồng; con thành niên không có khả năng lao động vẫn được hưởng phần di sản bằng hai phần ba suất của một người thừa kế theo pháp luật nếu họ không được cho hưởng hoặc được cho hưởng ít hơn, trừ trường hợp từ chối hoặc không có quyền hưởng.

## 8. Thanh toán và phân chia
Trước khi chia di sản phải thanh toán nghĩa vụ tài sản và chi phí theo thứ tự luật định. Phần còn lại mới được chia theo di chúc; nếu di chúc không định đoạt hết hoặc phần di chúc vô hiệu thì phần đó được chia theo pháp luật.

# Câu 12. Khái niệm Luật lao động, đối tượng và phương pháp điều chỉnh

## 1. Khái niệm
Luật lao động là ngành luật gồm các quy phạm điều chỉnh quan hệ lao động giữa người lao động làm việc có trả công, tiền lương và người sử dụng lao động, cùng các quan hệ liên quan trực tiếp đến quan hệ lao động.

## 2. Đối tượng điều chỉnh
### a) Quan hệ lao động cá nhân
Phát sinh từ việc tuyển dụng, giao kết và thực hiện hợp đồng lao động; việc làm, tiền lương, thời giờ làm việc và nghỉ ngơi; an toàn, vệ sinh lao động; kỷ luật, trách nhiệm vật chất; sửa đổi, tạm hoãn, chấm dứt hợp đồng.

### b) Quan hệ lao động tập thể
Gồm đối thoại tại nơi làm việc, thương lượng tập thể, thỏa ước lao động tập thể, quan hệ giữa tổ chức đại diện người lao động và người sử dụng lao động, giải quyết tranh chấp tập thể và đình công.

### c) Quan hệ liên quan trực tiếp
Bao gồm học nghề, đào tạo nghề; việc làm; bảo hiểm xã hội và bảo hiểm thất nghiệp trong mối liên hệ lao động; quản lý nhà nước, thanh tra lao động; giải quyết tranh chấp; đại diện và bảo vệ người lao động.

Quan hệ công vụ của cán bộ, công chức chủ yếu do pháp luật cán bộ, công chức và hành chính điều chỉnh. Quan hệ của thành viên hợp tác xã lao động cho chính hợp tác xã còn phụ thuộc tư cách thành viên và pháp luật hợp tác xã. Cần phân biệt với người được thuê theo hợp đồng lao động.

## 3. Phương pháp điều chỉnh
### a) Thỏa thuận
Các bên tự nguyện thỏa thuận về công việc, tiền lương, điều kiện lao động nhưng không được thấp hơn tiêu chuẩn tối thiểu hoặc trái quy định bắt buộc.

### b) Mệnh lệnh trong phạm vi quan hệ lao động
Người sử dụng lao động có quyền tổ chức, điều hành và kiểm tra lao động; người lao động phải tuân thủ sự điều hành hợp pháp, nội quy và kỷ luật. Quyền này bị giới hạn bởi pháp luật và hợp đồng.

### c) Bảo vệ người lao động
Do người lao động thường ở vị thế phụ thuộc về tổ chức và kinh tế, pháp luật đặt ra tiêu chuẩn tối thiểu, cơ chế đại diện, đối thoại, thương lượng, thanh tra và giải quyết tranh chấp.

### d) Sự tham gia của tổ chức đại diện
Tổ chức đại diện người lao động tham gia bảo vệ quyền lợi, đối thoại, thương lượng và giải quyết tranh chấp; tổ chức đại diện người sử dụng lao động tham gia trong các cơ chế ba bên và quan hệ lao động.

# Câu 13 (bổ sung do bản gốc bị thiếu). Hợp đồng lao động

## 1. Khái niệm
Hợp đồng lao động là sự thỏa thuận giữa người lao động và người sử dụng lao động về việc làm có trả công, tiền lương, điều kiện lao động, quyền và nghĩa vụ của mỗi bên. Thỏa thuận dù có tên gọi khác nhưng có nội dung về việc làm có trả công và sự quản lý, điều hành, giám sát của một bên vẫn có thể được coi là hợp đồng lao động.

## 2. Nguyên tắc giao kết
Các bên phải tự nguyện, bình đẳng, thiện chí, hợp tác, trung thực; được tự do giao kết nhưng không được trái pháp luật, thỏa ước lao động tập thể và đạo đức xã hội.

## 3. Hình thức
Hợp đồng về nguyên tắc phải lập bằng văn bản và giao cho mỗi bên một bản. Hợp đồng điện tử dưới hình thức thông điệp dữ liệu có giá trị như văn bản. Hai bên có thể giao kết bằng lời nói đối với hợp đồng dưới một tháng, trừ các trường hợp pháp luật bắt buộc bằng văn bản.

## 4. Các loại hợp đồng
Pháp luật hiện hành có hai loại:
- Hợp đồng không xác định thời hạn.
- Hợp đồng xác định thời hạn không quá 36 tháng.

Quy định cũ về hợp đồng theo mùa vụ hoặc công việc dưới 12 tháng đã được thay đổi. Khi hợp đồng xác định thời hạn hết hạn mà người lao động tiếp tục làm việc, các bên phải xử lý và ký mới trong thời hạn luật định; số lần ký hợp đồng xác định thời hạn bị giới hạn, trừ trường hợp đặc biệt.

## 5. Nội dung chủ yếu
Hợp đồng thường có tên, địa chỉ của người sử dụng lao động; thông tin người lao động; công việc và địa điểm; thời hạn; mức lương, hình thức và kỳ hạn trả lương, phụ cấp và khoản bổ sung; chế độ nâng lương; thời giờ làm việc, nghỉ ngơi; bảo hộ lao động; bảo hiểm; đào tạo.

## 6. Thực hiện, sửa đổi, tạm hoãn và chấm dứt
Các bên phải thực hiện đúng cam kết. Việc sửa đổi cần thỏa thuận; khi có căn cứ luật định có thể tạm hoãn. Hợp đồng chấm dứt do hết hạn, hoàn thành công việc, thỏa thuận, người lao động hoặc người sử dụng lao động đơn phương hợp pháp, sa thải, thay đổi cơ cấu hoặc các căn cứ khác.

Đơn phương chấm dứt phải có căn cứ và thời hạn báo trước theo luật, trừ trường hợp được miễn báo trước. Bên chấm dứt trái pháp luật phải chịu hậu quả như nhận người lao động trở lại, trả lương, đóng bảo hiểm, bồi thường hoặc bồi thường chi phí đào tạo tùy trường hợp.

## 7. Ý nghĩa
Hợp đồng lao động là căn cứ trung tâm xác lập quan hệ lao động, bảo vệ quyền lợi hai bên và là chứng cứ quan trọng khi giải quyết tranh chấp.

# Câu 14. Tranh chấp lao động, thẩm quyền và thủ tục giải quyết

## 1. Khái niệm và phân loại
Tranh chấp lao động là tranh chấp về quyền, nghĩa vụ và lợi ích phát sinh giữa các bên trong quá trình xác lập, thực hiện hoặc chấm dứt quan hệ lao động; tranh chấp giữa các tổ chức đại diện người lao động; hoặc tranh chấp từ quan hệ liên quan trực tiếp.

Tranh chấp gồm tranh chấp cá nhân và tranh chấp tập thể. Tranh chấp tập thể được chia thành tranh chấp về quyền và tranh chấp về lợi ích.

Tranh chấp tập thể về quyền phát sinh khi có khác nhau trong hiểu, thực hiện pháp luật, thỏa ước, nội quy hoặc thỏa thuận hợp pháp. Tranh chấp tập thể về lợi ích phát sinh trong thương lượng nhằm xác lập điều kiện lao động mới.

## 2. Nguyên tắc giải quyết
Tôn trọng quyền tự định đoạt; ưu tiên thương lượng, hòa giải; công khai, minh bạch, khách quan, kịp thời; có sự tham gia của đại diện các bên; chỉ cơ quan có thẩm quyền can thiệp khi được yêu cầu hoặc pháp luật quy định.

## 3. Chủ thể có thẩm quyền
- Hòa giải viên lao động.
- Hội đồng trọng tài lao động.
- Tòa án nhân dân.

Bản gốc nêu Hội đồng hòa giải cơ sở và Chủ tịch UBND cấp huyện là cơ quan giải quyết theo cơ chế cũ. Bộ luật Lao động năm 2019 đã thay đổi hệ thống này.

## 4. Tranh chấp lao động cá nhân
Thông thường, tranh chấp phải qua hòa giải viên lao động trước khi yêu cầu trọng tài hoặc tòa án. Một số tranh chấp được miễn hòa giải bắt buộc, như xử lý kỷ luật sa thải hoặc đơn phương chấm dứt hợp đồng; bồi thường, trợ cấp khi chấm dứt; giữa người giúp việc gia đình và người sử dụng; bảo hiểm; bồi thường giữa người lao động với doanh nghiệp đưa đi làm việc ở nước ngoài; cho thuê lại lao động.

Nếu hòa giải thành, các bên thực hiện biên bản hòa giải. Nếu không thành, hết thời hạn mà không hòa giải hoặc một bên không thực hiện thỏa thuận, các bên có thể yêu cầu Hội đồng trọng tài lao động hoặc Tòa án theo điều kiện luật định. Khi đã chọn trọng tài thì trong thời gian giải quyết không đồng thời yêu cầu Tòa án, trừ trường hợp luật cho phép.

## 5. Tranh chấp tập thể về quyền
Trước hết giải quyết qua hòa giải viên lao động. Nếu không thành hoặc hết thời hạn, các bên có thể lựa chọn Hội đồng trọng tài lao động hoặc Tòa án. Cơ quan có thẩm quyền có thể chuyển hồ sơ xử lý vi phạm nếu phát hiện hành vi vi phạm pháp luật.

## 6. Tranh chấp tập thể về lợi ích
Trước hết phải hòa giải. Nếu không thành hoặc hết thời hạn, các bên có thể yêu cầu Hội đồng trọng tài. Tổ chức đại diện người lao động có quyền tiến hành thủ tục lấy ý kiến và đình công khi đáp ứng điều kiện luật định, như hòa giải không thành, hết thời hạn hòa giải, trọng tài không được thành lập/không giải quyết hoặc một bên không thực hiện quyết định.

## 7. Lưu ý về đình công
Đình công chỉ được tiến hành đối với tranh chấp tập thể về lợi ích, do tổ chức đại diện người lao động có quyền thương lượng tập thể tổ chức và lãnh đạo, sau khi tuân thủ thủ tục. Không phải mọi tranh chấp tập thể đều được đình công.

# Câu 15. Khái niệm, đối tượng và phương pháp điều chỉnh của Luật hành chính

## 1. Khái niệm
Luật hành chính là ngành luật gồm các quy phạm điều chỉnh quan hệ xã hội phát sinh trong hoạt động quản lý hành chính nhà nước, tổ chức thực hiện quyền hành pháp và một số hoạt động quản lý nội bộ hoặc quản lý được Nhà nước trao quyền.

## 2. Đối tượng điều chỉnh
### a) Quan hệ trong hoạt động quản lý của cơ quan hành chính
Phát sinh khi Chính phủ, bộ, cơ quan ngang bộ, UBND và cơ quan chuyên môn thực hiện quản lý kinh tế, văn hóa, xã hội, an ninh, trật tự và các lĩnh vực khác.

### b) Quan hệ quản lý nội bộ
Phát sinh trong tổ chức bộ máy, phân công công tác, quản lý cán bộ, tài sản, kỷ luật và quy trình hoạt động nội bộ của cơ quan nhà nước.

### c) Quan hệ quản lý trong cơ quan lập pháp, tư pháp
Quốc hội, Hội đồng nhân dân, Tòa án, Viện kiểm sát khi thực hiện công tác hành chính nội bộ cũng làm phát sinh quan hệ mang tính quản lý, khác với hoạt động lập pháp, xét xử hoặc kiểm sát chuyên môn.

### d) Quan hệ do chủ thể được trao quyền quản lý
Một số tổ chức hoặc cá nhân không phải cơ quan hành chính có thể được pháp luật trao quyền thực hiện nhiệm vụ quản lý nhà nước cụ thể.

## 3. Phương pháp điều chỉnh
Phương pháp chủ yếu là quyền lực - phục tùng hoặc mệnh lệnh đơn phương. Một bên nhân danh Nhà nước có quyền ban hành quyết định, yêu cầu, kiểm tra và cưỡng chế; bên kia có nghĩa vụ chấp hành quyết định hợp pháp.

Tuy nhiên, không phải mọi quan hệ hành chính đều chỉ có mệnh lệnh cứng nhắc. Pháp luật hiện đại còn sử dụng phối hợp, tham vấn, cung cấp dịch vụ công, thỏa thuận hành chính trong phạm vi luật cho phép. Dù vậy, dấu hiệu quyền lực công vẫn là đặc trưng để phân biệt với quan hệ dân sự.

## 4. Đặc điểm của quyết định quản lý
Quyết định phải đúng thẩm quyền, căn cứ, trình tự, mục đích và hình thức; phải tôn trọng quyền, lợi ích hợp pháp. Cá nhân, tổ chức có quyền khiếu nại, khởi kiện hoặc yêu cầu xem xét khi quyết định, hành vi hành chính trái pháp luật.

# Câu 16. Khái niệm, đặc điểm và phân loại tội phạm

## 1. Khái niệm
Theo Bộ luật Hình sự năm 2015, được sửa đổi, bổ sung năm 2017, tội phạm là hành vi nguy hiểm cho xã hội được quy định trong Bộ luật Hình sự, do người có năng lực trách nhiệm hình sự hoặc pháp nhân thương mại thực hiện một cách cố ý hoặc vô ý, xâm phạm các quan hệ xã hội được luật hình sự bảo vệ và theo quy định phải bị xử lý hình sự.

Đề cương cũ trích Điều 8 Bộ luật Hình sự năm 1999 đã hết hiệu lực nên cần cập nhật.

## 2. Đặc điểm
### a) Tính nguy hiểm cho xã hội
Hành vi gây ra hoặc đe dọa gây ra thiệt hại đáng kể cho độc lập, chủ quyền, chế độ, quyền con người, quyền công dân, trật tự, an toàn xã hội và các lợi ích được luật hình sự bảo vệ. Đây là dấu hiệu nội dung quan trọng nhất.

Mức độ nguy hiểm được đánh giá qua khách thể bị xâm hại; tính chất, mức độ hành vi và hậu quả; lỗi, động cơ, mục đích; phương pháp, công cụ; nhân thân; hoàn cảnh phạm tội.

### b) Tính được quy định trong Bộ luật Hình sự
Không có tội nếu hành vi không được Bộ luật Hình sự quy định. Hành vi tuy có hại nhưng mức độ không đáng kể thì không phải tội phạm và có thể bị xử lý bằng biện pháp khác.

### c) Tính có lỗi
Chủ thể có thái độ cố ý hoặc vô ý. Cố ý trực tiếp là nhận thức rõ, thấy trước và mong muốn hậu quả; cố ý gián tiếp là thấy trước, không mong muốn nhưng có ý thức để mặc. Vô ý do quá tự tin là thấy trước nhưng tin hậu quả không xảy ra hoặc ngăn được; vô ý do cẩu thả là không thấy trước dù phải và có thể thấy trước.

### d) Do chủ thể có năng lực trách nhiệm hình sự thực hiện
Cá nhân phải đáp ứng điều kiện về tuổi và khả năng nhận thức, điều khiển hành vi. Người từ đủ 16 tuổi chịu trách nhiệm về mọi tội, trừ tội mà luật có quy định khác; người từ đủ 14 đến dưới 16 tuổi chỉ chịu trách nhiệm về các tội rất nghiêm trọng do cố ý hoặc đặc biệt nghiêm trọng được liệt kê trong luật.

Pháp nhân thương mại chỉ chịu trách nhiệm về những tội danh và khi đủ điều kiện luật định.

### e) Tính phải chịu hình phạt
Tội phạm gắn với khả năng bị áp dụng hình phạt. Tuy nhiên, người phạm tội có thể được miễn trách nhiệm hoặc miễn hình phạt trong trường hợp luật định; vì vậy nên hiểu đây là “tính chịu sự đe dọa của hình phạt”.

## 3. Phân loại
- Tội phạm ít nghiêm trọng: mức cao nhất của khung hình phạt là phạt tiền, cải tạo không giam giữ hoặc tù đến 3 năm.
- Tội phạm nghiêm trọng: mức cao nhất trên 3 năm đến 7 năm tù.
- Tội phạm rất nghiêm trọng: mức cao nhất trên 7 năm đến 15 năm tù.
- Tội phạm đặc biệt nghiêm trọng: mức cao nhất trên 15 năm đến 20 năm tù, tù chung thân hoặc tử hình.

Việc phân loại ảnh hưởng đến tuổi chịu trách nhiệm, thời hiệu, chuẩn bị phạm tội, thời hạn điều tra, truy tố, xét xử và nhiều hậu quả pháp lý khác.

# Câu 17. Khái niệm, đặc điểm, thẩm quyền và thủ tục giải quyết tố cáo

## 1. Khái niệm
Theo Luật Tố cáo năm 2018, tố cáo là việc cá nhân theo thủ tục luật định báo cho cơ quan, tổ chức, cá nhân có thẩm quyền biết về hành vi vi phạm pháp luật gây thiệt hại hoặc đe dọa gây thiệt hại đến lợi ích Nhà nước, quyền và lợi ích hợp pháp của cơ quan, tổ chức, cá nhân.

Tố cáo khác khiếu nại: khiếu nại nhằm bảo vệ trực tiếp quyền, lợi ích của người khiếu nại trước quyết định hoặc hành vi liên quan đến họ; tố cáo nhằm phát hiện, xử lý hành vi vi phạm, dù người tố cáo có thể đồng thời bị ảnh hưởng.

## 2. Đặc điểm
- Chủ thể thực hiện quyền tố cáo là cá nhân.
- Đối tượng là hành vi vi phạm trong thực hiện nhiệm vụ, công vụ hoặc hành vi vi phạm pháp luật về quản lý nhà nước trong các lĩnh vực.
- Nội dung có thể liên quan lợi ích công hoặc quyền, lợi ích của cơ quan, tổ chức, cá nhân.
- Người tố cáo phải trung thực, chịu trách nhiệm về nội dung; được giữ bí mật thông tin, bảo vệ và khen thưởng theo điều kiện luật định.
- Nghiêm cấm tố cáo sai sự thật, lợi dụng tố cáo để xúc phạm, gây rối hoặc trục lợi; đồng thời nghiêm cấm trả thù, trù dập người tố cáo.

## 3. Thẩm quyền
Nguyên tắc chung: người đứng đầu cơ quan, tổ chức có thẩm quyền quản lý cán bộ, công chức, viên chức giải quyết tố cáo hành vi vi phạm trong nhiệm vụ, công vụ của người thuộc quyền quản lý. Tố cáo người đứng đầu, cấp phó hoặc người do cấp trên bổ nhiệm thường thuộc thẩm quyền người đứng đầu cơ quan cấp trên trực tiếp hoặc chủ thể có quyền bổ nhiệm.

Đối với hành vi vi phạm pháp luật về quản lý nhà nước trong lĩnh vực, cơ quan có chức năng quản lý nhà nước lĩnh vực đó giải quyết. Tố cáo có dấu hiệu tội phạm phải được chuyển ngay cho cơ quan tiến hành tố tụng có thẩm quyền.

## 4. Hình thức và tiếp nhận
Tố cáo được thực hiện bằng đơn hoặc trình bày trực tiếp. Đơn phải có ngày tháng, họ tên, địa chỉ, cách liên hệ, nội dung, thông tin người bị tố cáo và chữ ký/điểm chỉ.

Thông tin không rõ họ tên, địa chỉ hoặc không xác định được người tố cáo thì không xử lý theo thủ tục giải quyết tố cáo thông thường. Tuy nhiên, nếu thông tin có nội dung cụ thể, tài liệu, chứng cứ rõ ràng và có cơ sở xác minh thì cơ quan tiếp nhận có thể thanh tra, kiểm tra hoặc chuyển cơ quan có thẩm quyền. Vì vậy, nhận định “tố cáo nặc danh là bất hợp pháp” trong bản cũ là không chính xác.

## 5. Trình tự giải quyết
### Bước 1: Thụ lý
Kiểm tra điều kiện, xác định thẩm quyền; ban hành quyết định thụ lý và thông báo theo quy định. Nếu không thuộc thẩm quyền thì chuyển đến cơ quan có thẩm quyền.

### Bước 2: Xác minh
Thu thập tài liệu, làm việc với người tố cáo, người bị tố cáo và chủ thể liên quan; áp dụng biện pháp bảo vệ, ngăn chặn khi cần thiết. Việc xác minh phải khách quan và bảo mật.

### Bước 3: Kết luận nội dung tố cáo
Người giải quyết kết luận từng nội dung đúng, sai hoặc đúng một phần; xác định vi phạm, thiệt hại, trách nhiệm và biện pháp xử lý.

### Bước 4: Xử lý kết luận
Áp dụng theo thẩm quyền hoặc kiến nghị xử lý kỷ luật, hành chính, thu hồi tài sản; nếu có dấu hiệu tội phạm thì chuyển hồ sơ; thông báo kết quả theo quy định.

## 6. Thời hạn
Thời hạn thông thường không quá 30 ngày kể từ ngày thụ lý. Vụ việc phức tạp có thể gia hạn một lần không quá 30 ngày; đặc biệt phức tạp có thể gia hạn hai lần, mỗi lần không quá 30 ngày. Mốc 60/90 ngày cố định trong bản cũ là quy định theo pháp luật trước đây.

# Câu 18. Luật Hôn nhân và gia đình, các nguyên tắc cơ bản và tính ưu việt

## 1. Khái niệm
Luật Hôn nhân và gia đình là ngành luật gồm các quy phạm điều chỉnh quan hệ nhân thân và tài sản phát sinh từ hôn nhân, quan hệ giữa vợ chồng, cha mẹ và con, các thành viên gia đình; việc xác lập, thực hiện, chấm dứt quan hệ và xử lý vi phạm.

Đề cương cũ dựa vào Luật năm 2000. Khi ôn hiện nay cần cập nhật Luật Hôn nhân và gia đình năm 2014.

## 2. Các nguyên tắc cơ bản
### a) Hôn nhân tự nguyện, tiến bộ, một vợ một chồng, vợ chồng bình đẳng
Nam, nữ tự quyết định kết hôn và ly hôn theo điều kiện luật định; nghiêm cấm cưỡng ép, lừa dối, cản trở kết hôn; người đang có vợ hoặc chồng không được kết hôn/chung sống như vợ chồng với người khác.

### b) Không phân biệt dân tộc, tôn giáo và yếu tố nước ngoài
Hôn nhân giữa công dân Việt Nam thuộc các dân tộc, tôn giáo; giữa người theo và không theo tôn giáo; giữa người Việt Nam và người nước ngoài được tôn trọng và bảo vệ nếu hợp pháp.

### c) Xây dựng gia đình ấm no, tiến bộ, hạnh phúc
Các thành viên có nghĩa vụ tôn trọng, quan tâm, chăm sóc, giúp đỡ nhau; không phân biệt đối xử giữa các con.

### d) Bảo vệ trẻ em, người cao tuổi, người khuyết tật và người yếu thế
Cha mẹ có nghĩa vụ nuôi dưỡng, giáo dục con; con có nghĩa vụ kính trọng, chăm sóc cha mẹ; Nhà nước và xã hội bảo vệ bà mẹ, trẻ em và hỗ trợ thực hiện chức năng gia đình.

### e) Kế thừa và phát huy truyền thống tốt đẹp
Pháp luật tôn trọng phong tục tốt đẹp nhưng loại bỏ tập quán lạc hậu, bạo lực, bất bình đẳng, tảo hôn và các hành vi xâm phạm quyền con người.

### f) Trách nhiệm của Nhà nước, xã hội và gia đình
Các chủ thể có trách nhiệm bảo vệ chế độ hôn nhân và gia đình; thực hiện chính sách dân số phù hợp; hòa giải mâu thuẫn nhưng không được dùng hòa giải để che giấu bạo lực hoặc vi phạm.

## 3. Tính ưu việt
- Đề cao tự do kết hôn, ly hôn có kiểm soát pháp lý và chống ép buộc.
- Xác lập bình đẳng thực chất giữa vợ và chồng về nhân thân, tài sản, lao động và quyết định gia đình.
- Bảo vệ trẻ em, không phân biệt con trai/con gái, con đẻ/con nuôi, con trong hay ngoài hôn nhân.
- Kết hợp giá trị truyền thống về tình nghĩa, hiếu thảo, tương trợ với chuẩn mực hiện đại về quyền con người.
- Bảo vệ người yếu thế và xử lý bạo lực gia đình, tảo hôn, cưỡng ép, mua bán người và lợi dụng hôn nhân.
- Thừa nhận và bảo vệ quan hệ hôn nhân có yếu tố nước ngoài trên cơ sở chủ quyền, bình đẳng và lợi ích hợp pháp.

## Kết luận
Chế độ hôn nhân và gia đình Việt Nam hướng đến gia đình dân chủ, bình đẳng, bền vững, vừa bảo tồn giá trị văn hóa tốt đẹp vừa bảo đảm quyền của từng thành viên.

# Câu 19. Nội dung cơ bản của Hiến pháp nước Cộng hòa xã hội chủ nghĩa Việt Nam năm 2013

## 1. Khái quát
Hiến pháp năm 2013 được Quốc hội thông qua ngày 28/11/2013, có hiệu lực từ ngày 01/01/2014, gồm Lời nói đầu, 11 chương và 120 điều. Đây là luật cơ bản, có hiệu lực pháp lý cao nhất; mọi văn bản khác phải phù hợp với Hiến pháp.

## 2. Nội dung theo các chương
### Chương I - Chế độ chính trị
Khẳng định Việt Nam là nước độc lập, có chủ quyền, thống nhất và toàn vẹn lãnh thổ; là Nhà nước pháp quyền xã hội chủ nghĩa của Nhân dân, do Nhân dân, vì Nhân dân. Quy định quyền lực thuộc về Nhân dân; sự lãnh đạo của Đảng Cộng sản Việt Nam; vai trò Mặt trận Tổ quốc; nguyên tắc tổ chức quyền lực; chính sách dân tộc, đối ngoại; Quốc kỳ, Quốc huy, Quốc ca, Quốc khánh và Thủ đô.

Điểm nổi bật là bổ sung yêu cầu kiểm soát quyền lực bên cạnh phân công và phối hợp; nhấn mạnh trách nhiệm của Đảng trước Nhân dân và việc tổ chức đảng, đảng viên hoạt động trong khuôn khổ Hiến pháp, pháp luật.

### Chương II - Quyền con người, quyền và nghĩa vụ cơ bản của công dân
Hiến pháp phân biệt quyền con người và quyền công dân; khẳng định các quyền được công nhận, tôn trọng, bảo vệ và bảo đảm. Quyền chỉ có thể bị hạn chế theo luật trong trường hợp cần thiết vì quốc phòng, an ninh quốc gia, trật tự, an toàn xã hội, đạo đức xã hội và sức khỏe cộng đồng.

Hiến pháp ghi nhận các quyền dân sự, chính trị, kinh tế, văn hóa, xã hội như quyền sống; bất khả xâm phạm thân thể, đời tư; tự do tín ngưỡng, ngôn luận; quyền bầu cử; sở hữu; kinh doanh; an sinh xã hội; học tập; hưởng thụ văn hóa; sống trong môi trường trong lành. Đồng thời quy định nghĩa vụ trung thành với Tổ quốc, tuân theo Hiến pháp, nộp thuế, bảo vệ Tổ quốc.

### Chương III - Kinh tế, xã hội, văn hóa, giáo dục, khoa học, công nghệ và môi trường
Quy định nền kinh tế thị trường định hướng xã hội chủ nghĩa với nhiều hình thức sở hữu, nhiều thành phần kinh tế; các chủ thể bình đẳng, hợp tác và cạnh tranh theo pháp luật; kinh tế nhà nước giữ vai trò chủ đạo.

Nhà nước quản lý, sử dụng hiệu quả tài nguyên, ngân sách và tài sản công; thực hiện chính sách an sinh, chăm sóc sức khỏe, giáo dục là quốc sách hàng đầu, phát triển khoa học - công nghệ và bảo vệ môi trường.

### Chương IV - Bảo vệ Tổ quốc
Bảo vệ Tổ quốc là sự nghiệp của toàn dân; lực lượng vũ trang tuyệt đối trung thành với Tổ quốc, Nhân dân, Đảng và Nhà nước. Nhà nước củng cố quốc phòng toàn dân, an ninh nhân dân, kết hợp quốc phòng, an ninh với kinh tế và đối ngoại.

### Chương V - Quốc hội
Quốc hội là cơ quan đại biểu cao nhất của Nhân dân, cơ quan quyền lực nhà nước cao nhất; thực hiện quyền lập hiến, lập pháp, quyết định vấn đề quan trọng và giám sát tối cao. Chương này quy định nhiệm vụ, quyền hạn, cơ cấu, đại biểu Quốc hội, Ủy ban thường vụ Quốc hội và các cơ quan của Quốc hội.

### Chương VI - Chủ tịch nước
Chủ tịch nước là người đứng đầu Nhà nước, thay mặt nước về đối nội và đối ngoại; công bố Hiến pháp, luật; thống lĩnh lực lượng vũ trang; thực hiện nhiệm vụ về nhân sự, quốc phòng, an ninh, đối ngoại theo Hiến pháp.

### Chương VII - Chính phủ
Chính phủ là cơ quan hành chính nhà nước cao nhất, thực hiện quyền hành pháp, là cơ quan chấp hành của Quốc hội. Chính phủ thống nhất quản lý các lĩnh vực và chịu trách nhiệm trước Quốc hội.

Điểm mới quan trọng là Hiến pháp chính thức xác định Chính phủ “thực hiện quyền hành pháp”, làm rõ hơn sự phân công quyền lực.

### Chương VIII - Tòa án nhân dân và Viện kiểm sát nhân dân
Tòa án là cơ quan xét xử, thực hiện quyền tư pháp, có nhiệm vụ bảo vệ công lý, quyền con người, quyền công dân, chế độ và lợi ích hợp pháp. Thẩm phán, hội thẩm xét xử độc lập và chỉ tuân theo pháp luật.

Viện kiểm sát thực hành quyền công tố, kiểm sát hoạt động tư pháp, góp phần bảo đảm pháp luật được chấp hành nghiêm chỉnh, thống nhất.

### Chương IX - Chính quyền địa phương
Quy định đơn vị hành chính, tổ chức chính quyền phù hợp đặc điểm nông thôn, đô thị, hải đảo và đơn vị hành chính - kinh tế đặc biệt; xác định Hội đồng nhân dân, Ủy ban nhân dân và nguyên tắc phân định thẩm quyền.

### Chương X - Hội đồng bầu cử quốc gia, Kiểm toán nhà nước
Hội đồng bầu cử quốc gia tổ chức bầu cử đại biểu Quốc hội, chỉ đạo và hướng dẫn bầu cử Hội đồng nhân dân. Kiểm toán nhà nước kiểm toán việc quản lý, sử dụng tài chính, tài sản công. Đây là hai thiết chế hiến định độc lập được quy định thành chương riêng.

### Chương XI - Hiệu lực và sửa đổi Hiến pháp
Khẳng định hiệu lực tối cao của Hiến pháp, trách nhiệm bảo vệ Hiến pháp; quy định chủ thể đề nghị sửa đổi và yêu cầu ít nhất hai phần ba tổng số đại biểu Quốc hội biểu quyết tán thành. Quốc hội có thể quyết định trưng cầu ý dân về Hiến pháp.

## 3. Những giá trị nổi bật
- Khẳng định mạnh mẽ chủ quyền Nhân dân và Nhà nước pháp quyền.
- Đề cao quyền con người, quyền công dân và nguyên tắc hạn chế quyền bằng luật.
- Làm rõ hơn phân công, phối hợp và kiểm soát quyền lực.
- Xác định rõ chức năng lập pháp, hành pháp, tư pháp.
- Tạo cơ sở hiến định cho cải cách bộ máy, tư pháp, kinh tế thị trường, hội nhập và bảo vệ môi trường.

## Kết luận
Hiến pháp năm 2013 là nền tảng chính trị - pháp lý của Nhà nước và xã hội, vừa tổ chức quyền lực nhà nước vừa giới hạn quyền lực, bảo vệ quyền con người và định hướng phát triển đất nước.

# Phụ lục. Cách trình bày để đạt điểm cao

## Công thức trả lời câu lý thuyết
- Mở đầu bằng khái niệm chuẩn, ngắn gọn.
- Chia đúng các ý mà đề yêu cầu; mỗi ý có giải thích, không chỉ liệt kê.
- Nếu có căn cứ pháp luật, nêu đúng tên văn bản và năm ban hành; không cần chép dài điều luật nếu đề không yêu cầu.
- Nêu một lưu ý phân biệt hoặc ví dụ ngắn ở phần dễ nhầm.
- Kết luận từ một đến hai câu, khẳng định ý nghĩa của vấn đề.

## Những lỗi cần tránh
- Dùng văn bản đã hết hiệu lực mà không ghi chú cập nhật.
- Nhầm khách thể với người bị thiệt hại; khách thể là quan hệ xã hội được pháp luật bảo vệ.
- Nhầm điều luật với quy phạm pháp luật.
- Cho rằng mọi vi phạm đều phải có hậu quả vật chất thực tế.
- Liệt kê đặc trưng của xã hội xã hội chủ nghĩa thay cho đặc trưng của Nhà nước.
- Khẳng định tố cáo nặc danh luôn “bất hợp pháp” hoặc mọi tranh chấp lao động đều có thể đình công.

# Tài liệu pháp luật chính được dùng để cập nhật
- Hiến pháp nước Cộng hòa xã hội chủ nghĩa Việt Nam năm 2013.
- Bộ luật Dân sự số 91/2015/QH13.
- Bộ luật Hình sự số 100/2015/QH13, được sửa đổi, bổ sung năm 2017.
- Bộ luật Lao động số 45/2019/QH14.
- Luật Tố cáo số 25/2018/QH14.
- Luật Hôn nhân và gia đình số 52/2014/QH13.
- Pháp luật hiện hành về ban hành văn bản quy phạm pháp luật và điều ước quốc tế.
"""


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, name="Arial", size=11, bold=None, italic=None, color="222222"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_font(run, size=9, color="666666")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def shade_paragraph(paragraph, fill="EAF1F8"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(34, 34, 34)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, color, before, after in (
    ("Heading 1", 16, "1F4E79", 16, 8),
    ("Heading 2", 13, "2E75B6", 11, 5),
    ("Heading 3", 11.5, "1F4E79", 7, 3),
):
    style = styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Bullet 2"):
    style = styles[name]
    style.font.name = "Arial"
    style.font.size = Pt(10.7)
    style.paragraph_format.left_indent = Inches(0.38 if name == "List Bullet" else 0.65)
    style.paragraph_format.first_line_indent = Inches(-0.18)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.15

if "Lead Note" not in styles:
    lead = styles.add_style("Lead Note", WD_STYLE_TYPE.PARAGRAPH)
else:
    lead = styles["Lead Note"]
lead.font.name = "Arial"
lead.font.size = Pt(10.5)
lead.font.italic = True
lead.font.color.rgb = RGBColor.from_string("4F4F4F")
lead.paragraph_format.left_indent = Inches(0.18)
lead.paragraph_format.right_indent = Inches(0.18)
lead.paragraph_format.space_before = Pt(4)
lead.paragraph_format.space_after = Pt(8)
lead.paragraph_format.line_spacing = 1.15

header = section.header
hp = header.paragraphs[0]
hp.text = "ĐỀ CƯƠNG ÔN THI PHÁP LUẬT ĐẠI CƯƠNG"
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(hp.runs[0], size=8.5, bold=True, color="6B7280")

footer = section.footer
add_page_number(footer.paragraphs[0])

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(90)
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ĐỀ CƯƠNG ÔN THI")
set_font(r, size=26, bold=True, color="17365D")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PHÁP LUẬT ĐẠI CƯƠNG")
set_font(r, size=20, bold=True, color="2E75B6")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(26)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("19 câu lý thuyết – bản sửa sai, bổ sung và cập nhật")
set_font(r, size=12.5, italic=True, color="555555")

note = doc.add_paragraph(style="Lead Note")
note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
note.add_run(
    "Lưu ý: File gốc bị thiếu Câu 13 (đánh số từ 12 sang 14). "
    "Tài liệu này bổ sung Câu 13 về hợp đồng lao động theo mạch kiến thức của chương Luật lao động. "
    "Các dẫn chiếu đã hết hiệu lực trong đề cương cũ được ghi chú và cập nhật."
)
shade_paragraph(note, "EAF1F8")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Tài liệu phục vụ ôn tập và trình bày bài thi")
set_font(r, size=10.5, color="666666")

doc.add_page_break()

# Compact contents list
p = doc.add_paragraph("DANH MỤC 19 CÂU", style="Heading 1")
for line in [x[2:].strip() for x in CONTENT.splitlines() if x.startswith("# ") and not x.startswith("# Phụ") and not x.startswith("# Tài")]:
    q = doc.add_paragraph(style="List Bullet")
    q.add_run(line)
doc.add_page_break()

first_question = True
for raw in CONTENT.strip().splitlines():
    line = raw.strip()
    if not line:
        continue
    if line.startswith("# "):
        title = line[2:].strip()
        p = doc.add_paragraph(title, style="Heading 1")
        if title.startswith("Câu "):
            p.paragraph_format.page_break_before = not first_question
            first_question = False
        continue
    if line.startswith("## "):
        doc.add_paragraph(line[3:].strip(), style="Heading 2")
        continue
    if line.startswith("### "):
        doc.add_paragraph(line[4:].strip(), style="Heading 3")
        continue
    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line[2:].strip())
        continue
    p = doc.add_paragraph()
    p.add_run(line)

# Keep final bibliography from becoming an orphaned fragment.
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True
    for run in p.runs:
        if not run.font.name:
            set_font(run, size=11)

doc.core_properties.title = "Đề cương Pháp luật đại cương - 19 câu hoàn chỉnh"
doc.core_properties.subject = "Tài liệu ôn thi đã sửa sai, bổ sung và cập nhật"
doc.core_properties.author = "Codex"
doc.core_properties.keywords = "Pháp luật đại cương, đề cương, ôn thi, 19 câu"
doc.save(OUTPUT)
print(OUTPUT)
