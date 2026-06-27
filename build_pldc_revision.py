from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_LEADER
from pathlib import Path


OUT = Path(r"D:\ProjectWeb\Project - Sao chép\DAP-AN-19-CAU-PLDC-RA-SOAT-TOAN-DIEN-VA-TU-DIEN-KHAI-NIEM.docx")

BLUE = "1F4E79"
DARK = "17365D"
LIGHT = "D9EAF7"
PALE = "EEF5FA"
GOLD = "FFF2CC"
RED = "C00000"
GRAY = "666666"
INK = "222222"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_inches):
    table.autofit = False
    total_dxa = sum(int(round(width * 1440)) for width in widths_inches)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width * 1440))))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(round(width * 1440))))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Arial", size=11, bold=None, color=INK, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.7)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 17, BLUE, 14, 7),
        ("Heading 2", 14, BLUE, 11, 5),
        ("Heading 3", 12, DARK, 8, 3),
    ]:
        st = doc.styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Key Point" not in [s.name for s in doc.styles]:
        st = doc.styles.add_style("Key Point", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.size = Pt(10.5)
        st.font.color.rgb = RGBColor.from_string(DARK)
        st.paragraph_format.left_indent = Inches(0.2)
        st.paragraph_format.right_indent = Inches(0.1)
        st.paragraph_format.space_before = Pt(4)
        st.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("ĐỀ CƯƠNG PHÁP LUẬT ĐẠI CƯƠNG • BẢN HIỆU ĐÍNH 2026")
        set_font(r, size=8, bold=True, color=GRAY)
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Tài liệu ôn tập – cần ưu tiên đề cương và hướng dẫn riêng của giảng viên")
        set_font(r, size=8, italic=True, color=GRAY)
        r = p.add_run("  |  Trang ")
        set_font(r, size=8, color=GRAY)
        run = p.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("ĐỀ CƯƠNG ÔN THI CUỐI KỲ")
    set_font(r, size=25, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PHÁP LUẬT ĐẠI CƯƠNG")
    set_font(r, size=22, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Bản bổ sung khái niệm, luận giải và ví dụ minh họa")
    set_font(r, size=13, italic=True, color=GRAY)
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    vals = [
        ("Mục tiêu", "Ôn tự luận, nhận diện vấn đề và giải bài tập tình huống"),
        ("Cấu trúc", "19 câu trả lời hoàn chỉnh + bảng ghi nhớ + bài tập tổng hợp nâng cao"),
        ("Căn cứ chính", "Các văn bản pháp luật hiện hành được liệt kê trong tài liệu"),
        ("Lưu ý", "Đây là tài liệu học tập, không thay thế tư vấn pháp lý cho vụ việc thực tế"),
    ]
    for row, (a, b) in zip(table.rows, vals):
        set_cell_shading(row.cells[0], LIGHT)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = row.cells[0].paragraphs[0].add_run(a)
        set_font(r, size=10, bold=True, color=DARK)
        r = row.cells[1].paragraphs[0].add_run(b)
        set_font(r, size=10)
    set_table_widths(table, [1.35, 5.35])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Biên soạn lại từ file tổng hợp của người học; đã loại bỏ hoặc đánh dấu các căn cứ hết hiệu lực.")
    set_font(r, size=9, italic=True, color=GRAY)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_contents(doc):
    doc.add_heading("MỤC LỤC NỘI DUNG", level=1)
    parts = [
        ("Hướng dẫn sử dụng và các điểm hiệu đính", "Cách học, văn bản nền tảng, nội dung đã thay thế"),
        ("Phần I. Đáp án 19 câu lý thuyết", "Mỗi câu gồm khái niệm, phân tích, ví dụ, lưu ý và dàn ý"),
        ("Phần II. Bảng ghi nhớ nhanh", "So sánh khái niệm và công thức giải bài"),
        ("Phần III. Bài tập", "Chỉ gồm các tình huống nâng cao, đa dạng và có lời giải chi tiết"),
        ("Phần IV. Từ điển khái niệm", "Liệt kê riêng từng thuật ngữ pháp lý xuất hiện trong tài liệu và định nghĩa"),
        ("Phần V. Checklist trước khi thi", "Kiểm tra kiến thức và khung trả lời tự luận"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ("Phần", "Nội dung")):
        set_cell_shading(cell, BLUE)
        r = cell.paragraphs[0].add_run(text)
        set_font(r, size=10, bold=True, color="FFFFFF")
    for a, b in parts:
        cells = table.add_row().cells
        r = cells[0].paragraphs[0].add_run(a)
        set_font(r, size=10, bold=True, color=DARK)
        r = cells[1].paragraphs[0].add_run(b)
        set_font(r, size=10)
    set_table_widths(table, [2.55, 4.15])
    set_repeat_table_header(table.rows[0])
    doc.add_heading("Danh sách 19 câu lý thuyết", level=2)
    qs = [
        "Nhà nước và bản chất Nhà nước",
        "Văn bản quy phạm pháp luật",
        "Nhà nước CHXHCN Việt Nam",
        "Nguyên tắc Luật Hình sự",
        "Quan hệ pháp luật",
        "Hình thức pháp luật",
        "Luật Dân sự",
        "Quy phạm pháp luật",
        "Quyền sở hữu",
        "Vi phạm pháp luật",
        "Thừa kế theo di chúc",
        "Luật Lao động",
        "Hợp đồng lao động",
        "Tranh chấp lao động",
        "Luật Hành chính",
        "Tội phạm",
        "Tố cáo",
        "Luật Hôn nhân và gia đình",
        "Hiến pháp năm 2013 và cập nhật 2025",
    ]
    for i, q in enumerate(qs, 1):
        add_number(doc, f"Câu {i}: {q}")
    doc.add_page_break()


def add_bullet(doc, text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)
    return p


def add_key(doc, text, label="Từ khóa ghi điểm"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [6.7])
    set_cell_shading(table.cell(0, 0), PALE)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run(label + ": ")
    set_font(r, size=10.5, bold=True, color=DARK)
    r = p.add_run(text)
    set_font(r, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


QUESTION_DETAILS = {
    1: {
        "basis": "Giáo trình Lý luận Nhà nước và pháp luật; các quy định nền tảng của Hiến pháp năm 2013.",
        "analysis": [
            "Nhà nước là quyền lực công cộng đặc biệt, tách tương đối khỏi dân cư và được tổ chức thành bộ máy gồm các cơ quan thực hiện quyền lực. Nhà nước khác các tổ chức xã hội ở quyền ban hành pháp luật, sử dụng cưỡng chế nhà nước, quản lý dân cư theo lãnh thổ, thực hiện chủ quyền và thu các khoản bắt buộc để duy trì bộ máy.",
            "Tính giai cấp thể hiện ở chỗ nhà nước ra đời trong xã hội có mâu thuẫn giai cấp và quyền lực nhà nước trước hết phản ánh ý chí, bảo vệ địa vị của lực lượng giữ ưu thế về kinh tế và chính trị. Tính giai cấp thể hiện thông qua chính sách kinh tế, tổ chức quyền lực, pháp luật và hệ tư tưởng.",
            "Tính xã hội thể hiện qua việc nhà nước phải giải quyết các công việc chung mà cá nhân hoặc tổ chức riêng lẻ không thể tự giải quyết có hiệu quả: bảo vệ trật tự, quốc phòng, hạ tầng, giáo dục, y tế, cứu trợ, môi trường và điều hòa lợi ích.",
            "Hai phương diện có quan hệ thống nhất. Nếu nhà nước không thực hiện chức năng xã hội thì trật tự chung và chính cơ sở tồn tại của quyền lực nhà nước cũng bị đe dọa. Ngược lại, chức năng xã hội được thực hiện trong những điều kiện chính trị và lợi ích nhất định.",
            "Ý nghĩa phương pháp luận là khi đánh giá một nhà nước phải xem xét điều kiện lịch sử, cơ sở kinh tế, lực lượng nắm quyền, cách tổ chức quyền lực và hoạt động thực tế; không được nhìn nhà nước như một hiện tượng phi giai cấp hoặc chỉ là bộ máy cưỡng chế.",
        ],
        "example": "Việc Nhà nước ban hành chính sách thuế vừa tạo nguồn lực duy trì bộ máy quyền lực, vừa cung cấp tài chính cho giáo dục, y tế và an sinh. Ví dụ này phản ánh đồng thời phương diện quyền lực và phương diện xã hội.",
        "pitfalls": ["Viết tính giai cấp là bảo vệ trước hết lợi ích toàn xã hội.", "Chỉ kể quyền lực kinh tế, chính trị, tư tưởng mà không giải thích.", "Nhầm bản chất Nhà nước với chức năng hoặc hình thức Nhà nước."],
        "outline": ["Nêu khái niệm và dấu hiệu Nhà nước.", "Phân tích tính giai cấp.", "Phân tích tính xã hội và mối quan hệ.", "Nêu ý nghĩa phương pháp luận và kết luận."],
    },
    2: {
        "basis": "Luật Ban hành văn bản quy phạm pháp luật số 64/2025/QH15 và Luật sửa đổi, bổ sung số 87/2025/QH15.",
        "analysis": [
            "Dấu hiệu cốt lõi của văn bản quy phạm pháp luật là có chứa quy phạm pháp luật, tức quy tắc xử sự chung, có hiệu lực bắt buộc chung và được áp dụng lặp lại. Văn bản giải quyết một cá nhân hoặc vụ việc xác định thường là văn bản áp dụng pháp luật, không phải văn bản QPPL.",
            "Văn bản phải được ban hành đúng chủ thể có thẩm quyền, đúng tên loại, nội dung, hình thức, trình tự và thủ tục do luật quy định. Việc một cơ quan nhà nước ban hành văn bản không đương nhiên làm văn bản đó trở thành văn bản QPPL.",
            "Tính hợp hiến yêu cầu mọi văn bản phù hợp Hiến pháp. Tính hợp pháp bao gồm phù hợp văn bản có hiệu lực cao hơn và đúng thẩm quyền, thủ tục. Tính thống nhất đòi hỏi quy định không mâu thuẫn, chồng chéo trong hệ thống.",
            "Công khai, minh bạch đòi hỏi quá trình xây dựng và nội dung văn bản có khả năng tiếp cận, lấy ý kiến khi pháp luật yêu cầu và công bố đúng quy định. Tính khả thi đòi hỏi quy định rõ ràng, có nguồn lực, có điều kiện thực hiện và không tạo chi phí tuân thủ bất hợp lý.",
            "Phải bảo đảm quyền con người, quyền công dân, lợi ích quốc gia, quốc phòng, an ninh và việc thực hiện điều ước quốc tế. Khi điều ước quốc tế và văn bản trong nước khác nhau phải áp dụng theo quy tắc của pháp luật về điều ước quốc tế; Hiến pháp vẫn có vị trí cao nhất.",
            "Quy trình ban hành không giống nhau cho mọi văn bản. Luật, nghị quyết của Quốc hội có quy trình khác nghị định, thông tư hoặc nghị quyết của HĐND. Vì vậy bài thi nên trình bày nguyên tắc chung, chỉ nêu quy trình cụ thể khi đề yêu cầu.",
        ],
        "example": "Quyết định xử phạt A 5 triệu đồng chỉ áp dụng một lần cho A nên là văn bản áp dụng pháp luật. Nghị định quy định mức phạt chung cho mọi người có hành vi tương ứng là văn bản QPPL.",
        "pitfalls": ["Gọi mọi công văn, thông báo là văn bản QPPL.", "Nêu quy trình xây dựng luật như quy trình chung của tất cả văn bản.", "Khẳng định điều ước quốc tế luôn cao hơn Hiến pháp."],
        "outline": ["Khái niệm.", "Bốn nhóm đặc điểm.", "Các nguyên tắc: hợp hiến–hợp pháp–thống nhất; đúng thẩm quyền; minh bạch; khả thi; bảo đảm quyền.", "Ví dụ phân biệt văn bản QPPL và cá biệt."],
    },
    3: {
        "basis": "Điều 2, Điều 3, Điều 4, Điều 6 và Điều 8 Hiến pháp năm 2013; nội dung sửa đổi Hiến pháp năm 2025.",
        "analysis": [
            "Bản chất của Nhà nước Việt Nam được thể hiện trước hết ở nguyên tắc Nhà nước pháp quyền xã hội chủ nghĩa của Nhân dân, do Nhân dân và vì Nhân dân. Nhân dân là chủ thể tối cao của quyền lực nhà nước, thực hiện quyền lực bằng dân chủ trực tiếp và dân chủ đại diện.",
            "Nền tảng xã hội của quyền lực nhà nước là liên minh giữa giai cấp công nhân với giai cấp nông dân và đội ngũ trí thức. Nhà nước đại diện cho lợi ích của Nhân dân và các dân tộc trên lãnh thổ Việt Nam.",
            "Quyền lực nhà nước là thống nhất nhưng có phân công, phối hợp và kiểm soát giữa các cơ quan thực hiện quyền lập pháp, hành pháp và tư pháp. “Thống nhất” không có nghĩa mọi quyền tập trung không giới hạn vào một cơ quan.",
            "Nhà nước tổ chức và hoạt động theo Hiến pháp và pháp luật; cơ quan, cán bộ, công chức phải tôn trọng Nhân dân, chịu sự giám sát và chịu trách nhiệm trước Nhân dân. Đây là biểu hiện quan trọng của tính pháp quyền.",
            "Nhà nước công nhận, tôn trọng, bảo vệ và bảo đảm quyền con người, quyền công dân; thực hiện mục tiêu dân giàu, nước mạnh, dân chủ, công bằng, văn minh và chính sách đại đoàn kết dân tộc.",
            "Đảng Cộng sản Việt Nam lãnh đạo Nhà nước và xã hội, đồng thời hoạt động trong khuôn khổ Hiến pháp, pháp luật, gắn bó với Nhân dân, chịu sự giám sát và chịu trách nhiệm trước Nhân dân.",
        ],
        "example": "Nhân dân bầu đại biểu Quốc hội và HĐND là dân chủ đại diện; biểu quyết khi Nhà nước tổ chức trưng cầu ý dân là dân chủ trực tiếp.",
        "pitfalls": ["Dùng tám đặc trưng của xã hội XHCN thay cho đặc trưng của Nhà nước.", "Bỏ cơ chế kiểm soát quyền lực.", "Chỉ viết “của dân, do dân, vì dân” mà không phân tích."],
        "outline": ["Trích bản chất hiến định.", "Phân tích chủ quyền Nhân dân.", "Phân tích Nhà nước pháp quyền và kiểm soát quyền lực.", "Quyền con người, đại đoàn kết, Đảng lãnh đạo.", "Kết luận."],
    },
    4: {
        "basis": "Bộ luật Hình sự năm 2015, Luật sửa đổi năm 2017 và Luật số 86/2025/QH15.",
        "analysis": [
            "Nguyên tắc pháp chế hình sự được thể hiện ở yêu cầu chỉ BLHS mới có thể quy định tội phạm và hình phạt; không được áp dụng tương tự pháp luật để kết tội; việc truy cứu và quyết định hình phạt phải đúng căn cứ và thủ tục.",
            "Bình đẳng trước pháp luật hình sự có nghĩa mọi người, pháp nhân thương mại khi đủ điều kiện đều bị xử lý theo luật, không phân biệt giới tính, dân tộc, tín ngưỡng, nghề nghiệp hoặc địa vị. Bình đẳng không có nghĩa áp dụng một mức hình phạt giống nhau cho mọi trường hợp.",
            "Trách nhiệm trên cơ sở có lỗi đòi hỏi chủ thể có khả năng nhận thức, điều khiển hành vi và có thái độ cố ý hoặc vô ý. Không thể chỉ dựa vào hậu quả để quy trách nhiệm hình sự.",
            "Cá thể hóa trách nhiệm và hình phạt yêu cầu xem xét tính chất, mức độ nguy hiểm, nhân thân, động cơ, mục đích, tình tiết tăng nặng, giảm nhẹ và khả năng giáo dục, phòng ngừa.",
            "Nguyên tắc công bằng, nhân đạo thể hiện ở việc không làm oan, không bỏ lọt tội phạm, hình phạt không nhằm hành hạ hoặc hạ thấp nhân phẩm, đồng thời tạo điều kiện để người phạm tội cải tạo và tái hòa nhập.",
            "Tài liệu cũ nói chỉ cá nhân chịu trách nhiệm là không còn chính xác. Pháp nhân thương mại có thể chịu trách nhiệm về các tội được BLHS liệt kê khi đủ điều kiện.",
        ],
        "example": "Hai người cùng gây thiệt hại nhưng một người chủ mưu, một người giúp sức hạn chế và thành khẩn khắc phục thì hình phạt không nhất thiết bằng nhau; đây là cá thể hóa chứ không vi phạm bình đẳng.",
        "pitfalls": ["Phủ nhận trách nhiệm hình sự của pháp nhân thương mại.", "Hiểu nhân đạo là không xử lý nghiêm.", "Nêu tên nguyên tắc nhưng không phân tích nội dung."],
        "outline": ["Khái niệm Luật Hình sự.", "Pháp chế.", "Bình đẳng.", "Có lỗi và cá thể hóa.", "Công bằng, nhân đạo.", "Lưu ý pháp nhân thương mại."],
    },
    5: {
        "basis": "Lý luận Nhà nước và pháp luật; Điều 16–24 và Điều 74 Bộ luật Dân sự 2015 về năng lực chủ thể và pháp nhân.",
        "analysis": [
            "Quan hệ pháp luật không tự phát sinh chỉ vì có quy phạm. Thông thường cần có quy phạm điều chỉnh, chủ thể có năng lực phù hợp và sự kiện pháp lý như giao kết hợp đồng, gây thiệt hại, sinh, chết hoặc quyết định của cơ quan có thẩm quyền.",
            "Chủ thể là cá nhân, pháp nhân, tổ chức hoặc Nhà nước tham gia quan hệ với quyền và nghĩa vụ xác định. Năng lực pháp luật là khả năng có quyền, nghĩa vụ; năng lực hành vi là khả năng bằng hành vi của mình xác lập, thực hiện quyền, nghĩa vụ.",
            "Năng lực hành vi không giống nhau ở mọi ngành luật. Một người có thể đủ năng lực thực hiện giao dịch dân sự thông thường nhưng chưa đủ tuổi kết hôn hoặc chịu trách nhiệm với một tội cụ thể.",
            "Nội dung quan hệ gồm quyền chủ thể và nghĩa vụ pháp lý tương ứng. Quyền có thể bao gồm tự xử sự, yêu cầu bên kia thực hiện và yêu cầu cơ quan có thẩm quyền bảo vệ.",
            "Khách thể là lợi ích vật chất, tinh thần hoặc kết quả hành vi mà các bên hướng tới. Khi phân tích tình huống không nên ghi khách thể là tên của một bên.",
            "Sự kiện pháp lý có thể là hành vi hợp pháp, hành vi vi phạm hoặc biến cố không phụ thuộc ý chí. Một sự kiện có thể làm phát sinh, thay đổi hoặc chấm dứt quan hệ.",
        ],
        "example": "Trong hợp đồng vay, chủ thể là bên vay và bên cho vay; nội dung là quyền, nghĩa vụ giao và hoàn trả tiền; khách thể là lợi ích đạt được qua khoản vay; việc giao kết và đến hạn là các sự kiện pháp lý.",
        "pitfalls": ["Dùng thuật ngữ “năng lực pháp lý” thay cho hai khái niệm chuẩn.", "Ghi khách thể là người bị hại.", "Không nêu sự kiện pháp lý."],
        "outline": ["Khái niệm.", "Đặc điểm.", "Chủ thể và năng lực.", "Nội dung quyền–nghĩa vụ.", "Khách thể.", "Sự kiện pháp lý và ví dụ."],
    },
    6: {
        "basis": "Giáo trình Lý luận Nhà nước và pháp luật; BLDS 2015 về tập quán; quy định về lựa chọn và áp dụng án lệ.",
        "analysis": [
            "Hình thức pháp luật là phương thức tồn tại và biểu hiện ra bên ngoài của các quy tắc pháp lý, giúp chủ thể nhận biết quyền, nghĩa vụ và căn cứ áp dụng.",
            "Tập quán pháp là tập quán được Nhà nước thừa nhận. Không phải mọi thói quen cộng đồng đều là tập quán pháp; tập quán không được trái các nguyên tắc cơ bản và chỉ áp dụng trong điều kiện pháp luật cho phép.",
            "Tiền lệ pháp theo nghĩa rộng là quyết định giải quyết vụ việc được thừa nhận làm mẫu. Ở Việt Nam cần dùng thuật ngữ án lệ thận trọng: chỉ lập luận, phán quyết đã được lựa chọn và công bố, không phải tất cả bản án.",
            "Văn bản QPPL là hình thức chủ yếu của pháp luật Việt Nam vì có tính thành văn, rõ thẩm quyền, hiệu lực và thuận lợi cho việc bảo đảm thống nhất.",
            "Điều ước quốc tế có vai trò quan trọng trong quan hệ đối ngoại và nhiều lĩnh vực nội luật. Việc áp dụng phải theo Hiến pháp và pháp luật về điều ước quốc tế.",
            "Các hình thức có ưu, nhược điểm khác nhau: tập quán linh hoạt nhưng có thể thiếu rõ ràng; án lệ góp phần thống nhất xét xử nhưng phụ thuộc vụ việc; văn bản rõ ràng nhưng có nguy cơ chậm so với thực tiễn.",
        ],
        "example": "Một tập quán thương mại chỉ được áp dụng khi các bên không thỏa thuận, pháp luật không quy định và tập quán không trái nguyên tắc cơ bản; không thể viện dẫn tập quán để hợp pháp hóa hành vi bị cấm.",
        "pitfalls": ["Cho rằng Việt Nam hoàn toàn không có án lệ.", "Gọi mọi tập quán là pháp luật.", "Trình bày lại toàn bộ Câu 2 mà không làm rõ khái niệm hình thức."],
        "outline": ["Khái niệm.", "Tập quán pháp.", "Án lệ/tiền lệ pháp.", "Văn bản QPPL.", "Điều ước quốc tế và nhận xét ưu nhược điểm."],
    },
    7: {
        "basis": "Bộ luật Dân sự năm 2015, đặc biệt Điều 1–6 và các nguyên tắc cơ bản tại Điều 3.",
        "analysis": [
            "Đối tượng điều chỉnh của luật dân sự gồm quan hệ tài sản và quan hệ nhân thân hình thành trên cơ sở bình đẳng, tự do ý chí, độc lập tài sản và tự chịu trách nhiệm.",
            "Quan hệ tài sản là quan hệ giữa các chủ thể thông qua tài sản, không phải quan hệ đơn thuần giữa người với vật. Tài sản gồm vật, tiền, giấy tờ có giá và quyền tài sản.",
            "Quan hệ nhân thân gắn với giá trị không thể chuyển giao tùy tiện như họ tên, danh dự, đời sống riêng tư; một số quyền nhân thân có thể làm phát sinh quyền tài sản như quyền tác giả.",
            "Bình đẳng có nghĩa không bên nào được dùng quyền lực nhà nước để áp đặt trong quan hệ dân sự. Cơ quan nhà nước khi mua bán, thuê tài sản cũng phải bình đẳng như chủ thể khác.",
            "Tự do, tự nguyện cam kết được tôn trọng nhưng bị giới hạn bởi điều cấm, đạo đức xã hội, quyền của người khác và lợi ích công cộng.",
            "Độc lập tài sản và tự chịu trách nhiệm tạo cơ sở để mỗi bên chịu hậu quả từ cam kết, hành vi của mình. Trách nhiệm dân sự chủ yếu hướng tới khôi phục quyền, bù đắp thiệt hại.",
        ],
        "example": "Cơ quan nhà nước thuê nhà của cá nhân thì quan hệ thuê là dân sự vì hai bên bình đẳng; việc cơ quan ra quyết định xử phạt cá nhân là quan hệ hành chính.",
        "pitfalls": ["Cho mọi quan hệ có tiền là dân sự.", "Xếp quyền kết hôn, ly hôn hoàn toàn vào quan hệ dân sự dù thuộc ngành hôn nhân gia đình.", "Bỏ phương pháp tự chịu trách nhiệm."],
        "outline": ["Khái niệm.", "Quan hệ tài sản.", "Quan hệ nhân thân.", "Bình đẳng.", "Tự do ý chí.", "Độc lập tài sản, tự chịu trách nhiệm."],
    },
    8: {
        "basis": "Luật Ban hành VBQPPL hiện hành và giáo trình Lý luận Nhà nước và pháp luật.",
        "analysis": [
            "Quy phạm pháp luật là quy tắc xử sự chung chứ không đồng nhất với một điều luật. Một điều có thể chứa nhiều quy phạm; một quy phạm có thể được thể hiện qua nhiều điều hoặc nhiều văn bản.",
            "Tính bắt buộc chung nghĩa là quy phạm áp dụng với mọi chủ thể thuộc điều kiện đã dự liệu, không phải áp dụng cho toàn bộ mọi người trong mọi hoàn cảnh.",
            "Giả định xác định chủ thể, thời gian, địa điểm, điều kiện hoặc hoàn cảnh. Giả định có thể đơn giản hoặc phức tạp, xác định hoặc tương đối xác định.",
            "Quy định nêu cách xử sự được phép, bắt buộc hoặc bị cấm. Đây là phần định hướng hành vi và thường được coi là trung tâm của quy phạm.",
            "Chế tài nêu hậu quả pháp lý bất lợi khi chủ thể không tuân thủ. Có thể là hình sự, hành chính, dân sự, kỷ luật; không nên chỉ phân loại bằng các cụm từ thiếu chuẩn như “phủ nhận pháp luật”.",
            "Trong quy phạm xử phạt, quy định về nghĩa vụ thường được thể hiện ngầm hoặc nằm ở văn bản giao thông, còn mức phạt nằm trong nghị định xử phạt.",
        ],
        "example": "Câu “phạt tiền đối với người đi xe máy không đội mũ” có giả định là người điều khiển/ngồi trên xe trong điều kiện tham gia giao thông; quy định ngầm là phải đội và cài đúng; chế tài là mức phạt.",
        "pitfalls": ["Ghi “quy phạm pháp luật đơn giản là điều luật”.", "Cố tìm đủ ba phần trong đúng một câu.", "Nhầm hành vi vi phạm trong giả định với phần quy định."],
        "outline": ["Khái niệm.", "Đặc điểm.", "Giả định.", "Quy định.", "Chế tài.", "Lưu ý cấu trúc logic và ví dụ."],
    },
    9: {
        "basis": "Bộ luật Dân sự 2015: Điều 158 và các quy định về chiếm hữu, sử dụng, định đoạt; Điều 221 và Điều 237 về căn cứ xác lập, chấm dứt.",
        "analysis": [
            "Sở hữu là quan hệ xã hội giữa người với người đối với tài sản. Quyền sở hữu là phạm trù pháp lý gồm quyền chiếm hữu, sử dụng và định đoạt.",
            "Chiếm hữu là nắm giữ, chi phối tài sản trực tiếp hoặc gián tiếp. Người thuê, mượn hoặc giữ tài sản có thể chiếm hữu hợp pháp nhưng không phải chủ sở hữu.",
            "Sử dụng là khai thác công dụng, hưởng hoa lợi và lợi tức. Việc sử dụng phải không gây thiệt hại, ảnh hưởng lợi ích quốc gia, công cộng hoặc quyền của người khác.",
            "Định đoạt là chuyển giao quyền sở hữu, từ bỏ, tiêu dùng hoặc tiêu hủy tài sản. Chủ thể định đoạt phải có năng lực và đúng thẩm quyền; người không phải chủ sở hữu chỉ định đoạt khi được ủy quyền hoặc luật cho phép.",
            "Căn cứ xác lập gồm lao động, sản xuất kinh doanh; chuyển giao; hoa lợi, lợi tức; tạo vật mới; thừa kế; chiếm hữu vật vô chủ, đánh rơi, bỏ quên; thời hiệu và căn cứ khác.",
            "Quyền sở hữu chấm dứt khi chuyển quyền, từ bỏ, tài sản tiêu hủy, bị xử lý nghĩa vụ, trưng mua, tịch thu hoặc người khác được xác lập quyền theo luật.",
        ],
        "example": "A thuê xe của B: A có quyền chiếm hữu, sử dụng trong phạm vi hợp đồng; B vẫn là chủ sở hữu; A không có quyền bán xe.",
        "pitfalls": ["Dẫn Điều 170, 171 BLDS 2005.", "Bỏ quyền sử dụng khi liệt kê nội dung quyền sở hữu.", "Cho rằng người đang giữ tài sản là chủ sở hữu."],
        "outline": ["Khái niệm sở hữu và quyền sở hữu.", "Ba quyền năng.", "Giới hạn.", "Nhóm căn cứ xác lập.", "Nhóm căn cứ chấm dứt.", "Ví dụ."],
    },
    10: {
        "basis": "Lý luận Nhà nước và pháp luật; các luật chuyên ngành quy định từng loại vi phạm.",
        "analysis": [
            "Vi phạm pháp luật phải là hành vi biểu hiện ra bên ngoài bằng hành động hoặc không hành động. Suy nghĩ hoặc ý định chưa biểu hiện thành hành vi không phải vi phạm.",
            "Hành vi phải trái quy định pháp luật và xâm hại hoặc đe dọa xâm hại quan hệ được bảo vệ. Không phải hành vi gây phản cảm nào cũng là vi phạm nếu pháp luật không điều chỉnh.",
            "Chủ thể phải có lỗi. Lỗi là thái độ tâm lý đối với hành vi và hậu quả trong điều kiện có khả năng lựa chọn cách xử sự phù hợp.",
            "Chủ thể phải có năng lực trách nhiệm pháp lý theo ngành luật, phụ thuộc độ tuổi, khả năng nhận thức và tư cách tổ chức.",
            "Vi phạm hình sự có tính nguy hiểm cao và được BLHS quy định; vi phạm hành chính xâm phạm quản lý nhà nước nhưng không phải tội; vi phạm dân sự xâm phạm tài sản, nhân thân hoặc nghĩa vụ; vi phạm kỷ luật xâm phạm trật tự nội bộ.",
            "Một hành vi có thể làm phát sinh nhiều trách nhiệm: người lao động cố ý phá tài sản có thể bị kỷ luật, bồi thường và bị xử lý hình sự nếu đủ dấu hiệu.",
        ],
        "example": "Người lái xe vượt sai gây thương tích có thể bị xử phạt hành chính, bồi thường; nếu hậu quả và các dấu hiệu đạt ngưỡng luật hình sự thì còn bị truy cứu.",
        "pitfalls": ["Đòi hỏi mọi vi phạm phải có thiệt hại thực tế.", "Dùng “năng lực hành vi” thay “năng lực trách nhiệm pháp lý”.", "Cho rằng mỗi hành vi chỉ chịu một trách nhiệm."],
        "outline": ["Khái niệm.", "Hành vi.", "Trái pháp luật.", "Có lỗi.", "Năng lực trách nhiệm.", "Bốn loại vi phạm và ví dụ."],
    },
    11: {
        "basis": "Phần thứ tư Bộ luật Dân sự năm 2015, đặc biệt quy định về di chúc, hiệu lực, người thừa kế bắt buộc và thừa kế theo pháp luật.",
        "analysis": [
            "Thừa kế là sự chuyển dịch tài sản của người chết. Di sản chỉ gồm tài sản riêng và phần của người chết trong tài sản chung, sau khi xác định nghĩa vụ tài sản.",
            "Di chúc hợp pháp đòi hỏi người lập minh mẫn, tự nguyện; nội dung không vi phạm điều cấm, đạo đức xã hội; hình thức phù hợp. Di chúc công chứng vẫn có thể vô hiệu nếu bị cưỡng ép.",
            "Người từ đủ 15 đến dưới 18 tuổi phải lập di chúc bằng văn bản và được cha mẹ hoặc người giám hộ đồng ý về việc lập; không nên viết bắt buộc cha mẹ ký thay vào di chúc.",
            "Di chúc miệng chỉ khi tính mạng bị đe dọa và không thể lập văn bản; phải có ít nhất hai người làm chứng, ghi chép, ký/điểm chỉ và được xác nhận theo thời hạn luật định. Sau ba tháng người lập còn sống, minh mẫn thì di chúc miệng mặc nhiên bị hủy bỏ.",
            "Người thừa kế bắt buộc gồm cha, mẹ, vợ/chồng, con chưa thành niên, con thành niên không có khả năng lao động; được ít nhất hai phần ba suất theo pháp luật, trừ trường hợp từ chối hoặc không có quyền hưởng.",
            "BLDS 2015 không còn chế định riêng về di chúc chung của vợ chồng. Mỗi người chỉ định đoạt phần tài sản thuộc quyền của mình.",
            "Khi giải bài phải theo thứ tự: thời điểm mở thừa kế – di sản – nghĩa vụ – di chúc – người hưởng – chia phần – kiểm tra tổng.",
        ],
        "example": "Tài sản chung vợ chồng 900 triệu và người chết có 100 triệu riêng: di sản là 450 + 100 = 550 triệu, không phải một tỷ.",
        "pitfalls": ["Chia toàn bộ tài sản chung.", "Bỏ người thừa kế bắt buộc.", "Dùng di chúc chung vợ chồng của luật cũ.", "Bỏ cha mẹ người chết khi đề chưa cho dữ kiện."],
        "outline": ["Khái niệm và di sản.", "Di chúc, điều kiện hợp pháp.", "Hình thức.", "Quyền người lập.", "Người thừa kế bắt buộc.", "Hiệu lực và cách chia."],
    },
    12: {
        "basis": "Bộ luật Lao động năm 2019, có hiệu lực từ ngày 01/01/2021.",
        "analysis": [
            "Đối tượng trung tâm là quan hệ lao động có việc làm, trả công/tiền lương và sự quản lý, điều hành, giám sát của người sử dụng lao động.",
            "Các quan hệ liên quan gồm việc làm, học nghề, an toàn vệ sinh, đối thoại, thương lượng tập thể, tổ chức đại diện, bảo hiểm và giải quyết tranh chấp.",
            "Phương pháp thỏa thuận thể hiện trong giao kết hợp đồng và thương lượng nhưng không tuyệt đối; các bên không được thỏa thuận thấp hơn tiêu chuẩn tối thiểu có lợi cho người lao động.",
            "Phương pháp quyền uy trong nội bộ thể hiện ở quyền tổ chức, phân công, kiểm tra và xử lý kỷ luật của người sử dụng lao động, nhưng phải trong giới hạn pháp luật.",
            "Sự tham gia của tổ chức đại diện người lao động giúp cân bằng vị thế, đặc biệt trong thương lượng tập thể, kỷ luật và tranh chấp.",
            "Quan hệ công vụ của cán bộ, công chức chủ yếu do pháp luật cán bộ, công chức và hành chính điều chỉnh, không nên xếp toàn bộ vào luật lao động.",
        ],
        "example": "Doanh nghiệp có quyền bố trí ca làm theo hợp đồng và nội quy, nhưng không thể buộc người lao động làm quá giới hạn thời giờ hoặc cắt lương làm hình thức kỷ luật.",
        "pitfalls": ["Cho mọi người làm việc trong cơ quan nhà nước đều thuộc BLLĐ.", "Chỉ nêu phương pháp thỏa thuận.", "Bỏ vai trò tổ chức đại diện người lao động."],
        "outline": ["Khái niệm.", "Quan hệ lao động.", "Quan hệ liên quan.", "Phương pháp thỏa thuận.", "Quyền quản lý.", "Tổ chức đại diện."],
    },
    13: {
        "basis": "Chương III Bộ luật Lao động 2019 về hợp đồng lao động.",
        "analysis": [
            "Tên gọi không quyết định bản chất. Thỏa thuận mang tên cộng tác viên, khoán việc hoặc dịch vụ vẫn là hợp đồng lao động nếu có việc làm có trả công và chịu quản lý, điều hành, giám sát.",
            "Thông thường hợp đồng phải bằng văn bản hoặc thông điệp dữ liệu. Hợp đồng dưới một tháng có thể bằng lời nói, trừ các trường hợp luật bắt buộc văn bản.",
            "Có hai loại: không xác định thời hạn và xác định thời hạn không quá 36 tháng. Luật giới hạn việc ký nối tiếp hợp đồng xác định thời hạn.",
            "Nội dung chủ yếu gồm công việc, địa điểm, thời hạn, lương, hình thức trả, nâng lương, thời giờ, nghỉ ngơi, bảo hộ, bảo hiểm, đào tạo và nội dung khác.",
            "Chấm dứt có thể do hết hạn, thỏa thuận, hoàn thành công việc, đơn phương hợp pháp hoặc căn cứ khác. Người lao động thường có quyền đơn phương khi báo trước; một số trường hợp không phải báo trước.",
            "Đơn phương trái pháp luật dẫn đến nghĩa vụ nhận lại, trả lương, bảo hiểm, bồi thường hoặc nghĩa vụ của người lao động tùy bên vi phạm.",
        ],
        "example": "Hợp đồng ghi “dịch vụ” nhưng người làm chấm công hằng ngày, nhận lương tháng và chịu quản lý trực tiếp có thể được xác định là hợp đồng lao động.",
        "pitfalls": ["Cho người lao động nghỉ phải được công ty đồng ý.", "Không phân biệt hai loại hợp đồng.", "Không kiểm tra thời hạn báo trước."],
        "outline": ["Khái niệm và nhận diện.", "Hình thức.", "Loại hợp đồng.", "Nội dung.", "Các căn cứ chấm dứt.", "Đơn phương và hậu quả."],
    },
    14: {
        "basis": "Chương XIV Bộ luật Lao động 2019 về giải quyết tranh chấp lao động.",
        "analysis": [
            "Tranh chấp cá nhân phát sinh giữa người lao động với người sử dụng lao động hoặc chủ thể liên quan về quyền, nghĩa vụ và lợi ích của cá nhân.",
            "Tranh chấp tập thể về quyền liên quan việc giải thích, thực hiện quy định đã có trong luật, thỏa ước, nội quy hoặc thỏa thuận hợp pháp. Tranh chấp lợi ích nhằm xác lập điều kiện mới.",
            "Chủ thể giải quyết gồm hòa giải viên lao động, Hội đồng trọng tài lao động và Tòa án. Không còn mô hình Hội đồng hòa giải cơ sở và Chủ tịch UBND huyện như tài liệu cũ.",
            "Tranh chấp cá nhân thường phải hòa giải trước, nhưng các tranh chấp như sa thải, đơn phương chấm dứt, bồi thường khi chấm dứt, bảo hiểm và một số nhóm khác được miễn.",
            "Tranh chấp tập thể về quyền có thể được giải quyết qua hòa giải, trọng tài hoặc Tòa án theo quy định. Tranh chấp lợi ích qua hòa giải, trọng tài và có thể dẫn tới đình công nếu đủ điều kiện.",
            "Đình công chỉ là biện pháp của tập thể người lao động đối với tranh chấp lợi ích, phải do tổ chức đại diện có quyền tổ chức và tuân thủ trình tự.",
        ],
        "example": "Yêu cầu trả khoản phụ cấp đã ghi trong thỏa ước là tranh chấp về quyền; yêu cầu tạo một phụ cấp hoàn toàn mới là tranh chấp về lợi ích.",
        "pitfalls": ["Dẫn cơ chế cũ.", "Cho mọi tranh chấp đều bắt buộc hòa giải.", "Cho tranh chấp về quyền được đình công."],
        "outline": ["Khái niệm.", "Ba loại tranh chấp.", "Cơ quan giải quyết.", "Thủ tục tranh chấp cá nhân.", "Tập thể về quyền.", "Tập thể về lợi ích và đình công."],
    },
    15: {
        "basis": "Giáo trình Luật Hành chính; Luật Xử lý vi phạm hành chính; Luật Khiếu nại và Luật Tố tụng hành chính.",
        "analysis": [
            "Luật Hành chính điều chỉnh quan hệ phát sinh trong quản lý hành chính nhà nước và các hoạt động chấp hành – điều hành.",
            "Nhóm điển hình là quan hệ giữa cơ quan hành chính hoặc người có thẩm quyền với cá nhân, tổ chức; quan hệ tổ chức nội bộ; và quan hệ khi chủ thể khác được trao quyền quản lý.",
            "Không phải mọi quan hệ có cơ quan nhà nước đều là hành chính. Cơ quan ký hợp đồng mua bán bình đẳng là quan hệ dân sự.",
            "Phương pháp quyền lực – phục tùng cho phép một bên ra quyết định đơn phương trong phạm vi thẩm quyền. Bên kia phải chấp hành nhưng có quyền khiếu nại, khởi kiện và yêu cầu kiểm soát.",
            "Quyết định hành chính hợp pháp phải đúng thẩm quyền, căn cứ, nội dung, hình thức, trình tự, thủ tục và mục đích. Một hành vi vi phạm có thật không làm hợp pháp hóa quyết định sai thẩm quyền.",
            "Cưỡng chế hành chính chỉ được áp dụng bởi chủ thể có thẩm quyền, đúng trường hợp và thủ tục, bảo đảm quyền con người.",
        ],
        "example": "Chủ tịch có thẩm quyền phạt ra quyết định xử phạt là quan hệ hành chính; UBND thuê dịch vụ sửa máy tính là quan hệ hợp đồng.",
        "pitfalls": ["Viết “chấp hành và diễu hành”.", "Cho bên quyền lực được tùy ý áp đặt.", "Không phân biệt quyết định QPPL và quyết định cá biệt."],
        "outline": ["Khái niệm.", "Các nhóm đối tượng.", "Phương pháp quyền lực–phục tùng.", "Giới hạn thẩm quyền.", "Ví dụ và quyền kiểm soát."],
    },
    16: {
        "basis": "Điều 8, Điều 9, Điều 12 và các quy định liên quan của BLHS 2015 đã được sửa đổi.",
        "analysis": [
            "Tính nguy hiểm cho xã hội là dấu hiệu nội dung cơ bản. Hành vi phải gây hoặc đe dọa gây thiệt hại đáng kể cho quan hệ được BLHS bảo vệ.",
            "Tính có lỗi yêu cầu chủ thể có thái độ cố ý hoặc vô ý. Người gây hậu quả trong sự kiện bất ngờ, không thể thấy trước hoặc không buộc phải thấy trước, không có lỗi.",
            "Tính trái BLHS đòi hỏi hành vi được mô tả trong một tội danh. Không được suy diễn hoặc áp dụng tương tự để buộc tội.",
            "Tính chịu hình phạt nghĩa là tội phạm bị đe dọa áp dụng hình phạt, nhưng người phạm tội thực tế có thể được miễn trách nhiệm hoặc miễn hình phạt khi đủ điều kiện.",
            "Cấu thành gồm khách thể; mặt khách quan; chủ thể; mặt chủ quan. Phải phân biệt khách thể với đối tượng tác động và bị hại.",
            "Từ đủ 16 tuổi chịu trách nhiệm về mọi tội trừ quy định khác; từ 14 đến dưới 16 chỉ chịu trách nhiệm về tội cụ thể được liệt kê tại Điều 12.",
            "Bốn loại tội căn cứ tính chất, mức độ nguy hiểm và mức cao nhất của khung: đến 3 năm; trên 3 đến 7; trên 7 đến 15; trên 15, tù chung thân hoặc tử hình.",
        ],
        "example": "Người 15 tuổi không đương nhiên chịu trách nhiệm chỉ vì hành vi được xếp loại rất nghiêm trọng; phải kiểm tra tội danh có trong danh mục Điều 12.",
        "pitfalls": ["Dùng định nghĩa BLHS 1999.", "Áp dụng quy tắc tuổi cũ.", "Ghi mức cao nhất là “tù giam” thay vì hình phạt tù.", "Nhầm bị hại với khách thể."],
        "outline": ["Khái niệm.", "Nguy hiểm.", "Có lỗi.", "Trái BLHS.", "Chịu hình phạt.", "Chủ thể/tuổi.", "Phân loại."],
    },
    17: {
        "basis": "Luật Tố cáo số 25/2018/QH14.",
        "analysis": [
            "Chủ thể tố cáo là cá nhân, không chỉ công dân. Đối tượng là hành vi vi phạm pháp luật của cơ quan, tổ chức, cá nhân trong thực hiện nhiệm vụ, công vụ hoặc quản lý nhà nước.",
            "Tố cáo khác khiếu nại ở đối tượng và mục đích. Khiếu nại yêu cầu xem xét quyết định/hành vi xâm phạm trực tiếp quyền người khiếu nại; tố cáo báo hành vi vi phạm để xử lý.",
            "Nguyên tắc chung về thẩm quyền là người đứng đầu cơ quan, tổ chức có thẩm quyền quản lý người bị tố cáo giải quyết; trường hợp đặc biệt theo luật.",
            "Trình tự gồm tiếp nhận, xử lý ban đầu; thụ lý; xác minh; kết luận nội dung; xử lý kết luận và thông báo theo luật.",
            "Thời hạn thông thường không quá 30 ngày từ ngày thụ lý; vụ phức tạp có thể gia hạn một lần, đặc biệt phức tạp có thể gia hạn hai lần, mỗi lần không quá 30 ngày.",
            "Đơn không rõ tên, địa chỉ không giải quyết theo thủ tục tố cáo. Tuy nhiên nếu thông tin rõ, có tài liệu cụ thể và cơ sở kiểm tra thì cơ quan có thể thanh tra, kiểm tra hoặc chuyển cơ quan có thẩm quyền.",
            "Người tố cáo được bảo mật và có thể được bảo vệ tính mạng, sức khỏe, tài sản, danh dự, vị trí việc làm khi có nguy cơ bị xâm hại.",
        ],
        "example": "A phản đối quyết định phạt mình thì khiếu nại/khởi kiện; B phát hiện cán bộ nhận hối lộ thì tố cáo hoặc báo tin tội phạm.",
        "pitfalls": ["Dùng Luật Khiếu nại, tố cáo cũ.", "Ghi thời hạn 60/90 ngày như tài liệu cũ.", "Nói đơn nặc danh luôn phải tiêu hủy.", "Cho tổ chức là người tố cáo."],
        "outline": ["Khái niệm.", "Đặc điểm.", "So sánh khiếu nại.", "Thẩm quyền.", "Trình tự.", "Thời hạn.", "Bảo vệ người tố cáo."],
    },
    18: {
        "basis": "Điều 2 và các quy định liên quan của Luật Hôn nhân và gia đình năm 2014.",
        "analysis": [
            "Hôn nhân tự nguyện, tiến bộ đòi hỏi nam nữ tự quyết định, đủ điều kiện, không bị cưỡng ép hoặc lừa dối và đăng ký tại cơ quan có thẩm quyền.",
            "Một vợ một chồng là nguyên tắc bắt buộc; sự đồng ý của người vợ hoặc chồng hiện tại không làm hợp pháp quan hệ với người thứ ba.",
            "Vợ chồng bình đẳng về nhân thân, tài sản, lựa chọn nơi cư trú, nghề nghiệp, chăm sóc con và đại diện. Bình đẳng không phủ nhận việc thỏa thuận phân công phù hợp.",
            "Nhà nước tôn trọng hôn nhân giữa các dân tộc, tôn giáo và có yếu tố nước ngoài nếu đáp ứng điều kiện pháp luật.",
            "Gia đình có nghĩa vụ tôn trọng, chăm sóc, giúp đỡ; bảo vệ trẻ em, người cao tuổi, người khuyết tật; không phân biệt giữa các con.",
            "Nam nữ chung sống mà không đăng ký không phát sinh quyền nghĩa vụ vợ chồng; quan hệ con chung vẫn được bảo vệ; tài sản giải quyết theo thỏa thuận, công sức đóng góp và pháp luật dân sự.",
            "Các hành vi bị cấm gồm tảo hôn, cưỡng ép, kết hôn giả, vi phạm một vợ một chồng, bạo lực, lợi dụng hôn nhân để mua bán người hoặc trục lợi.",
        ],
        "example": "Ông A đang có vợ mà sống với bà B: B không trở thành vợ hợp pháp dù vợ A đồng ý; nhưng con của A và B có quyền như các con khác khi quan hệ cha con được xác định.",
        "pitfalls": ["Dùng Luật năm 2000.", "Cho vợ đồng ý thì được lấy vợ hai.", "Mặc nhiên chia đôi tài sản người chung sống không đăng ký.", "Dùng thuật ngữ phân biệt con."],
        "outline": ["Khái niệm.", "Tự nguyện–tiến bộ.", "Một vợ một chồng–bình đẳng.", "Tôn trọng dân tộc/tôn giáo.", "Bảo vệ gia đình và thành viên.", "Chung sống không đăng ký."],
    },
    19: {
        "basis": "Hiến pháp năm 2013 và Nghị quyết sửa đổi, bổ sung một số điều được Quốc hội thông qua ngày 16/6/2025.",
        "analysis": [
            "Hiến pháp là luật cơ bản, có hiệu lực cao nhất, quy định nền tảng chế độ chính trị, quyền con người, tổ chức quyền lực, kinh tế – xã hội và bảo vệ Tổ quốc.",
            "Chương về chế độ chính trị xác định bản chất Nhà nước, chủ quyền Nhân dân, nguyên tắc quyền lực thống nhất có phân công, phối hợp, kiểm soát, vai trò lãnh đạo của Đảng và Mặt trận.",
            "Chương quyền con người, quyền công dân ghi nhận các quyền dân sự, chính trị, kinh tế, xã hội, văn hóa và nghĩa vụ. Quyền con người gắn với mọi cá nhân; quyền công dân gắn với quốc tịch.",
            "Quyền chỉ có thể bị hạn chế theo luật trong trường hợp cần thiết vì các mục đích hiến định. Văn bản dưới luật không thể tự đặt giới hạn quyền vượt quá luật.",
            "Quốc hội thực hiện quyền lập hiến, lập pháp và giám sát tối cao; Chính phủ thực hiện quyền hành pháp; Tòa án thực hiện quyền tư pháp; Viện kiểm sát công tố và kiểm sát tư pháp; Chủ tịch nước đứng đầu Nhà nước.",
            "Các chương khác quy định kinh tế, xã hội, văn hóa, giáo dục, khoa học, công nghệ, môi trường; bảo vệ Tổ quốc; chính quyền địa phương; Hội đồng bầu cử quốc gia; Kiểm toán nhà nước.",
            "Sửa đổi năm 2025 liên quan sắp xếp tổ chức bộ máy, Mặt trận Tổ quốc và mô hình chính quyền địa phương hai cấp. Khi thi cần căn cứ yêu cầu giảng viên để trình bày nội dung hiện hành hoặc so sánh bản năm 2013.",
        ],
        "example": "Nếu một thông tư hạn chế quyền công dân mà không có căn cứ luật thì không chỉ có nguy cơ trái luật mà còn đặt ra vấn đề bảo đảm quyền theo Hiến pháp.",
        "pitfalls": ["Chỉ liệt kê tên 11 chương.", "Nhầm quyền con người và quyền công dân.", "Gọi Chính phủ là cơ quan quyền lực cao nhất.", "Bỏ cập nhật 2025."],
        "outline": ["Vị trí và hiệu lực.", "Chế độ chính trị.", "Quyền con người/công dân.", "Kinh tế–xã hội.", "Bộ máy nhà nước.", "Bảo vệ Tổ quốc.", "Sửa đổi 2025 và kết luận."],
    },
}


def add_question_detail(doc, number):
    detail = QUESTION_DETAILS.get(number)
    if not detail:
        return
    doc.add_heading("Cơ sở pháp lý", level=3)
    p = doc.add_paragraph()
    r = p.add_run(detail["basis"])
    set_font(r)
    doc.add_heading("Phân tích chi tiết", level=3)
    for item in detail["analysis"]:
        add_bullet(doc, item)
    doc.add_heading("Ví dụ minh họa", level=3)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(detail["example"])
    set_font(r)
    doc.add_heading("Những lỗi dễ mất điểm", level=3)
    for item in detail["pitfalls"]:
        add_bullet(doc, item)
    doc.add_heading("Dàn ý trả lời khi thi", level=3)
    for item in detail["outline"]:
        add_number(doc, item)


DETAIL_SECTION_MAP = {
    1: [[0], [1, 2, 3], [4]],
    2: [[0], [1, 2], [3, 4, 5]],
    3: [[0, 1], [2, 3, 4, 5]],
    4: [[], [0, 1, 2, 3, 4, 5]],
    5: [[0, 5], [], [1, 2, 3, 4]],
    6: [[0], [], [1, 2, 3, 4, 5]],
    7: [[0], [1, 2], [3, 4, 5]],
    8: [[0], [1], [2, 3, 4, 5]],
    9: [[0], [1, 2, 3], [4, 5]],
    10: [[0], [1, 2, 3], [4, 5]],
    11: [[0, 1], [2, 3, 4, 5, 6]],
    12: [[0], [1], [2, 3, 4, 5]],
    13: [[0], [1, 2, 3], [4, 5]],
    14: [[0, 1], [2], [3, 4, 5]],
    15: [[0], [1, 2], [3, 4, 5]],
    16: [[0], [1, 2, 3, 4, 5], [6]],
    17: [[0], [1], [2], [3, 4, 5, 6]],
    18: [[0], [1, 2, 3, 4, 5], []],
    19: [[0, 1, 2, 3, 4, 5, 6]],
}


QUESTION_ENRICHMENTS = {
    1: {
        0: [
            "Các dấu hiệu nhận diện Nhà nước gồm: có quyền lực công cộng đặc biệt; phân chia và quản lý dân cư theo lãnh thổ; có chủ quyền quốc gia; ban hành pháp luật; quy định và thu các khoản bắt buộc để duy trì bộ máy, thực hiện nhiệm vụ công.",
            "Quyền lực công cộng đặc biệt không đồng nhất với ảnh hưởng xã hội thông thường: quyền lực nhà nước có tính bắt buộc, được thực hiện bằng một bộ máy chuyên nghiệp và có khả năng áp dụng cưỡng chế hợp pháp.",
        ],
        1: [
            "Quyền lực kinh tế tạo cơ sở vật chất; quyền lực chính trị bảo đảm khả năng tổ chức và cưỡng chế; quyền lực tư tưởng góp phần tạo sự thừa nhận đối với trật tự xã hội. Ba phương diện này liên hệ với nhau nhưng không phải ba “bản chất” tách biệt.",
            "Ví dụ bổ sung: Trong tình trạng thiên tai, Nhà nước có thể huy động nguồn lực và áp dụng biện pháp bắt buộc để bảo đảm an toàn. Việc này vừa thể hiện quyền lực công, vừa nhằm giải quyết nhu cầu chung của xã hội.",
        ],
        2: [
            "Phương pháp đúng là phân tích Nhà nước trong hoàn cảnh lịch sử cụ thể, thông qua pháp luật, cơ cấu quyền lực, chính sách và kết quả hoạt động thực tế; tránh kết luận chỉ từ khẩu hiệu hoặc tên gọi của một thiết chế.",
        ],
    },
    2: {
        0: [
            "Cần phân biệt: quy phạm pháp luật là quy tắc xử sự chung; văn bản quy phạm pháp luật là hình thức văn bản chứa một hoặc nhiều quy phạm; văn bản áp dụng pháp luật là quyết định cá biệt giải quyết một vụ việc hoặc một chủ thể xác định.",
        ],
        1: [
            "Tính áp dụng nhiều lần nghĩa là quy tắc tiếp tục tồn tại sau mỗi lần được áp dụng, cho đến khi hết hiệu lực, bị bãi bỏ hoặc được thay thế. Đây là điểm khác với quyết định xử phạt hay quyết định bổ nhiệm chỉ nhằm giải quyết trường hợp cụ thể.",
            "Văn bản QPPL còn phải được công bố, đăng tải và xác định hiệu lực theo luật; việc có câu chữ mang tính mệnh lệnh nhưng không đúng loại văn bản và thẩm quyền không làm phát sinh một văn bản QPPL hợp pháp.",
        ],
        2: [
            "Bảo đảm tính ổn định không có nghĩa pháp luật không được sửa đổi; yêu cầu là chính sách phải được đánh giá kỹ, dự báo tác động và tránh thay đổi tùy tiện làm giảm khả năng dự liệu của người dân, doanh nghiệp.",
            "Ví dụ bổ sung: Một công văn của sở yêu cầu doanh nghiệp thực hiện nghĩa vụ mới mà luật và nghị định không quy định không thể tự trở thành căn cứ đặt thêm nghĩa vụ chung, dù công văn được gửi cho nhiều doanh nghiệp.",
        ],
    },
    3: {
        0: [
            "“Của Nhân dân” xác định chủ thể của quyền lực; “do Nhân dân” thể hiện nguồn gốc và cách hình thành quyền lực thông qua bầu cử, dân chủ trực tiếp và đại diện; “vì Nhân dân” yêu cầu mục tiêu, chính sách và hoạt động công quyền hướng tới quyền, lợi ích hợp pháp và đời sống của Nhân dân.",
        ],
        1: [
            "Tính pháp quyền không chỉ là quản lý xã hội bằng pháp luật mà còn là chính quyền bị ràng buộc bởi Hiến pháp và pháp luật, quyền lực phải được kiểm soát, quyết định công phải có căn cứ và người dân có cơ chế bảo vệ quyền.",
            "Ví dụ bổ sung: Cá nhân có quyền khiếu nại hoặc khởi kiện quyết định hành chính trái pháp luật. Cơ chế này cho thấy cơ quan công quyền không đứng trên pháp luật.",
        ],
    },
    4: {
        0: [
            "Luật Hình sự đồng thời xác định giới hạn của quyền trừng phạt nhà nước: chỉ hành vi được BLHS quy định và do chủ thể có lỗi thực hiện mới có thể làm phát sinh trách nhiệm hình sự.",
        ],
        1: [
            "Nguyên tắc không tránh khỏi trách nhiệm đòi hỏi tội phạm phải được phát hiện, xử lý kịp thời và đúng pháp luật; tuy nhiên xử lý nghiêm không đồng nghĩa áp dụng hình phạt nặng nhất trong mọi trường hợp.",
            "Trách nhiệm hình sự của pháp nhân thương mại không thay thế trách nhiệm của cá nhân. Nếu cá nhân và pháp nhân đều đủ điều kiện chịu trách nhiệm thì có thể xem xét độc lập đối với từng chủ thể.",
            "Ví dụ bổ sung: Người gây hậu quả vì một sự kiện bất ngờ mà họ không thể thấy trước và cũng không buộc phải thấy trước thì không có lỗi; chỉ có hậu quả nghiêm trọng chưa đủ để kết luận có tội.",
        ],
    },
    5: {
        0: [
            "Ba tiền đề thường được kiểm tra khi xác định một quan hệ pháp luật cụ thể là: có quy phạm điều chỉnh; có chủ thể đủ năng lực tương ứng; có sự kiện pháp lý làm phát sinh, thay đổi hoặc chấm dứt quan hệ.",
        ],
        1: [
            "Tính ý chí thể hiện ở ý chí Nhà nước được ghi nhận trong quy phạm và, trong nhiều quan hệ như hợp đồng, còn có ý chí của các bên. Với quan hệ phát sinh từ vi phạm, quan hệ trách nhiệm vẫn có thể hình thành dù người vi phạm không mong muốn.",
            "Ví dụ bổ sung: Khi một người chết, quan hệ nhân thân của họ chấm dứt nhưng sự kiện chết lại làm phát sinh quan hệ thừa kế giữa những người có quyền hưởng di sản.",
        ],
        2: [
            "Quyền chủ thể thường bao gồm khả năng tự thực hiện hành vi được phép, yêu cầu chủ thể có nghĩa vụ thực hiện đúng cam kết và yêu cầu cơ quan có thẩm quyền bảo vệ. Nghĩa vụ có thể là phải làm, không được làm hoặc phải chịu hậu quả pháp lý.",
            "Cần phân biệt khách thể với đối tượng vật chất. Trong quan hệ mua bán xe, chiếc xe là đối tượng của giao dịch; khách thể rộng hơn là lợi ích tài sản mà các bên hướng tới thông qua việc chuyển giao xe và thanh toán.",
        ],
    },
    6: {
        0: [
            "Hình thức pháp luật còn được gọi là nguồn pháp luật theo nghĩa hình thức, tức nơi chứa đựng hoặc phương thức ghi nhận các quy tắc được sử dụng làm căn cứ giải quyết vụ việc.",
        ],
        1: [
            "Hình thức pháp luật chịu ảnh hưởng của điều kiện kinh tế, chính trị, lịch sử và truyền thống pháp lý; không phải sản phẩm hoàn toàn tùy ý của người làm luật.",
            "Mỗi hình thức phải có cách nhận diện và điều kiện áp dụng nhất định. Chủ thể không thể viện dẫn một thói quen, một bản án bất kỳ hoặc một văn bản nội bộ như nguồn pháp luật nếu Nhà nước chưa thừa nhận giá trị pháp lý của nó.",
            "Các hình thức cùng tồn tại và bổ trợ nhau: văn bản QPPL tạo tính rõ ràng, án lệ góp phần thống nhất cách hiểu, tập quán giúp xử lý khoảng trống trong giới hạn luật cho phép.",
        ],
        2: [
            "Ví dụ bổ sung: Một bản án đã có hiệu lực nhưng chưa được lựa chọn và công bố thành án lệ không đương nhiên là án lệ để áp dụng cho mọi vụ việc tương tự.",
        ],
    },
    7: {
        0: [
            "Luật Dân sự giữ vai trò luật chung của lĩnh vực luật tư. Khi luật chuyên ngành không có quy định khác, các nguyên tắc và quy định dân sự có thể được sử dụng để điều chỉnh quan hệ tài sản, nhân thân tương ứng.",
        ],
        1: [
            "Quan hệ tài sản thường có thể quy đổi thành tiền nhưng không phải mọi quan hệ liên quan tiền đều là dân sự; thu thuế và xử phạt tiền là quan hệ hành chính vì một bên nhân danh quyền lực nhà nước.",
            "Quan hệ nhân thân không gắn với tài sản gồm họ tên, hình ảnh, danh dự, đời sống riêng tư. Quan hệ nhân thân gắn với tài sản có thể thấy ở quyền công bố tác phẩm và quyền hưởng nhuận bút của tác giả.",
        ],
        2: [
            "Các bên bình đẳng về địa vị pháp lý nhưng quyền và nghĩa vụ cụ thể không nhất thiết giống hệt nhau; bên bán có nghĩa vụ giao tài sản, bên mua có nghĩa vụ trả tiền.",
            "Ví dụ bổ sung: A đăng ảnh riêng tư của B không được đồng ý. Dù không có hợp đồng hay thiệt hại tài sản ngay lập tức, B vẫn có thể yêu cầu chấm dứt, xin lỗi, cải chính và bồi thường nếu đủ căn cứ.",
        ],
    },
    8: {
        0: [
            "Quy phạm pháp luật khác quy tắc đạo đức ở nguồn hình thành và cơ chế bảo đảm: quy phạm pháp luật do Nhà nước ban hành hoặc thừa nhận và có thể được bảo đảm bằng cưỡng chế nhà nước.",
        ],
        1: [
            "Quy phạm có tính khuôn mẫu: nó dự liệu một nhóm tình huống và chủ thể có đặc điểm chung, chứ không mô tả riêng một cá nhân đã được xác định bằng tên.",
        ],
        2: [
            "Giả định có thể chứa nhiều điều kiện đồng thời; chỉ khi các điều kiện luật định xuất hiện thì phần quy định hoặc chế tài mới được đặt ra.",
            "Chế tài không chỉ là hình phạt. Buộc bồi thường, khôi phục tình trạng ban đầu, hủy quyết định trái luật hoặc xử lý kỷ luật đều có thể là hậu quả pháp lý trong các ngành luật khác nhau.",
            "Ví dụ bổ sung: Quy định “người gây thiệt hại do lỗi của mình phải bồi thường” có giả định là có hành vi, thiệt hại, quan hệ nhân quả và lỗi; quy định là phải bồi thường; các hậu quả cụ thể được xác định ở những điều liên quan.",
        ],
    },
    9: {
        0: [
            "Sở hữu theo nghĩa kinh tế là quan hệ giữa người với người trong việc chiếm hữu của cải; quyền sở hữu theo nghĩa pháp lý là phạm vi quyền năng được pháp luật công nhận và bảo vệ cho chủ sở hữu.",
            "Không đồng nhất quyền sở hữu với việc đang cầm giữ tài sản. Người trông xe đang chiếm hữu trực tiếp chiếc xe nhưng chủ xe mới là người có quyền sở hữu.",
        ],
        1: [
            "Ba quyền năng có thể được chủ sở hữu tự thực hiện hoặc giao cho người khác trong giới hạn nhất định. Việc giao quyền sử dụng hoặc chiếm hữu không mặc nhiên làm chuyển quyền sở hữu.",
            "Quyền sở hữu không tuyệt đối: chủ sở hữu phải tuân thủ quy định về môi trường, trật tự, quyền liền kề, lợi ích công cộng và không được lạm dụng quyền gây thiệt hại cho người khác.",
        ],
        2: [
            "Có thể nhận diện căn cứ xác lập theo hai nhóm: nguyên sinh, khi quyền không bắt nguồn trực tiếp từ chủ sở hữu trước; và phái sinh, khi quyền được chuyển từ chủ thể trước qua hợp đồng, thừa kế hoặc căn cứ khác.",
            "Ví dụ bổ sung: A bán xe cho B theo hợp đồng hợp pháp thì quyền của A chấm dứt và quyền của B được xác lập theo căn cứ chuyển giao; nếu xe bị tịch thu đúng luật thì quyền sở hữu chấm dứt không phụ thuộc ý chí của A.",
        ],
    },
    10: {
        0: [
            "Khái niệm này cho thấy không thể chỉ căn cứ vào thiệt hại để kết luận vi phạm; phải đồng thời kiểm tra tính trái pháp luật, lỗi và năng lực chịu trách nhiệm của chủ thể.",
        ],
        1: [
            "Hành vi gây thiệt hại nhưng được thực hiện trong phòng vệ chính đáng, tình thế cấp thiết hoặc trường hợp pháp luật cho phép có thể không mang tính trái pháp luật.",
            "Lỗi cố ý gồm trường hợp chủ thể mong muốn hậu quả hoặc tuy không mong muốn nhưng có ý thức để mặc; lỗi vô ý gồm quá tự tin và cẩu thả. Việc xác định lỗi phải dựa trên khả năng nhận thức, điều khiển và hoàn cảnh cụ thể.",
            "Ví dụ bổ sung: Một em nhỏ chưa đạt độ tuổi chịu trách nhiệm hành chính thực hiện hành vi bị cấm thì hành vi khách quan có thể trái quy định, nhưng không đủ điều kiện xác định em là chủ thể vi phạm hành chính.",
        ],
        2: [
            "Phân loại vi phạm dựa trước hết vào ngành luật bảo vệ quan hệ bị xâm hại và tính chất, mức độ nguy hiểm. Không được căn cứ duy nhất vào tên cơ quan đang xử lý.",
            "Một hành vi chỉ bị áp dụng một lần đối với cùng một loại trách nhiệm, nhưng có thể đồng thời phát sinh các trách nhiệm khác nhau, chẳng hạn hình sự và bồi thường dân sự.",
        ],
    },
    11: {
        0: [
            "Thời điểm mở thừa kế là thời điểm người có tài sản chết; địa điểm mở thừa kế thường là nơi cư trú cuối cùng của người để lại di sản, hoặc nơi có toàn bộ hay phần lớn di sản nếu không xác định được nơi cư trú cuối cùng.",
            "Di sản không phải toàn bộ khối tài sản đang thấy trong gia đình. Phải xác định tài sản riêng, phần của người chết trong tài sản chung và khấu trừ nghĩa vụ tài sản trước khi chia.",
        ],
        1: [
            "Di chúc có hiệu lực từ thời điểm mở thừa kế. Một phần di chúc vô hiệu không nhất thiết làm toàn bộ di chúc vô hiệu nếu phần còn lại độc lập và vẫn có thể thực hiện.",
            "Người lập di chúc chỉ được định đoạt tài sản thuộc quyền của mình. Nếu di chúc định đoạt cả phần của người khác trong tài sản chung thì phần vượt quá quyền định đoạt không có hiệu lực.",
            "Ví dụ bổ sung: Ông A có 1/2 căn nhà chung trị giá 2 tỷ đồng và 200 triệu đồng riêng. A chỉ có thể định đoạt di sản 1,2 tỷ đồng; phần 1 tỷ đồng của người đồng sở hữu không phải di sản.",
            "Ví dụ về người thừa kế bắt buộc: Nếu A để toàn bộ di sản cho bạn nhưng còn con 10 tuổi, người con vẫn có thể được hưởng phần bằng hai phần ba một suất theo pháp luật, trừ trường hợp từ chối hoặc không có quyền hưởng.",
        ],
    },
    12: {
        0: [
            "Quan hệ lao động được nhận diện bằng ba dấu hiệu trung tâm: có công việc được thực hiện; có trả công hoặc tiền lương; có sự quản lý, điều hành, giám sát của người sử dụng lao động.",
        ],
        1: [
            "Đối tượng còn bao gồm quan hệ lao động tập thể giữa tổ chức đại diện người lao động với người sử dụng lao động hoặc tổ chức đại diện của họ, đặc biệt trong đối thoại, thương lượng và thỏa ước tập thể.",
            "Không phải mọi hoạt động có trả tiền đều là lao động: người nhận khoán độc lập, tự quyết định cách làm và chịu rủi ro kinh doanh có thể thuộc quan hệ dịch vụ dân sự.",
        ],
        2: [
            "Sự kết hợp giữa thỏa thuận và quyền quản lý phản ánh tính đặc thù: khi giao kết các bên bình đẳng về mặt pháp lý, nhưng trong quá trình làm việc người lao động phải tuân theo sự điều hành hợp pháp của người sử dụng lao động.",
            "Ví dụ bổ sung: Doanh nghiệp có thể ban hành nội quy và bố trí ca làm đúng luật, nhưng không thể phạt tiền hoặc cắt lương thay cho xử lý kỷ luật lao động.",
        ],
    },
    13: {
        0: [
            "Dấu hiệu quản lý được thể hiện qua chấm công, ca làm cố định, phân công nhiệm vụ, kiểm tra, đánh giá, nội quy và quyền xử lý vi phạm. Không cần hợp đồng mang đúng tên “hợp đồng lao động” mới được pháp luật lao động bảo vệ.",
        ],
        1: [
            "Hợp đồng điện tử dưới hình thức thông điệp dữ liệu có giá trị như hợp đồng bằng văn bản nếu đáp ứng điều kiện luật định.",
            "Khi hợp đồng xác định thời hạn hết hạn mà người lao động vẫn tiếp tục làm việc, các bên phải xử lý việc ký mới trong thời hạn luật định; nếu không, loại hợp đồng có thể chuyển đổi theo quy định.",
        ],
        2: [
            "Cần phân biệt đơn phương chấm dứt hợp đồng với xử lý kỷ luật sa thải: đây là hai căn cứ, điều kiện và thủ tục khác nhau. Không được dùng một thông báo “cho nghỉ ngay” để thay thế thủ tục sa thải.",
            "Ví dụ bổ sung: Người lao động nghỉ vì không được trả đủ lương hoặc bị quấy rối tại nơi làm việc có thể thuộc trường hợp luật cho phép đơn phương không cần báo trước nếu đủ dữ kiện.",
        ],
    },
    14: {
        0: [
            "Tranh chấp chỉ trở thành tranh chấp lao động khi bất đồng gắn với quyền, nghĩa vụ hoặc lợi ích phát sinh từ quan hệ lao động hoặc quan hệ liên quan trực tiếp đến lao động.",
            "Điểm mấu chốt để phân biệt tranh chấp tập thể về quyền và về lợi ích là: quyền đã có căn cứ trong pháp luật, thỏa ước, nội quy hay chưa. Nếu tập thể yêu cầu thiết lập điều kiện mới thì đó là tranh chấp về lợi ích.",
        ],
        1: [
            "Thẩm quyền phải được xác định theo đúng loại tranh chấp. Hội đồng trọng tài chỉ giải quyết khi luật cho phép và các bên lựa chọn; Tòa án không phải lúc nào cũng là bước đầu tiên.",
        ],
        2: [
            "Đối với tranh chấp cá nhân thuộc diện phải hòa giải, hòa giải viên tiến hành hòa giải trong thời hạn luật định. Khi hòa giải không thành, hết thời hạn mà không hòa giải hoặc thuộc trường hợp được miễn, bên tranh chấp mới lựa chọn cơ chế tiếp theo.",
            "Khi các bên đã thống nhất yêu cầu Hội đồng trọng tài giải quyết thì trong thời gian trọng tài đang giải quyết không đồng thời yêu cầu Tòa án, trừ trường hợp luật cho phép chuyển sang Tòa án do cơ chế trọng tài không được thực hiện.",
            "Ví dụ bổ sung: Người lao động bị sa thải yêu cầu hủy quyết định sa thải có thể khởi kiện mà không bắt buộc qua hòa giải; còn tranh chấp về tiền thưởng thông thường phải kiểm tra có thuộc trường hợp miễn hòa giải hay không.",
        ],
    },
    15: {
        0: [
            "Hoạt động chấp hành - điều hành là việc tổ chức thi hành Hiến pháp, luật và quản lý thường xuyên các lĩnh vực đời sống; vì vậy Luật Hành chính gắn chặt với thẩm quyền, thủ tục và trách nhiệm công vụ.",
        ],
        1: [
            "Đối tượng có thể chia thành ba nhóm: quan hệ quản lý phát sinh giữa cấp trên với cấp dưới; giữa cơ quan có thẩm quyền với cá nhân, tổ chức; và quan hệ khi tổ chức hoặc cá nhân được Nhà nước trao quyền thực hiện nhiệm vụ quản lý.",
            "Ví dụ bổ sung: Hiệu trưởng trường công lập xử lý kỷ luật viên chức theo thẩm quyền là quan hệ quản lý; trường ký hợp đồng mua bàn ghế với doanh nghiệp lại là quan hệ dân sự.",
        ],
        2: [
            "Sự không bình đẳng về ý chí không có nghĩa bên quản lý được tùy tiện. Quyết định đơn phương chỉ hợp pháp khi đúng thẩm quyền, đúng căn cứ, đúng mục đích và tuân thủ trình tự, thủ tục.",
        ],
    },
    16: {
        0: [
            "Không phải mọi hành vi có vẻ nguy hiểm đều là tội phạm. Nếu mức độ nguy hiểm không đáng kể thì có thể không bị coi là tội phạm và được xử lý bằng biện pháp khác theo luật.",
        ],
        1: [
            "Mặt khách quan có thể gồm hành vi, hậu quả, quan hệ nhân quả, công cụ, phương tiện, thời gian, địa điểm; nhưng dấu hiệu nào bắt buộc phụ thuộc cấu thành của từng tội.",
            "Mặt chủ quan ngoài lỗi còn có thể đòi hỏi động cơ hoặc mục đích. Chẳng hạn nhiều tội chiếm đoạt yêu cầu mục đích chiếm đoạt, nên không thể chỉ thấy tài sản bị mất mà kết luận tội danh.",
            "Ví dụ bổ sung: A mượn xe rồi trả chậm vì tai nạn, vẫn liên lạc và có khả năng hoàn trả. Chỉ dữ kiện trả chậm chưa chứng minh mục đích chiếm đoạt để cấu thành tội phạm.",
        ],
        2: [
            "Việc phân loại căn cứ vào mức cao nhất của khung hình phạt do BLHS quy định cho tội danh, không căn cứ vào mức án Tòa án thực tế tuyên cho bị cáo.",
        ],
    },
    17: {
        0: [
            "Tố cáo có thể nhằm bảo vệ lợi ích Nhà nước, cộng đồng hoặc quyền của người khác; người tố cáo không bắt buộc phải là người trực tiếp bị thiệt hại.",
        ],
        1: [
            "Người tố cáo phải trình bày trung thực, cung cấp thông tin mình có và chịu trách nhiệm về nội dung cố ý tố cáo sai; đồng thời được bảo mật thông tin và bảo vệ khi có nguy cơ bị trả thù, trù dập.",
            "Ví dụ bổ sung: Nhân viên phát hiện trưởng phòng lập hồ sơ khống để rút ngân sách có thể tố cáo dù tiền bị chiếm đoạt là tài sản Nhà nước chứ không phải tài sản của người tố cáo.",
        ],
        2: [
            "Nếu người bị tố cáo thuộc nhiều cơ quan quản lý hoặc hành vi liên quan nhiều cấp, phải căn cứ quan hệ quản lý và quy định chuyên ngành; cơ quan nhận đơn sai thẩm quyền có trách nhiệm xử lý, chuyển theo luật chứ không tự giải quyết.",
        ],
        3: [
            "Kết luận nội dung tố cáo phải xác định nội dung đúng, sai hoặc đúng một phần, trách nhiệm của từng chủ thể và biện pháp xử lý; đây là cơ sở cho quyết định xử lý chứ không phải bản thân đơn tố cáo.",
        ],
    },
    18: {
        0: [
            "Quan hệ nhân thân là nội dung trung tâm và thường không thể chuyển giao, như quan hệ vợ chồng, cha mẹ - con; quan hệ tài sản trong gia đình chịu ảnh hưởng của quan hệ nhân thân và mục tiêu bảo vệ gia đình.",
        ],
        1: [
            "Hôn nhân tiến bộ còn thể hiện ở việc pháp luật công nhận quyền yêu cầu ly hôn khi mục đích hôn nhân không đạt được; tự nguyện kết hôn không đồng nghĩa phải duy trì hôn nhân bằng mọi giá.",
            "Bình đẳng giữa vợ và chồng bao gồm bình đẳng trong lựa chọn nghề nghiệp, nơi cư trú, đại diện, chăm sóc con và tạo lập, chiếm hữu, sử dụng, định đoạt tài sản chung.",
            "Ví dụ bổ sung: Vợ chồng có thể thỏa thuận một người chăm sóc con, một người tạo thu nhập; sự phân công tự nguyện này không trái bình đẳng. Trái lại, việc cấm vợ đi làm chỉ vì định kiến giới có thể xâm phạm quyền bình đẳng.",
        ],
        2: [
            "Tính ưu việt còn thể hiện ở việc quyền của con được bảo vệ không phụ thuộc tình trạng hôn nhân của cha mẹ; mọi người con đều bình đẳng về chăm sóc, cấp dưỡng và thừa kế khi quan hệ cha mẹ - con được xác định.",
        ],
    },
    19: {
        0: [
            "Hiến pháp năm 2013 gồm Lời nói đầu và các chương quy định những vấn đề nền tảng nhất; vì vậy khi phân tích không nên chỉ kể tên chương mà phải nêu tư tưởng và nguyên tắc chi phối toàn bộ hệ thống pháp luật.",
            "Về chế độ chính trị, nội dung trọng tâm là chủ quyền Nhân dân, Nhà nước pháp quyền xã hội chủ nghĩa, quyền lực thống nhất có phân công, phối hợp, kiểm soát, vai trò lãnh đạo của Đảng và vị trí của Mặt trận Tổ quốc Việt Nam.",
            "Về quyền con người, Hiến pháp chuyển mạnh từ cách chỉ ghi nhận sang yêu cầu công nhận, tôn trọng, bảo vệ và bảo đảm. Nghĩa vụ của Nhà nước không chỉ là không xâm phạm mà còn phải xây dựng cơ chế để quyền được thực hiện.",
            "Cần phân biệt quyền con người với quyền công dân: quyền con người về nguyên tắc dành cho mọi cá nhân; quyền công dân gắn với quốc tịch Việt Nam, như quyền bầu cử và ứng cử.",
            "Nguyên tắc hạn chế quyền đòi hỏi phải do luật quy định và chỉ trong trường hợp cần thiết vì quốc phòng, an ninh quốc gia, trật tự, an toàn xã hội, đạo đức xã hội hoặc sức khỏe cộng đồng.",
            "Về kinh tế - xã hội, Hiến pháp ghi nhận nhiều hình thức sở hữu, thành phần kinh tế, vai trò quản lý của Nhà nước, chính sách an sinh, văn hóa, giáo dục, khoa học, công nghệ và bảo vệ môi trường.",
            "Về bộ máy, mỗi cơ quan có vị trí và chức năng hiến định khác nhau nhưng đều nằm trong cơ chế quyền lực nhà nước thống nhất và chịu các hình thức kiểm soát, giám sát theo Hiến pháp.",
            "Ví dụ bổ sung: Quyền tự do kinh doanh được bảo đảm trong những ngành nghề mà pháp luật không cấm; cơ quan dưới luật không thể tùy ý tạo thêm điều kiện kinh doanh nếu không có căn cứ luật.",
            "Cập nhật 2025 cần được trình bày như phần sửa đổi của Hiến pháp hiện hành, tập trung vào tổ chức bộ máy, Mặt trận Tổ quốc và mô hình chính quyền địa phương hai cấp; không nên hiểu là ban hành một bản Hiến pháp hoàn toàn mới.",
        ],
    },
}


def add_question(doc, number, title, sections, keywords):
    doc.add_heading(f"Câu {number}. {title}", level=2)
    detail = QUESTION_DETAILS.get(number, {})
    section_map = DETAIL_SECTION_MAP.get(number, [])
    enrichments = QUESTION_ENRICHMENTS.get(number, {})
    for section_index, (heading, content) in enumerate(sections):
        doc.add_heading(heading, level=3)
        if isinstance(content, str):
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(content)
            set_font(r)
        else:
            for item in content:
                add_bullet(doc, item)
        if section_index < len(section_map):
            for detail_index in section_map[section_index]:
                if detail_index < len(detail.get("analysis", [])):
                    add_bullet(doc, detail["analysis"][detail_index])
        for item in enrichments.get(section_index, []):
            add_bullet(doc, item)
        if section_index == len(sections) - 1 and detail.get("example"):
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run("Ví dụ: ")
            set_font(r, bold=True, color=DARK)
            r = p.add_run(detail["example"])
            set_font(r)
    add_key(doc, keywords)


def add_intro(doc):
    doc.add_heading("HƯỚNG DẪN SỬ DỤNG", level=1)
    add_bullet(doc, "Học theo cấu trúc: khái niệm → đặc điểm/nguyên tắc → phân loại/nội dung → ví dụ → kết luận.")
    add_bullet(doc, "Khi làm bài tình huống, không suy đoán dữ kiện. Nêu giả định nếu đề bài chưa cho đủ thông tin.")
    add_bullet(doc, "Nếu đề thi yêu cầu theo giáo trình của trường, ưu tiên thuật ngữ và phạm vi do giảng viên đã dạy.")
    add_bullet(doc, "Không học thuộc các điều luật đã hết hiệu lực trong bản tổng hợp cũ.")

    doc.add_heading("NHỮNG ĐIỂM ĐÃ HIỆU ĐÍNH SO VỚI FILE CŨ", level=1)
    rows = [
        ("BLDS 2005", "Dùng Bộ luật Dân sự 2015."),
        ("BLHS 1999", "Dùng BLHS 2015, sửa đổi 2017 và các sửa đổi có hiệu lực hiện hành."),
        ("Luật HN&GĐ 2000", "Dùng Luật Hôn nhân và gia đình 2014."),
        ("Luật Khiếu nại, tố cáo cũ", "Dùng Luật Tố cáo 2018 cho nội dung tố cáo."),
        ("Tranh chấp lao động", "Dùng Bộ luật Lao động 2019; bỏ mô hình Hội đồng hòa giải cơ sở cũ."),
        ("Hiến pháp 2013", "Đọc cùng nội dung sửa đổi, bổ sung năm 2025 về Mặt trận Tổ quốc và chính quyền địa phương."),
        ("Thiếu Câu 13", "Bổ sung câu trả lời về hợp đồng lao động."),
        ("Bài tập hình sự", "Sửa lỗi nhầm giữa cướp, cướp giật, công nhiên chiếm đoạt, trộm cắp, lừa đảo và lạm dụng tín nhiệm."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(("Nội dung cũ", "Hướng sử dụng hiện nay")):
        set_cell_shading(table.rows[0].cells[i], BLUE)
        r = table.rows[0].cells[i].paragraphs[0].add_run(text)
        set_font(r, size=10, bold=True, color="FFFFFF")
    for a, b in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (a, b)):
            r = cell.paragraphs[0].add_run(text)
            set_font(r, size=10)
    set_table_widths(table, [2.05, 4.65])
    set_repeat_table_header(table.rows[0])

    doc.add_heading("DANH MỤC VĂN BẢN NỀN TẢNG", level=1)
    laws = [
        "Hiến pháp năm 2013 và Nghị quyết sửa đổi, bổ sung một số điều của Hiến pháp năm 2025.",
        "Luật Ban hành văn bản quy phạm pháp luật số 64/2025/QH15 và Luật sửa đổi, bổ sung số 87/2025/QH15.",
        "Bộ luật Dân sự năm 2015.",
        "Bộ luật Hình sự năm 2015, Luật sửa đổi năm 2017 và Luật sửa đổi, bổ sung số 86/2025/QH15.",
        "Bộ luật Lao động năm 2019.",
        "Luật Hôn nhân và gia đình năm 2014.",
        "Luật Tố cáo năm 2018; Luật Khiếu nại năm 2011.",
        "Luật Xử lý vi phạm hành chính năm 2012 và các luật sửa đổi, bổ sung hiện hành.",
    ]
    for law in laws:
        add_bullet(doc, law)
    p = doc.add_paragraph()
    r = p.add_run("Tra cứu bản chính thức tại: ")
    set_font(r, size=10)
    r = p.add_run("https://vanban.chinhphu.vn/  |  https://vbpl.vn/")
    set_font(r, size=10, color=BLUE)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_theory(doc):
    doc.add_heading("PHẦN I. ĐÁP ÁN ĐẦY ĐỦ 19 CÂU LÝ THUYẾT", level=1)
    add_question(doc, 1, "Khái niệm và bản chất của Nhà nước; ý nghĩa phương pháp luận", [
        ("Khái niệm", "Nhà nước là tổ chức đặc biệt của quyền lực chính trị, có bộ máy chuyên làm nhiệm vụ quản lý và cưỡng chế, quản lý dân cư theo lãnh thổ, ban hành pháp luật và thu thuế nhằm duy trì trật tự, tổ chức đời sống chung của xã hội."),
        ("Bản chất", [
            "Tính giai cấp: trong xã hội có giai cấp, nhà nước trước hết là công cụ quyền lực của lực lượng giữ địa vị thống trị về kinh tế và chính trị; tính giai cấp thể hiện qua quyền lực kinh tế, chính trị và tư tưởng.",
            "Tính xã hội: nhà nước phải giải quyết những công việc chung như bảo đảm an ninh, trật tự, hạ tầng, giáo dục, y tế, bảo vệ môi trường và điều hòa lợi ích xã hội.",
            "Hai phương diện có quan hệ thống nhất. Không được viết rằng nhà nước chỉ có tính giai cấp hoặc chỉ phục vụ lợi ích chung một cách phi giai cấp.",
        ]),
        ("Ý nghĩa phương pháp luận", "Khi phân tích một nhà nước cụ thể phải đặt trong điều kiện kinh tế – xã hội, tương quan giai cấp, tổ chức quyền lực và chức năng thực tế; đồng thời xem xét cả tính giai cấp và vai trò xã hội."),
    ], "quyền lực công cộng đặc biệt; quản lý dân cư; chủ quyền; pháp luật; thuế; tính giai cấp; tính xã hội")

    add_question(doc, 2, "Văn bản quy phạm pháp luật và nguyên tắc xây dựng, ban hành", [
        ("Khái niệm", "Văn bản quy phạm pháp luật là văn bản có chứa quy phạm pháp luật, được ban hành đúng thẩm quyền, hình thức, trình tự và thủ tục theo luật. Nếu ban hành không đúng thẩm quyền, hình thức hoặc thủ tục luật định thì không được coi là văn bản quy phạm pháp luật."),
        ("Đặc điểm", [
            "Chứa quy tắc xử sự chung, áp dụng nhiều lần đối với nhiều chủ thể.",
            "Do chủ thể có thẩm quyền ban hành và có tên loại văn bản do pháp luật quy định.",
            "Có hiệu lực bắt buộc và được Nhà nước bảo đảm thực hiện.",
            "Phải phù hợp với Hiến pháp, văn bản có hiệu lực pháp lý cao hơn và bảo đảm tính thống nhất của hệ thống pháp luật.",
        ]),
        ("Nguyên tắc trọng tâm", [
            "Bảo đảm sự lãnh đạo của Đảng, chủ quyền nhân dân và kiểm soát quyền lực nhà nước.",
            "Bảo đảm tính hợp hiến, hợp pháp, thống nhất, công khai, minh bạch, khả thi và ổn định.",
            "Đúng thẩm quyền, nội dung, hình thức, trình tự và thủ tục.",
            "Tôn trọng, bảo vệ quyền con người, quyền công dân; bảo đảm quốc phòng, an ninh và thực hiện điều ước quốc tế mà Việt Nam là thành viên.",
            "Quy trình cụ thể phụ thuộc từng loại văn bản; không dùng một quy trình của luật để áp dụng cho mọi văn bản.",
        ]),
    ], "quy tắc xử sự chung; áp dụng nhiều lần; đúng thẩm quyền; hợp hiến; hợp pháp; thống nhất; minh bạch; khả thi")

    add_question(doc, 3, "Bản chất và đặc trưng của Nhà nước CHXHCN Việt Nam", [
        ("Bản chất hiến định", "Nhà nước Cộng hòa xã hội chủ nghĩa Việt Nam là Nhà nước pháp quyền xã hội chủ nghĩa của Nhân dân, do Nhân dân, vì Nhân dân. Tất cả quyền lực nhà nước thuộc về Nhân dân; nền tảng là liên minh giữa giai cấp công nhân với giai cấp nông dân và đội ngũ trí thức."),
        ("Đặc trưng", [
            "Quyền lực nhà nước là thống nhất, có sự phân công, phối hợp và kiểm soát giữa các cơ quan thực hiện quyền lập pháp, hành pháp và tư pháp.",
            "Nhà nước tổ chức và hoạt động theo Hiến pháp, pháp luật; quản lý xã hội bằng Hiến pháp và pháp luật.",
            "Công nhận, tôn trọng, bảo vệ và bảo đảm quyền con người, quyền công dân.",
            "Đảng Cộng sản Việt Nam lãnh đạo Nhà nước và xã hội, gắn bó và chịu sự giám sát của Nhân dân, chịu trách nhiệm trước Nhân dân về quyết định của mình.",
            "Thực hiện đại đoàn kết dân tộc, bình đẳng giữa các dân tộc và đường lối đối ngoại hòa bình, hợp tác.",
            "Các đặc trưng trên là đặc trưng của Nhà nước; không dùng tám đặc trưng của xã hội xã hội chủ nghĩa để thay thế toàn bộ nội dung này.",
        ]),
    ], "Nhân dân; pháp quyền XHCN; quyền lực thống nhất; phân công–phối hợp–kiểm soát; quyền con người; Đảng lãnh đạo")

    add_question(doc, 4, "Khái niệm và các nguyên tắc cơ bản của Luật Hình sự", [
        ("Khái niệm", "Luật Hình sự là ngành luật gồm các quy phạm xác định hành vi nào là tội phạm, quy định hình phạt, biện pháp tư pháp, căn cứ trách nhiệm hình sự và các vấn đề liên quan."),
        ("Nguyên tắc", [
            "Pháp chế: chỉ Bộ luật Hình sự mới quy định tội phạm và hình phạt; xử lý đúng căn cứ, thẩm quyền và thủ tục.",
            "Bình đẳng trước pháp luật hình sự.",
            "Có lỗi: không có lỗi thì không có trách nhiệm hình sự.",
            "Cá thể hóa trách nhiệm và hình phạt theo tính chất hành vi, nhân thân và tình tiết vụ án.",
            "Công bằng, nhân đạo; hình phạt không nhằm trả thù hoặc hạ thấp nhân phẩm.",
            "Trách nhiệm hình sự không chỉ áp dụng với cá nhân; pháp nhân thương mại có thể chịu trách nhiệm đối với các tội luật định.",
        ]),
    ], "tội phạm; hình phạt; pháp chế; bình đẳng; có lỗi; cá thể hóa; công bằng; nhân đạo; pháp nhân thương mại")

    add_question(doc, 5, "Quan hệ pháp luật: khái niệm, đặc điểm và các yếu tố cấu thành", [
        ("Khái niệm", "Quan hệ pháp luật là quan hệ xã hội được quy phạm pháp luật điều chỉnh, trong đó các bên có quyền và nghĩa vụ pháp lý được Nhà nước bảo đảm thực hiện."),
        ("Đặc điểm", [
            "Phát sinh trên cơ sở quy phạm pháp luật và sự kiện pháp lý.",
            "Mang tính ý chí; có chủ thể xác định.",
            "Các bên có quyền chủ thể và nghĩa vụ pháp lý tương ứng.",
            "Được Nhà nước bảo đảm thực hiện.",
        ]),
        ("Các yếu tố", [
            "Chủ thể: cá nhân, pháp nhân, tổ chức hoặc Nhà nước có năng lực chủ thể phù hợp. Pháp nhân phải đáp ứng Điều 74 BLDS 2015.",
            "Nội dung: quyền chủ thể và nghĩa vụ pháp lý.",
            "Khách thể: lợi ích vật chất, tinh thần hoặc kết quả hành vi mà các chủ thể hướng tới.",
        ]),
    ], "quy phạm pháp luật + sự kiện pháp lý; chủ thể; năng lực pháp luật; năng lực hành vi; nội dung; khách thể")

    add_question(doc, 6, "Hình thức pháp luật và các loại hình thức pháp luật", [
        ("Khái niệm", "Hình thức pháp luật là cách thức biểu hiện ra bên ngoài của pháp luật, qua đó các quy tắc pháp lý được nhận biết và áp dụng."),
        ("Đặc điểm", [
            "Là phương thức tồn tại và biểu hiện chính thức của pháp luật, giúp xác định căn cứ được dùng để giải quyết vụ việc.",
            "Được hình thành và thừa nhận trong những điều kiện kinh tế, chính trị, lịch sử và truyền thống pháp lý nhất định.",
            "Mỗi hình thức có điều kiện nhận diện, giá trị và phạm vi áp dụng riêng; không phải mọi tập quán, bản án hoặc văn bản đều là nguồn pháp luật.",
        ]),
        ("Các loại", [
            "Tập quán pháp: tập quán được Nhà nước thừa nhận và bảo đảm thực hiện.",
            "Tiền lệ pháp/án lệ: cách giải quyết trong bản án, quyết định được lựa chọn và công bố để nghiên cứu, áp dụng cho vụ việc tương tự theo điều kiện pháp luật.",
            "Văn bản quy phạm pháp luật: hình thức chủ yếu trong hệ thống pháp luật Việt Nam.",
            "Ngoài ra, điều ước quốc tế và các nguyên tắc chung có vai trò theo quy định của từng lĩnh vực.",
        ]),
    ], "hình thức bên ngoài; tập quán pháp; án lệ; văn bản quy phạm pháp luật; điều ước quốc tế")

    add_question(doc, 7, "Luật Dân sự: đối tượng và phương pháp điều chỉnh", [
        ("Khái niệm", "Luật Dân sự là ngành luật điều chỉnh quan hệ tài sản và quan hệ nhân thân được hình thành trên cơ sở bình đẳng, tự do ý chí, độc lập về tài sản và tự chịu trách nhiệm."),
        ("Đối tượng", [
            "Quan hệ tài sản: quan hệ giữa các chủ thể thông qua tài sản, thường có tính giá trị và trao đổi.",
            "Quan hệ nhân thân: quan hệ gắn với giá trị nhân thân; có thể liên quan hoặc không liên quan đến tài sản.",
        ]),
        ("Phương pháp", [
            "Bình đẳng về địa vị pháp lý.",
            "Tự do, tự nguyện cam kết, thỏa thuận nhưng không được vi phạm điều cấm hoặc trái đạo đức xã hội.",
            "Độc lập về tài sản và tự chịu trách nhiệm.",
            "Tòa án chỉ giải quyết khi có yêu cầu và trong phạm vi yêu cầu, trừ trường hợp luật quy định khác.",
        ]),
    ], "quan hệ tài sản; quan hệ nhân thân; bình đẳng; tự do ý chí; độc lập tài sản; tự chịu trách nhiệm")

    add_question(doc, 8, "Quy phạm pháp luật: khái niệm, đặc điểm và cấu trúc", [
        ("Khái niệm", "Quy phạm pháp luật là quy tắc xử sự chung, có hiệu lực bắt buộc chung, do Nhà nước ban hành hoặc thừa nhận và bảo đảm thực hiện nhằm điều chỉnh quan hệ xã hội."),
        ("Đặc điểm", [
            "Tính bắt buộc chung và áp dụng nhiều lần.",
            "Thể hiện ý chí nhà nước; xác định quyền, nghĩa vụ hoặc hậu quả pháp lý.",
            "Có hình thức chặt chẽ và được bảo đảm bằng quyền lực nhà nước.",
        ]),
        ("Cấu trúc logic", [
            "Giả định: chủ thể, điều kiện, hoàn cảnh áp dụng – trả lời ai, khi nào, trong trường hợp nào.",
            "Quy định: cách xử sự được phép, bắt buộc hoặc bị cấm.",
            "Chế tài: hậu quả pháp lý bất lợi khi không thực hiện đúng quy định.",
            "Một quy phạm không nhất thiết có đủ ba bộ phận trong cùng một điều luật; một điều luật cũng có thể chứa nhiều quy phạm.",
        ]),
    ], "quy tắc xử sự chung; giả định; quy định; chế tài; quy phạm không đồng nhất với điều luật")

    add_question(doc, 9, "Sở hữu và quyền sở hữu theo Bộ luật Dân sự 2015", [
        ("Khái niệm", "Sở hữu là quan hệ xã hội về việc chiếm hữu của cải. Quyền sở hữu là quyền của chủ sở hữu đối với tài sản, gồm quyền chiếm hữu, quyền sử dụng và quyền định đoạt."),
        ("Nội dung", [
            "Chiếm hữu: nắm giữ, chi phối tài sản trực tiếp hoặc gián tiếp.",
            "Sử dụng: khai thác công dụng, hưởng hoa lợi, lợi tức.",
            "Định đoạt: chuyển giao quyền sở hữu, từ bỏ quyền sở hữu, tiêu dùng hoặc tiêu hủy tài sản.",
        ]),
        ("Căn cứ xác lập và chấm dứt", [
            "Xác lập từ lao động, sản xuất kinh doanh; chuyển giao; thu hoa lợi, lợi tức; tạo vật mới; thừa kế; chiếm hữu theo luật; bản án/quyết định và các căn cứ khác.",
            "Chấm dứt khi chuyển quyền, từ bỏ, tài sản bị tiêu hủy, xử lý để thực hiện nghĩa vụ, trưng mua, tịch thu hoặc người khác được xác lập quyền theo luật.",
        ]),
    ], "BLDS 2015; chiếm hữu; sử dụng; định đoạt; căn cứ xác lập; căn cứ chấm dứt")

    add_question(doc, 10, "Vi phạm pháp luật: khái niệm, dấu hiệu và phân loại", [
        ("Khái niệm", "Vi phạm pháp luật là hành vi trái pháp luật, có lỗi, do chủ thể có năng lực trách nhiệm pháp lý thực hiện, xâm hại quan hệ xã hội được pháp luật bảo vệ."),
        ("Dấu hiệu", [
            "Là hành vi xác định của con người hoặc tổ chức, dưới dạng hành động hoặc không hành động.",
            "Trái pháp luật và xâm hại quan hệ xã hội được bảo vệ.",
            "Có lỗi cố ý hoặc vô ý.",
            "Do chủ thể có năng lực trách nhiệm pháp lý thực hiện.",
        ]),
        ("Phân loại", [
            "Vi phạm hình sự (tội phạm) là hành vi nguy hiểm cho xã hội được quy định trong Bộ luật Hình sự, do chủ thể có năng lực trách nhiệm hình sự thực hiện có lỗi và phải chịu trách nhiệm hình sự.",
            "Vi phạm hành chính là hành vi có lỗi của cá nhân hoặc tổ chức, vi phạm quy định quản lý nhà nước, không phải tội phạm và theo pháp luật phải bị xử phạt vi phạm hành chính.",
            "Vi phạm dân sự là hành vi trái pháp luật xâm phạm quan hệ tài sản, quan hệ nhân thân hoặc không thực hiện, thực hiện không đúng nghĩa vụ dân sự, làm phát sinh trách nhiệm dân sự khi đủ căn cứ.",
            "Vi phạm kỷ luật là hành vi có lỗi vi phạm quy tắc, nội quy hoặc nghĩa vụ trong quan hệ công vụ, lao động, học tập hay nội bộ tổ chức, làm phát sinh trách nhiệm kỷ luật.",
            "Một hành vi có thể đồng thời làm phát sinh nhiều loại trách nhiệm pháp lý nếu đủ căn cứ.",
        ]),
    ], "hành vi; trái pháp luật; có lỗi; năng lực trách nhiệm pháp lý; hình sự; hành chính; dân sự; kỷ luật")

    add_question(doc, 11, "Thừa kế và thừa kế theo di chúc", [
        ("Khái niệm", "Thừa kế là việc chuyển dịch tài sản của người chết cho người còn sống theo di chúc hoặc theo pháp luật. Di chúc là sự thể hiện ý chí của cá nhân nhằm chuyển tài sản của mình cho người khác sau khi chết."),
        ("Thừa kế theo di chúc", [
            "Thừa kế theo di chúc là việc chuyển dịch di sản của người chết cho người còn sống theo ý chí của người để lại di sản được thể hiện trong di chúc hợp pháp.",
            "Người lập minh mẫn, sáng suốt; không bị lừa dối, đe dọa hoặc cưỡng ép.",
            "Nội dung không vi phạm điều cấm, không trái đạo đức xã hội; hình thức phù hợp luật.",
            "Người từ đủ 15 đến dưới 18 tuổi lập bằng văn bản và được cha, mẹ hoặc người giám hộ đồng ý về việc lập di chúc.",
            "Di chúc miệng chỉ trong tình trạng tính mạng bị đe dọa và phải tuân thủ điều kiện người làm chứng, ghi chép, xác nhận theo luật.",
            "Quyền chỉ định, truất quyền, phân chia di sản, di tặng, dành tài sản thờ cúng và giao nghĩa vụ.",
            "Người thừa kế không phụ thuộc nội dung di chúc: cha, mẹ, vợ/chồng, con chưa thành niên và con thành niên không có khả năng lao động, trừ trường hợp luật loại trừ.",
            "BLDS 2015 không còn chế định riêng về di chúc chung của vợ chồng như BLDS 2005.",
            "Trước khi chia phải xác định tài sản của người chết, tách tài sản chung và thanh toán nghĩa vụ tài sản.",
        ]),
    ], "di sản; thời điểm mở thừa kế; di chúc hợp pháp; người thừa kế bắt buộc; tách tài sản chung; nghĩa vụ tài sản")

    add_question(doc, 12, "Luật Lao động: khái niệm, đối tượng và phương pháp điều chỉnh", [
        ("Khái niệm", "Luật Lao động là ngành luật điều chỉnh quan hệ lao động giữa người lao động và người sử dụng lao động cùng các quan hệ liên quan trực tiếp."),
        ("Đối tượng", [
            "Quan hệ phát sinh từ việc làm có trả công, tiền lương và sự quản lý, điều hành, giám sát của một bên.",
            "Quan hệ việc làm, học nghề, đối thoại, thương lượng tập thể, an toàn vệ sinh lao động, bảo hiểm, đại diện người lao động và giải quyết tranh chấp.",
        ]),
        ("Phương pháp", [
            "Thỏa thuận trên cơ sở tự nguyện nhưng phải tuân thủ chuẩn tối thiểu bảo vệ người lao động.",
            "Quyền quản lý, điều hành của người sử dụng lao động trong giới hạn pháp luật.",
            "Sự tham gia của tổ chức đại diện người lao động và cơ chế thương lượng tập thể.",
        ]),
    ], "việc làm có trả công; tiền lương; quản lý–điều hành–giám sát; thỏa thuận; chuẩn tối thiểu; đại diện người lao động")

    add_question(doc, 13, "Hợp đồng lao động: nhận diện, hình thức và chấm dứt", [
        ("Khái niệm và nhận diện", "Hợp đồng lao động là thỏa thuận về việc làm có trả công, tiền lương, điều kiện lao động, quyền và nghĩa vụ của mỗi bên. Dù mang tên khác, nếu có việc làm có trả công và sự quản lý, điều hành, giám sát thì vẫn được coi là hợp đồng lao động."),
        ("Hình thức và loại hợp đồng", [
            "Thông thường phải bằng văn bản hoặc thông điệp dữ liệu; một số hợp đồng dưới một tháng có thể bằng lời nói, trừ trường hợp luật bắt buộc bằng văn bản.",
            "Hợp đồng lao động không xác định thời hạn là hợp đồng trong đó các bên không xác định thời hạn và thời điểm chấm dứt hiệu lực.",
            "Hợp đồng lao động xác định thời hạn là hợp đồng trong đó các bên xác định thời hạn và thời điểm chấm dứt hiệu lực trong thời gian không quá 36 tháng kể từ khi hợp đồng có hiệu lực.",
        ]),
        ("Chấm dứt", [
            "Theo thỏa thuận, hết hạn, hoàn thành công việc, đơn phương hợp pháp và các căn cứ luật định.",
            "Khi đơn phương phải có căn cứ hoặc tuân thủ thời hạn báo trước tùy chủ thể và trường hợp.",
            "Đơn phương trái pháp luật có thể dẫn đến nghĩa vụ nhận lại làm việc, trả lương, đóng bảo hiểm, bồi thường hoặc hoàn trả chi phí.",
        ]),
    ], "việc làm có trả công; quản lý; hợp đồng bằng văn bản; xác định thời hạn; không xác định thời hạn; đơn phương; báo trước")

    add_question(doc, 14, "Tranh chấp lao động và cơ chế giải quyết theo Bộ luật Lao động 2019", [
        ("Khái niệm và phân loại", [
            "Tranh chấp lao động cá nhân là tranh chấp giữa người lao động với người sử dụng lao động hoặc giữa người lao động với doanh nghiệp, tổ chức đưa người lao động đi làm việc ở nước ngoài hay chủ thể khác theo luật về quyền, nghĩa vụ và lợi ích của cá nhân.",
            "Tranh chấp lao động tập thể về quyền là tranh chấp giữa một hay nhiều tổ chức đại diện người lao động với người sử dụng lao động hoặc tổ chức của người sử dụng lao động về việc giải thích, thực hiện quyền đã được xác lập trong pháp luật, thỏa ước, nội quy hoặc thỏa thuận hợp pháp.",
            "Tranh chấp lao động tập thể về lợi ích là tranh chấp phát sinh trong quá trình thương lượng nhằm xác lập điều kiện lao động mới so với quy định, thỏa ước hoặc thỏa thuận đang có.",
        ]),
        ("Chủ thể có thẩm quyền", [
            "Hòa giải viên lao động.",
            "Hội đồng trọng tài lao động trong trường hợp luật quy định và các bên lựa chọn.",
            "Tòa án nhân dân.",
        ]),
        ("Trình tự khái quát", [
            "Tranh chấp cá nhân thường phải qua hòa giải viên trước khi yêu cầu trọng tài hoặc Tòa án, trừ các trường hợp được miễn hòa giải như xử lý kỷ luật sa thải, đơn phương chấm dứt, bồi thường khi chấm dứt, bảo hiểm và một số tranh chấp khác.",
            "Tranh chấp tập thể về quyền: hòa giải → trọng tài hoặc Tòa án theo lựa chọn và điều kiện luật định.",
            "Tranh chấp tập thể về lợi ích: hòa giải → trọng tài; đình công chỉ có thể đặt ra theo đúng điều kiện và trình tự.",
            "Không còn dùng mô hình Hội đồng hòa giải cơ sở hoặc Chủ tịch UBND cấp huyện như tài liệu cũ.",
        ]),
    ], "cá nhân; tập thể về quyền; tập thể về lợi ích; hòa giải viên; trọng tài lao động; Tòa án; đình công")

    add_question(doc, 15, "Luật Hành chính: khái niệm, đối tượng và phương pháp điều chỉnh", [
        ("Khái niệm", "Luật Hành chính là ngành luật điều chỉnh các quan hệ xã hội phát sinh trong hoạt động quản lý hành chính nhà nước và những quan hệ quản lý khác được pháp luật giao."),
        ("Đối tượng", [
            "Quan hệ giữa cơ quan hành chính với cá nhân, tổ chức.",
            "Quan hệ tổ chức và hoạt động nội bộ của cơ quan nhà nước.",
            "Quan hệ khi cơ quan, tổ chức hoặc cá nhân được trao quyền quản lý hành chính.",
        ]),
        ("Phương pháp", "Phương pháp quyền lực – phục tùng là chủ yếu: một bên nhân danh quyền lực nhà nước có quyền đưa ra quyết định đơn phương trong phạm vi thẩm quyền; bên kia có nghĩa vụ chấp hành và có quyền khiếu nại, khởi kiện khi quyết định hoặc hành vi trái pháp luật."),
    ], "quản lý hành chính; quyền lực–phục tùng; quyết định đơn phương; thẩm quyền; khiếu nại; khởi kiện")

    add_question(doc, 16, "Tội phạm: khái niệm, dấu hiệu và phân loại", [
        ("Khái niệm", "Tội phạm là hành vi nguy hiểm cho xã hội được quy định trong Bộ luật Hình sự, do người có năng lực trách nhiệm hình sự hoặc pháp nhân thương mại thực hiện một cách cố ý hoặc vô ý, xâm phạm các quan hệ được luật hình sự bảo vệ và phải bị xử lý hình sự."),
        ("Dấu hiệu", [
            "Tính nguy hiểm đáng kể cho xã hội.",
            "Tính có lỗi.",
            "Tính trái pháp luật hình sự.",
            "Tính phải chịu hình phạt hoặc biện pháp hình sự.",
            "Chủ thể có năng lực trách nhiệm hình sự; pháp nhân thương mại chỉ chịu trách nhiệm về các tội được liệt kê.",
            "Về độ tuổi: người từ đủ 16 tuổi chịu trách nhiệm về mọi tội, trừ trường hợp luật quy định khác; người từ đủ 14 đến dưới 16 tuổi chỉ chịu trách nhiệm về những tội được liệt kê tại Điều 12 BLHS.",
        ]),
        ("Phân loại tội phạm", [
            "Tội phạm ít nghiêm trọng là tội có tính chất và mức độ nguy hiểm cho xã hội không lớn, mức cao nhất của khung hình phạt là phạt tiền, cải tạo không giam giữ hoặc tù đến 03 năm.",
            "Tội phạm nghiêm trọng là tội có tính chất và mức độ nguy hiểm cho xã hội lớn, mức cao nhất của khung hình phạt trên 03 năm đến 07 năm tù.",
            "Tội phạm rất nghiêm trọng là tội có tính chất và mức độ nguy hiểm cho xã hội rất lớn, mức cao nhất của khung hình phạt trên 07 năm đến 15 năm tù.",
            "Tội phạm đặc biệt nghiêm trọng là tội có tính chất và mức độ nguy hiểm cho xã hội đặc biệt lớn, mức cao nhất của khung hình phạt trên 15 năm đến 20 năm tù, tù chung thân hoặc tử hình.",
            "Việc phân loại căn cứ vào mức cao nhất của khung hình phạt do luật quy định, không căn cứ vào mức án cụ thể Tòa án đã tuyên.",
        ]),
    ], "nguy hiểm cho xã hội; có lỗi; trái BLHS; chịu hình phạt; năng lực TNHS; Điều 12; bốn loại tội phạm")

    add_question(doc, 17, "Tố cáo: khái niệm, đặc điểm, thẩm quyền và thủ tục", [
        ("Khái niệm", "Tố cáo là việc cá nhân theo thủ tục luật định báo cho cơ quan, tổ chức hoặc cá nhân có thẩm quyền biết về hành vi vi phạm pháp luật gây thiệt hại hoặc đe dọa gây thiệt hại đến lợi ích Nhà nước, quyền và lợi ích hợp pháp của cơ quan, tổ chức, cá nhân."),
        ("Đặc điểm", [
            "Chủ thể tố cáo là cá nhân.",
            "Đối tượng là hành vi vi phạm pháp luật trong thực hiện nhiệm vụ, công vụ hoặc hành vi vi phạm về quản lý nhà nước trong các lĩnh vực.",
            "Khác khiếu nại: khiếu nại nhằm bảo vệ trực tiếp quyền, lợi ích của người khiếu nại trước quyết định/hành vi hành chính hoặc kỷ luật.",
        ]),
        ("Thẩm quyền giải quyết", [
            "Nguyên tắc chung: người đứng đầu cơ quan, tổ chức có thẩm quyền quản lý người bị tố cáo giải quyết; trường hợp khác theo luật chuyên ngành.",
        ]),
        ("Thủ tục giải quyết", [
            "Trình tự: tiếp nhận, xử lý ban đầu → thụ lý → xác minh → kết luận nội dung → xử lý kết luận.",
            "Thời hạn thông thường không quá 30 ngày từ ngày thụ lý; vụ việc phức tạp có thể gia hạn theo luật.",
            "Thông tin không rõ họ tên, địa chỉ không xử lý theo thủ tục tố cáo; nhưng nếu nội dung rõ, có tài liệu cụ thể và cơ sở kiểm tra thì có thể được chuyển để thanh tra, kiểm tra.",
        ]),
    ], "cá nhân; hành vi vi phạm; phân biệt khiếu nại; thụ lý; xác minh; kết luận; 30 ngày; bảo vệ người tố cáo")

    add_question(doc, 18, "Khái niệm, nguyên tắc và tính ưu việt của Luật Hôn nhân và gia đình", [
        ("Khái niệm", "Luật Hôn nhân và gia đình là ngành luật điều chỉnh quan hệ nhân thân và tài sản giữa vợ chồng, cha mẹ và con, cùng các thành viên gia đình."),
        ("Nguyên tắc theo Luật năm 2014", [
            "Hôn nhân tự nguyện, tiến bộ, một vợ một chồng, vợ chồng bình đẳng.",
            "Hôn nhân giữa công dân thuộc các dân tộc, tôn giáo; giữa người có và không có tôn giáo; giữa công dân Việt Nam với người nước ngoài được tôn trọng và bảo vệ.",
            "Xây dựng gia đình ấm no, tiến bộ, hạnh phúc; các thành viên tôn trọng, quan tâm, chăm sóc, giúp đỡ nhau.",
            "Không phân biệt đối xử giữa các con; bảo vệ trẻ em, người cao tuổi, người khuyết tật và người mẹ.",
            "Kế thừa, phát huy truyền thống tốt đẹp; xóa bỏ phong tục, tập quán lạc hậu.",
            "Nam, nữ chung sống như vợ chồng mà không đăng ký thì không phát sinh quyền, nghĩa vụ giữa vợ và chồng; tài sản và con chung được giải quyết theo quy định tương ứng.",
        ]),
        ("Tính ưu việt", [
            "Bảo đảm quyền tự do kết hôn, ly hôn trên cơ sở tự nguyện và bình đẳng.",
            "Xóa bỏ chế độ đa thê, cưỡng ép, trọng nam khinh nữ và phân biệt giữa các con.",
            "Bảo vệ phụ nữ, trẻ em và thành viên yếu thế; đề cao trách nhiệm chăm sóc, phụng dưỡng và tương trợ.",
            "Kết hợp giá trị tiến bộ với việc kế thừa truyền thống gia đình tốt đẹp của Việt Nam.",
        ]),
    ], "tự nguyện; tiến bộ; một vợ một chồng; bình đẳng; đăng ký kết hôn; bảo vệ thành viên yếu thế")

    add_question(doc, 19, "Phân tích nội dung cơ bản của Hiến pháp năm 2013", [
        ("Nội dung cơ bản", [
            "Hiến pháp là luật cơ bản, có hiệu lực pháp lý cao nhất. Mọi văn bản pháp luật khác phải phù hợp với Hiến pháp; mọi hành vi vi phạm Hiến pháp phải bị xử lý.",
            "Chế độ chính trị; bản chất Nhà nước; chủ quyền Nhân dân; vai trò lãnh đạo của Đảng.",
            "Quyền con người, quyền và nghĩa vụ cơ bản của công dân.",
            "Kinh tế, xã hội, văn hóa, giáo dục, khoa học, công nghệ và môi trường.",
            "Bảo vệ Tổ quốc.",
            "Tổ chức bộ máy nhà nước: Quốc hội, Chủ tịch nước, Chính phủ, Tòa án, Viện kiểm sát, chính quyền địa phương, Hội đồng bầu cử quốc gia và Kiểm toán nhà nước.",
            "Hiệu lực và thủ tục sửa đổi Hiến pháp.",
            "Năm 2025, một số điều của Hiến pháp được sửa đổi, bổ sung để phù hợp việc sắp xếp tổ chức bộ máy, củng cố vai trò của Mặt trận Tổ quốc Việt Nam và tổ chức chính quyền địa phương theo mô hình hai cấp.",
        ]),
    ], "luật cơ bản; hiệu lực cao nhất; quyền con người; tổ chức bộ máy; kiểm soát quyền lực; sửa đổi 2025; chính quyền hai cấp")


def add_analysis_block(doc, title, legal_basis, thesis, analyses, example, mistakes, outline):
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    r = p.add_run("Cơ sở pháp lý và học thuật: ")
    set_font(r, bold=True, color=DARK)
    r = p.add_run(legal_basis)
    set_font(r)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Luận điểm trung tâm: ")
    set_font(r, bold=True, color=DARK)
    r = p.add_run(thesis)
    set_font(r)
    doc.add_heading("Phân tích", level=3)
    for item in analyses:
        add_bullet(doc, item)
    doc.add_heading("Ví dụ minh họa", level=3)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(example)
    set_font(r)
    doc.add_heading("Lỗi thường gặp", level=3)
    for item in mistakes:
        add_bullet(doc, item)
    doc.add_heading("Dàn ý viết bài tự luận", level=3)
    for item in outline:
        add_number(doc, item)
    add_key(doc, "Không chỉ liệt kê. Mỗi ý nên có một câu giải thích bản chất và một câu liên hệ hoặc ví dụ.", label="Mẹo trình bày")


def add_deep_analysis(doc):
    doc.add_page_break()
    doc.add_heading("PHẦN II. PHÂN TÍCH CHUYÊN SÂU THEO CHUYÊN ĐỀ", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(
        "Phần này mở rộng 19 câu hỏi thành các chuyên đề hoàn chỉnh. Người học có thể dùng phần đề cương ngắn ở trên "
        "để ghi nhớ, sau đó dùng phần này để hiểu bản chất, viết bài tự luận dài và xử lý câu hỏi biến đổi."
    )
    set_font(r)

    add_analysis_block(
        doc,
        "Chuyên đề 1. Nhà nước, bản chất Nhà nước và Nhà nước pháp quyền XHCN Việt Nam (Câu 1 và Câu 3)",
        "Hiến pháp năm 2013, đặc biệt các điều 2, 3, 4, 6, 8; nội dung sửa đổi Hiến pháp được Quốc hội thông qua năm 2025; giáo trình Lý luận Nhà nước và pháp luật.",
        "Nhà nước vừa là tổ chức quyền lực công đặc biệt, vừa là thiết chế tổ chức và giải quyết công việc chung. Khi phân tích Nhà nước Việt Nam phải xuất phát từ chủ quyền Nhân dân, tính pháp quyền và cơ chế phân công, phối hợp, kiểm soát quyền lực.",
        [
            "Nhà nước chỉ xuất hiện khi xã hội phát triển đến một giai đoạn nhất định, có phân hóa lợi ích và cần một quyền lực công tách tương đối khỏi cộng đồng để duy trì trật tự. Vì vậy, Nhà nước khác thị tộc, bộ lạc ở bộ máy quyền lực chuyên nghiệp, quản lý dân cư theo lãnh thổ, chủ quyền, pháp luật và thuế.",
            "Tính giai cấp trả lời câu hỏi quyền lực nhà nước phản ánh và bảo vệ trước hết lợi ích của lực lượng xã hội nào. Tính xã hội trả lời câu hỏi Nhà nước thực hiện những công việc chung nào để xã hội tồn tại và phát triển. Hai mặt không loại trừ nhau.",
            "Trong Nhà nước pháp quyền, quyền lực không được thực hiện tùy tiện. Cơ quan và người có chức vụ chỉ được làm trong phạm vi thẩm quyền; quyền con người, quyền công dân được công nhận, tôn trọng, bảo vệ và bảo đảm.",
            "Cụm từ “quyền lực nhà nước là thống nhất” không đồng nghĩa mọi quyền tập trung vào một cơ quan. Hiến pháp đồng thời yêu cầu phân công, phối hợp và kiểm soát giữa các cơ quan thực hiện quyền lập pháp, hành pháp và tư pháp.",
            "Nguyên tắc Nhân dân là chủ thể quyền lực được thực hiện thông qua dân chủ trực tiếp và dân chủ đại diện: bầu cử, ứng cử, trưng cầu ý dân, tham gia quản lý, kiến nghị, khiếu nại, tố cáo và giám sát.",
            "Vai trò lãnh đạo của Đảng phải được trình bày cùng yêu cầu hoạt động trong khuôn khổ Hiến pháp, pháp luật, gắn bó với Nhân dân, chịu sự giám sát của Nhân dân và chịu trách nhiệm trước Nhân dân.",
        ],
        "Quốc hội ban hành luật nhưng luật vẫn phải phù hợp Hiến pháp; Chính phủ thực hiện quyền hành pháp nhưng chịu giám sát; Tòa án xét xử độc lập và chỉ tuân theo pháp luật. Ba biểu hiện này cho thấy quyền lực thống nhất nhưng có phân công và kiểm soát.",
        [
            "Đồng nhất bản chất Nhà nước với chức năng Nhà nước.",
            "Chỉ nói tính giai cấp mà bỏ tính xã hội, hoặc ngược lại.",
            "Dùng đặc trưng của xã hội xã hội chủ nghĩa để thay hoàn toàn đặc trưng của Nhà nước.",
            "Hiểu Nhà nước pháp quyền đơn giản là Nhà nước ban hành nhiều pháp luật; cốt lõi còn là giới hạn quyền lực và bảo vệ quyền.",
        ],
        [
            "Nêu khái niệm Nhà nước và các dấu hiệu đặc trưng.",
            "Phân tích tính giai cấp, tính xã hội và mối quan hệ giữa hai phương diện.",
            "Trình bày bản chất hiến định của Nhà nước Việt Nam.",
            "Phân tích các đặc trưng: chủ quyền Nhân dân, pháp quyền, kiểm soát quyền lực, quyền con người, Đảng lãnh đạo.",
            "Kết luận về ý nghĩa đối với xây dựng và hoàn thiện bộ máy nhà nước.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 2. Hình thức pháp luật, quy phạm pháp luật và văn bản QPPL (Câu 2, Câu 6 và Câu 8)",
        "Luật Ban hành văn bản quy phạm pháp luật số 64/2025/QH15; Luật sửa đổi, bổ sung số 87/2025/QH15; Bộ luật Dân sự 2015 về tập quán và áp dụng tương tự pháp luật; quy định về án lệ.",
        "Cần phân biệt ba tầng: hình thức pháp luật là cách biểu hiện của pháp luật; văn bản QPPL là một hình thức pháp luật; quy phạm pháp luật là quy tắc xử sự chung nằm trong hoặc được thể hiện qua các nguồn pháp luật.",
        [
            "Tập quán pháp hình thành từ quy tắc xử sự tồn tại lâu dài trong cộng đồng và được Nhà nước thừa nhận. Không phải mọi tập quán đều là pháp luật; tập quán trái nguyên tắc cơ bản hoặc điều cấm không được áp dụng.",
            "Án lệ ở Việt Nam không đồng nhất với toàn bộ bản án. Chỉ lập luận, phán quyết trong bản án hoặc quyết định đã được lựa chọn, công bố theo thủ tục mới có giá trị để nghiên cứu và áp dụng cho vụ việc tương tự.",
            "Văn bản QPPL chứa quy tắc áp dụng chung và nhiều lần. Văn bản cá biệt chỉ giải quyết một trường hợp xác định, như quyết định xử phạt một người cụ thể hoặc quyết định bổ nhiệm.",
            "Tính hợp hiến yêu cầu văn bản phù hợp Hiến pháp; tính hợp pháp còn bao gồm đúng thẩm quyền, nội dung, hình thức, trình tự và thủ tục; tính thống nhất yêu cầu không mâu thuẫn, chồng chéo trong toàn hệ thống.",
            "Cấu trúc giả định – quy định – chế tài là cấu trúc logic, không phải khuôn mẫu bắt buộc của từng câu chữ. Chế tài có thể nằm ở điều khác, chương khác hoặc văn bản khác.",
            "Hiệu lực văn bản phải xem xét theo thời gian, không gian và đối tượng. Khi hai văn bản có quy định khác nhau cần xác định hiệu lực pháp lý, thời điểm ban hành, quy định chuyên ngành và quy tắc áp dụng cụ thể.",
            "Luật Ban hành VBQPPL năm 2025 nhấn mạnh chất lượng chính sách, tính khả thi, phân định rõ thẩm quyền và đổi mới quy trình xây dựng pháp luật. Không nên học thuộc quy trình duy nhất cho mọi loại văn bản.",
        ],
        "Quy định “người đủ 18 tuổi có năng lực hành vi dân sự đầy đủ, trừ trường hợp luật định” là quy phạm. Điều luật chứa quy định đó là đơn vị trình bày. Bộ luật Dân sự là văn bản QPPL. Ba khái niệm liên quan nhưng không đồng nhất.",
        [
            "Viết “quy phạm pháp luật chính là điều luật”.",
            "Cho rằng mọi công văn, thông báo, quyết định của cơ quan nhà nước đều là văn bản QPPL.",
            "Tách máy móc một câu luật thành ba bộ phận dù phần quy định hoặc chế tài nằm ở nơi khác.",
            "Chỉ nêu hợp hiến mà bỏ đúng thẩm quyền, trình tự, thủ tục và tính khả thi.",
        ],
        [
            "Định nghĩa từng khái niệm.",
            "Phân tích đặc điểm và mối quan hệ giữa nguồn luật, văn bản và quy phạm.",
            "Trình bày ba bộ phận logic của quy phạm và lưu ý về cách thể hiện.",
            "Phân tích nguyên tắc xây dựng, ban hành và hiệu lực văn bản.",
            "Minh họa bằng một văn bản quy phạm và một văn bản cá biệt.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 3. Quan hệ pháp luật, sự kiện pháp lý và vi phạm pháp luật (Câu 5 và Câu 10)",
        "Giáo trình Lý luận Nhà nước và pháp luật; Bộ luật Dân sự 2015 và các luật chuyên ngành về năng lực chủ thể, quyền, nghĩa vụ và trách nhiệm pháp lý.",
        "Quan hệ pháp luật là hình thức pháp lý của quan hệ xã hội. Vi phạm pháp luật là một loại sự kiện pháp lý làm phát sinh quan hệ trách nhiệm pháp lý, nhưng không phải mọi sự kiện pháp lý đều là vi phạm.",
        [
            "Một quan hệ pháp luật thường hình thành khi có quy phạm điều chỉnh, chủ thể có năng lực phù hợp và sự kiện pháp lý xảy ra. Ví dụ, quy định về hợp đồng đã tồn tại nhưng quan hệ mua bán cụ thể chỉ phát sinh khi các bên giao kết.",
            "Năng lực pháp luật là khả năng có quyền, nghĩa vụ; năng lực hành vi là khả năng bằng hành vi của mình xác lập, thực hiện quyền, nghĩa vụ. Mức độ năng lực hành vi phụ thuộc ngành luật, độ tuổi và trạng thái nhận thức.",
            "Chủ thể không chỉ là cá nhân và pháp nhân. Tổ chức không có tư cách pháp nhân vẫn có thể tham gia một số quan hệ thông qua thành viên hoặc người đại diện theo quy định.",
            "Khách thể không nên đồng nhất đơn giản với người bị thiệt hại. Khách thể là lợi ích hoặc quan hệ xã hội mà chủ thể hướng tới hay pháp luật bảo vệ.",
            "Mặt khách quan của vi phạm gồm hành vi, hậu quả, quan hệ nhân quả và các dấu hiệu hoàn cảnh nếu luật yêu cầu. Một số vi phạm có cấu thành hình thức, không bắt buộc hậu quả thiệt hại thực tế.",
            "Mặt chủ quan gồm lỗi; trong một số trường hợp còn có động cơ, mục đích. Lỗi phản ánh thái độ tâm lý với hành vi và hậu quả, không thể suy ra chỉ từ việc hậu quả đã xảy ra.",
            "Trách nhiệm pháp lý có mục đích khác nhau: hình sự trừng trị và giáo dục; hành chính bảo vệ trật tự quản lý; dân sự khôi phục, bù đắp; kỷ luật bảo đảm trật tự nội bộ.",
        ],
        "Người lái xe tuân thủ luật nhưng sự cố kỹ thuật bất khả kháng gây thiệt hại có thể không có lỗi và không cấu thành vi phạm pháp luật, dù trách nhiệm bồi thường trong một số trường hợp vẫn cần phân tích theo chế định nguồn nguy hiểm cao độ.",
        [
            "Dùng “năng lực pháp lý” thay cho thuật ngữ chuẩn “năng lực pháp luật”.",
            "Ghi khách thể là tên người bị hại.",
            "Cho rằng vi phạm nào cũng phải gây hậu quả vật chất.",
            "Nhầm động cơ với mục đích và nhầm cố ý trực tiếp với cố ý gián tiếp.",
        ],
        [
            "Nêu khái niệm và điều kiện phát sinh quan hệ pháp luật.",
            "Phân tích chủ thể, nội dung, khách thể và sự kiện pháp lý.",
            "Nêu khái niệm vi phạm pháp luật và bốn dấu hiệu.",
            "Phân tích bốn yếu tố cấu thành.",
            "Phân loại vi phạm và trách nhiệm pháp lý tương ứng.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 4. Luật Dân sự, quyền sở hữu và nghĩa vụ dân sự (Câu 7 và Câu 9)",
        "Bộ luật Dân sự năm 2015, đặc biệt các nguyên tắc cơ bản, quy định về tài sản, quyền sở hữu, giao dịch dân sự, nghĩa vụ và hợp đồng.",
        "Pháp luật dân sự lấy bình đẳng, tự do ý chí, độc lập tài sản và tự chịu trách nhiệm làm nền tảng; quyền dân sự được tôn trọng nhưng việc thực hiện quyền không được xâm phạm lợi ích công cộng và quyền hợp pháp của người khác.",
        [
            "Quan hệ tài sản trong luật dân sự rất rộng: sở hữu, nghĩa vụ, hợp đồng, thừa kế, bồi thường và quyền tài sản. Không phải mọi quan hệ có tiền đều thuộc luật dân sự; thuế và xử phạt tiền mang tính quyền lực hành chính.",
            "Quan hệ nhân thân có thể không gắn với tài sản như họ tên, danh dự, đời sống riêng tư; hoặc gắn với tài sản như quyền công bố tác phẩm gắn với quyền hưởng nhuận bút.",
            "Quyền chiếm hữu là nắm giữ, chi phối tài sản; quyền sử dụng là khai thác công dụng và hưởng hoa lợi, lợi tức; quyền định đoạt là quyết định số phận pháp lý hoặc thực tế của tài sản.",
            "Người không phải chủ sở hữu vẫn có thể chiếm hữu, sử dụng hoặc định đoạt trong phạm vi được giao, ủy quyền hoặc luật quy định. Vì vậy không được suy luận ai đang giữ tài sản thì người đó là chủ sở hữu.",
            "Căn cứ xác lập quyền sở hữu có thể nguyên sinh, như tạo vật mới hoặc chiếm hữu theo thời hiệu, và phái sinh, như mua bán, tặng cho, thừa kế.",
            "Giao dịch dân sự có hiệu lực khi chủ thể có năng lực phù hợp, tự nguyện, mục đích và nội dung không vi phạm điều cấm, không trái đạo đức xã hội, đồng thời tuân thủ hình thức nếu luật quy định hình thức là điều kiện.",
            "Vi phạm nghĩa vụ có thể dẫn đến buộc thực hiện đúng, phạt vi phạm nếu có thỏa thuận hoặc luật định, bồi thường thiệt hại, lãi chậm trả và chấm dứt/hủy bỏ hợp đồng theo điều kiện.",
        ],
        "A thuê xe của B thì A là người chiếm hữu và sử dụng hợp pháp trong thời hạn thuê nhưng B vẫn là chủ sở hữu. A không được tự ý bán xe vì không có quyền định đoạt.",
        [
            "Liệt kê quyền sở hữu nhưng bỏ điều kiện, giới hạn thực hiện quyền.",
            "Nhầm quyền sử dụng với quyền định đoạt.",
            "Cho rằng tài sản chỉ gồm vật và tiền; BLDS còn ghi nhận giấy tờ có giá và quyền tài sản.",
            "Dùng các điều 170, 171 BLDS 2005 thay cho hệ thống quy định của BLDS 2015.",
        ],
        [
            "Nêu đối tượng và phương pháp điều chỉnh của luật dân sự.",
            "Phân tích nguyên tắc bình đẳng, tự do ý chí, độc lập tài sản.",
            "Định nghĩa sở hữu và quyền sở hữu.",
            "Phân tích chiếm hữu, sử dụng, định đoạt.",
            "Trình bày nhóm căn cứ xác lập, chấm dứt và ví dụ.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 5. Thừa kế theo di chúc và theo pháp luật (Câu 11)",
        "Phần thứ tư Bộ luật Dân sự 2015 về thừa kế, đặc biệt quy định về di sản, người thừa kế, di chúc, người thừa kế không phụ thuộc nội dung di chúc và thừa kế theo pháp luật.",
        "Bài thừa kế phải giải quyết theo thứ tự: xác định thời điểm mở thừa kế – xác định di sản – thanh toán nghĩa vụ – kiểm tra di chúc – xác định người hưởng – chia phần. Sai ở bước xác định di sản sẽ làm toàn bộ phép chia sai.",
        [
            "Di sản gồm tài sản riêng của người chết và phần của người chết trong tài sản chung. Tài sản chung phải được xác định và tách trước, không được lấy toàn bộ tài sản gia đình làm di sản.",
            "Thời điểm mở thừa kế là thời điểm người có tài sản chết. Thời điểm này quyết định pháp luật áp dụng, người còn sống để hưởng, phạm vi di sản và thời hiệu.",
            "Di chúc hợp pháp cần đáp ứng điều kiện về chủ thể, ý chí, nội dung và hình thức. Di chúc có công chứng không đương nhiên hợp pháp nếu người lập bị cưỡng ép hoặc nội dung vi phạm điều cấm.",
            "Di chúc chỉ định đoạt phần tài sản thuộc người lập. Phần định đoạt vượt quá quyền sở hữu không có hiệu lực tương ứng.",
            "Người thừa kế bắt buộc được bảo vệ bằng phần bằng hai phần ba suất theo pháp luật, trừ người từ chối hoặc không có quyền hưởng. Phải xác định đúng ai thuộc nhóm được bảo vệ.",
            "Thừa kế theo pháp luật áp dụng khi không có di chúc, di chúc không hợp pháp, người được chỉ định chết trước, không có quyền hưởng, từ chối hoặc phần di sản không được định đoạt.",
            "Thừa kế thế vị xảy ra khi con của người để lại di sản chết trước hoặc cùng thời điểm; cháu được hưởng phần mà cha hoặc mẹ của cháu đáng lẽ được hưởng. Không nhầm với thừa kế chuyển tiếp.",
            "Quan hệ vợ chồng không hợp pháp thường không tạo tư cách thừa kế với danh nghĩa vợ/chồng; nhưng con chung vẫn có quyền thừa kế nếu quan hệ cha mẹ con được xác định.",
        ],
        "Vợ chồng có nhà chung 1,2 tỷ đồng; chồng có tài sản riêng 300 triệu và chết. Phần của vợ trong nhà là 600 triệu. Di sản của chồng là 600 + 300 = 900 triệu, chưa trừ nghĩa vụ. Chỉ 900 triệu mới được đưa vào chia.",
        [
            "Lấy toàn bộ tài sản chung làm di sản.",
            "Chia ngay mà không trừ chi phí mai táng và nghĩa vụ tài sản.",
            "Cho người chung sống không đăng ký kết hôn hưởng với tư cách vợ/chồng.",
            "Áp dụng quy định di chúc chung vợ chồng của BLDS 2005 như quy định hiện hành.",
        ],
        [
            "Vẽ sơ đồ gia đình và ghi thời điểm chết.",
            "Tách tài sản chung, xác định di sản ròng.",
            "Kiểm tra hiệu lực di chúc và phạm vi định đoạt.",
            "Xác định người thừa kế bắt buộc, người không có quyền hưởng, người từ chối.",
            "Chia theo di chúc và phần còn lại theo pháp luật; kiểm tra tổng số.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 6. Luật Hình sự, tội phạm và cấu thành tội phạm (Câu 4 và Câu 16)",
        "Bộ luật Hình sự 2015; Luật sửa đổi, bổ sung năm 2017; Luật số 86/2025/QH15; các nghị quyết hướng dẫn áp dụng có liên quan.",
        "Tội phạm không chỉ là hành vi gây hậu quả xấu. Hành vi phải được BLHS quy định, có tính nguy hiểm đáng kể, do chủ thể có năng lực thực hiện với lỗi và đủ các dấu hiệu cấu thành cụ thể.",
        [
            "Nguyên tắc pháp chế hình sự thể hiện ở yêu cầu không có tội nếu không có luật, không áp dụng tương tự pháp luật để buộc tội và chỉ áp dụng hình phạt theo BLHS.",
            "Tính nguy hiểm cho xã hội được đánh giá tổng hợp qua khách thể, phương thức, hậu quả, lỗi, động cơ, nhân thân và hoàn cảnh. Hành vi có dấu hiệu nhưng tính nguy hiểm không đáng kể có thể không phải tội phạm.",
            "Khách thể của tội phạm là quan hệ xã hội được luật hình sự bảo vệ; đối tượng tác động là bộ phận vật chất cụ thể, ví dụ chiếc xe hoặc khoản tiền.",
            "Mặt khách quan có hành vi, hậu quả, quan hệ nhân quả và các dấu hiệu bắt buộc khác. Cần phân biệt tội có cấu thành vật chất và cấu thành hình thức.",
            "Cố ý trực tiếp: thấy trước và mong muốn hậu quả; cố ý gián tiếp: thấy trước, không mong muốn nhưng có ý thức để mặc. Vô ý quá tự tin: thấy trước nhưng tin hậu quả không xảy ra hoặc ngăn được; vô ý cẩu thả: không thấy trước dù phải và có thể thấy trước.",
            "Chủ thể tội phạm gồm cá nhân có năng lực và đủ tuổi; pháp nhân thương mại chỉ chịu trách nhiệm về các tội được BLHS liệt kê và khi đủ điều kiện.",
            "Người từ 14 đến dưới 16 tuổi chỉ chịu trách nhiệm về các tội được liệt kê tại Điều 12. Phải kiểm tra tội danh cụ thể, không sử dụng công thức cũ.",
            "Phân biệt nhóm chiếm đoạt bằng phương thức: vũ lực – nhanh chóng công khai – công nhiên – lén lút – gian dối trước khi giao tài sản – nhận hợp pháp rồi chiếm đoạt.",
        ],
        "A nói dối mình bán điện thoại để B chuyển tiền ngay từ đầu rồi chiếm đoạt: hướng lừa đảo. A mượn điện thoại hợp pháp, sau đó nảy sinh ý định bán lấy tiền: hướng lạm dụng tín nhiệm nếu đủ điều kiện. Thời điểm và vai trò của gian dối là chìa khóa.",
        [
            "Định tội theo ngôn ngữ đời thường, như gọi mọi hành vi lấy tài sản là “cướp”.",
            "Ghi khách thể là người bị hại.",
            "Cho rằng động cơ khó khăn tài chính làm mất trách nhiệm hình sự.",
            "Khẳng định tội danh khi đề thiếu tuổi, giá trị, hậu quả hoặc phương thức thực hiện.",
        ],
        [
            "Nêu khái niệm, nguyên tắc của luật hình sự.",
            "Phân tích bốn đặc điểm của tội phạm và chủ thể.",
            "Phân tích bốn yếu tố cấu thành tội phạm.",
            "Xác định lỗi và giai đoạn thực hiện.",
            "Nếu là tội chiếm đoạt, so sánh phương thức với các tội gần nhau rồi kết luận có điều kiện.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 7. Luật Lao động, hợp đồng và tranh chấp lao động (Câu 12, Câu 13 và Câu 14)",
        "Bộ luật Lao động năm 2019, có hiệu lực từ ngày 01/01/2021; các nghị định hướng dẫn về quan hệ lao động, tiền lương và giải quyết tranh chấp.",
        "Quan hệ lao động có sự kết hợp giữa thỏa thuận và quyền quản lý. Pháp luật tôn trọng tự do hợp đồng nhưng đặt ra chuẩn tối thiểu để bảo vệ người lao động, là bên thường yếu thế hơn.",
        [
            "Dấu hiệu nhận diện hợp đồng lao động quan trọng hơn tên gọi: có việc làm, có trả công/tiền lương và có sự quản lý, điều hành, giám sát. Hợp đồng “cộng tác viên” vẫn có thể là hợp đồng lao động nếu đủ dấu hiệu.",
            "Hai loại hợp đồng cơ bản là xác định thời hạn không quá 36 tháng và không xác định thời hạn. Việc ký nối tiếp hợp đồng xác định thời hạn bị giới hạn để tránh kéo dài tình trạng việc làm bấp bênh.",
            "Thử việc phải phù hợp loại công việc, thời gian tối đa và mức lương tối thiểu theo luật. Không phải công việc nào cũng được yêu cầu thử việc.",
            "Tiền lương là khoản người sử dụng lao động trả theo thỏa thuận để thực hiện công việc, gồm mức lương theo công việc/chức danh, phụ cấp và khoản bổ sung; không được thấp hơn mức tối thiểu áp dụng.",
            "Kỷ luật lao động phải có căn cứ, chứng minh lỗi, bảo đảm quyền tham gia của tổ chức đại diện trong trường hợp luật định, đúng thời hiệu và thủ tục; cấm phạt tiền, cắt lương thay kỷ luật.",
            "Người lao động có quyền đơn phương chấm dứt và thường chỉ cần tuân thủ thời hạn báo trước; một số trường hợp được nghỉ không cần báo trước. Người sử dụng lao động cần có căn cứ luật định và thủ tục chặt chẽ hơn.",
            "Tranh chấp cá nhân thường qua hòa giải viên, nhưng nhiều nhóm như sa thải, đơn phương chấm dứt, bồi thường khi chấm dứt và bảo hiểm được miễn thủ tục hòa giải bắt buộc.",
            "Tranh chấp tập thể về quyền nhằm giải thích, thực hiện quy định đã có; tranh chấp về lợi ích nhằm xác lập điều kiện mới. Chỉ tranh chấp lợi ích mới có thể dẫn tới đình công hợp pháp theo trình tự.",
        ],
        "Công ty ký “hợp đồng dịch vụ” với nhân viên làm đủ giờ, nhận lương tháng, chịu chấm công và quản lý trực tiếp. Tòa án có thể căn cứ bản chất để xác định đây là quan hệ lao động, không phụ thuộc nhãn hợp đồng.",
        [
            "Cho rằng cán bộ, công chức thuộc quan hệ lao động thông thường giống người làm trong doanh nghiệp.",
            "Dùng mô hình Hội đồng hòa giải cơ sở và Chủ tịch UBND huyện của pháp luật cũ.",
            "Cho rằng người lao động muốn nghỉ phải được công ty đồng ý.",
            "Kết luận sa thải hợp pháp chỉ vì có hành vi vi phạm mà không kiểm tra thủ tục.",
        ],
        [
            "Nhận diện quan hệ lao động qua ba dấu hiệu.",
            "Xác định loại hợp đồng, thời hạn và hình thức.",
            "Phân tích quyền, nghĩa vụ và hành vi vi phạm.",
            "Kiểm tra căn cứ, báo trước, thời hiệu, thủ tục.",
            "Xác định loại tranh chấp, cơ quan giải quyết và thời hiệu yêu cầu.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 8. Luật Hành chính, quyết định hành chính, khiếu nại và tố cáo (Câu 15 và Câu 17)",
        "Luật Tố cáo 2018; Luật Khiếu nại 2011; Luật Tố tụng hành chính 2015; Luật Xử lý vi phạm hành chính và các luật sửa đổi; pháp luật tổ chức chính quyền địa phương hiện hành.",
        "Quản lý hành chính dựa trên quyền lực công nhưng quyền lực đó bị giới hạn bởi thẩm quyền, trình tự, mục đích và quyền kiểm soát của cá nhân, tổ chức thông qua khiếu nại, tố cáo, khởi kiện và giám sát.",
        [
            "Quan hệ hành chính thường có một bên nhân danh quyền lực nhà nước, nhưng không phải cứ cơ quan nhà nước tham gia thì quan hệ là hành chính. Khi cơ quan thuê văn phòng, quan hệ có thể là dân sự.",
            "Quyết định hành chính cá biệt áp dụng một lần cho đối tượng cụ thể, như quyết định xử phạt. Văn bản QPPL đặt ra quy tắc chung. Phân biệt này quyết định cơ chế kiểm tra và khởi kiện.",
            "Một quyết định hợp pháp phải đúng thẩm quyền, căn cứ, nội dung, hình thức, trình tự, thủ tục và mục đích. Sai thẩm quyền là lỗi nghiêm trọng dù nội dung có vẻ hợp lý.",
            "Khiếu nại do người cho rằng quyết định/hành vi xâm phạm trực tiếp quyền của mình thực hiện, yêu cầu xem xét lại. Người khiếu nại có thể lựa chọn khởi kiện hành chính theo điều kiện luật định.",
            "Tố cáo do cá nhân báo hành vi vi phạm pháp luật. Mục tiêu trước hết là phát hiện, ngăn chặn và xử lý vi phạm, đồng thời bảo vệ lợi ích hợp pháp.",
            "Thời hạn giải quyết tố cáo thông thường là 30 ngày từ ngày thụ lý; vụ việc phức tạp hoặc đặc biệt phức tạp có thể được gia hạn theo số lần và thời gian luật định.",
            "Thông tin nặc danh không được xử lý theo đầy đủ thủ tục tố cáo, nhưng thông tin rõ ràng, có tài liệu cụ thể và cơ sở kiểm tra có thể được dùng cho hoạt động thanh tra, kiểm tra.",
            "Người tố cáo được bảo mật thông tin và có thể được áp dụng biện pháp bảo vệ vị trí công tác, việc làm, tính mạng, sức khỏe, tài sản, danh dự khi có căn cứ.",
        ],
        "A bị phạt hành chính và cho rằng quyết định sai: A khiếu nại hoặc khởi kiện. B biết cán bộ nhận tiền để bỏ qua vi phạm: B tố cáo hoặc báo tin tội phạm tùy tính chất. Hai cơ chế không thể dùng lẫn chỉ vì đều gửi đơn đến cơ quan nhà nước.",
        [
            "Nói phương pháp hành chính là “không bình đẳng” nhưng không giải thích giới hạn pháp luật.",
            "Nhầm tố cáo với khiếu nại.",
            "Dùng thời hạn 60 ngày của tài liệu cũ cho mọi vụ tố cáo.",
            "Cho rằng mọi đơn không ghi tên đều phải bỏ qua hoàn toàn, dù có chứng cứ rõ ràng.",
        ],
        [
            "Nêu khái niệm, đối tượng, phương pháp của luật hành chính.",
            "Phân tích tính quyền lực – phục tùng và giới hạn thẩm quyền.",
            "So sánh khiếu nại, tố cáo và khởi kiện hành chính.",
            "Trình bày thẩm quyền, trình tự, thời hạn tố cáo.",
            "Liên hệ cơ chế kiểm soát quyền lực và bảo vệ công dân.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 9. Hôn nhân và gia đình (Câu 18)",
        "Luật Hôn nhân và gia đình năm 2014; Bộ luật Dân sự 2015; các văn bản hướng dẫn về đăng ký hộ tịch, tài sản vợ chồng và giải quyết việc nam nữ chung sống như vợ chồng.",
        "Hôn nhân hợp pháp được xây dựng trên tự nguyện, một vợ một chồng, bình đẳng và đăng ký. Quan hệ nhân thân, tài sản và quan hệ cha mẹ con phải được tách riêng khi giải quyết tình huống.",
        [
            "Kết hôn có hiệu lực khi nam nữ đáp ứng điều kiện về tuổi, tự nguyện, năng lực, không thuộc trường hợp cấm và thực hiện đăng ký tại cơ quan có thẩm quyền.",
            "Tổ chức đám cưới hoặc chung sống lâu dài không tự động thay thế đăng ký, trừ một số quan hệ lịch sử được pháp luật chuyển tiếp công nhận.",
            "Tài sản chung vợ chồng về nguyên tắc gồm thu nhập, tài sản tạo lập trong thời kỳ hôn nhân và tài sản khác theo luật; tài sản riêng gồm tài sản có trước, được thừa kế riêng, tặng cho riêng và tài sản được chia riêng.",
            "Việc định đoạt tài sản chung quan trọng, đặc biệt bất động sản hoặc nguồn tạo thu nhập chủ yếu, thường cần sự thỏa thuận của cả hai.",
            "Nam nữ chung sống không đăng ký không phát sinh quyền, nghĩa vụ giữa vợ chồng; nhưng quyền và nghĩa vụ với con được giải quyết như cha mẹ con, tài sản giải quyết theo thỏa thuận và công sức đóng góp.",
            "Con sinh trong hay ngoài hôn nhân, con đẻ hay con nuôi đều được pháp luật bảo vệ bình đẳng. Cần tránh thuật ngữ mang tính phân biệt như “con ngoài giá thú” trong lập luận hiện đại.",
            "Bạo lực gia đình, cưỡng ép kết hôn, tảo hôn, kết hôn giả tạo và vi phạm chế độ một vợ một chồng là những hành vi bị cấm hoặc xử lý theo luật.",
        ],
        "Ông A có vợ hợp pháp nhưng sống với bà B và mua nhà bằng tiền cả hai cùng góp. B không phải người thừa kế với tư cách vợ, nhưng có thể chứng minh phần sở hữu dựa trên đóng góp; con chung có quyền thừa kế bình đẳng.",
        [
            "Cho rằng vợ cả đồng ý thì người chồng được kết hôn hoặc chung sống như vợ chồng với người khác.",
            "Mặc nhiên chia đôi tài sản của người chung sống không đăng ký.",
            "Đồng nhất quan hệ hôn nhân không hợp pháp với việc con chung không có quyền.",
            "Dùng nguyên tắc của Luật Hôn nhân và gia đình 2000 thay cho luật năm 2014.",
        ],
        [
            "Kiểm tra điều kiện và đăng ký kết hôn.",
            "Xác định quan hệ nhân thân vợ chồng.",
            "Phân loại tài sản chung, riêng và nghĩa vụ.",
            "Giải quyết quan hệ cha mẹ con độc lập.",
            "Kết luận quyền thừa kế hoặc chia tài sản dựa trên tư cách pháp lý từng người.",
        ],
    )

    add_analysis_block(
        doc,
        "Chuyên đề 10. Hiến pháp, quyền con người và tổ chức bộ máy nhà nước (Câu 19)",
        "Hiến pháp năm 2013; Nghị quyết sửa đổi, bổ sung một số điều của Hiến pháp được Quốc hội thông qua ngày 16/6/2025; pháp luật tổ chức Quốc hội, Chính phủ, Tòa án, Viện kiểm sát và chính quyền địa phương.",
        "Hiến pháp vừa xác lập nền tảng chính trị – pháp lý của Nhà nước, vừa giới hạn quyền lực và ghi nhận địa vị pháp lý cơ bản của con người, công dân. Không nên học Câu 19 như bản liệt kê tên chương.",
        [
            "Hiến pháp có hiệu lực pháp lý cao nhất và là căn cứ đánh giá tính hợp hiến của các văn bản khác. Quyền lập hiến thuộc về Nhân dân và được thực hiện thông qua cơ chế hiến định.",
            "Quyền con người gắn với mọi cá nhân; quyền công dân gắn với quốc tịch. Hiến pháp 2013 phân biệt rõ hơn hai nhóm nhưng nhiều quyền có phạm vi chủ thể và điều kiện khác nhau.",
            "Quyền chỉ có thể bị hạn chế theo quy định của luật trong trường hợp cần thiết vì quốc phòng, an ninh quốc gia, trật tự, an toàn xã hội, đạo đức xã hội hoặc sức khỏe cộng đồng. Hạn chế quyền không thể tùy tiện bằng văn bản cấp thấp.",
            "Quốc hội thực hiện quyền lập hiến, lập pháp và giám sát tối cao; Chính phủ là cơ quan hành chính nhà nước cao nhất, thực hiện quyền hành pháp; Tòa án thực hiện quyền tư pháp; Viện kiểm sát thực hành quyền công tố và kiểm sát hoạt động tư pháp.",
            "Chủ tịch nước là người đứng đầu Nhà nước, thay mặt nước về đối nội và đối ngoại; vị trí này không đồng nhất với người đứng đầu Chính phủ.",
            "Sửa đổi năm 2025 gắn với sắp xếp tổ chức bộ máy, vai trò của Mặt trận Tổ quốc và mô hình chính quyền địa phương hai cấp. Khi ôn thi cần dùng nội dung hiện hành nhưng cũng nhận biết cấu trúc nguyên bản năm 2013 nếu đề yêu cầu so sánh.",
            "Kiểm soát quyền lực được thực hiện bằng nhiều cơ chế: phân công giữa các cơ quan, giám sát của Quốc hội, xét xử, kiểm sát, kiểm toán, thanh tra, giám sát xã hội và quyền khiếu nại, tố cáo của Nhân dân.",
        ],
        "Nếu một nghị định hạn chế một quyền cơ bản vượt quá phạm vi luật cho phép, vấn đề không chỉ là nghị định trái luật mà còn có thể liên quan yêu cầu bảo đảm quyền theo Hiến pháp.",
        [
            "Chỉ kể 11 chương mà không phân tích nội dung cốt lõi.",
            "Đồng nhất quyền con người và quyền công dân.",
            "Cho rằng mọi quyền đều tuyệt đối, hoặc ngược lại có thể bị hạn chế bởi bất kỳ văn bản nào.",
            "Nhầm Chính phủ là cơ quan quyền lực nhà nước cao nhất; vị trí đó thuộc Quốc hội.",
        ],
        [
            "Nêu vị trí, hiệu lực và chức năng của Hiến pháp.",
            "Phân tích chế độ chính trị và chủ quyền Nhân dân.",
            "Phân tích nguyên tắc về quyền con người, quyền công dân.",
            "Trình bày vị trí, chức năng các thiết chế nhà nước chủ yếu.",
            "Nêu cập nhật năm 2025 và ý nghĩa với kiểm soát quyền lực.",
        ],
    )


def add_memory_tables(doc):
    doc.add_page_break()
    doc.add_heading("PHẦN II. BẢNG GHI NHỚ NHANH", level=1)
    doc.add_heading("1. Phân biệt quy phạm pháp luật – điều luật – văn bản QPPL", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    heads = ("Tiêu chí", "Quy phạm pháp luật", "Điều luật", "Văn bản QPPL")
    for cell, text in zip(table.rows[0].cells, heads):
        set_cell_shading(cell, BLUE)
        r = cell.paragraphs[0].add_run(text)
        set_font(r, size=9, bold=True, color="FFFFFF")
    rows = [
        ("Bản chất", "Quy tắc xử sự chung", "Đơn vị trình bày trong văn bản", "Văn bản chứa quy phạm"),
        ("Quan hệ", "Có thể nằm ở nhiều điều", "Có thể chứa một hoặc nhiều quy phạm", "Gồm nhiều điều, chương"),
        ("Ví dụ", "Cấm chiếm đoạt tài sản", "Một điều quy định tội danh", "Bộ luật, luật, nghị định…"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for c, text in zip(cells, row):
            r = c.paragraphs[0].add_run(text)
            set_font(r, size=9)
    set_table_widths(table, [1.05, 1.85, 1.85, 1.95])
    set_repeat_table_header(table.rows[0])

    doc.add_heading("2. Phân biệt các loại vi phạm pháp luật", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    heads = ("Loại", "Quan hệ bị xâm hại", "Căn cứ chính", "Chủ thể xử lý", "Hậu quả điển hình")
    for c, text in zip(table.rows[0].cells, heads):
        set_cell_shading(c, BLUE)
        r = c.paragraphs[0].add_run(text)
        set_font(r, size=8.5, bold=True, color="FFFFFF")
    rows = [
        ("Hình sự", "Quan hệ được BLHS bảo vệ", "Bộ luật Hình sự", "Tòa án", "Hình phạt, biện pháp tư pháp"),
        ("Hành chính", "Trật tự quản lý nhà nước", "Luật XLVPHC và luật chuyên ngành", "Người/cơ quan có thẩm quyền", "Cảnh cáo, phạt tiền, biện pháp khắc phục"),
        ("Dân sự", "Tài sản, nhân thân", "BLDS, hợp đồng", "Tòa án/trọng tài", "Bồi thường, buộc thực hiện nghĩa vụ"),
        ("Kỷ luật", "Trật tự nội bộ", "Nội quy, luật công vụ/lao động", "Người sử dụng lao động/cơ quan", "Khiển trách, kéo dài nâng lương, sa thải…"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for c, text in zip(cells, row):
            r = c.paragraphs[0].add_run(text)
            set_font(r, size=8.5)
    set_table_widths(table, [0.75, 1.55, 1.55, 1.35, 1.5])
    set_repeat_table_header(table.rows[0])

    doc.add_heading("3. Công thức giải bài tình huống", level=2)
    steps = [
        "Tóm tắt đúng dữ kiện pháp lý; xác định câu hỏi đề bài.",
        "Xác định quan hệ pháp luật và ngành luật điều chỉnh.",
        "Nêu quy tắc/căn cứ pháp lý cần áp dụng.",
        "Đối chiếu từng dấu hiệu với dữ kiện; không bỏ qua lỗi, năng lực chủ thể và quan hệ nhân quả.",
        "Nêu dữ kiện còn thiếu hoặc các trường hợp cần phân nhánh.",
        "Kết luận ngắn, rõ, không tuyệt đối hóa khi chưa đủ chứng cứ.",
    ]
    for step in steps:
        add_number(doc, step)

    doc.add_heading("4. Nhận diện nhanh các tội chiếm đoạt", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for c, text in zip(table.rows[0].cells, ("Tội danh", "Dấu hiệu nhận diện chính", "Thời điểm hình thành ý định chiếm đoạt")):
        set_cell_shading(c, BLUE)
        r = c.paragraphs[0].add_run(text)
        set_font(r, size=9, bold=True, color="FFFFFF")
    rows = [
        ("Cướp", "Dùng/đe dọa dùng ngay tức khắc vũ lực hoặc thủ đoạn làm nạn nhân không thể chống cự", "Trước hoặc trong khi chiếm đoạt"),
        ("Cướp giật", "Công khai, nhanh chóng giật lấy rồi tẩu thoát", "Trước khi chiếm đoạt"),
        ("Công nhiên chiếm đoạt", "Công khai lấy tài sản khi chủ tài sản không thể ngăn cản", "Trước khi chiếm đoạt"),
        ("Trộm cắp", "Lén lút chiếm đoạt", "Trước khi chiếm đoạt"),
        ("Lừa đảo", "Gian dối có trước làm chủ tài sản tự nguyện giao tài sản", "Có trước việc giao tài sản"),
        ("Lạm dụng tín nhiệm", "Nhận tài sản hợp pháp rồi dùng thủ đoạn/điều kiện luật định để chiếm đoạt", "Thường phát sinh sau khi nhận tài sản"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for c, text in zip(cells, row):
            r = c.paragraphs[0].add_run(text)
            set_font(r, size=8.7)
    set_table_widths(table, [1.1, 4.0, 1.6])
    set_repeat_table_header(table.rows[0])


def add_exercise(doc, n, title, prompt, answer_parts, note=None):
    doc.add_heading(f"Bài {n}. {title}", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Đề bài: ")
    set_font(r, bold=True, color=DARK)
    r = p.add_run(prompt)
    set_font(r)
    doc.add_heading("Gợi ý giải", level=3)
    for label, text in answer_parts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(label + ": ")
        set_font(r, bold=True)
        r = p.add_run(text)
        set_font(r)
    if note:
        add_key(doc, note, label="Lưu ý chấm điểm")


def add_exercises(doc):
    doc.add_page_break()
    doc.add_heading("PHẦN III. BÀI TẬP ĐIỂN HÌNH VÀ NÂNG CAO", level=1)
    p = doc.add_paragraph()
    r = p.add_run("Nguyên tắc: ")
    set_font(r, bold=True, color=RED)
    r = p.add_run("các bài có yếu tố hình sự/hành chính chỉ kết luận trong phạm vi dữ kiện đề cho; nếu thiếu tuổi, giá trị tài sản, lỗi hoặc hậu quả thì phải nêu điều kiện.")
    set_font(r)

    add_exercise(doc, 1, "Xác định quan hệ pháp luật dân sự",
        "A vay B 100 triệu đồng trong 12 tháng, lãi theo thỏa thuận hợp pháp. B đã giao tiền nhưng đến hạn A không trả. Hãy xác định các yếu tố của quan hệ pháp luật.",
        [
            ("Quan hệ pháp luật", "Quan hệ hợp đồng vay tài sản thuộc pháp luật dân sự."),
            ("Chủ thể", "A là bên vay; B là bên cho vay. Cần giả định cả hai có năng lực hành vi dân sự phù hợp."),
            ("Khách thể", "Lợi ích mà các bên hướng tới thông qua khoản tiền vay và việc hoàn trả; không nên chỉ ghi máy móc “100 triệu đồng”."),
            ("Nội dung", "B có nghĩa vụ giao tiền và quyền yêu cầu trả nợ; A có quyền nhận, sử dụng tiền và nghĩa vụ trả nợ gốc, lãi hợp pháp đúng hạn."),
            ("Sự kiện pháp lý", "Giao kết hợp đồng, giao tiền và việc đến hạn không trả."),
            ("Hậu quả", "B có quyền yêu cầu thực hiện nghĩa vụ, lãi chậm trả và bồi thường nếu đủ căn cứ."),
        ])

    add_exercise(doc, 2, "Giả định – quy định – chế tài",
        "Quy tắc: “Người điều khiển xe mô tô khi tham gia giao thông phải đội mũ bảo hiểm đúng quy cách; vi phạm bị xử phạt theo quy định.” Hãy xác định cấu trúc logic.",
        [
            ("Giả định", "Người điều khiển xe mô tô khi tham gia giao thông."),
            ("Quy định", "Phải đội mũ bảo hiểm đúng quy cách."),
            ("Chế tài", "Bị xử phạt theo quy định nếu không thực hiện."),
            ("Nhận xét", "Trong văn bản thực tế, mức phạt có thể nằm ở điều khoản hoặc văn bản khác; cấu trúc quy phạm không nhất thiết nằm trọn trong một câu."),
        ])

    add_exercise(doc, 3, "Phân tích cấu thành vi phạm pháp luật",
        "A, 23 tuổi, lén lấy 50 triệu đồng trong tủ của H rồi dùng tiền mua xe và đóng học phí. Phân tích bốn yếu tố cấu thành.",
        [
            ("Mặt khách quan", "Hành vi lén lút lấy 50 triệu; hậu quả H mất tài sản; có quan hệ nhân quả trực tiếp."),
            ("Mặt chủ quan", "Lỗi cố ý trực tiếp: A nhận thức rõ và mong muốn chiếm đoạt. Động cơ đóng học phí không loại trừ lỗi; mục đích là chiếm đoạt tài sản."),
            ("Chủ thể", "A 23 tuổi, giả định có năng lực trách nhiệm hình sự."),
            ("Khách thể", "Quan hệ sở hữu tài sản, không phải bản thân H."),
            ("Định hướng tội danh", "Có dấu hiệu trộm cắp tài sản vì hành vi chiếm đoạt được thực hiện lén lút; kết luận cụ thể cần đối chiếu đầy đủ BLHS."),
        ],
        "Không được ghi “không có hậu quả”, “cố ý gián tiếp” hoặc “khách thể là anh H”.")

    add_exercise(doc, 4, "Phân biệt tội chiếm đoạt qua tình huống sửa xe",
        "A nhờ H sửa chiếc xe chết máy. Sau khi sửa, H bất ngờ nổ máy chạy đi và bán xe. H không dùng vũ lực hoặc đe dọa A. Hành vi có phải tội cướp không?",
        [
            ("Loại trừ tội cướp", "Không có hành vi dùng vũ lực, đe dọa dùng ngay tức khắc vũ lực hoặc thủ đoạn làm A không thể chống cự."),
            ("Hướng định tội", "Cần làm rõ H chiếm đoạt công khai trong hoàn cảnh A có khả năng hay không có khả năng ngăn cản; từ đó xem xét công nhiên chiếm đoạt hoặc tội danh phù hợp khác."),
            ("Kết luận", "Không đủ căn cứ gọi là tội cướp chỉ vì H công khai lấy xe."),
        ],
        "Điểm quan trọng là phương thức chiếm đoạt, không phải cách gọi đời thường “cướp xe”.")

    add_exercise(doc, 5, "Vé số trúng thưởng được giao giữ",
        "C mua vé số rồi đồng ý để Đ cầm giữ. Khi biết trúng 150 triệu đồng, Đ bí mật nhận thưởng và nói dối C rằng vé không trúng. Xác định hướng xử lý.",
        [
            ("Quan hệ ban đầu", "Đ nhận giữ vé hợp pháp theo sự đồng ý của C."),
            ("Thời điểm gian dối", "Thủ đoạn gian dối xuất hiện sau khi Đ đã nhận vé, nhằm che giấu việc chiếm đoạt."),
            ("Hướng tội danh", "Dấu hiệu phù hợp hơn với lạm dụng tín nhiệm chiếm đoạt tài sản so với lừa đảo; tuy nhiên cần làm rõ quyền sở hữu vé, thỏa thuận giữa các bên và đầy đủ điều kiện luật định."),
            ("Lỗi", "Cố ý trực tiếp; mục đích chiếm đoạt tiền thưởng."),
        ])

    add_exercise(doc, 6, "Người giữ tài sản do phạm tội mà có",
        "Sau khi lấy xe của A, H gửi xe tại nhà B. Phân tích trách nhiệm của B trong ba trường hợp: không biết; biết sau khi nhận; biết và hứa hẹn trước.",
        [
            ("Không biết", "B không có lỗi nên không chịu trách nhiệm hình sự về việc chứa chấp."),
            ("Biết sau khi nhận", "Nếu sau đó biết rõ tài sản do phạm tội mà có nhưng vẫn cất giấu, tiêu thụ và đủ dấu hiệu luật định, có thể bị xem xét về tội chứa chấp hoặc tiêu thụ tài sản do người khác phạm tội mà có."),
            ("Hứa hẹn trước", "Nếu B đã hứa hẹn trước và hỗ trợ H theo kế hoạch chung, có thể được xem xét với vai trò đồng phạm của tội chiếm đoạt tương ứng."),
        ])

    add_exercise(doc, 7, "Tai nạn giao thông và dữ kiện còn thiếu",
        "K điều khiển xe lấn trái để vượt, va chạm làm T bị thương tích 60%. Có thể kết luận K phạm tội vi phạm quy định về tham gia giao thông đường bộ không?",
        [
            ("Hành vi", "Có dấu hiệu vi phạm quy tắc giao thông và lỗi vô ý."),
            ("Hậu quả", "Thương tích 60% cho một người, nếu chỉ có dữ kiện này, chưa đạt ngưỡng thông thường 61% của trường hợp cơ bản về thương tích một người."),
            ("Dữ kiện cần bổ sung", "Có chết người không; có thêm người bị thương; thiệt hại tài sản; nồng độ cồn, ma túy; giấy phép; bỏ chạy hoặc không cứu giúp; các tình tiết định khung khác."),
            ("Kết luận", "Chưa đủ dữ kiện để khẳng định trách nhiệm hình sự; vẫn có thể phát sinh xử phạt hành chính và trách nhiệm bồi thường dân sự."),
        ])

    add_exercise(doc, 8, "Chia thừa kế theo pháp luật",
        "Ông D và bà M có tài sản chung 900 triệu đồng. Ông D có tài sản riêng 100 triệu đồng. Ông D chết không để lại di chúc; cha mẹ ông đã mất, hàng thừa kế thứ nhất còn bà M và hai người con A, B. Không có nghĩa vụ tài sản.",
        [
            ("Bước 1 – tách tài sản chung", "Phần của bà M là 450 triệu; phần của ông D trong tài sản chung là 450 triệu."),
            ("Bước 2 – xác định di sản", "450 triệu + 100 triệu tài sản riêng = 550 triệu."),
            ("Bước 3 – xác định người thừa kế", "Bà M, A và B cùng hàng thừa kế thứ nhất."),
            ("Bước 4 – chia", "Mỗi người hưởng 550/3 = 183,33 triệu đồng."),
            ("Kết quả cuối", "Bà M có 450 + 183,33 = 633,33 triệu; A và B mỗi người 183,33 triệu."),
        ],
        "Không được lấy toàn bộ 1 tỷ đồng làm di sản; luôn tách tài sản chung trước.")

    add_exercise(doc, 9, "Di chúc và người thừa kế bắt buộc",
        "Ông X có di sản 900 triệu đồng, lập di chúc để lại toàn bộ cho bạn Y. Ông X còn vợ, một con 10 tuổi và một con 25 tuổi khỏe mạnh. Giả định di chúc hợp pháp và không ai từ chối nhận.",
        [
            ("Suất theo pháp luật", "Nếu chia theo pháp luật, hàng thứ nhất gồm vợ và hai người con: mỗi suất 300 triệu."),
            ("Người hưởng bắt buộc", "Vợ và con 10 tuổi thuộc diện hưởng không phụ thuộc nội dung di chúc; người con 25 tuổi khỏe mạnh không thuộc diện này."),
            ("Phần bắt buộc", "Mỗi người thuộc diện bắt buộc hưởng ít nhất 2/3 suất pháp luật = 200 triệu."),
            ("Phần còn lại", "Y nhận 900 – 200 – 200 = 500 triệu, nếu không có nghĩa vụ và tình tiết khác."),
        ])

    add_exercise(doc, 10, "Tài sản khi chung sống không đăng ký kết hôn",
        "Ông P đang có vợ hợp pháp nhưng chung sống với bà Q và cùng góp tiền mua nhà 1,2 tỷ đồng. Ông P chết. Có thể tự động coi bà Q là vợ và chia đôi nhà hay không?",
        [
            ("Quan hệ hôn nhân", "Bà Q không đương nhiên là vợ hợp pháp và không thuộc hàng thừa kế với tư cách vợ."),
            ("Quan hệ tài sản", "Phải xác định công sức đóng góp và thỏa thuận để xác định phần sở hữu của từng người; không mặc nhiên 50/50 nếu không có căn cứ."),
            ("Di sản", "Chỉ phần tài sản thuộc ông P sau khi xác định đồng sở hữu mới là di sản."),
            ("Con chung", "Con của ông P, nếu quan hệ cha con được xác định, có quyền thừa kế bình đẳng với các con khác."),
        ])

    add_exercise(doc, 11, "Đơn phương chấm dứt hợp đồng lao động",
        "Người lao động L làm theo hợp đồng không xác định thời hạn, gửi thông báo nghỉ việc trước 45 ngày và bàn giao công việc. Công ty không đồng ý nên giữ lương tháng cuối.",
        [
            ("Quyền đơn phương", "Người lao động có quyền đơn phương chấm dứt hợp đồng khi tuân thủ thời hạn báo trước, trừ trường hợp được nghỉ không cần báo trước."),
            ("Đánh giá", "Nếu L thông báo hợp lệ trước 45 ngày và không có cam kết đặc biệt hợp pháp khác, việc công ty “không đồng ý” không làm mất quyền nghỉ."),
            ("Nghĩa vụ công ty", "Phải thanh toán các khoản liên quan trong thời hạn luật định, hoàn thành thủ tục bảo hiểm và xác nhận giấy tờ; không được giữ lương chỉ vì không đồng ý."),
        ])

    add_exercise(doc, 12, "Tranh chấp lao động có bắt buộc hòa giải?",
        "Người lao động bị công ty sa thải và yêu cầu hủy quyết định, nhận trở lại làm việc. Có bắt buộc phải qua hòa giải viên trước khi khởi kiện không?",
        [
            ("Loại tranh chấp", "Tranh chấp lao động cá nhân về xử lý kỷ luật sa thải."),
            ("Thủ tục", "Đây là nhóm tranh chấp không bắt buộc phải qua hòa giải viên lao động trước khi yêu cầu Tòa án giải quyết."),
            ("Lưu ý", "Người lao động vẫn có thể lựa chọn hòa giải nếu thấy phù hợp, nhưng phải chú ý thời hiệu yêu cầu giải quyết."),
        ])

    add_exercise(doc, 13, "Phân biệt tố cáo và khiếu nại",
        "A bị Chủ tịch UBND ban hành quyết định xử phạt mà A cho là trái luật. B phát hiện cán bộ nhận hối lộ trong khi thi hành công vụ. A và B nên sử dụng cơ chế nào?",
        [
            ("Trường hợp A", "Khiếu nại quyết định hành chính hoặc khởi kiện vụ án hành chính để bảo vệ quyền, lợi ích trực tiếp của mình."),
            ("Trường hợp B", "Tố cáo hành vi vi phạm pháp luật; nếu có dấu hiệu tội phạm thì thông tin có thể được chuyển đến cơ quan tiến hành tố tụng."),
            ("Phân biệt cốt lõi", "Khiếu nại hướng tới quyết định/hành vi xâm phạm trực tiếp quyền của người khiếu nại; tố cáo hướng tới hành vi vi phạm pháp luật."),
        ])

    add_exercise(doc, 14, "Hiệu lực của văn bản quy phạm pháp luật",
        "Một thông tư quy định nghĩa vụ mới trái với luật do Quốc hội ban hành. Cơ quan áp dụng nên xử lý nhận thức pháp lý như thế nào?",
        [
            ("Nguyên tắc", "Văn bản cấp dưới phải phù hợp văn bản có hiệu lực pháp lý cao hơn."),
            ("Áp dụng", "Không thể lấy thông tư trái luật làm căn cứ hợp pháp để hạn chế quyền hoặc đặt nghĩa vụ trái luật."),
            ("Xử lý văn bản", "Cơ quan, người có thẩm quyền phải kiểm tra, kiến nghị đình chỉ, bãi bỏ hoặc xử lý theo cơ chế kiểm tra văn bản; người học cần nêu đúng thẩm quyền cụ thể nếu đề cung cấp chủ thể ban hành."),
        ])

    add_exercise(doc, 15, "Bài tổng hợp trách nhiệm pháp lý",
        "Nhân viên N cố ý làm hỏng máy của công ty trị giá 20 triệu đồng vì mâu thuẫn với quản lý. N đồng thời vi phạm nội quy lao động. Hãy xác định các loại trách nhiệm có thể phát sinh.",
        [
            ("Kỷ luật lao động", "Có thể bị xử lý theo nội quy và Bộ luật Lao động nếu bảo đảm căn cứ, nguyên tắc, thời hiệu và thủ tục."),
            ("Dân sự/vật chất", "Có thể phải bồi thường thiệt hại theo căn cứ và phạm vi trách nhiệm phù hợp."),
            ("Hành chính hoặc hình sự", "Tùy giá trị, tính chất, hậu quả và các dấu hiệu luật định, hành vi hủy hoại tài sản có thể bị xử phạt hành chính hoặc truy cứu hình sự."),
            ("Nguyên tắc", "Một hành vi có thể làm phát sinh nhiều loại trách nhiệm có mục đích khác nhau; không được kết luận hình sự nếu chưa đủ dấu hiệu."),
        ])

    add_exercise(doc, 16, "Năng lực hành vi dân sự của người chưa thành niên",
        "M, 16 tuổi, tự mua xe máy trị giá 45 triệu đồng bằng tiền tiết kiệm mà không có sự đồng ý của cha mẹ. Giao dịch có đương nhiên có hiệu lực không?",
        [
            ("Vấn đề pháp lý", "M là người từ đủ 15 đến dưới 18 tuổi; có thể tự mình xác lập nhiều giao dịch nhưng giao dịch liên quan bất động sản, động sản phải đăng ký và giao dịch khác theo luật có thể cần người đại diện đồng ý."),
            ("Đánh giá", "Xe máy là động sản phải đăng ký. Cần sự đồng ý của người đại diện theo pháp luật; thiếu sự đồng ý có thể dẫn đến giao dịch vô hiệu theo yêu cầu của chủ thể có quyền."),
            ("Lưu ý", "Không được kết luận mọi giao dịch của người dưới 18 tuổi đều vô hiệu; phải xét nhóm tuổi, loại tài sản và nguồn tài sản."),
        ])

    add_exercise(doc, 17, "Giao dịch dân sự do bị lừa dối",
        "B biết mảnh đất đang tranh chấp nhưng cố ý che giấu, cung cấp giấy tờ làm A tin đất không có tranh chấp và ký hợp đồng. A phát hiện sau khi đã trả tiền.",
        [
            ("Hành vi", "B cố ý đưa thông tin sai hoặc che giấu thông tin quan trọng làm A hiểu sai và giao kết."),
            ("Hậu quả pháp lý", "A có thể yêu cầu Tòa án tuyên giao dịch vô hiệu do bị lừa dối trong thời hiệu luật định."),
            ("Giải quyết hậu quả", "Các bên hoàn trả cho nhau những gì đã nhận; bên có lỗi gây thiệt hại phải bồi thường."),
            ("Phân biệt", "Không phải mọi vi phạm thông tin đều là tội lừa đảo; trách nhiệm hình sự chỉ đặt ra khi đủ dấu hiệu tội phạm và ý định chiếm đoạt."),
        ])

    add_exercise(doc, 18, "Bồi thường thiệt hại ngoài hợp đồng",
        "C vô ý làm rơi chậu cây từ ban công, làm hỏng xe của D đang đỗ đúng nơi quy định. Chi phí sửa xe 18 triệu đồng.",
        [
            ("Căn cứ", "Có thiệt hại thực tế, hành vi gây thiệt hại, quan hệ nhân quả và lỗi vô ý."),
            ("Trách nhiệm", "C phải bồi thường thiệt hại hợp lý để khôi phục tình trạng tài sản; các bên có thể thỏa thuận mức và phương thức."),
            ("Thiệt hại", "Chi phí sửa chữa hợp lý, thiệt hại do không khai thác được tài sản nếu chứng minh được và khoản khác theo luật."),
            ("Giảm trách nhiệm", "Chỉ xem xét nếu D cũng có lỗi hoặc có căn cứ khác theo luật; dữ kiện cho thấy D đỗ đúng nơi."),
        ])

    add_exercise(doc, 19, "Chiếm hữu ngay tình",
        "A mua chiếc đồng hồ từ B tại cửa hàng, có hóa đơn và không biết B không phải chủ sở hữu. Chủ thật sự C yêu cầu A trả lại.",
        [
            ("Vấn đề", "A có thể là người chiếm hữu ngay tình; tuy nhiên ngay tình không tự động làm phát sinh quyền sở hữu."),
            ("Phân tích", "Cần xác định tài sản bị mất, bị trộm hay được chủ sở hữu giao cho người khác; giao dịch có đền bù hay không; thời hiệu và trường hợp bảo vệ người thứ ba ngay tình."),
            ("Kết luận", "Không đủ dữ kiện để khẳng định A được giữ hay phải trả; bài đúng phải phân nhánh theo nguồn gốc tài sản."),
        ])

    add_exercise(doc, 20, "Thừa kế thế vị",
        "Ông X chết không để lại di chúc. X có vợ Y, con A và con B. A đã chết trước X và có hai con A1, A2. Cha mẹ X đã mất.",
        [
            ("Hàng thừa kế", "Y, A và B thuộc hàng thứ nhất; do A chết trước, A1 và A2 được thừa kế thế vị phần A đáng lẽ được hưởng."),
            ("Cách chia", "Di sản chia thành ba suất: Y một suất, B một suất, nhánh A một suất."),
            ("Trong nhánh A", "A1 và A2 chia đều suất của A."),
            ("Lưu ý", "Không tính A1, A2 thành hai suất ngang độc lập với Y và B."),
        ])

    add_exercise(doc, 21, "Người không được quyền hưởng di sản",
        "K cố ý xâm phạm nghiêm trọng tính mạng của cha là ông P để sớm được nhận thừa kế. Sau đó P vẫn lập di chúc cho K và ghi rõ biết hành vi của K.",
        [
            ("Nguyên tắc", "Người có hành vi cố ý xâm phạm tính mạng người để lại di sản thuộc nhóm không được quyền hưởng."),
            ("Ngoại lệ", "Nếu người để lại di sản đã biết hành vi mà vẫn cho hưởng theo di chúc, người đó vẫn có thể được hưởng theo ý chí này."),
            ("Yêu cầu chứng minh", "Phải chứng minh P biết rõ hành vi và ý chí cho hưởng là tự nguyện, hợp pháp."),
        ])

    add_exercise(doc, 22, "Kết hôn trái pháp luật",
        "Nam 19 tuổi và nữ 17 tuổi tổ chức cưới, sống chung nhưng không đăng ký. Một năm sau họ có con. Xác định quan hệ pháp lý.",
        [
            ("Hôn nhân", "Tại thời điểm chung sống, nữ chưa đủ tuổi kết hôn và hai người không đăng ký; không phát sinh quan hệ vợ chồng hợp pháp."),
            ("Con chung", "Quyền, nghĩa vụ cha mẹ con vẫn được xác định và bảo vệ; cần đăng ký khai sinh, xác định cha nếu cần."),
            ("Tài sản", "Giải quyết theo thỏa thuận; nếu không thỏa thuận thì theo pháp luật dân sự, có tính đến công sức đóng góp và bảo vệ phụ nữ, trẻ em."),
            ("Sau khi đủ điều kiện", "Nếu muốn xác lập hôn nhân, hai người phải thực hiện đăng ký; thời điểm có hiệu lực tính từ đăng ký."),
        ])

    add_exercise(doc, 23, "Tài sản chung và tài sản riêng vợ chồng",
        "Trước khi kết hôn, A có căn nhà. Trong thời kỳ hôn nhân, vợ chồng dùng tiền chung sửa chữa lớn và cho thuê nhà, thu 20 triệu đồng mỗi tháng.",
        [
            ("Căn nhà", "Về nguyên tắc là tài sản riêng của A nếu không có thỏa thuận nhập vào tài sản chung."),
            ("Tiền thuê", "Hoa lợi, lợi tức phát sinh từ tài sản riêng trong thời kỳ hôn nhân thường là tài sản chung, trừ trường hợp pháp luật hoặc thỏa thuận hợp pháp quy định khác."),
            ("Chi phí sửa chữa", "Cần xác định nguồn tiền và thỏa thuận; bên còn lại có thể có yêu cầu hoàn trả hoặc tính công sức khi giải quyết tài sản."),
        ])

    add_exercise(doc, 24, "Kỷ luật sa thải không đúng thủ tục",
        "N tự ý nghỉ việc 6 ngày cộng dồn trong 30 ngày nhưng công ty lập tức ra quyết định sa thải mà không mời N tham dự cuộc họp xử lý kỷ luật.",
        [
            ("Căn cứ nội dung", "Hành vi có thể thuộc trường hợp xem xét sa thải nếu không có lý do chính đáng và đủ số ngày theo luật."),
            ("Thủ tục", "Người sử dụng lao động phải chứng minh lỗi và thực hiện trình tự họp, thông báo, thành phần tham dự, thời hiệu theo luật."),
            ("Kết luận", "Dù có thể có căn cứ về hành vi, quyết định vẫn có nguy cơ trái pháp luật do vi phạm thủ tục."),
            ("Hậu quả", "Công ty có thể phải nhận người lao động trở lại, trả lương, đóng bảo hiểm, bồi thường tùy yêu cầu và phán quyết."),
        ])

    add_exercise(doc, 25, "Phân biệt tranh chấp tập thể về quyền và lợi ích",
        "Tập thể người lao động yêu cầu công ty trả phụ cấp đã ghi trong thỏa ước. Ở tình huống khác, họ yêu cầu bổ sung phụ cấp mới chưa từng được thỏa thuận.",
        [
            ("Tình huống 1", "Tranh chấp tập thể về quyền vì yêu cầu thực hiện điều kiện đã được xác lập trong thỏa ước."),
            ("Tình huống 2", "Tranh chấp tập thể về lợi ích vì yêu cầu xác lập điều kiện lao động mới."),
            ("Ý nghĩa", "Phân loại quyết định cơ chế giải quyết và khả năng tiến hành đình công."),
        ])

    add_exercise(doc, 26, "Xử phạt hành chính sai thẩm quyền",
        "Một người không có thẩm quyền theo luật ký quyết định phạt doanh nghiệp 100 triệu đồng, dù doanh nghiệp thực sự có hành vi vi phạm.",
        [
            ("Tính hợp pháp", "Quyết định hành chính phải đúng cả nội dung và thẩm quyền. Vi phạm thực tế không hợp thức hóa quyết định do người không có thẩm quyền ban hành."),
            ("Quyền của doanh nghiệp", "Có thể khiếu nại hoặc khởi kiện hành chính, yêu cầu hủy quyết định và khôi phục quyền lợi."),
            ("Xử lý vi phạm gốc", "Cơ quan có thẩm quyền vẫn có thể xử lý hành vi vi phạm nếu còn thời hiệu và thực hiện đúng thủ tục."),
        ])

    add_exercise(doc, 27, "Tố cáo nặc danh có tài liệu rõ ràng",
        "Cơ quan nhận email không ghi tên người gửi nhưng nêu rõ cán bộ, thời gian, địa điểm nhận tiền và kèm video.",
        [
            ("Thủ tục tố cáo", "Không xử lý như một tố cáo đầy đủ do không xác định được người tố cáo."),
            ("Xử lý thông tin", "Vì nội dung rõ, có tài liệu cụ thể và cơ sở kiểm tra, cơ quan có thể sử dụng để thanh tra, kiểm tra hoặc chuyển cơ quan điều tra."),
            ("Kết luận", "Sai nếu nói mọi thông tin nặc danh đều phải hủy bỏ và không được xem xét."),
        ])

    add_exercise(doc, 28, "Độ tuổi chịu trách nhiệm hình sự",
        "T, 15 tuổi, thực hiện một hành vi có dấu hiệu tội nghiêm trọng. Có thể kết luận T phải chịu trách nhiệm hình sự chỉ dựa vào phân loại “nghiêm trọng” không?",
        [
            ("Quy tắc", "Người từ đủ 14 đến dưới 16 tuổi chỉ chịu trách nhiệm về các tội cụ thể được Điều 12 BLHS liệt kê."),
            ("Kết luận", "Không thể chỉ dựa vào nhãn ít nghiêm trọng, nghiêm trọng, rất nghiêm trọng hay đặc biệt nghiêm trọng."),
            ("Cách làm", "Phải xác định tội danh dự kiến, kiểm tra tội đó có nằm trong danh mục và đánh giá năng lực nhận thức, điều khiển hành vi."),
        ])

    add_exercise(doc, 29, "Phòng vệ chính đáng và vượt quá giới hạn",
        "A bị B dùng dao tấn công. A giằng được gậy và đánh B để chấm dứt tấn công. Khi B đã bỏ chạy, A tiếp tục đuổi theo đánh gây thương tích nặng.",
        [
            ("Giai đoạn chống trả", "Hành vi cần thiết nhằm chống lại cuộc tấn công đang xảy ra có thể là phòng vệ chính đáng."),
            ("Giai đoạn đuổi theo", "Khi nguy cơ đã chấm dứt, việc tiếp tục tấn công không còn nằm trong phạm vi phòng vệ; cần xem xét lỗi và tội danh tương ứng."),
            ("Bài học", "Phải tách diễn biến theo từng thời điểm, không đánh giá toàn bộ chuỗi hành vi bằng một kết luận duy nhất."),
        ])

    add_exercise(doc, 30, "Hiệu lực và giới hạn quyền con người",
        "Một cơ quan cấp sở ban hành công văn nội bộ yêu cầu mọi người dân trong địa phương phải cung cấp dữ liệu đời tư, nếu không sẽ bị từ chối dịch vụ công.",
        [
            ("Quyền liên quan", "Quyền về đời sống riêng tư và bảo vệ dữ liệu/thông tin cá nhân."),
            ("Giới hạn quyền", "Việc hạn chế quyền phải có căn cứ luật, nhằm mục đích hiến định cần thiết và bảo đảm phạm vi phù hợp; công văn không thể tự đặt nghĩa vụ và chế tài vượt luật."),
            ("Đánh giá", "Cần kiểm tra thẩm quyền, mục đích, loại dữ liệu, mức cần thiết và cơ chế bảo vệ. Quy định chung chung có nguy cơ trái pháp luật và xâm phạm quyền."),
        ])

    doc.add_page_break()
    doc.add_heading("B. BÀI TẬP TỔNG HỢP DÀI – DẠNG ĐỀ THI", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Các bài dưới đây được trình bày theo dạng tình huống dài. Đáp án không chỉ đưa ra kết luận mà phải phân tích "
        "phương thức thực hiện, thời điểm phát sinh ý định, bốn yếu tố cấu thành, dữ kiện còn thiếu và căn cứ pháp luật."
    )
    set_font(r)

    add_exercise(doc, 31, "Vé số trúng thưởng – phân biệt lừa đảo và lạm dụng tín nhiệm",
        "Chiều ngày 20/9/2018, Đ cùng C đang ngồi uống nước thì có một em bé mời mua vé số. C bỏ 15.000 đồng mua 3 tờ. Trong khi C trả tiền, Đ nhận 3 tờ vé số từ người bán, cất vào túi và nói: “Để tôi cầm cho may mắn, nếu trúng thưởng thì tối nay lại nhậu”. C chỉ cười, không phản đối. Sáng hôm sau, biết vé trúng 150 triệu đồng, Đ đi nhận thưởng rồi gọi điện nói dối C rằng cả 3 vé đều trượt. Đ dùng tiền mua xe máy. Khi C biết sự việc và yêu cầu trả tiền, Đ tiếp tục nói dối rằng vé không trúng và đã xé bỏ. Hành vi của Đ cấu thành tội gì? Tại sao?",
        [
            ("Xác định quyền đối với tài sản", "C là người bỏ tiền mua và là chủ sở hữu các tờ vé cùng quyền nhận thưởng. Đ chỉ là người được C đồng ý cho cầm giữ; việc C cười, không phản đối cho thấy Đ nhận tài sản một cách hợp pháp ban đầu."),
            ("Loại trừ tội trộm cắp", "Đ không lén lút lấy vé khỏi sự quản lý của C. Việc cầm vé diễn ra công khai và được C chấp nhận."),
            ("Loại trừ hướng lừa đảo", "Đ không dùng thủ đoạn gian dối trước khi C giao vé. Lời nói dối “vé không trúng” xuất hiện sau khi Đ đã nhận giữ vé và đã biết kết quả. Gian dối không phải nguyên nhân làm C giao tài sản."),
            ("Hướng định tội phù hợp", "Đ nhận tài sản hợp pháp thông qua thỏa thuận giữ hộ, sau đó nảy sinh ý định chiếm đoạt, đi nhận thưởng, dùng thủ đoạn gian dối che giấu và từ chối hoàn trả. Hành vi có dấu hiệu của tội lạm dụng tín nhiệm chiếm đoạt tài sản theo Điều 175 BLHS, nếu cơ quan tiến hành tố tụng xác định đầy đủ quan hệ giao tài sản và các điều kiện cấu thành."),
            ("Khách thể", "Quan hệ sở hữu của C đối với vé số và khoản tiền thưởng."),
            ("Mặt khách quan", "Nhận giữ vé hợp pháp; sau khi vé trúng đã tự nhận 150 triệu, nói dối, sử dụng tiền và kiên quyết không trả. Hậu quả là C mất quyền hưởng 150 triệu; hành vi chiếm đoạt là nguyên nhân trực tiếp."),
            ("Chủ thể", "Đ phải là người có năng lực trách nhiệm hình sự và đủ tuổi. Đề không cho tuổi nên bài làm phải nêu giả định Đ đủ điều kiện chủ thể."),
            ("Mặt chủ quan", "Lỗi cố ý trực tiếp: Đ biết tiền thưởng thuộc C nhưng mong muốn chiếm đoạt. Mục đích là chiếm đoạt 150 triệu; động cơ tư lợi."),
            ("Giá trị tài sản", "Giá trị chiếm đoạt là 150 triệu đồng. Khi xác định khung trách nhiệm cần đối chiếu khoản tương ứng của Điều 175 và pháp luật áp dụng tại thời điểm hành vi; bài Pháp luật đại cương chỉ nên kết luận tội danh khi không yêu cầu định khung."),
            ("Kết luận", "Có cơ sở xác định theo hướng lạm dụng tín nhiệm chiếm đoạt tài sản, không phải lừa đảo, vì tài sản được giao hợp pháp trước và thủ đoạn gian dối xuất hiện sau để chiếm đoạt."),
        ],
        "Điểm quyết định là thời điểm xuất hiện thủ đoạn gian dối và cách Đ có được tài sản. Không được thấy có nói dối rồi kết luận ngay là lừa đảo.")

    add_exercise(doc, 32, "Sinh viên lấy 50 triệu – phân tích đầy đủ cấu thành vi phạm",
        "Nguyễn Văn A, 23 tuổi, là sinh viên năm 2. Do không có tiền đóng học phí, ngày 01/02/2020 A đến chơi và ở lại nhà anh Hùng. Ngày 02/02/2020, lợi dụng lúc anh Hùng đi vắng và tủ không khóa, A lấy 50 triệu đồng. A dùng 20 triệu mua xe máy, số còn lại nộp học phí. Hãy phân tích các yếu tố cấu thành vi phạm pháp luật.",
        [
            ("Xác định hành vi", "A lén lút lấy 50 triệu thuộc sở hữu của Hùng khi Hùng vắng nhà. Đây là hành vi xác định, trái pháp luật và xâm phạm quan hệ sở hữu."),
            ("Mặt khách quan – hành vi", "Thủ đoạn là lợi dụng chủ tài sản vắng mặt và tủ không khóa để bí mật chiếm đoạt. Tủ không khóa không làm mất quyền sở hữu và không phải sự đồng ý giao tài sản."),
            ("Mặt khách quan – hậu quả", "Hùng bị mất 50 triệu đồng. Vì vậy đáp án cũ ghi “không có hậu quả” là sai."),
            ("Quan hệ nhân quả", "Hậu quả mất tài sản phát sinh trực tiếp từ hành vi A lấy và mang tiền đi sử dụng."),
            ("Dấu hiệu khác", "Thời gian là 02/02/2020; địa điểm nhà Hùng; tiền 50 triệu là đối tượng tác động. Không nhất thiết phải xác định “phương tiện phạm tội” nếu không có."),
            ("Mặt chủ quan – lỗi", "A nhận thức tiền thuộc Hùng, biết hành vi lấy là trái pháp luật và mong muốn chiếm đoạt để sử dụng nên là cố ý trực tiếp, không phải cố ý gián tiếp."),
            ("Động cơ và mục đích", "Động cơ là cần tiền đóng học phí và mua xe; mục đích pháp lý trực tiếp vẫn là chiếm đoạt tài sản. Hoàn cảnh khó khăn không loại trừ lỗi, có thể chỉ được xem xét khi lượng hình nếu pháp luật cho phép."),
            ("Chủ thể", "A 23 tuổi, giả định có khả năng nhận thức và điều khiển hành vi, nên có năng lực trách nhiệm pháp lý và trách nhiệm hình sự."),
            ("Khách thể", "Quan hệ sở hữu tài sản được pháp luật bảo vệ; không ghi khách thể là “anh Hùng”."),
            ("Phân loại vi phạm", "Hành vi có dấu hiệu tội trộm cắp tài sản vì phương thức chiếm đoạt là lén lút. Với giá trị 50 triệu đồng, nếu định khung cần đối chiếu khoản tương ứng của Điều 173 BLHS hiện hành."),
            ("Trách nhiệm khác", "A phải hoàn trả tài sản, bồi thường thiệt hại và có thể bị xử lý kỷ luật theo quy định của trường nếu hành vi ảnh hưởng tư cách sinh viên, nhưng các trách nhiệm này không thay thế trách nhiệm hình sự."),
        ],
        "Bốn lỗi thường gặp: hậu quả “không có”; lỗi cố ý gián tiếp; khách thể là Hùng; động cơ đóng học phí làm mất tính phạm tội.")

    add_exercise(doc, 33, "Xác định giả định – quy định – chế tài trong quy phạm xử phạt",
        "Phân tích quy phạm: “Phạt tiền từ 100.000 đồng đến 200.000 đồng đối với người điều khiển, người ngồi trên xe máy không đội mũ bảo hiểm hoặc đội mũ bảo hiểm không cài đúng quy cách khi tham gia giao thông trên đường bộ” (trích quy định tại Nghị định 71/2012/NĐ-CP sửa đổi Nghị định 34/2010/NĐ-CP).",
        [
            ("Lưu ý về hiệu lực", "Hai nghị định được dẫn trong đề đã hết hiệu lực. Tuy nhiên nếu đề chỉ yêu cầu nhận diện cấu trúc quy phạm thì vẫn có thể phân tích chính câu chữ được cung cấp; không dùng mức phạt này để tư vấn pháp luật hiện hành."),
            ("Giả định", "Chủ thể là người điều khiển hoặc người ngồi trên xe máy; hoàn cảnh là khi tham gia giao thông trên đường bộ; điều kiện là không đội mũ hoặc đội nhưng không cài đúng quy cách."),
            ("Quy định", "Phần quy định không được viết trực tiếp dưới dạng mệnh lệnh mà được hiểu ngầm: người điều khiển và người ngồi trên xe máy khi tham gia giao thông phải đội mũ bảo hiểm và cài đúng quy cách."),
            ("Chế tài", "Phạt tiền từ 100.000 đến 200.000 đồng đối với chủ thể vi phạm."),
            ("Giải thích cấu trúc", "Cụm mô tả người và hành vi trong điều khoản xử phạt vừa xác định phạm vi áp dụng của chế tài; quy tắc phải đội mũ có thể được quy định rõ ở luật giao thông và được dẫn chiếu ngầm."),
            ("Kết luận", "Giả định trả lời ai, khi nào, trong hoàn cảnh nào; quy định trả lời phải làm gì; chế tài trả lời hậu quả bất lợi khi không thực hiện."),
        ],
        "Không lấy cụm “không đội mũ” làm phần quy định cấm một cách máy móc mà bỏ việc diễn đạt quy tắc tích cực “phải đội và cài đúng”.")

    add_exercise(doc, 34, "Hai lần mở thừa kế – ông Thành, bà An và bà Mai",
        "Ông Thành và bà An kết hôn hợp pháp, có hai con Hà (1975), Hồng (1977), tạo lập ngôi nhà số 30 trị giá 540 triệu đồng. Năm 2000, dù đang có vợ và được bà An đồng ý, ông Thành chung sống với bà Mai, có hai con Hằng (2001), Thái (2002). Ông Thành và bà Mai mua ngôi nhà số 14 trị giá 120 triệu đồng để ba mẹ con bà Mai ở. Năm 2018 bà An chết không để lại di chúc; năm 2020 ông Thành chết không để lại di chúc. Xác định di sản và chia thừa kế. Giả định cha mẹ của ông Thành, bà An đều đã chết, không có nghĩa vụ tài sản và công sức mua nhà số 14 của ông Thành, bà Mai ngang nhau.",
        [
            ("Tư cách của bà Mai", "Việc bà An đồng ý không làm quan hệ giữa ông Thành và bà Mai hợp pháp, vì ông Thành đang có vợ và pháp luật bảo vệ chế độ một vợ một chồng. Bà Mai không phải vợ hợp pháp, không hưởng thừa kế với tư cách vợ."),
            ("Quyền của Hằng và Thái", "Con sinh từ quan hệ không hợp pháp vẫn có quyền bình đẳng. Khi quan hệ cha con được xác định, Hằng và Thái thuộc hàng thừa kế thứ nhất của ông Thành."),
            ("Nhà số 14", "Không gọi là tài sản chung vợ chồng Thành–Mai. Đây là tài sản đồng sở hữu theo công sức đóng góp. Theo giả định đóng góp ngang nhau, ông Thành và bà Mai mỗi người có 60 triệu."),
            ("Lần mở thừa kế thứ nhất", "Nhà số 30 là tài sản chung Thành–An: mỗi người 270 triệu. Di sản của bà An là 270 triệu."),
            ("Người hưởng di sản bà An", "Hàng thứ nhất gồm ông Thành, Hà và Hồng theo giả định cha mẹ bà An đã chết. Mỗi người hưởng 270/3 = 90 triệu."),
            ("Tài sản ông Thành sau lần thứ nhất", "Ông Thành có 270 triệu là phần của mình trong nhà số 30 và 90 triệu thừa kế từ bà An, tổng 360 triệu."),
            ("Lần mở thừa kế thứ hai", "Di sản ông Thành năm 2020 gồm 360 triệu và 60 triệu phần sở hữu trong nhà số 14, tổng 420 triệu."),
            ("Người hưởng di sản ông Thành", "Bà An đã chết; bà Mai không phải vợ hợp pháp; bốn người con Hà, Hồng, Hằng, Thái cùng hàng thứ nhất. Mỗi người hưởng 420/4 = 105 triệu."),
            ("Kết quả cuối", "Hà: 90 + 105 = 195 triệu; Hồng: 195 triệu; Hằng: 105 triệu; Thái: 105 triệu. Bà Mai giữ 60 triệu là phần sở hữu của mình trong nhà số 14, không phải phần thừa kế."),
            ("Kiểm tra", "Tổng quyền lợi từ di sản bà An là 270 triệu; tổng di sản ông Thành là 420 triệu. Các con được hưởng bình đẳng, không phân biệt con trong quan hệ hôn nhân hợp pháp hay không."),
            ("Phân nhánh cần nêu", "Nếu cha mẹ của người chết còn sống, có nghĩa vụ tài sản hoặc tỷ lệ đóng góp vào nhà số 14 không bằng nhau thì kết quả thay đổi. Bài thi nên ghi rõ các giả định."),
        ],
        "Điểm khó là phải tách hai lần mở thừa kế, tách tài sản chung và không coi bà Mai là vợ hợp pháp.")

    add_exercise(doc, 35, "Chia thừa kế – ông Dũng, bà Dung và bà Thơm",
        "Ông Dũng và bà Dung kết hôn hợp pháp, có hai con Hoa, Hà và tài sản chung là nhà trị giá 850 triệu đồng. Ông Dũng có tài sản riêng 9 triệu. Năm 1980, trong khi hôn nhân với bà Dung còn tồn tại, ông Dũng chung sống với bà Thơm và có con Thu. Năm 1997 bà Dung chết không di chúc; năm 2000 ông Dũng chết không di chúc. Sau khi bà Dung chết, ông Dũng và bà Thơm tiếp tục chung sống nhưng không đăng ký. Giả định cha mẹ hai người đã chết và không có nghĩa vụ tài sản.",
        [
            ("Di sản bà Dung", "Nhà 850 triệu là tài sản chung nên phần của bà Dung là 425 triệu. Đây mới là di sản lần thứ nhất."),
            ("Người thừa kế bà Dung", "Ông Dũng, Hoa và Hà cùng hàng thứ nhất; mỗi người hưởng 425/3 = 141,666 triệu đồng."),
            ("Tài sản của ông Dũng", "Phần của ông trong nhà: 425 triệu; tài sản riêng: 9 triệu; thừa kế từ bà Dung: 141,666 triệu. Tổng khoảng 575,666 triệu."),
            ("Tư cách bà Thơm", "Quan hệ năm 1980 hình thành khi ông Dũng đang có vợ, vi phạm một vợ một chồng; sau khi bà Dung chết hai người vẫn không đăng ký. Theo dữ kiện, bà Thơm không được xác định là vợ hợp pháp của ông Dũng và không hưởng với tư cách vợ."),
            ("Quyền của Thu", "Thu là con ông Dũng nên cùng hàng thừa kế thứ nhất với Hoa và Hà."),
            ("Chia di sản ông Dũng", "Hoa, Hà, Thu mỗi người hưởng 575,666/3 ≈ 191,889 triệu đồng."),
            ("Kết quả cộng dồn", "Hoa và Hà mỗi người có khoảng 141,666 + 191,889 = 333,555 triệu; Thu khoảng 191,889 triệu."),
            ("Sửa lỗi đáp án cũ", "Không ghi di sản ngay từ đầu là 859 triệu; đó là tổng giá trị nhà và tài sản riêng, chưa tách tài sản chung. Cũng cần tính nhất quán, tránh đổi 141,67 thành 141,76."),
        ],
        "Nếu đề có dữ kiện về quan hệ hôn nhân thực tế được pháp luật chuyển tiếp công nhận thì phải phân tích thêm; nhưng quan hệ vi phạm hôn nhân đang tồn tại không được hợp pháp hóa chỉ vì có sự đồng ý.")

    add_exercise(doc, 36, "Lấy xe sau khi được nhờ sửa – định tội và trách nhiệm người giữ xe",
        "Chị A đi xe máy thì xe chết máy. A nhờ H sửa. Sau khi sửa, H ngồi lên xe, nổ máy và phóng đi dù A hô giữ lại; sau đó H gửi xe ở nhà B rồi bán được 12 triệu. H phạm tội gì? B có chịu trách nhiệm không?",
        [
            ("Không phải tội cướp", "H không dùng vũ lực, đe dọa dùng ngay tức khắc vũ lực hoặc thủ đoạn làm A lâm vào tình trạng không thể chống cự."),
            ("Không kết luận máy móc", "Cần xác định A có giao quyền quản lý xe cho H đến mức nào, thời điểm H nảy sinh ý định và khả năng ngăn cản của A."),
            ("Hướng thứ nhất", "Nếu H chỉ được tiếp cận để sửa, chưa được giao tài sản theo một quan hệ tín nhiệm có ý nghĩa pháp lý và công khai lấy xe trong tình trạng A không thể ngăn cản, có thể xem xét công nhiên chiếm đoạt."),
            ("Hướng thứ hai", "Nếu xác định A đã giao xe cho H quản lý/sửa chữa hợp pháp rồi H lợi dụng việc được giao để chiếm đoạt, có thể xem xét lạm dụng tín nhiệm khi đủ dấu hiệu."),
            ("Không phải cứ phóng xe là cướp giật", "Cướp giật đòi hỏi hành vi công khai, nhanh chóng chiếm lấy tài sản từ sự quản lý của chủ thể. Việc H đang trực tiếp sửa và nắm xe làm tình huống cần đánh giá thận trọng."),
            ("Trường hợp B không biết", "B không có lỗi, không chịu trách nhiệm về việc chứa chấp."),
            ("B biết sau khi nhận", "Nếu biết rõ xe do phạm tội mà có nhưng vẫn cất giấu, tiêu thụ và đủ điều kiện, B có thể chịu trách nhiệm về tội chứa chấp hoặc tiêu thụ tài sản do người khác phạm tội mà có."),
            ("B hứa hẹn trước", "Nếu B đã thống nhất, hứa giúp cất hoặc bán xe trước khi H chiếm đoạt, B có thể là đồng phạm của tội chiếm đoạt tương ứng."),
            ("Kết luận", "Đáp án “tội cướp của” là sai. Bài tốt phải loại trừ tội cướp và phân nhánh tội danh theo bản chất việc giao, quản lý tài sản."),
        ])


def add_advanced_exercises_only(doc):
    doc.add_page_break()
    doc.add_heading("PHẦN III. BÀI TẬP NÂNG CAO CÓ LỜI GIẢI CHI TIẾT", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Mỗi bài được giải theo trình tự: xác định quan hệ pháp luật hoặc vấn đề cần giải quyết; nêu quy tắc áp dụng; "
        "đối chiếu từng dữ kiện; phân biệt với trường hợp gần giống; kết luận và chỉ rõ dữ kiện có thể làm thay đổi đáp án."
    )
    set_font(r)

    add_exercise(doc, 1, "Vé số trúng thưởng – thời điểm phát sinh ý định chiếm đoạt",
        "C bỏ tiền mua 3 tờ vé số và đồng ý để Đ cầm giữ. Sau khi biết vé trúng 150 triệu đồng, Đ tự nhận thưởng, nói dối rằng vé không trúng, dùng tiền mua xe và không trả C. Xác định tội danh và phân tích cấu thành.",
        [
            ("Quan hệ ban đầu", "C là chủ sở hữu vé và quyền nhận thưởng. Đ nhận vé công khai, được C chấp nhận nên việc có tài sản ban đầu là hợp pháp, mang tính giữ hộ."),
            ("Phân biệt lừa đảo", "Lừa đảo đòi hỏi gian dối có trước, làm chủ tài sản tin tưởng và tự nguyện giao tài sản. Trong tình huống, lời nói dối xuất hiện sau khi Đ đã nhận vé; vì vậy gian dối không phải nguyên nhân làm C giao vé."),
            ("Phân biệt trộm cắp", "Không có hành vi lén lút lấy vé khỏi sự quản lý của C."),
            ("Hướng định tội", "Đ nhận tài sản hợp pháp rồi nảy sinh ý định chiếm đoạt, nhận thưởng, che giấu và từ chối hoàn trả. Hành vi có dấu hiệu lạm dụng tín nhiệm chiếm đoạt tài sản theo Điều 175 BLHS nếu đủ điều kiện cấu thành."),
            ("Khách thể và đối tượng", "Khách thể là quan hệ sở hữu; đối tượng tác động là vé số và khoản tiền thưởng 150 triệu đồng."),
            ("Mặt khách quan", "Hành vi nhận giữ, tự nhận thưởng, gian dối và sử dụng tiền; hậu quả C mất 150 triệu; có quan hệ nhân quả trực tiếp."),
            ("Mặt chủ quan", "Cố ý trực tiếp, mục đích chiếm đoạt, động cơ tư lợi."),
            ("Chủ thể và kết luận", "Phải giả định Đ đủ tuổi và có năng lực trách nhiệm hình sự. Kết luận phù hợp hơn là lạm dụng tín nhiệm, không phải cứ có lời nói dối là lừa đảo."),
        ])

    add_exercise(doc, 2, "Lén lấy 50 triệu để đóng học phí",
        "A, 23 tuổi, lợi dụng chủ nhà đi vắng và tủ không khóa để lấy 50 triệu đồng, dùng 20 triệu mua xe và số còn lại đóng học phí. Phân tích các yếu tố cấu thành vi phạm pháp luật và hướng tội danh.",
        [
            ("Mặt khách quan", "A lén lút lấy tài sản; hậu quả chủ sở hữu mất 50 triệu; quan hệ nhân quả trực tiếp. Việc tủ không khóa không phải là từ bỏ quyền sở hữu."),
            ("Mặt chủ quan", "A biết tiền không thuộc mình, biết hành vi trái luật và mong muốn chiếm đoạt nên là cố ý trực tiếp. Động cơ khó khăn học phí không loại trừ lỗi."),
            ("Chủ thể", "A 23 tuổi và giả định có khả năng nhận thức, điều khiển hành vi nên có năng lực trách nhiệm hình sự."),
            ("Khách thể", "Quan hệ sở hữu, không phải bản thân người bị mất tiền."),
            ("Phân loại vi phạm", "Có dấu hiệu tội trộm cắp tài sản theo Điều 173 BLHS vì phương thức là lén lút; giá trị 50 triệu được dùng để xác định khung nếu đề yêu cầu."),
            ("Trách nhiệm bổ sung", "A phải hoàn trả, bồi thường; trách nhiệm dân sự hoặc kỷ luật không thay thế trách nhiệm hình sự."),
        ])

    add_exercise(doc, 3, "Cấu trúc quy phạm xử phạt đội mũ bảo hiểm",
        "Xác định giả định, quy định, chế tài trong câu: “Phạt tiền từ 100.000 đồng đến 200.000 đồng đối với người điều khiển, người ngồi trên xe máy không đội mũ bảo hiểm hoặc đội mũ bảo hiểm không cài đúng quy cách khi tham gia giao thông trên đường bộ”.",
        [
            ("Giả định", "Người điều khiển hoặc người ngồi trên xe máy; trong hoàn cảnh tham gia giao thông đường bộ; có hành vi không đội hoặc cài không đúng."),
            ("Quy định", "Được thể hiện ngầm: các chủ thể trên phải đội mũ bảo hiểm và cài đúng quy cách khi tham gia giao thông."),
            ("Chế tài", "Phạt tiền từ 100.000 đến 200.000 đồng."),
            ("Lưu ý hiệu lực", "Nghị định 71/2012 và Nghị định 34/2010 trong đề cũ đã hết hiệu lực. Ta chỉ phân tích cấu trúc câu được cung cấp, không dùng mức phạt này như pháp luật hiện hành."),
            ("Nhận xét", "Ba bộ phận là cấu trúc logic; quy định về nghĩa vụ và chế tài có thể nằm ở các văn bản khác nhau."),
        ])

    add_exercise(doc, 4, "Hai lần mở thừa kế – Thành, An, Mai và bốn người con",
        "Thành và An kết hôn hợp pháp, có nhà chung 540 triệu và hai con Hà, Hồng. Khi hôn nhân còn tồn tại, Thành chung sống với Mai, có Hằng, Thái; Thành và Mai cùng mua nhà 120 triệu. An chết năm 2018, Thành chết năm 2020, đều không để lại di chúc. Giả định cha mẹ họ đã chết, không có nghĩa vụ và Thành–Mai đóng góp nhà 120 triệu ngang nhau.",
        [
            ("Tư cách Mai", "Mai không phải vợ hợp pháp vì quan hệ hình thành khi Thành đang có vợ; sự đồng ý của An không làm hợp pháp hóa việc vi phạm một vợ một chồng."),
            ("Di sản của An", "Nhà 540 triệu chia phần sở hữu vợ chồng: An 270, Thành 270. Di sản An là 270 triệu."),
            ("Chia lần một", "Hàng thứ nhất của An gồm Thành, Hà, Hồng. Mỗi người hưởng 90 triệu. Thành lúc này có 270 + 90 = 360 triệu."),
            ("Nhà của Thành và Mai", "Không phải tài sản chung vợ chồng mà là đồng sở hữu theo đóng góp. Theo giả định, mỗi người 60 triệu."),
            ("Di sản của Thành", "360 + 60 = 420 triệu."),
            ("Chia lần hai", "Bốn con Hà, Hồng, Hằng, Thái bình đẳng, mỗi người 105 triệu; Mai không hưởng với tư cách vợ."),
            ("Kết quả", "Hà và Hồng mỗi người 195 triệu sau hai lần; Hằng, Thái mỗi người 105 triệu; Mai giữ 60 triệu là tài sản của mình."),
            ("Dữ kiện làm đổi đáp án", "Cha mẹ người chết còn sống, nghĩa vụ tài sản, tỷ lệ đóng góp khác hoặc bằng chứng về tài sản riêng."),
        ])

    add_exercise(doc, 5, "Hai lần mở thừa kế – Dũng, Dung, Thơm",
        "Dũng và Dung có nhà chung 850 triệu, Dũng có 9 triệu riêng và hai con Hoa, Hà. Dũng chung sống với Thơm khi hôn nhân với Dung còn tồn tại và có con Thu. Dung chết năm 1997, Dũng chết năm 2000, không di chúc. Giả định cha mẹ đã chết, không có nghĩa vụ.",
        [
            ("Di sản Dung", "Phần của Dung trong nhà là 425 triệu."),
            ("Chia di sản Dung", "Dũng, Hoa, Hà mỗi người hưởng 425/3 = 141,666 triệu."),
            ("Di sản Dũng", "425 triệu phần nhà + 9 triệu riêng + 141,666 triệu thừa kế = 575,666 triệu."),
            ("Tư cách Thơm", "Không phải vợ hợp pháp theo dữ kiện; không hưởng với tư cách vợ."),
            ("Chia di sản Dũng", "Hoa, Hà, Thu là ba con cùng hàng, mỗi người khoảng 191,889 triệu."),
            ("Kết quả", "Hoa và Hà mỗi người khoảng 333,555 triệu; Thu khoảng 191,889 triệu."),
            ("Sai lầm cần tránh", "Không lấy 859 triệu làm di sản ngay từ đầu; phải tách tài sản chung và giải quyết từng thời điểm mở thừa kế."),
        ])

    add_exercise(doc, 6, "Lấy xe khi được nhờ sửa và trách nhiệm người cất giấu",
        "A nhờ H sửa xe. Sau khi sửa, H nổ máy phóng đi, gửi xe ở nhà B và bán được 12 triệu. H không dùng vũ lực. Phân tích tội danh của H và trách nhiệm của B.",
        [
            ("Loại trừ cướp", "Không có vũ lực, đe dọa dùng ngay tức khắc vũ lực hoặc thủ đoạn làm A không thể chống cự."),
            ("Phân nhánh đối với H", "Nếu A chỉ cho tiếp cận để sửa và H công khai lấy trong hoàn cảnh A không thể ngăn cản, xem xét công nhiên chiếm đoạt; nếu A đã giao xe cho H quản lý/sửa chữa và H lợi dụng sự tín nhiệm, xem xét lạm dụng tín nhiệm. Cần làm rõ thời điểm nảy sinh ý định."),
            ("Không kết luận theo ngôn ngữ đời thường", "“Phóng xe đi” không đương nhiên là cướp giật; phải xác định phương thức chiếm đoạt và trạng thái quản lý tài sản."),
            ("B không biết", "Không có lỗi nên không chịu trách nhiệm hình sự."),
            ("B biết sau khi nhận", "Nếu biết rõ tài sản do phạm tội mà có nhưng vẫn cất giấu hoặc bán, có thể phạm tội chứa chấp hoặc tiêu thụ tài sản do người khác phạm tội mà có."),
            ("B hứa hẹn trước", "Nếu thống nhất hỗ trợ từ trước, B có thể là đồng phạm của tội H thực hiện."),
        ])

    add_exercise(doc, 7, "Bán hàng qua mạng – lừa đảo hay tranh chấp hợp đồng?",
        "M đăng bán máy ảnh 40 triệu, nhận đủ tiền từ N nhưng không giao hàng. M khai do nguồn hàng bị chậm và muốn hoàn tiền; N cho rằng M lừa đảo. Điều tra xác định trước khi nhận tiền M có hoặc không có máy ảnh trong hai giả thiết khác nhau.",
        [
            ("Giả thiết M không có hàng và dùng ảnh giả", "Nếu ngay từ đầu M dựng thông tin gian dối nhằm làm N chuyển tiền rồi cắt liên lạc, hành vi có dấu hiệu lừa đảo chiếm đoạt tài sản."),
            ("Giả thiết M có hàng nhưng phát sinh vi phạm sau", "Nếu giao dịch được xác lập thật, M có khả năng thực hiện nhưng sau đó chậm giao, tranh chấp trước hết là dân sự. Chỉ chuyển sang hình sự khi chứng minh ý định chiếm đoạt và hành vi thỏa mãn cấu thành."),
            ("Điểm phân biệt", "Thời điểm có ý định chiếm đoạt, tính giả tạo của thông tin trước khi giao tiền, cách sử dụng tiền và hành vi sau khi nhận."),
            ("Trách nhiệm dân sự", "Hoàn trả tiền, lãi, phạt vi phạm và bồi thường nếu có căn cứ."),
            ("Kết luận", "Không hình sự hóa quan hệ dân sự chỉ vì một bên không thực hiện đúng hạn; đồng thời không được dùng hợp đồng giả để che giấu chiếm đoạt."),
        ])

    add_exercise(doc, 8, "Người chưa thành niên, đồng phạm và tội chiếm đoạt",
        "P 17 tuổi rủ Q 15 tuổi giật điện thoại trị giá 30 triệu từ tay người đi đường. P lái xe, Q ngồi sau trực tiếp giật. Sau đó R 20 tuổi biết rõ nguồn gốc nhưng vẫn bán giúp và lấy 2 triệu tiền công.",
        [
            ("Tội danh dự kiến", "Hành vi công khai, nhanh chóng giật tài sản rồi tẩu thoát có dấu hiệu cướp giật tài sản."),
            ("Vai trò P và Q", "P tổ chức, điều khiển xe tạo điều kiện; Q trực tiếp thực hành. Cả hai có sự thống nhất ý chí nên có thể là đồng phạm."),
            ("Tuổi của P", "P 17 tuổi phải chịu trách nhiệm về tội phạm theo quy định chung nếu có năng lực."),
            ("Tuổi của Q", "Q 15 tuổi chỉ chịu trách nhiệm nếu tội cướp giật thuộc danh mục Điều 12 BLHS và hành vi đạt điều kiện luật định; không chỉ dựa vào tên loại tội."),
            ("Trách nhiệm R", "Nếu không hứa hẹn trước, R không là đồng phạm của vụ cướp giật nhưng có thể phạm tội chứa chấp hoặc tiêu thụ tài sản do người khác phạm tội mà có."),
            ("Nếu R hứa bán từ trước", "R có thể bị xem xét là người giúp sức trong đồng phạm."),
        ])

    add_exercise(doc, 9, "Hợp đồng cộng tác viên, tiền lương và sa thải",
        "Công ty ký “hợp đồng cộng tác viên” 24 tháng với L, yêu cầu làm đủ 8 giờ/ngày, chấm công, nhận lương cố định và chịu quản lý trực tiếp. L đi muộn nhiều lần, công ty cắt nửa tháng lương và thông báo nghỉ ngay không họp kỷ luật.",
        [
            ("Nhận diện quan hệ", "Có việc làm, trả lương và quản lý, điều hành nên dù mang tên cộng tác viên vẫn có dấu hiệu hợp đồng lao động."),
            ("Loại hợp đồng", "Thời hạn 24 tháng là hợp đồng xác định thời hạn nếu quan hệ lao động được công nhận."),
            ("Cắt lương", "Pháp luật cấm phạt tiền, cắt lương thay xử lý kỷ luật."),
            ("Chấm dứt", "Công ty phải xác định căn cứ đơn phương hoặc căn cứ kỷ luật sa thải; đi muộn nhiều lần không tự động thuộc trường hợp sa thải."),
            ("Thủ tục kỷ luật", "Phải chứng minh lỗi, thông báo họp, bảo đảm thành phần, thời hiệu và lập biên bản theo luật."),
            ("Tranh chấp", "L có thể yêu cầu giải quyết tranh chấp về chấm dứt và tiền lương; tranh chấp sa thải/đơn phương thuộc nhóm không bắt buộc hòa giải trước khi khởi kiện."),
        ])

    add_exercise(doc, 10, "Quyết định xử phạt sai thẩm quyền và hành vi nhận hối lộ",
        "Trưởng phòng cấp huyện ký quyết định phạt doanh nghiệp 100 triệu dù không có thẩm quyền. Doanh nghiệp đồng thời có video cho thấy cán bộ nhận tiền để bỏ qua một vi phạm khác. Doanh nghiệp phải sử dụng cơ chế nào?",
        [
            ("Quyết định xử phạt", "Sai thẩm quyền là căn cứ xem xét tính bất hợp pháp dù doanh nghiệp có hành vi vi phạm thật."),
            ("Cơ chế bảo vệ", "Doanh nghiệp có thể khiếu nại quyết định hoặc khởi kiện vụ án hành chính theo điều kiện, thời hiệu luật định."),
            ("Hậu quả hủy quyết định", "Việc hủy quyết định sai không làm mất quyền xử lý hành vi gốc bởi người có thẩm quyền nếu còn thời hiệu và đúng thủ tục."),
            ("Hành vi nhận tiền", "Cá nhân biết sự việc có thể tố cáo; nếu có dấu hiệu tội phạm, cung cấp tin báo và chứng cứ cho cơ quan điều tra, viện kiểm sát."),
            ("Phân biệt", "Khiếu nại hướng tới quyết định xâm phạm trực tiếp doanh nghiệp; tố cáo/tin báo hướng tới hành vi vi phạm của cán bộ."),
        ])

    add_exercise(doc, 11, "Hôn nhân không đăng ký, tài sản và thừa kế",
        "A đang có vợ hợp pháp là B nhưng chung sống với C trong 15 năm, có con D. A và C cùng mua nhà 2 tỷ, C chứng minh đóng góp 70%. A chết không di chúc, còn mẹ ruột, B, hai con với B và D.",
        [
            ("Quan hệ A–C", "Không phát sinh tư cách vợ chồng hợp pháp; C không hưởng thừa kế với tư cách vợ."),
            ("Nhà 2 tỷ", "Giải quyết theo đồng sở hữu và công sức đóng góp. Nếu chứng minh C 70%, C có 1,4 tỷ; phần A là 600 triệu và được đưa vào di sản."),
            ("Di sản khác", "Phải cộng tài sản riêng, phần trong tài sản chung với B và trừ nghĩa vụ trước khi chia."),
            ("Người thừa kế", "Mẹ A, B và tất cả các con, gồm D, cùng hàng thứ nhất. Các con bình đẳng, không phụ thuộc tình trạng hôn nhân của cha mẹ."),
            ("Chia", "Sau khi xác định tổng di sản ròng, chia đều cho mẹ, B và ba người con nếu không có người khác và không có tình tiết loại trừ."),
            ("Vi phạm hôn nhân", "Việc A vi phạm một vợ một chồng là vấn đề riêng, không làm D mất quyền."),
        ])

    add_exercise(doc, 12, "Tai nạn giao thông, hình sự, hành chính và bồi thường",
        "K lái xe tải lấn làn vượt xe, va chạm làm T bị thương 60%, hỏng xe 80 triệu. K có giấy phép, không dùng rượu bia, ở lại cứu người. Xác định các trách nhiệm có thể phát sinh.",
        [
            ("Hành vi và lỗi", "K vi phạm quy tắc giao thông, có lỗi vô ý; hậu quả thương tích và tài sản có quan hệ nhân quả."),
            ("Trách nhiệm hình sự", "Nếu chỉ một người bị thương 60% và thiệt hại 80 triệu, dữ kiện thường chưa đạt các ngưỡng cơ bản tương ứng của tội vi phạm quy định về tham gia giao thông. Cần kiểm tra đầy đủ hậu quả và tình tiết định tội hiện hành."),
            ("Trách nhiệm hành chính", "Nếu chưa đến mức tội phạm, K có thể bị xử phạt về hành vi lấn làn, vượt sai và các vi phạm liên quan."),
            ("Trách nhiệm dân sự", "Phải bồi thường chi phí chữa trị, thu nhập mất/giảm, người chăm sóc, tổn thất tinh thần, thiệt hại xe và khoản hợp lý khác."),
            ("Chủ sở hữu xe", "Nếu K lái xe trong nhiệm vụ của doanh nghiệp hoặc xe là nguồn nguy hiểm cao độ của chủ thể khác, phải phân tích trách nhiệm của người sử dụng lao động/chủ sở hữu và quyền hoàn trả."),
            ("Kết luận", "Không được thấy thương tích cao rồi tự động kết luận hình sự; phải đối chiếu ngưỡng và tình tiết."),
        ])

    add_exercise(doc, 13, "Di chúc vô hiệu một phần, người thừa kế bắt buộc và thế vị",
        "Ông X có di sản ròng 1,2 tỷ. Di chúc hợp pháp để toàn bộ cho bạn Y, nhưng 300 triệu trong khối tài sản thực tế thuộc vợ X. X còn vợ, mẹ, con 10 tuổi; một người con A chết trước X và có hai con A1, A2.",
        [
            ("Xác định di sản", "Loại 300 triệu thuộc vợ; di sản của X chỉ 900 triệu."),
            ("Suất theo pháp luật", "Nếu chia theo pháp luật, hàng thứ nhất gồm vợ, mẹ, con 10 tuổi và nhánh A. Có bốn suất; mỗi suất 225 triệu."),
            ("Thế vị", "A1, A2 cùng hưởng suất 225 triệu của A, mỗi người 112,5 triệu nếu phải chia theo pháp luật."),
            ("Người thừa kế bắt buộc", "Vợ, mẹ và con 10 tuổi thuộc diện bắt buộc, mỗi người ít nhất 2/3 × 225 = 150 triệu."),
            ("A1, A2", "Cháu thế vị không thuộc nhóm thừa kế bắt buộc chỉ vì là cháu; do di chúc đã định đoạt toàn bộ, họ không đương nhiên hưởng nếu không có căn cứ khác."),
            ("Phần Y", "Y hưởng 900 – 450 = 450 triệu sau khi bảo đảm ba phần bắt buộc, giả định không có tình tiết khác."),
            ("Vô hiệu một phần", "Di chúc không có hiệu lực đối với 300 triệu không thuộc X; điều này không làm toàn bộ di chúc vô hiệu."),
        ])

    add_exercise(doc, 14, "Xung đột văn bản và bảo vệ quyền công dân",
        "Một thông tư yêu cầu công dân cung cấp toàn bộ lịch sử liên lạc để được giải quyết thủ tục, trong khi luật chỉ yêu cầu thông tin định danh cần thiết. Cơ quan từ chối hồ sơ vì công dân không cung cấp lịch sử liên lạc.",
        [
            ("Thứ bậc hiệu lực", "Thông tư phải phù hợp luật và Hiến pháp; không được đặt thêm điều kiện trái luật."),
            ("Quyền liên quan", "Đời sống riêng tư, bí mật cá nhân và quyền được giải quyết thủ tục theo quy định."),
            ("Giới hạn quyền", "Chỉ được hạn chế quyền theo luật, trong trường hợp cần thiết và với phạm vi tương xứng; văn bản dưới luật không tự tạo giới hạn vượt luật."),
            ("Xử lý vụ việc", "Công dân có thể khiếu nại hành vi/quyết định từ chối hoặc khởi kiện hành chính nếu đủ điều kiện."),
            ("Xử lý văn bản", "Cơ quan có thẩm quyền phải kiểm tra, đình chỉ, bãi bỏ hoặc kiến nghị xử lý thông tư trái luật."),
            ("Kết luận", "Áp dụng văn bản cấp dưới trái văn bản cấp trên không trở thành hợp pháp chỉ vì thông tư chưa bị bãi bỏ tại thời điểm tranh chấp; cần vận dụng cơ chế xử lý và bảo vệ quyền phù hợp."),
        ])


GLOSSARY_GROUPS = [
    ("1. Lý luận chung về Nhà nước và pháp luật", [
        ("Nhà nước", "Tổ chức đặc biệt của quyền lực chính trị, có bộ máy quản lý và cưỡng chế, quản lý dân cư theo lãnh thổ, có chủ quyền, ban hành pháp luật và tổ chức thực hiện các công việc chung của xã hội."),
        ("Quyền lực chính trị", "Khả năng của một lực lượng xã hội tổ chức, định hướng và chi phối các quyết định chung liên quan đến việc quản lý xã hội."),
        ("Quyền lực công cộng đặc biệt", "Quyền lực tách tương đối khỏi cộng đồng dân cư, được thực hiện bằng bộ máy chuyên nghiệp và có tính bắt buộc đối với các chủ thể thuộc phạm vi quản lý."),
        ("Bộ máy nhà nước", "Hệ thống các cơ quan nhà nước được tổ chức và hoạt động theo pháp luật để thực hiện quyền lực và nhiệm vụ của Nhà nước."),
        ("Cưỡng chế nhà nước", "Biện pháp bắt buộc do chủ thể có thẩm quyền áp dụng theo pháp luật nhằm bảo đảm trật tự và việc thực hiện nghĩa vụ pháp lý."),
        ("Quản lý dân cư theo lãnh thổ", "Việc Nhà nước tổ chức quyền lực và quản lý mọi cá nhân, tổ chức trong phạm vi lãnh thổ, không phụ thuộc quan hệ huyết thống hay nghề nghiệp."),
        ("Chủ quyền quốc gia", "Quyền tối cao của Nhà nước trong phạm vi lãnh thổ và quyền độc lập của quốc gia trong quan hệ quốc tế."),
        ("Thuế", "Khoản nộp bắt buộc vào ngân sách nhà nước do luật quy định đối với tổ chức, cá nhân khi có đủ điều kiện chịu thuế."),
        ("Bản chất Nhà nước", "Những thuộc tính cơ bản phản ánh cơ sở xã hội, mục đích và lực lượng mà quyền lực nhà nước trước hết đại diện, bảo vệ."),
        ("Tính giai cấp của Nhà nước", "Phương diện thể hiện quyền lực nhà nước gắn với lợi ích và ý chí của lực lượng giữ địa vị chi phối về kinh tế, chính trị trong xã hội có giai cấp."),
        ("Tính xã hội của Nhà nước", "Phương diện thể hiện Nhà nước phải giải quyết công việc chung, duy trì trật tự và đáp ứng nhu cầu tồn tại, phát triển của xã hội."),
        ("Quyền lực kinh tế", "Khả năng chi phối dựa trên việc nắm giữ, kiểm soát nguồn lực và những điều kiện vật chất quan trọng của xã hội."),
        ("Quyền lực tư tưởng", "Khả năng định hướng nhận thức, giá trị và hệ tư tưởng được thừa nhận trong đời sống xã hội."),
        ("Ý nghĩa phương pháp luận", "Giá trị định hướng về cách tiếp cận, phân tích và đánh giá một hiện tượng; đối với Nhà nước là phải đặt trong điều kiện lịch sử, kinh tế, xã hội và xem xét cả tính giai cấp lẫn tính xã hội."),
        ("Pháp luật", "Hệ thống quy tắc xử sự chung do Nhà nước ban hành hoặc thừa nhận và bảo đảm thực hiện để điều chỉnh các quan hệ xã hội."),
        ("Hệ thống pháp luật", "Tổng thể các quy phạm pháp luật có quan hệ thống nhất, được sắp xếp thành chế định pháp luật và ngành luật."),
        ("Quy phạm pháp luật", "Quy tắc xử sự chung, có hiệu lực bắt buộc chung, do Nhà nước ban hành hoặc thừa nhận và bảo đảm thực hiện."),
        ("Chế định pháp luật", "Nhóm quy phạm pháp luật điều chỉnh một nhóm quan hệ xã hội cùng loại, có liên hệ chặt chẽ trong phạm vi một ngành luật hoặc liên ngành."),
        ("Ngành luật", "Tổng thể quy phạm pháp luật điều chỉnh một lĩnh vực quan hệ xã hội tương đối độc lập bằng phương pháp điều chỉnh đặc thù."),
        ("Đối tượng điều chỉnh", "Nhóm quan hệ xã hội mà một ngành luật tác động và điều chỉnh."),
        ("Phương pháp điều chỉnh", "Cách thức pháp luật tác động lên quan hệ xã hội, xác định địa vị, quyền, nghĩa vụ và cách hình thành quan hệ giữa các chủ thể."),
    ]),
    ("2. Văn bản, nguồn và cấu trúc pháp luật", [
        ("Văn bản quy phạm pháp luật", "Văn bản có chứa quy phạm pháp luật, được ban hành đúng thẩm quyền, hình thức, trình tự và thủ tục do pháp luật quy định."),
        ("Văn bản áp dụng pháp luật", "Văn bản cá biệt do chủ thể có thẩm quyền ban hành để giải quyết một vụ việc hoặc áp dụng pháp luật đối với chủ thể xác định."),
        ("Quy tắc xử sự chung", "Khuôn mẫu hành vi áp dụng cho mọi chủ thể thuộc điều kiện đã được dự liệu và có thể được áp dụng nhiều lần."),
        ("Hiệu lực bắt buộc chung", "Tính chất buộc các chủ thể thuộc phạm vi điều chỉnh phải tuân thủ quy phạm pháp luật."),
        ("Thẩm quyền", "Phạm vi quyền hạn và nhiệm vụ mà pháp luật giao cho một cơ quan, tổ chức hoặc cá nhân."),
        ("Hình thức văn bản", "Tên loại và cách thức thể hiện chính thức của văn bản theo quy định pháp luật."),
        ("Trình tự ban hành", "Thứ tự các giai đoạn phải thực hiện trong quá trình xây dựng và ban hành văn bản."),
        ("Thủ tục ban hành", "Các yêu cầu, công việc và cách thức cụ thể phải tuân theo ở từng giai đoạn ban hành văn bản."),
        ("Tính hợp hiến", "Yêu cầu văn bản và hoạt động pháp lý phải phù hợp với Hiến pháp."),
        ("Tính hợp pháp", "Yêu cầu văn bản hoặc hành vi đúng thẩm quyền, căn cứ, nội dung, hình thức, trình tự, thủ tục và phù hợp văn bản có hiệu lực cao hơn."),
        ("Tính thống nhất của pháp luật", "Yêu cầu các quy định trong hệ thống pháp luật không mâu thuẫn, chồng chéo và có sự liên kết nhất quán."),
        ("Tính công khai", "Yêu cầu thông tin pháp luật phải được công bố, đăng tải hoặc cung cấp để chủ thể có thể tiếp cận theo luật."),
        ("Tính minh bạch", "Yêu cầu quy định và quá trình ban hành rõ ràng, có thể hiểu, kiểm tra và dự liệu được."),
        ("Tính khả thi", "Khả năng quy định được thực hiện trong thực tế với nội dung rõ ràng, điều kiện và nguồn lực phù hợp."),
        ("Tính ổn định", "Mức độ duy trì tương đối bền vững của quy định để các chủ thể có thể dự liệu, đồng thời vẫn cho phép sửa đổi khi thực tiễn đòi hỏi."),
        ("Hình thức pháp luật", "Phương thức tồn tại và biểu hiện chính thức ra bên ngoài của pháp luật, được sử dụng làm căn cứ nhận biết và áp dụng quy tắc pháp lý."),
        ("Nguồn pháp luật", "Nơi chứa đựng hoặc phương thức được Nhà nước thừa nhận để xác định quy tắc pháp lý áp dụng cho vụ việc."),
        ("Tập quán", "Quy tắc xử sự hình thành và được lặp lại lâu dài trong cộng đồng nhưng chưa đương nhiên có giá trị pháp luật."),
        ("Tập quán pháp", "Tập quán được Nhà nước thừa nhận và bảo đảm thực hiện như một quy tắc pháp luật trong điều kiện nhất định."),
        ("Tiền lệ pháp", "Cách giải quyết một vụ việc cụ thể được Nhà nước thừa nhận làm khuôn mẫu để giải quyết vụ việc tương tự về sau."),
        ("Án lệ", "Lập luận và phán quyết trong bản án, quyết định đã có hiệu lực được lựa chọn, công bố để nghiên cứu và áp dụng trong xét xử vụ việc tương tự."),
        ("Điều ước quốc tế", "Thỏa thuận bằng văn bản được ký kết giữa các chủ thể của luật quốc tế, làm phát sinh quyền và nghĩa vụ quốc tế theo luật điều ước."),
        ("Nguyên tắc chung của pháp luật", "Tư tưởng và chuẩn mực nền tảng định hướng việc xây dựng, giải thích và áp dụng pháp luật."),
        ("Điều luật", "Đơn vị cấu trúc của văn bản pháp luật; một điều luật có thể chứa một hoặc nhiều quy phạm, hoặc một quy phạm có thể được thể hiện qua nhiều điều."),
        ("Giả định", "Bộ phận của quy phạm nêu chủ thể, điều kiện, hoàn cảnh, thời gian hoặc địa điểm mà quy phạm được áp dụng."),
        ("Quy định", "Bộ phận của quy phạm nêu cách xử sự được phép, bắt buộc hoặc bị cấm."),
        ("Chế tài", "Bộ phận nêu hậu quả pháp lý áp dụng khi chủ thể không thực hiện đúng phần quy định hoặc khi có sự kiện luật định."),
        ("Hiệu lực theo thời gian", "Khoảng thời gian văn bản hoặc quy phạm bắt đầu, tiếp tục và chấm dứt giá trị áp dụng."),
        ("Hiệu lực theo không gian", "Phạm vi lãnh thổ mà văn bản hoặc quy phạm được áp dụng."),
        ("Hiệu lực theo đối tượng", "Phạm vi cá nhân, tổ chức hoặc chủ thể mà văn bản, quy phạm điều chỉnh."),
        ("Xung đột văn bản", "Tình trạng hai hay nhiều quy định có nội dung khác nhau hoặc mâu thuẫn khi cùng điều chỉnh một vấn đề."),
    ]),
    ("3. Nhà nước pháp quyền và Hiến pháp", [
        ("Nhà nước pháp quyền xã hội chủ nghĩa", "Nhà nước tổ chức và hoạt động trên cơ sở Hiến pháp, pháp luật, bảo đảm chủ quyền Nhân dân, quyền con người và kiểm soát quyền lực theo định hướng xã hội chủ nghĩa."),
        ("Chủ quyền Nhân dân", "Nguyên tắc xác định Nhân dân là chủ thể tối cao và nguồn gốc của quyền lực nhà nước."),
        ("Dân chủ trực tiếp", "Hình thức Nhân dân trực tiếp quyết định hoặc tham gia quyết định vấn đề công, chẳng hạn biểu quyết trưng cầu ý dân."),
        ("Dân chủ đại diện", "Hình thức Nhân dân thực hiện quyền lực thông qua cơ quan và người đại diện do mình bầu ra."),
        ("Quyền lực nhà nước thống nhất", "Nguyên tắc quyền lực nhà nước thuộc về Nhân dân và được thực hiện như một chỉnh thể, không bị chia cắt thành các quyền lực đối lập tuyệt đối."),
        ("Phân công quyền lực", "Việc Hiến pháp, pháp luật giao chức năng, nhiệm vụ khác nhau cho các cơ quan thực hiện quyền lập pháp, hành pháp và tư pháp."),
        ("Phối hợp quyền lực", "Sự liên hệ, hỗ trợ và thực hiện trách nhiệm giữa các cơ quan nhà nước trong quá trình thực hiện quyền lực."),
        ("Kiểm soát quyền lực", "Cơ chế ngăn ngừa, phát hiện và xử lý việc lạm quyền, vượt quyền hoặc thực hiện quyền lực trái pháp luật."),
        ("Quyền lập pháp", "Quyền làm Hiến pháp, sửa đổi Hiến pháp và làm luật theo phạm vi hiến định."),
        ("Quyền hành pháp", "Quyền tổ chức thi hành Hiến pháp, luật và quản lý, điều hành các lĩnh vực của đời sống nhà nước, xã hội."),
        ("Quyền tư pháp", "Quyền xét xử và thực hiện hoạt động tư pháp nhằm bảo vệ công lý, quyền con người, pháp luật và trật tự pháp lý."),
        ("Hiến pháp", "Luật cơ bản của Nhà nước, có hiệu lực pháp lý cao nhất, quy định nền tảng của chế độ chính trị, quyền con người, quyền công dân và tổ chức quyền lực nhà nước."),
        ("Hiệu lực pháp lý cao nhất", "Vị trí pháp lý buộc mọi văn bản và hành vi pháp lý khác phải phù hợp với Hiến pháp."),
        ("Chế độ chính trị", "Tổng thể nguyên tắc nền tảng về bản chất Nhà nước, chủ quyền, tổ chức và thực hiện quyền lực chính trị."),
        ("Quyền con người", "Những quyền tự nhiên, vốn có của mọi cá nhân, được Hiến pháp và pháp luật công nhận, tôn trọng, bảo vệ và bảo đảm."),
        ("Quyền công dân", "Những quyền gắn với tư cách công dân của một quốc gia, được Hiến pháp và pháp luật quy định."),
        ("Nghĩa vụ cơ bản của công dân", "Những cách xử sự bắt buộc có ý nghĩa nền tảng mà Hiến pháp đặt ra đối với công dân."),
        ("Hạn chế quyền", "Việc thu hẹp phạm vi thực hiện quyền theo luật, chỉ trong trường hợp cần thiết vì các mục đích hiến định."),
        ("Quốc hội", "Cơ quan đại biểu cao nhất của Nhân dân, cơ quan quyền lực nhà nước cao nhất, thực hiện quyền lập hiến, lập pháp và giám sát tối cao."),
        ("Chủ tịch nước", "Người đứng đầu Nhà nước, thay mặt nước Cộng hòa xã hội chủ nghĩa Việt Nam về đối nội và đối ngoại."),
        ("Chính phủ", "Cơ quan hành chính nhà nước cao nhất, thực hiện quyền hành pháp và là cơ quan chấp hành của Quốc hội."),
        ("Tòa án nhân dân", "Cơ quan xét xử của nước Cộng hòa xã hội chủ nghĩa Việt Nam, thực hiện quyền tư pháp."),
        ("Viện kiểm sát nhân dân", "Cơ quan thực hành quyền công tố và kiểm sát hoạt động tư pháp."),
        ("Chính quyền địa phương", "Hệ thống cơ quan nhà nước được tổ chức tại các đơn vị hành chính để thực hiện quyền lực và quản lý địa phương theo luật."),
        ("Hội đồng bầu cử quốc gia", "Cơ quan do Quốc hội thành lập, có nhiệm vụ tổ chức bầu cử đại biểu Quốc hội và chỉ đạo, hướng dẫn bầu cử đại biểu Hội đồng nhân dân."),
        ("Kiểm toán nhà nước", "Cơ quan do Quốc hội thành lập, hoạt động độc lập và chỉ tuân theo pháp luật, thực hiện kiểm toán việc quản lý, sử dụng tài chính, tài sản công."),
        ("Mặt trận Tổ quốc Việt Nam", "Tổ chức liên minh chính trị, liên hiệp tự nguyện của các tổ chức và cá nhân tiêu biểu, là cơ sở chính trị của chính quyền Nhân dân."),
        ("Chính quyền địa phương hai cấp", "Mô hình tổ chức chính quyền địa phương gồm hai cấp theo quy định hiện hành sau việc sắp xếp tổ chức bộ máy."),
    ]),
    ("4. Quan hệ pháp luật và chủ thể", [
        ("Quan hệ xã hội", "Mối liên hệ phát sinh giữa người với người hoặc giữa các chủ thể trong đời sống xã hội."),
        ("Quan hệ pháp luật", "Quan hệ xã hội được quy phạm pháp luật điều chỉnh, trong đó các bên có quyền và nghĩa vụ pháp lý được Nhà nước bảo đảm thực hiện."),
        ("Chủ thể quan hệ pháp luật", "Cá nhân, pháp nhân, tổ chức hoặc Nhà nước tham gia quan hệ pháp luật và có quyền, nghĩa vụ xác định."),
        ("Cá nhân", "Con người cụ thể tham gia quan hệ pháp luật với năng lực chủ thể theo quy định."),
        ("Người chưa thành niên", "Người chưa đủ 18 tuổi, có phạm vi năng lực hành vi và trách nhiệm pháp lý được xác định theo độ tuổi và từng ngành luật."),
        ("Người giám hộ", "Cá nhân hoặc pháp nhân thực hiện việc chăm sóc, bảo vệ quyền, lợi ích của người được giám hộ và đại diện trong trường hợp luật định."),
        ("Pháp nhân", "Tổ chức được thành lập hợp pháp, có cơ cấu, có tài sản độc lập và tự chịu trách nhiệm, nhân danh mình tham gia quan hệ pháp luật một cách độc lập."),
        ("Tổ chức không có tư cách pháp nhân", "Tổ chức không đáp ứng đầy đủ điều kiện pháp nhân nhưng vẫn có thể tham gia một số quan hệ theo cơ chế đại diện hoặc trách nhiệm của thành viên."),
        ("Năng lực chủ thể", "Tổng hợp năng lực pháp luật và năng lực hành vi cần thiết để một chủ thể tham gia quan hệ pháp luật cụ thể."),
        ("Năng lực pháp luật", "Khả năng được pháp luật công nhận có quyền và nghĩa vụ pháp lý."),
        ("Năng lực hành vi", "Khả năng bằng hành vi của mình xác lập, thực hiện quyền và nghĩa vụ pháp lý."),
        ("Năng lực trách nhiệm pháp lý", "Khả năng của chủ thể nhận thức, điều khiển hành vi và chịu hậu quả pháp lý khi thực hiện vi phạm."),
        ("Nội dung quan hệ pháp luật", "Tổng thể quyền chủ thể và nghĩa vụ pháp lý của các bên trong quan hệ."),
        ("Quyền chủ thể", "Khả năng xử sự được pháp luật cho phép, bao gồm tự hành động, yêu cầu bên có nghĩa vụ và yêu cầu Nhà nước bảo vệ."),
        ("Nghĩa vụ pháp lý", "Cách xử sự bắt buộc mà chủ thể phải thực hiện, không được thực hiện hoặc phải gánh chịu theo pháp luật."),
        ("Khách thể quan hệ pháp luật", "Lợi ích vật chất, tinh thần, giá trị xã hội hoặc kết quả hành vi mà các chủ thể hướng tới."),
        ("Sự kiện pháp lý", "Sự kiện thực tế mà pháp luật gắn với hậu quả làm phát sinh, thay đổi hoặc chấm dứt quan hệ pháp luật."),
        ("Hành vi pháp lý", "Sự kiện pháp lý phát sinh từ hành động hoặc không hành động có ý chí của chủ thể."),
        ("Hành vi hợp pháp", "Hành vi phù hợp quy định pháp luật và có thể làm phát sinh hậu quả pháp lý được pháp luật thừa nhận."),
        ("Hành vi vi phạm pháp luật", "Hành vi trái pháp luật, có lỗi, do chủ thể có năng lực trách nhiệm pháp lý thực hiện."),
        ("Biến cố pháp lý", "Sự kiện xảy ra không phụ thuộc hoặc không hoàn toàn phụ thuộc ý chí con người nhưng làm phát sinh hậu quả pháp lý, như sinh, chết hoặc thiên tai."),
    ]),
    ("5. Luật Dân sự, tài sản và quyền sở hữu", [
        ("Luật Dân sự", "Ngành luật điều chỉnh quan hệ tài sản và quan hệ nhân thân hình thành trên cơ sở bình đẳng, tự do ý chí, độc lập tài sản và tự chịu trách nhiệm."),
        ("Quan hệ tài sản", "Quan hệ giữa các chủ thể thông qua tài sản, thường mang tính giá trị và trao đổi."),
        ("Quan hệ nhân thân", "Quan hệ gắn với giá trị nhân thân của cá nhân hoặc tổ chức, có thể gắn hoặc không gắn với tài sản."),
        ("Tài sản", "Vật, tiền, giấy tờ có giá và quyền tài sản theo Bộ luật Dân sự."),
        ("Vật", "Bộ phận của thế giới vật chất có thể đáp ứng nhu cầu và trở thành đối tượng của quyền dân sự."),
        ("Tiền", "Phương tiện thanh toán và thước đo giá trị được Nhà nước phát hành, thừa nhận."),
        ("Giấy tờ có giá", "Chứng chỉ hoặc dữ liệu xác nhận nghĩa vụ trả nợ giữa tổ chức phát hành với người sở hữu trong điều kiện pháp luật quy định."),
        ("Quyền tài sản", "Quyền trị giá được bằng tiền, bao gồm quyền tài sản đối với đối tượng sở hữu trí tuệ, quyền sử dụng đất và quyền tài sản khác."),
        ("Bình đẳng dân sự", "Nguyên tắc các bên có địa vị pháp lý ngang nhau, không bên nào được dùng quyền lực công để áp đặt ý chí trong quan hệ dân sự."),
        ("Tự do, tự nguyện cam kết, thỏa thuận", "Nguyên tắc các bên tự quyết định việc xác lập và nội dung giao dịch trong giới hạn điều cấm, đạo đức xã hội và quyền của người khác."),
        ("Điều cấm của luật", "Quy định của luật không cho phép chủ thể thực hiện hành vi nhất định."),
        ("Đạo đức xã hội", "Chuẩn mực ứng xử chung được cộng đồng thừa nhận, tôn trọng và được pháp luật sử dụng làm giới hạn trong một số quan hệ."),
        ("Độc lập về tài sản", "Khả năng chủ thể có và chịu trách nhiệm bằng khối tài sản độc lập của mình."),
        ("Tự chịu trách nhiệm", "Nguyên tắc chủ thể phải gánh chịu hậu quả pháp lý từ cam kết hoặc hành vi của chính mình."),
        ("Giao dịch dân sự", "Hợp đồng hoặc hành vi pháp lý đơn phương làm phát sinh, thay đổi hoặc chấm dứt quyền, nghĩa vụ dân sự."),
        ("Hợp đồng", "Sự thỏa thuận giữa các bên về việc xác lập, thay đổi hoặc chấm dứt quyền, nghĩa vụ."),
        ("Tranh chấp dân sự", "Bất đồng về quyền, nghĩa vụ hoặc lợi ích phát sinh từ quan hệ dân sự."),
        ("Trách nhiệm dân sự", "Hậu quả pháp lý bất lợi chủ yếu nhằm khôi phục quyền, buộc thực hiện nghĩa vụ hoặc bồi thường thiệt hại."),
        ("Sở hữu", "Quan hệ xã hội giữa người với người đối với việc chiếm hữu của cải, tài sản."),
        ("Quyền sở hữu", "Quyền của chủ sở hữu đối với tài sản, bao gồm quyền chiếm hữu, sử dụng và định đoạt."),
        ("Chủ sở hữu", "Chủ thể có quyền sở hữu hợp pháp đối với tài sản."),
        ("Chiếm hữu", "Việc chủ thể nắm giữ, chi phối tài sản trực tiếp hoặc gián tiếp như chủ thể có quyền."),
        ("Chiếm hữu hợp pháp", "Việc nắm giữ, chi phối tài sản dựa trên căn cứ pháp luật hoặc sự giao tài sản hợp pháp."),
        ("Sử dụng tài sản", "Việc khai thác công dụng và hưởng hoa lợi, lợi tức từ tài sản."),
        ("Hoa lợi", "Sản vật tự nhiên mà tài sản mang lại."),
        ("Lợi tức", "Khoản lợi thu được từ việc khai thác tài sản."),
        ("Định đoạt tài sản", "Việc chuyển giao, từ bỏ quyền sở hữu, tiêu dùng hoặc tiêu hủy tài sản."),
        ("Căn cứ xác lập quyền sở hữu", "Sự kiện được pháp luật công nhận làm phát sinh quyền sở hữu cho một chủ thể."),
        ("Căn cứ nguyên sinh", "Căn cứ xác lập quyền sở hữu không trực tiếp dựa trên việc chuyển quyền từ một chủ sở hữu trước."),
        ("Căn cứ phái sinh", "Căn cứ xác lập quyền sở hữu thông qua việc chuyển quyền từ chủ thể có quyền trước đó."),
        ("Chấm dứt quyền sở hữu", "Việc chủ thể không còn quyền sở hữu do chuyển quyền, từ bỏ, tài sản bị tiêu hủy, xử lý, tịch thu hoặc căn cứ luật định khác."),
        ("Bồi thường thiệt hại", "Việc bên có trách nhiệm bù đắp tổn thất vật chất, tinh thần cho bên bị thiệt hại theo căn cứ và phạm vi luật định."),
        ("Bồi thường thiệt hại ngoài hợp đồng", "Trách nhiệm bồi thường phát sinh từ hành vi gây thiệt hại không dựa trên việc vi phạm một nghĩa vụ hợp đồng giữa các bên."),
        ("Nguồn nguy hiểm cao độ", "Phương tiện, hệ thống hoặc vật có khả năng gây thiệt hại lớn theo danh mục và tính chất do pháp luật dân sự xác định."),
        ("Sự kiện bất khả kháng", "Sự kiện khách quan, không thể lường trước và không thể khắc phục dù đã áp dụng mọi biện pháp cần thiết, khả năng cho phép."),
    ]),
    ("6. Vi phạm pháp luật, lỗi và trách nhiệm pháp lý", [
        ("Vi phạm pháp luật", "Hành vi trái pháp luật, có lỗi, do chủ thể có năng lực trách nhiệm pháp lý thực hiện, xâm hại quan hệ xã hội được pháp luật bảo vệ."),
        ("Hành động", "Dạng hành vi thể hiện bằng việc chủ thể thực hiện một việc mà pháp luật cấm hoặc thực hiện sai nghĩa vụ."),
        ("Không hành động", "Dạng hành vi thể hiện bằng việc chủ thể không thực hiện việc pháp luật buộc phải thực hiện dù có điều kiện thực hiện."),
        ("Tính trái pháp luật", "Tính chất không phù hợp yêu cầu, nghĩa vụ hoặc điều cấm do pháp luật đặt ra."),
        ("Hậu quả", "Thiệt hại hoặc biến đổi bất lợi do hành vi gây ra đối với quan hệ được pháp luật bảo vệ."),
        ("Quan hệ nhân quả", "Mối liên hệ khách quan trong đó hành vi là nguyên nhân dẫn đến hậu quả pháp lý đang được xem xét."),
        ("Lỗi", "Thái độ tâm lý của chủ thể đối với hành vi và hậu quả trong điều kiện có khả năng nhận thức, lựa chọn cách xử sự phù hợp."),
        ("Lỗi cố ý", "Lỗi khi chủ thể nhận thức rõ tính chất hành vi và có thái độ mong muốn hoặc có ý thức để mặc hậu quả."),
        ("Cố ý trực tiếp", "Trường hợp chủ thể nhận thức rõ hành vi nguy hiểm, thấy trước hậu quả và mong muốn hậu quả xảy ra."),
        ("Cố ý gián tiếp", "Trường hợp chủ thể nhận thức rõ hành vi nguy hiểm, thấy trước hậu quả, không mong muốn nhưng có ý thức để mặc hậu quả xảy ra."),
        ("Lỗi vô ý", "Lỗi khi chủ thể không mong muốn hậu quả nhưng hậu quả xảy ra do quá tự tin hoặc cẩu thả."),
        ("Vô ý vì quá tự tin", "Trường hợp chủ thể thấy trước khả năng hậu quả nhưng cho rằng hậu quả sẽ không xảy ra hoặc có thể ngăn ngừa."),
        ("Vô ý do cẩu thả", "Trường hợp chủ thể không thấy trước hậu quả dù phải thấy trước và có thể thấy trước."),
        ("Động cơ", "Nguyên nhân bên trong thúc đẩy chủ thể thực hiện hành vi."),
        ("Mục đích", "Kết quả chủ quan mà chủ thể hướng tới khi thực hiện hành vi."),
        ("Vi phạm hình sự", "Tội phạm: hành vi nguy hiểm cho xã hội được Bộ luật Hình sự quy định, do chủ thể có năng lực trách nhiệm hình sự thực hiện có lỗi."),
        ("Vi phạm hành chính", "Hành vi có lỗi của cá nhân hoặc tổ chức vi phạm quy định quản lý nhà nước, không phải tội phạm và phải bị xử phạt hành chính theo luật."),
        ("Vi phạm dân sự", "Hành vi trái pháp luật xâm phạm quan hệ tài sản, nhân thân hoặc vi phạm nghĩa vụ dân sự, làm phát sinh trách nhiệm dân sự khi đủ căn cứ."),
        ("Vi phạm kỷ luật", "Hành vi có lỗi vi phạm nội quy, quy chế hoặc nghĩa vụ trong quan hệ công vụ, lao động, học tập hay tổ chức."),
        ("Trách nhiệm pháp lý", "Hậu quả bất lợi do Nhà nước hoặc chủ thể có thẩm quyền áp dụng đối với chủ thể vi phạm theo pháp luật."),
        ("Trách nhiệm hình sự", "Trách nhiệm pháp lý nghiêm khắc của cá nhân hoặc pháp nhân thương mại phạm tội trước Nhà nước."),
        ("Trách nhiệm hành chính", "Hậu quả pháp lý do chủ thể có thẩm quyền áp dụng đối với cá nhân, tổ chức vi phạm hành chính."),
        ("Trách nhiệm dân sự do vi phạm", "Trách nhiệm khôi phục quyền, buộc thực hiện nghĩa vụ, bồi thường hoặc chịu hậu quả dân sự khác khi có căn cứ phát sinh."),
        ("Trách nhiệm kỷ luật", "Hậu quả bất lợi áp dụng trong quan hệ nội bộ đối với người vi phạm kỷ luật."),
        ("Phòng vệ chính đáng", "Hành vi chống trả cần thiết người đang có hành vi xâm phạm quyền, lợi ích hợp pháp; không phải tội phạm."),
        ("Tình thế cấp thiết", "Tình thế buộc phải gây thiệt hại nhỏ hơn để tránh một nguy cơ thực tế đang đe dọa lợi ích được pháp luật bảo vệ."),
        ("Sự kiện bất ngờ", "Trường hợp chủ thể không thể thấy trước hoặc không buộc phải thấy trước hậu quả nên không có lỗi hình sự."),
    ]),
    ("7. Thừa kế", [
        ("Thừa kế", "Việc chuyển dịch tài sản của người chết cho người còn sống theo di chúc hoặc theo pháp luật."),
        ("Di sản", "Tài sản riêng của người chết và phần tài sản của người chết trong tài sản chung với người khác."),
        ("Thời điểm mở thừa kế", "Thời điểm người có tài sản chết hoặc bị Tòa án tuyên bố là đã chết."),
        ("Địa điểm mở thừa kế", "Nơi cư trú cuối cùng của người để lại di sản; nếu không xác định được thì là nơi có toàn bộ hoặc phần lớn di sản."),
        ("Người để lại di sản", "Cá nhân có tài sản được chuyển giao cho người khác sau khi chết."),
        ("Người thừa kế", "Cá nhân hoặc tổ chức có quyền hưởng di sản theo di chúc hoặc theo pháp luật; người thừa kế là cá nhân phải còn sống hoặc được thành thai theo điều kiện luật định."),
        ("Thừa kế theo di chúc", "Việc chuyển dịch di sản theo ý chí của người chết được thể hiện trong di chúc hợp pháp."),
        ("Thừa kế theo pháp luật", "Việc chia di sản theo hàng thừa kế, điều kiện và trình tự do pháp luật quy định khi có căn cứ áp dụng."),
        ("Di chúc", "Sự thể hiện ý chí của cá nhân nhằm chuyển tài sản của mình cho người khác sau khi chết."),
        ("Di chúc hợp pháp", "Di chúc đáp ứng điều kiện về người lập, sự tự nguyện, nội dung và hình thức theo pháp luật."),
        ("Di chúc bằng văn bản", "Di chúc được thể hiện bằng chữ viết và đáp ứng điều kiện về chữ ký, người làm chứng, công chứng hoặc chứng thực khi luật yêu cầu."),
        ("Người làm chứng", "Người chứng kiến việc lập di chúc hoặc sự kiện pháp lý và xác nhận theo điều kiện luật định, không thuộc trường hợp bị cấm làm chứng."),
        ("Công chứng", "Việc công chứng viên chứng nhận tính xác thực, hợp pháp của hợp đồng, giao dịch hoặc bản dịch theo pháp luật công chứng."),
        ("Chứng thực", "Việc cơ quan, người có thẩm quyền xác nhận bản sao, chữ ký, hợp đồng hoặc giao dịch theo phạm vi pháp luật quy định."),
        ("Di chúc miệng", "Di chúc được thể hiện bằng lời nói trong tình trạng tính mạng bị đe dọa và không thể lập bằng văn bản, phải tuân thủ điều kiện đặc biệt."),
        ("Di chúc vô hiệu", "Di chúc không làm phát sinh hiệu lực toàn bộ hoặc một phần vì không đáp ứng điều kiện pháp luật hoặc đối tượng định đoạt không còn, không thuộc quyền."),
        ("Di chúc vô hiệu một phần", "Trường hợp chỉ một phần nội dung di chúc không có hiệu lực nhưng phần còn lại độc lập vẫn được thực hiện."),
        ("Người thừa kế không phụ thuộc nội dung di chúc", "Cha, mẹ, vợ, chồng, con chưa thành niên hoặc con thành niên không có khả năng lao động được hưởng phần bắt buộc, trừ trường hợp luật loại trừ."),
        ("Suất thừa kế theo pháp luật", "Phần di sản mà một người trong cùng hàng thừa kế được hưởng nếu di sản được chia theo pháp luật."),
        ("Hàng thừa kế", "Nhóm người được pháp luật xếp theo thứ tự ưu tiên hưởng di sản; người ở hàng sau chỉ hưởng khi không còn ai ở hàng trước đủ điều kiện hưởng."),
        ("Thừa kế thế vị", "Việc cháu hưởng phần di sản mà cha hoặc mẹ của cháu đáng lẽ được hưởng nếu còn sống, trong điều kiện luật định."),
        ("Di tặng", "Việc người lập di chúc dành một phần di sản để tặng cho người khác."),
        ("Tài sản dùng vào việc thờ cúng", "Phần di sản được người lập di chúc dành để thờ cúng và giao cho người quản lý theo ý chí hợp pháp."),
        ("Truất quyền hưởng di sản", "Việc người lập di chúc thể hiện ý chí không cho một người có thể thuộc diện thừa kế hưởng di sản, trong giới hạn phần thừa kế bắt buộc."),
        ("Từ chối nhận di sản", "Việc người thừa kế tuyên bố không nhận phần di sản theo điều kiện, hình thức và thời hạn luật định."),
        ("Người không được quyền hưởng di sản", "Người có hành vi thuộc trường hợp luật quy định bị loại khỏi quyền hưởng, trừ khi người để lại di sản biết mà vẫn cho hưởng theo di chúc."),
        ("Nghĩa vụ tài sản do người chết để lại", "Khoản nợ và nghĩa vụ tài sản được thanh toán bằng di sản trước khi phân chia cho người thừa kế."),
    ]),
    ("8. Luật Lao động và tranh chấp lao động", [
        ("Luật Lao động", "Ngành luật điều chỉnh quan hệ lao động giữa người lao động, người sử dụng lao động và các quan hệ liên quan trực tiếp."),
        ("Người lao động", "Người làm việc cho người sử dụng lao động theo thỏa thuận, được trả lương và chịu sự quản lý, điều hành, giám sát."),
        ("Người sử dụng lao động", "Doanh nghiệp, cơ quan, tổ chức, hợp tác xã, hộ gia đình hoặc cá nhân thuê mướn, sử dụng người lao động theo thỏa thuận."),
        ("Quan hệ lao động", "Quan hệ phát sinh trong việc thuê mướn, sử dụng lao động và trả lương giữa người lao động với người sử dụng lao động."),
        ("Học nghề", "Việc người học được đào tạo kiến thức, kỹ năng nghề nghiệp để làm việc theo chương trình và điều kiện pháp luật."),
        ("An toàn, vệ sinh lao động", "Hệ thống biện pháp phòng ngừa yếu tố nguy hiểm, có hại nhằm bảo đảm sức khỏe, tính mạng người lao động."),
        ("Bảo hiểm xã hội", "Sự bảo đảm thay thế hoặc bù đắp một phần thu nhập của người lao động khi bị giảm hoặc mất thu nhập do sự kiện luật định."),
        ("Việc làm có trả công", "Công việc do một người thực hiện để nhận tiền lương hoặc khoản trả công theo thỏa thuận."),
        ("Tiền lương", "Số tiền người sử dụng lao động trả cho người lao động theo thỏa thuận để thực hiện công việc, gồm mức lương, phụ cấp và khoản bổ sung theo luật."),
        ("Quản lý, điều hành, giám sát lao động", "Quyền hợp pháp của người sử dụng lao động trong tổ chức công việc, phân công, kiểm tra và đánh giá việc thực hiện."),
        ("Quan hệ lao động tập thể", "Quan hệ giữa tổ chức đại diện người lao động với người sử dụng lao động hoặc tổ chức đại diện của họ."),
        ("Tổ chức đại diện người lao động", "Tổ chức được thành lập hợp pháp để đại diện, bảo vệ quyền và lợi ích của người lao động trong quan hệ lao động."),
        ("Đối thoại tại nơi làm việc", "Việc chia sẻ thông tin, tham khảo và trao đổi giữa người sử dụng lao động với người lao động hoặc tổ chức đại diện."),
        ("Thương lượng tập thể", "Quá trình đàm phán giữa các bên đại diện nhằm xác lập điều kiện lao động và điều chỉnh quan hệ lao động tập thể."),
        ("Thỏa ước lao động tập thể", "Thỏa thuận đạt được thông qua thương lượng tập thể và được các bên ký kết bằng văn bản."),
        ("Hợp đồng lao động", "Thỏa thuận về việc làm có trả công, tiền lương, điều kiện lao động, quyền và nghĩa vụ giữa người lao động với người sử dụng lao động."),
        ("Hợp đồng lao động xác định thời hạn", "Hợp đồng xác định thời hạn và thời điểm chấm dứt hiệu lực trong thời gian không quá 36 tháng."),
        ("Hợp đồng lao động không xác định thời hạn", "Hợp đồng không xác định thời hạn và thời điểm chấm dứt hiệu lực."),
        ("Thông điệp dữ liệu", "Thông tin được tạo ra, gửi, nhận và lưu trữ bằng phương tiện điện tử, có thể được sử dụng làm hình thức hợp đồng theo luật."),
        ("Chấm dứt hợp đồng lao động", "Việc quan hệ hợp đồng lao động kết thúc theo thỏa thuận, hết hạn, đơn phương hoặc căn cứ luật định khác."),
        ("Đơn phương chấm dứt hợp đồng lao động", "Việc một bên tự mình chấm dứt hợp đồng theo quyền và điều kiện pháp luật quy định."),
        ("Thời hạn báo trước", "Khoảng thời gian tối thiểu một bên phải thông báo cho bên kia trước khi đơn phương chấm dứt trong trường hợp luật yêu cầu."),
        ("Đơn phương chấm dứt trái pháp luật", "Việc đơn phương không đúng căn cứ, thời hạn báo trước hoặc điều kiện luật định, làm phát sinh nghĩa vụ khắc phục, bồi thường."),
        ("Kỷ luật lao động", "Quy định về việc tuân theo thời gian, công nghệ, điều hành sản xuất, kinh doanh và hình thức xử lý hành vi vi phạm trong quan hệ lao động."),
        ("Nội quy lao động", "Văn bản do người sử dụng lao động ban hành theo luật, quy định trật tự, kỷ luật, trách nhiệm và hành vi vi phạm tại nơi làm việc."),
        ("Xử lý kỷ luật lao động", "Việc áp dụng hình thức kỷ luật đối với người lao động có lỗi theo đúng căn cứ, nguyên tắc, thẩm quyền và thủ tục."),
        ("Sa thải", "Hình thức xử lý kỷ luật lao động nghiêm khắc làm chấm dứt quan hệ lao động trong trường hợp và theo thủ tục luật định."),
        ("Tranh chấp lao động", "Tranh chấp về quyền, nghĩa vụ và lợi ích phát sinh giữa các bên trong việc xác lập, thực hiện hoặc chấm dứt quan hệ lao động và quan hệ liên quan."),
        ("Tranh chấp lao động cá nhân", "Tranh chấp về quyền, nghĩa vụ và lợi ích của một người lao động với người sử dụng lao động hoặc chủ thể liên quan theo luật."),
        ("Tranh chấp lao động tập thể về quyền", "Tranh chấp về việc giải thích, thực hiện quyền đã được xác lập trong pháp luật, thỏa ước, nội quy hoặc thỏa thuận hợp pháp."),
        ("Tranh chấp lao động tập thể về lợi ích", "Tranh chấp phát sinh trong thương lượng nhằm xác lập điều kiện lao động mới."),
        ("Hòa giải viên lao động", "Người được cơ quan có thẩm quyền bổ nhiệm để hòa giải tranh chấp lao động và hỗ trợ phát triển quan hệ lao động."),
        ("Hội đồng trọng tài lao động", "Thiết chế giải quyết tranh chấp lao động theo lựa chọn và điều kiện do Bộ luật Lao động quy định."),
        ("Hòa giải lao động", "Phương thức trong đó hòa giải viên hỗ trợ các bên thương lượng để tự giải quyết tranh chấp."),
        ("Đình công", "Sự ngừng việc tạm thời, tự nguyện và có tổ chức của người lao động nhằm đạt yêu cầu trong quá trình giải quyết tranh chấp tập thể về lợi ích."),
    ]),
    ("9. Luật Hành chính, tố cáo và khiếu nại", [
        ("Luật Hành chính", "Ngành luật điều chỉnh các quan hệ phát sinh trong hoạt động quản lý hành chính nhà nước và các quan hệ quản lý khác được pháp luật giao."),
        ("Quản lý hành chính nhà nước", "Hoạt động chấp hành, điều hành mang tính quyền lực nhà nước nhằm tổ chức thi hành pháp luật và quản lý thường xuyên các lĩnh vực xã hội."),
        ("Hoạt động chấp hành - điều hành", "Hoạt động tổ chức thực hiện luật và đưa ra quyết định quản lý cụ thể để điều hành đời sống hành chính."),
        ("Phương pháp quyền lực - phục tùng", "Phương pháp trong đó một bên nhân danh quyền lực nhà nước ra quyết định đơn phương trong phạm vi thẩm quyền, bên kia có nghĩa vụ chấp hành."),
        ("Quyết định hành chính", "Văn bản hoặc quyết định cá biệt của cơ quan, người có thẩm quyền quản lý hành chính áp dụng một lần đối với đối tượng cụ thể."),
        ("Hành vi hành chính", "Hành vi thực hiện hoặc không thực hiện nhiệm vụ, công vụ của cơ quan hành chính hoặc người có thẩm quyền quản lý hành chính."),
        ("Cưỡng chế hành chính", "Biện pháp bắt buộc do chủ thể có thẩm quyền áp dụng theo thủ tục hành chính để bảo đảm quản lý hoặc thi hành quyết định."),
        ("Xử phạt vi phạm hành chính", "Việc người có thẩm quyền áp dụng hình thức xử phạt và biện pháp khắc phục hậu quả đối với cá nhân, tổ chức vi phạm hành chính."),
        ("Khiếu nại", "Việc cá nhân, cơ quan, tổ chức yêu cầu xem xét lại quyết định hoặc hành vi hành chính, quyết định kỷ luật khi cho rằng quyết định, hành vi đó xâm phạm quyền, lợi ích hợp pháp của mình."),
        ("Khởi kiện hành chính", "Việc chủ thể yêu cầu Tòa án giải quyết vụ án hành chính để bảo vệ quyền, lợi ích hợp pháp bị quyết định hoặc hành vi hành chính xâm phạm."),
        ("Tố cáo", "Việc cá nhân báo cho chủ thể có thẩm quyền biết về hành vi vi phạm pháp luật gây hoặc đe dọa gây thiệt hại đến lợi ích Nhà nước, quyền, lợi ích hợp pháp của chủ thể khác."),
        ("Người tố cáo", "Cá nhân thực hiện việc tố cáo theo thủ tục luật định."),
        ("Người bị tố cáo", "Cơ quan, tổ chức hoặc cá nhân có hành vi bị người tố cáo cho là vi phạm pháp luật."),
        ("Thụ lý tố cáo", "Việc người giải quyết tố cáo ra quyết định chính thức tiếp nhận vụ việc để giải quyết khi đủ điều kiện."),
        ("Xác minh nội dung tố cáo", "Hoạt động thu thập, kiểm tra, đánh giá thông tin, tài liệu để làm rõ nội dung tố cáo."),
        ("Kết luận nội dung tố cáo", "Văn bản xác định nội dung tố cáo đúng, sai hoặc đúng một phần, trách nhiệm và kiến nghị xử lý."),
        ("Xử lý kết luận tố cáo", "Việc áp dụng hoặc kiến nghị biện pháp xử lý hành vi vi phạm, khắc phục hậu quả sau kết luận."),
        ("Thông tin tố cáo không rõ danh tính", "Thông tin không xác định được tên, địa chỉ người cung cấp; không giải quyết theo thủ tục tố cáo nhưng có thể được dùng để thanh tra, kiểm tra khi cụ thể, có căn cứ."),
        ("Bảo vệ người tố cáo", "Biện pháp bảo vệ bí mật thông tin, vị trí công tác, tính mạng, sức khỏe, tài sản, danh dự của người tố cáo và người liên quan khi có nguy cơ bị xâm hại."),
        ("Báo tin về tội phạm", "Việc cá nhân, cơ quan, tổ chức cung cấp thông tin về hành vi có dấu hiệu tội phạm cho cơ quan có thẩm quyền tố tụng."),
    ]),
    ("10. Luật Hình sự, tội phạm và các tội chiếm đoạt", [
        ("Luật Hình sự", "Ngành luật xác định hành vi nào là tội phạm, căn cứ trách nhiệm hình sự, hình phạt, biện pháp tư pháp và vấn đề liên quan."),
        ("Tội danh", "Tên pháp lý của tội phạm được quy định trong Bộ luật Hình sự."),
        ("Định tội danh", "Hoạt động xác định hành vi thực tế phù hợp với cấu thành của tội nào trong Bộ luật Hình sự."),
        ("Khung hình phạt", "Phạm vi loại và mức hình phạt được điều luật quy định cho một trường hợp phạm tội có các dấu hiệu nhất định."),
        ("Định khung hình phạt", "Hoạt động xác định khoản hoặc khung hình phạt tương ứng với các dấu hiệu định khung của hành vi phạm tội."),
        ("Tội phạm", "Hành vi nguy hiểm cho xã hội được Bộ luật Hình sự quy định, do chủ thể có năng lực trách nhiệm hình sự thực hiện có lỗi và phải chịu trách nhiệm hình sự."),
        ("Hình phạt", "Biện pháp cưỡng chế nghiêm khắc nhất của Nhà nước do Tòa án áp dụng đối với người hoặc pháp nhân thương mại phạm tội."),
        ("Phạt tiền", "Hình phạt buộc người hoặc pháp nhân thương mại bị kết án nộp một khoản tiền theo quyết định của Tòa án."),
        ("Cải tạo không giam giữ", "Hình phạt không cách ly người bị kết án khỏi xã hội, đặt họ dưới sự giám sát, giáo dục và thực hiện nghĩa vụ theo luật."),
        ("Tù có thời hạn", "Hình phạt buộc người bị kết án chấp hành tại cơ sở giam giữ trong một thời hạn do Tòa án quyết định trong giới hạn luật định."),
        ("Tù chung thân", "Hình phạt tù không thời hạn áp dụng đối với người phạm tội đặc biệt nghiêm trọng trong trường hợp luật quy định."),
        ("Tử hình", "Hình phạt đặc biệt nghiêm khắc tước bỏ tính mạng người bị kết án, chỉ áp dụng đối với một số tội đặc biệt nghiêm trọng và chủ thể luật định."),
        ("Miễn trách nhiệm hình sự", "Việc không buộc người hoặc pháp nhân thương mại phải chịu trách nhiệm hình sự dù hành vi đã có dấu hiệu tội phạm, khi đủ căn cứ luật định."),
        ("Miễn hình phạt", "Việc Tòa án không buộc người phạm tội phải chấp hành hình phạt khi đủ căn cứ, dù họ vẫn bị xác định là có tội."),
        ("Biện pháp tư pháp", "Biện pháp hình sự do cơ quan có thẩm quyền áp dụng nhằm hỗ trợ xử lý tội phạm, khắc phục hậu quả hoặc giáo dục, phòng ngừa."),
        ("Nguyên tắc pháp chế hình sự", "Nguyên tắc chỉ Bộ luật Hình sự được quy định tội phạm, hình phạt và việc xử lý phải đúng căn cứ, thẩm quyền, thủ tục."),
        ("Bình đẳng trước pháp luật hình sự", "Nguyên tắc mọi chủ thể đủ điều kiện đều bị xử lý theo pháp luật, không phân biệt địa vị, giới tính, dân tộc, tín ngưỡng hay tài sản."),
        ("Trách nhiệm trên cơ sở có lỗi", "Nguyên tắc chỉ truy cứu trách nhiệm hình sự khi chủ thể có lỗi đối với hành vi nguy hiểm."),
        ("Cá thể hóa trách nhiệm hình sự", "Yêu cầu xác định trách nhiệm và hình phạt phù hợp vai trò, tính chất hành vi, nhân thân và tình tiết của từng chủ thể."),
        ("Nguyên tắc công bằng", "Yêu cầu xử lý tương xứng, khách quan, không làm oan người vô tội và không bỏ lọt tội phạm."),
        ("Nguyên tắc nhân đạo", "Nguyên tắc tôn trọng nhân phẩm, không nhằm trả thù, tạo điều kiện cải tạo và tái hòa nhập cho người phạm tội."),
        ("Tính nguy hiểm cho xã hội", "Dấu hiệu nội dung thể hiện hành vi gây hoặc đe dọa gây thiệt hại đáng kể cho quan hệ được luật hình sự bảo vệ."),
        ("Tính trái pháp luật hình sự", "Dấu hiệu hành vi phải được Bộ luật Hình sự quy định là tội phạm."),
        ("Tính phải chịu hình phạt", "Dấu hiệu thể hiện tội phạm bị đe dọa áp dụng hình phạt hoặc biện pháp hình sự, dù trường hợp cụ thể có thể được miễn theo luật."),
        ("Năng lực trách nhiệm hình sự", "Khả năng của cá nhân nhận thức, điều khiển hành vi và chịu trách nhiệm hình sự; pháp nhân thương mại chịu trách nhiệm theo điều kiện riêng."),
        ("Pháp nhân thương mại", "Pháp nhân có mục tiêu chính là tìm kiếm lợi nhuận và lợi nhuận được chia cho thành viên, có thể chịu trách nhiệm hình sự về tội luật định."),
        ("Cấu thành tội phạm", "Tổng hợp dấu hiệu pháp lý đặc trưng của một tội, gồm khách thể, mặt khách quan, chủ thể và mặt chủ quan."),
        ("Khách thể của tội phạm", "Quan hệ xã hội được luật hình sự bảo vệ và bị tội phạm xâm hại."),
        ("Đối tượng tác động của tội phạm", "Bộ phận cụ thể mà hành vi tác động tới, như tài sản, con người hoặc vật thể."),
        ("Bị hại", "Cá nhân trực tiếp bị thiệt hại về thể chất, tinh thần, tài sản hoặc cơ quan, tổ chức bị thiệt hại do tội phạm gây ra theo luật tố tụng."),
        ("Mặt khách quan của tội phạm", "Biểu hiện bên ngoài gồm hành vi, hậu quả, quan hệ nhân quả và dấu hiệu khác do cấu thành tội quy định."),
        ("Chủ thể của tội phạm", "Người có năng lực trách nhiệm hình sự, đủ tuổi luật định hoặc pháp nhân thương mại đủ điều kiện thực hiện tội phạm."),
        ("Mặt chủ quan của tội phạm", "Diễn biến tâm lý bên trong gồm lỗi và, khi luật yêu cầu, động cơ, mục đích."),
        ("Tội phạm ít nghiêm trọng", "Tội có mức độ nguy hiểm không lớn, mức cao nhất của khung hình phạt là phạt tiền, cải tạo không giam giữ hoặc tù đến 03 năm."),
        ("Tội phạm nghiêm trọng", "Tội có mức độ nguy hiểm lớn, mức cao nhất của khung hình phạt trên 03 năm đến 07 năm tù."),
        ("Tội phạm rất nghiêm trọng", "Tội có mức độ nguy hiểm rất lớn, mức cao nhất của khung hình phạt trên 07 năm đến 15 năm tù."),
        ("Tội phạm đặc biệt nghiêm trọng", "Tội có mức độ nguy hiểm đặc biệt lớn, mức cao nhất của khung hình phạt trên 15 năm đến 20 năm tù, tù chung thân hoặc tử hình."),
        ("Đồng phạm", "Trường hợp có từ hai người trở lên cố ý cùng thực hiện một tội phạm."),
        ("Người tổ chức", "Người chủ mưu, cầm đầu hoặc chỉ huy việc thực hiện tội phạm."),
        ("Người thực hành", "Người trực tiếp thực hiện tội phạm."),
        ("Người xúi giục", "Người kích động, dụ dỗ hoặc thúc đẩy người khác thực hiện tội phạm."),
        ("Người giúp sức", "Người tạo điều kiện tinh thần hoặc vật chất cho việc thực hiện tội phạm."),
        ("Tội trộm cắp tài sản", "Tội chiếm đoạt tài sản bằng thủ đoạn lén lút trong điều kiện luật định."),
        ("Tội lừa đảo chiếm đoạt tài sản", "Tội dùng thủ đoạn gian dối có trước để làm chủ tài sản tin tưởng giao tài sản rồi chiếm đoạt."),
        ("Tội lạm dụng tín nhiệm chiếm đoạt tài sản", "Tội nhận tài sản hợp pháp thông qua giao dịch hoặc tín nhiệm rồi thực hiện hành vi luật định nhằm chiếm đoạt."),
        ("Tội cướp tài sản", "Tội dùng vũ lực, đe dọa dùng ngay tức khắc vũ lực hoặc thủ đoạn khác làm người bị tấn công không thể chống cự nhằm chiếm đoạt."),
        ("Tội cướp giật tài sản", "Tội công khai, nhanh chóng giật lấy tài sản rồi tẩu thoát."),
        ("Tội công nhiên chiếm đoạt tài sản", "Tội công khai lấy tài sản khi chủ thể quản lý biết nhưng do hoàn cảnh không thể ngăn cản."),
        ("Tội chứa chấp hoặc tiêu thụ tài sản do người khác phạm tội mà có", "Tội không hứa hẹn trước nhưng biết rõ tài sản do người khác phạm tội mà có vẫn cất giữ, sử dụng, chuyển nhượng hoặc tiêu thụ."),
        ("Mục đích chiếm đoạt", "Ý chí chuyển tài sản của người khác thành của mình hoặc làm chủ tài sản trái pháp luật."),
    ]),
    ("11. Hôn nhân và gia đình", [
        ("Luật Hôn nhân và gia đình", "Ngành luật điều chỉnh quan hệ nhân thân và tài sản giữa vợ chồng, cha mẹ và con, cùng các thành viên gia đình."),
        ("Hôn nhân", "Quan hệ giữa vợ và chồng sau khi kết hôn hợp pháp."),
        ("Gia đình", "Tập hợp người gắn bó do hôn nhân, huyết thống, nuôi dưỡng, làm phát sinh quyền và nghĩa vụ giữa họ."),
        ("Kết hôn", "Việc nam và nữ xác lập quan hệ vợ chồng theo điều kiện kết hôn và đăng ký kết hôn."),
        ("Hôn nhân tự nguyện", "Hôn nhân được xác lập trên ý chí tự do của các bên, không bị cưỡng ép, lừa dối hoặc cản trở."),
        ("Hôn nhân tiến bộ", "Hôn nhân bảo đảm tự nguyện, bình đẳng, tôn trọng quyền nhân thân và quyền yêu cầu ly hôn theo luật."),
        ("Nguyên tắc một vợ một chồng", "Nguyên tắc một người đang có vợ hoặc chồng không được kết hôn hay chung sống như vợ chồng với người khác."),
        ("Bình đẳng vợ chồng", "Nguyên tắc vợ và chồng có quyền, nghĩa vụ ngang nhau về nhân thân, tài sản, gia đình và chăm sóc con."),
        ("Đăng ký kết hôn", "Thủ tục hộ tịch tại cơ quan có thẩm quyền để Nhà nước công nhận quan hệ vợ chồng."),
        ("Chung sống như vợ chồng", "Việc nam và nữ tổ chức cuộc sống chung và coi nhau là vợ chồng nhưng có thể chưa đăng ký kết hôn."),
        ("Tài sản chung vợ chồng", "Tài sản được tạo lập trong thời kỳ hôn nhân và tài sản khác thuộc sở hữu chung theo luật hoặc thỏa thuận."),
        ("Tài sản riêng vợ chồng", "Tài sản có trước hôn nhân, được thừa kế riêng, tặng cho riêng, được chia riêng hoặc tài sản khác theo luật."),
        ("Con chung", "Người con có quan hệ cha mẹ - con được xác định với cả hai bên, không phụ thuộc cha mẹ có đăng ký kết hôn hay không."),
        ("Cấp dưỡng", "Nghĩa vụ đóng góp tiền hoặc tài sản để đáp ứng nhu cầu thiết yếu của người không sống chung nhưng có quan hệ gia đình theo luật."),
        ("Ly hôn", "Việc chấm dứt quan hệ vợ chồng theo bản án hoặc quyết định có hiệu lực của Tòa án."),
        ("Kết hôn trái pháp luật", "Việc đã đăng ký kết hôn nhưng một hoặc cả hai bên vi phạm điều kiện kết hôn."),
        ("Tảo hôn", "Việc lấy vợ, lấy chồng khi một hoặc cả hai bên chưa đủ tuổi kết hôn theo luật."),
        ("Cưỡng ép kết hôn", "Hành vi đe dọa, uy hiếp, hành hạ hoặc dùng thủ đoạn buộc người khác kết hôn trái ý muốn."),
        ("Kết hôn giả tạo", "Việc lợi dụng kết hôn không nhằm xây dựng gia đình mà để xuất nhập cảnh, cư trú, nhập quốc tịch, hưởng chế độ hoặc trục lợi khác."),
        ("Bạo lực gia đình", "Hành vi cố ý gây tổn hại hoặc có khả năng gây tổn hại thể chất, tinh thần, tình dục hoặc kinh tế giữa các thành viên gia đình."),
    ]),
    ("12. Khái niệm bổ sung trong các bài tập", [
        ("Hợp đồng vay tài sản", "Thỏa thuận bên cho vay giao tài sản cho bên vay; đến hạn bên vay phải hoàn trả tài sản cùng loại và lãi nếu có thỏa thuận hoặc luật quy định."),
        ("Lãi", "Khoản lợi ích bên vay phải trả thêm trên số tiền hoặc tài sản vay theo thỏa thuận trong giới hạn pháp luật."),
        ("Giao dịch dân sự vô hiệu", "Giao dịch không làm phát sinh, thay đổi hoặc chấm dứt quyền, nghĩa vụ như các bên mong muốn do vi phạm điều kiện có hiệu lực."),
        ("Lừa dối trong giao dịch dân sự", "Hành vi cố ý làm bên kia hiểu sai lệch về chủ thể, tính chất đối tượng hoặc nội dung để họ xác lập giao dịch."),
        ("Chiếm hữu ngay tình", "Việc chiếm hữu mà người chiếm hữu có căn cứ tin rằng mình có quyền đối với tài sản và không biết, không thể biết việc chiếm hữu là không có căn cứ."),
        ("Vật vô chủ", "Vật mà chủ sở hữu đã từ bỏ quyền sở hữu hoặc chưa xác định được chủ sở hữu theo pháp luật."),
        ("Vật bị đánh rơi, bỏ quên", "Tài sản rời khỏi sự quản lý của chủ sở hữu ngoài ý chí hoặc do sơ suất, phải được thông báo, giao nộp và xử lý theo luật."),
        ("Thời hiệu", "Thời hạn do luật quy định mà khi kết thúc làm phát sinh hậu quả pháp lý đối với chủ thể theo điều kiện nhất định."),
        ("Tai nạn giao thông", "Sự kiện xảy ra trong hoạt động giao thông gây thiệt hại về người, tài sản hoặc lợi ích khác."),
        ("Thương tích", "Tổn hại cơ thể được xác định về tính chất và tỷ lệ theo giám định trong trường hợp pháp luật yêu cầu."),
        ("Quyết định xử phạt sai thẩm quyền", "Quyết định áp dụng xử phạt do người không có quyền hoặc vượt quá phạm vi quyền hạn luật định ban hành."),
        ("Nhận hối lộ", "Hành vi của người có chức vụ, quyền hạn trực tiếp hoặc qua trung gian nhận hoặc sẽ nhận lợi ích để làm hoặc không làm việc vì lợi ích của người đưa."),
        ("Đời sống riêng tư", "Phạm vi thông tin và hoạt động cá nhân được pháp luật bảo vệ khỏi việc thu thập, sử dụng, công khai trái phép."),
        ("Bí mật cá nhân", "Thông tin gắn với cá nhân được pháp luật bảo vệ và chỉ được thu thập, sử dụng, công khai theo căn cứ hợp pháp."),
        ("Quyền đối với hình ảnh", "Quyền của cá nhân quyết định việc sử dụng hình ảnh của mình, trừ trường hợp luật quy định."),
        ("Giao dịch qua mạng", "Giao dịch được xác lập, thực hiện bằng phương tiện điện tử và chịu các điều kiện pháp luật về giao dịch, chứng cứ, bảo vệ chủ thể."),
        ("Hàng hóa giả tạo trong lừa đảo", "Thông tin, hình ảnh hoặc đối tượng không có thật được dựng lên nhằm làm người khác tin và chuyển giao tài sản."),
        ("Tình tiết tăng nặng trách nhiệm", "Tình tiết luật định làm tăng mức độ trách nhiệm khi quyết định xử lý, nhưng không được dùng lặp lại nếu đã là dấu hiệu định tội hoặc định khung."),
        ("Tình tiết giảm nhẹ trách nhiệm", "Tình tiết luật định hoặc được xem xét làm giảm mức độ trách nhiệm, như tự nguyện khắc phục, thành khẩn trong điều kiện luật định."),
    ]),
]


def add_glossary(doc):
    doc.add_page_break()
    doc.add_heading("PHẦN IV. TỪ ĐIỂN TOÀN BỘ KHÁI NIỆM, THUẬT NGỮ PHÁP LÝ", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(
        "Phần này liệt kê riêng từng khái niệm pháp lý được sử dụng trong 19 câu lý thuyết, "
        "bảng ghi nhớ và 14 bài tập. Các thuật ngữ gần nhau vẫn được tách thành từng mục, "
        "không gộp thành một định nghĩa chung."
    )
    set_font(r)
    for group_title, entries in GLOSSARY_GROUPS:
        doc.add_heading(group_title, level=2)
        for term, definition in entries:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(term + ": ")
            set_font(r, bold=True, color=DARK)
            r = p.add_run(definition)
            set_font(r)
    add_key(
        doc,
        f"Tổng số thuật ngữ/khái niệm được định nghĩa riêng: {sum(len(entries) for _, entries in GLOSSARY_GROUPS)}.",
        label="Kết quả rà soát",
    )


def add_checklist(doc):
    doc.add_page_break()
    doc.add_heading("PHẦN V. CHECKLIST TRƯỚC KHI ĐI THI", level=1)
    checks = [
        "Tôi phân biệt được quy phạm pháp luật với điều luật và văn bản QPPL.",
        "Tôi nhớ ba yếu tố của quan hệ pháp luật: chủ thể – nội dung – khách thể.",
        "Tôi nhớ bốn dấu hiệu của vi phạm pháp luật.",
        "Tôi phân tích được mặt khách quan – mặt chủ quan – chủ thể – khách thể.",
        "Tôi phân biệt được sáu tội chiếm đoạt tài sản thường gặp.",
        "Tôi luôn tách tài sản chung trước khi chia thừa kế.",
        "Tôi biết xác định người thừa kế bắt buộc.",
        "Tôi phân biệt tố cáo với khiếu nại.",
        "Tôi biết tranh chấp lao động nào không bắt buộc hòa giải.",
        "Tôi không dùng BLDS 2005, BLHS 1999, Luật HN&GĐ 2000 hoặc cơ chế lao động cũ nếu đề hỏi pháp luật hiện hành.",
    ]
    for item in checks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run("☐  ")
        set_font(r, size=13, bold=True, color=BLUE)
        r = p.add_run(item)
        set_font(r)

    doc.add_heading("MẪU KHUNG TRẢ LỜI TỰ LUẬN", level=2)
    add_number(doc, "Nêu khái niệm chính xác, không vòng vo.")
    add_number(doc, "Liệt kê đặc điểm hoặc nguyên tắc theo từng ý.")
    add_number(doc, "Phân tích ngắn từng ý và gắn một ví dụ.")
    add_number(doc, "Nếu có căn cứ pháp luật, nêu tên văn bản và điều luật khi chắc chắn.")
    add_number(doc, "Kết luận hoặc nêu ý nghĩa thực tiễn.")
    add_key(doc, "Một bài có bố cục rõ, thuật ngữ đúng và biết phân nhánh dữ kiện thường an toàn hơn một bài dài nhưng khẳng định sai.", label="Lời nhắc cuối")


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_contents(doc)
    add_intro(doc)
    add_theory(doc)
    add_memory_tables(doc)
    add_advanced_exercises_only(doc)
    add_glossary(doc)
    add_checklist(doc)

    core = doc.core_properties
    core.title = "Đáp án hoàn chỉnh 19 câu Pháp luật đại cương"
    core.subject = "Rà soát toàn diện 19 câu, bổ sung định nghĩa phân loại và từ điển đầy đủ khái niệm pháp lý"
    core.author = "Bản hiệu đính phục vụ học tập"
    core.keywords = "Pháp luật đại cương, đề cương, bài tập, Việt Nam"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
