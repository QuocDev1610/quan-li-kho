from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def extract(path: Path) -> dict:
    doc = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            paragraphs.append(
                {
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else "",
                    "text": text,
                }
            )

    tables = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row_index, row in enumerate(table.rows, start=1):
            cells = []
            for cell_index, cell in enumerate(row.cells, start=1):
                cells.append(
                    {
                        "index": cell_index,
                        "text": "\n".join(
                            p.text.strip() for p in cell.paragraphs if p.text.strip()
                        ),
                    }
                )
            rows.append({"index": row_index, "cells": cells})
        tables.append({"index": table_index, "rows": rows})

    return {
        "file": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def main() -> None:
    if len(sys.argv) not in (2, 3):
        raise SystemExit("Usage: extract_docx.py INPUT.docx [OUTPUT.json]")
    payload = extract(Path(sys.argv[1]))
    if len(sys.argv) == 3:
        Path(sys.argv[2]).write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    else:
        json.dump(payload, sys.stdout, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()
