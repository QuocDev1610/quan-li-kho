from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\ProjectWeb\Project - Sao chép")
OUT = ROOT / "output" / "docx" / "cam-nang-on-thi-ky-thuat-vi-xu-ly.docx"
MEDIA = ROOT / "tmp" / "study_sources" / "pptx_media"

NAVY = "17324D"
BLUE = "2E75B6"
LIGHT_BLUE = "DCEAF7"
PALE_BLUE = "EEF5FB"
GOLD = "C9972B"
PALE_GOLD = "FFF4D6"
RED = "A12A2A"
PALE_RED = "FCE8E6"
GREEN = "38755B"
PALE_GREEN = "E7F3EC"
GRAY = "5B6570"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "111111"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B7C1CC", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_keep(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def set_font(run, name="Arial", size=10.5, bold=None, italic=None, color=BLACK) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_run(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_font(run, **kwargs)
    return run


def add_body(doc, text, *, bold_prefix=None, italic=False, after=5, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix):], italic=italic)
    else:
        add_run(p, text, italic=italic)
    if keep:
        set_repeat_keep(p)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        add_run(p, item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        add_run(p, item)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 14, 2: 10, 3: 7}[level])
    p.paragraph_format.space_after = Pt({1: 7, 2: 5, 3: 4}[level])
    p.clear()
    size = {1: 17, 2: 13.5, 3: 11.5}[level]
    color = {1: NAVY, 2: BLUE, 3: GREEN}[level]
    add_run(p, text, size=size, bold=True, color=color)
    return p


def add_callout(doc, label, text, kind="note"):
    fills = {"note": PALE_BLUE, "trap": PALE_RED, "formula": PALE_GOLD, "tip": PALE_GREEN}
    colors = {"note": BLUE, "trap": RED, "formula": GOLD, "tip": GREEN}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.3)
    set_cell_shading(cell, fills[kind])
    set_cell_margins(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"{label}: ", bold=True, color=colors[kind])
    add_run(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(header), size=font_size, bold=True, color=WHITE)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if r_idx % 2:
                set_cell_shading(cell, LIGHT_GRAY)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i == 0 and len(headers) > 1:
                add_run(p, str(value), size=font_size, bold=True)
            else:
                add_run(p, str(value), size=font_size)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_run(p, caption, size=9.5, bold=True, color=GRAY)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7F9")
    set_cell_margins(cell, 100, 150, 100, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, code, name="Consolas", size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_picture(doc, filename, width, caption, source):
    path = MEDIA / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    p.paragraph_format.space_after = Pt(2)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(6)
    add_run(cp, caption, size=8.5, italic=True, color=GRAY)
    add_run(cp, f" ({source})", size=8.5, color=GRAY)


def add_source(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, f"Nguồn trong bộ tài liệu: {text}", size=8.3, italic=True, color=GRAY)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        styles[name].font.name = "Arial"
        styles[name]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        styles[name]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        styles[name].font.size = Pt(10.3)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "CẨM NANG ÔN THI KỸ THUẬT VI XỬ LÝ", size=8.2, bold=True, color=GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Biên soạn hoàn toàn từ ma trận đề, VXL_CNTT.pptx và vxl.docx", size=8, color=GRAY)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "CẨM NANG ÔN THI", size=18, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "KỸ THUẬT VI XỬ LÝ", size=31, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Trọng tâm 89C51 - Timer/Counter - UART - Ngắt - ARM", size=14, color=BLUE)
    doc.add_paragraph()
    add_callout(
        doc,
        "Phạm vi",
        "Nội dung chuyên môn trong tài liệu này chỉ được tổng hợp từ ba tệp người học cung cấp. "
        "Các liên hệ hiện đại chỉ dùng để gợi hình, không bổ sung kiến thức ngoài nguồn.",
        "note",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    add_run(p, "Dành cho sinh viên năm 2 ngành CNTT", size=12, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Bản hệ thống hóa theo ma trận 40 câu / 60 phút", size=11, color=GRAY)
    add_page_break(doc)

    # How to use
    add_heading(doc, "Cách dùng cẩm nang", 1)
    add_numbered(
        doc,
        [
            "Đọc phần Bản đồ đề thi để biết mỗi câu đang kiểm tra nhóm kiến thức nào.",
            "Học bảng tra nhanh trước, sau đó làm các ví dụ truy vết lệnh và tính Timer/UART.",
            "Khi gặp câu trắc nghiệm, ghi lại trạng thái A, B, Rn, ô nhớ, CY và các cờ sau từng lệnh.",
            "Đối chiếu mục Bẫy đề thi trước khi chốt đáp án, đặc biệt với dấu #, ký hiệu @, MOV/MOVC/MOVX và RET/RETI.",
        ],
    )
    add_callout(
        doc,
        "Chiến lược 60 phút",
        "40 câu tương đương trung bình 90 giây/câu. Làm trước câu nhận biết về thanh ghi, chân, vector và chế độ; "
        "đánh dấu câu truy vết dài hoặc tính Timer để quay lại sau.",
        "tip",
    )

    add_heading(doc, "Bản đồ đề thi 40 câu", 1)
    add_table(
        doc,
        ["Cụm", "Nội dung", "Số câu theo ma trận", "Mức cần đạt"],
        [
            ("CLO1", "Tổng quan hệ VXL, cấu trúc CPU, nguyên lý hoạt động, cấu trúc bộ nhớ", "4", "Nhận biết"),
            ("CLO2", "Kiến trúc 89C51, chân, SFR, cờ, bộ nhớ, địa chỉ, ASM, phân loại lệnh", "11", "Nhận biết + thông hiểu"),
            ("CLO3-A", "Chức năng lệnh, giá trị thanh ghi/cờ, đoạn lệnh, chương trình con", "10", "Thông hiểu + vận dụng"),
            ("CLO3-B", "Timer/Counter: nguyên lý, chế độ, giá trị đếm, khởi tạo, lập trình", "5", "Thông hiểu + vận dụng"),
            ("CLO3-C", "UART: nguyên lý, thanh ghi, tốc độ baud, lập trình", "4", "Thông hiểu + vận dụng"),
            ("CLO3-D", "Ngắt: hoạt động, vector, IE/IP, Timer, ngoài, nối tiếp", "6", "Thông hiểu + vận dụng"),
        ],
        widths=[0.8, 3.5, 0.9, 1.2],
    )
    add_source(doc, "ma-tran-de-ky-thuat-vi-xu-ly.pdf, trang 6-8.")

    add_heading(doc, "Vai trò của ngân hàng 100 câu", 2)
    add_body(
        doc,
        "vxl.docx cho thấy đề luyện tập tập trung dày vào: cú pháp lệnh, truy vết RAM nội, cờ PSW, "
        "MOV/MOVC/MOVX, bộ nhớ ngoài, TMOD/TCON, Timer chế độ 1-2, SCON/SBUF, vector ngắt và IE/IP. "
        "Vì vậy các phần này được trình bày chi tiết hơn ARM."
    )
    add_source(doc, "vxl.docx, bộ 100 câu hỏi 89C51.")

    # Chapter 1
    add_page_break(doc)
    add_heading(doc, "1. Tổng quan bộ vi xử lý và hệ vi xử lý", 1)
    add_heading(doc, "1.1. Cấu trúc hệ vi xử lý", 2)
    add_body(
        doc,
        "Một hệ vi xử lý gồm CPU, ROM, RAM, khối vào/ra (I/O), thiết bị ngoại vi, nguồn xung CLK và ba hệ bus: "
        "bus địa chỉ, bus dữ liệu, bus điều khiển."
    )
    add_table(
        doc,
        ["Thành phần", "Chức năng trọng tâm", "Câu hỏi dễ gặp"],
        [
            ("CPU", "Nhận lệnh, giải mã, thực hiện; điều khiển toàn hệ thống", "Thứ tự chu trình lệnh"),
            ("ROM", "Lưu chương trình/dữ liệu cố định", "Phân biệt với RAM"),
            ("RAM", "Lưu dữ liệu và chương trình đang hoạt động", "Không gian và dung lượng"),
            ("I/O", "Trung gian trao đổi với ngoại vi", "Khi nào cổng được mở"),
            ("Bus địa chỉ", "Chọn ô nhớ/cổng cần truy cập", "Một chiều từ CPU trong mô hình cơ bản"),
            ("Bus dữ liệu", "Mang dữ liệu/lệnh", "Hai chiều"),
            ("Bus điều khiển", "Mang tín hiệu đọc, ghi, cho phép...", "Phân biệt RD, WR, PSEN, ALE"),
        ],
        widths=[1.0, 3.4, 1.9],
    )
    add_source(doc, "VXL_CNTT.pptx, slide 8-10.")

    add_heading(doc, "1.2. Chu trình hoạt động", 2)
    add_numbered(
        doc,
        [
            "Nhận lệnh: PC tạo địa chỉ trên bus địa chỉ; CPU phát tín hiệu đọc; byte lệnh được đưa về CPU; PC tăng để chuẩn bị nhận byte tiếp theo.",
            "Giải mã: byte đầu là mã lệnh; CPU xác định chức năng và cách lấy toán hạng.",
            "Thực hiện: CPU thao tác nội bộ hoặc phát địa chỉ/tín hiệu điều khiển ra bên ngoài; kết quả được ghi về đích.",
            "Lặp lại với lệnh kế tiếp. Các giai đoạn có thể tuần tự hoặc song song tùy cấu trúc bộ xử lý.",
        ],
    )
    add_callout(
        doc,
        "Dạng thi",
        "Ngân hàng 100 câu hỏi yêu cầu chọn đúng thứ tự: nhận lệnh -> giải mã lệnh -> nhận dữ liệu -> xử lý dữ liệu -> ghi dữ liệu.",
        "note",
    )

    add_heading(doc, "1.3. Vi xử lý và vi điều khiển", 2)
    add_table(
        doc,
        ["Tiêu chí", "Vi xử lý đa năng", "Vi điều khiển"],
        [
            ("Mức tích hợp", "Khối tính toán trên chip; bộ nhớ và I/O thường ở ngoài", "Tích hợp CPU, bộ nhớ và I/O"),
            ("Ví dụ trong slide", "8086/8088, 6800", "8051"),
            ("Định hướng", "Hệ thống tính toán đa năng", "Điều khiển nhúng"),
        ],
        widths=[1.2, 2.55, 2.55],
    )
    add_source(doc, "VXL_CNTT.pptx, slide 11-14.")

    # Chapter 2
    add_page_break(doc)
    add_heading(doc, "2. Kiến trúc vi điều khiển 89C51", 1)
    add_heading(doc, "2.1. Thông số phải nhớ", 2)
    add_table(
        doc,
        ["Hạng mục", "Giá trị nêu trong slide"],
        [
            ("Đóng gói", "40 chân"),
            ("ROM trong", "4 KB"),
            ("RAM trong", "128 byte"),
            ("Cổng I/O", "4 cổng 8 bit"),
            ("Timer/Counter", "2 bộ 16 bit"),
            ("Cổng nối tiếp", "1 cổng, full duplex"),
            ("Bộ nhớ chương trình ngoài", "Tối đa 64 KB"),
            ("Bộ nhớ dữ liệu ngoài", "Tối đa 64 KB"),
            ("Vị trí định địa chỉ bit", "210 vị trí"),
            ("Tần số được nêu", "11.0592, 12, 16, 24 MHz"),
        ],
        widths=[2.4, 3.9],
    )
    add_callout(
        doc,
        "Mâu thuẫn trong nguồn",
        "Slide tổng quan ghi 6 nguồn ngắt, nhưng chương ngắt và bảng vector liệt kê 5 nguồn chuẩn của nội dung môn học. "
        "Khi làm câu vector/IE/IP, bám bảng 5 nguồn ở slide 111-116.",
        "trap",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 18, 111-116.")

    add_heading(doc, "2.2. Chân và cổng", 2)
    add_picture(doc, "image7.png", 3.1, "Sơ đồ chân 8051", "slide 19 và 23")
    add_table(
        doc,
        ["Nhóm chân", "Chức năng"],
        [
            ("P0.0-P0.7", "I/O hoặc AD0-AD7 khi ghép bộ nhớ ngoài; cực máng hở nên cần điện trở kéo ngoài"),
            ("P1.0-P1.7", "I/O"),
            ("P2.0-P2.7", "I/O hoặc A8-A15 khi ghép bộ nhớ ngoài"),
            ("P3.0/RXD", "Nhận nối tiếp"),
            ("P3.1/TXD", "Phát nối tiếp"),
            ("P3.2/INT0", "Ngắt ngoài 0"),
            ("P3.3/INT1", "Ngắt ngoài 1"),
            ("P3.4/T0", "Đầu vào xung đếm Timer/Counter 0"),
            ("P3.5/T1", "Đầu vào xung đếm Timer/Counter 1"),
            ("P3.6/WR", "Ghi bộ nhớ dữ liệu ngoài"),
            ("P3.7/RD", "Đọc bộ nhớ dữ liệu ngoài"),
            ("ALE", "Chốt byte địa chỉ thấp khi P0 ghép kênh địa chỉ/dữ liệu"),
            ("PSEN", "Cho phép đọc bộ nhớ chương trình ngoài"),
            ("EA", "Chọn cách truy xuất bộ nhớ chương trình"),
            ("RST", "Mức cao làm khởi động lại; PC về 0000H"),
            ("XTAL1/XTAL2", "Kết nối mạch dao động"),
        ],
        widths=[1.35, 4.95],
        font_size=8.9,
    )
    add_callout(
        doc,
        "Bẫy EA",
        "Theo ngân hàng câu hỏi: EA=0 -> chỉ lấy lệnh từ bộ nhớ chương trình ngoài. "
        "EA=1 -> ưu tiên ROM trong; khi PC vượt phạm vi ROM trong thì truy xuất ngoài.",
        "trap",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 19-24; vxl.docx, câu 68-69.")

    add_heading(doc, "2.3. Tổ chức RAM trong", 2)
    add_picture(doc, "image12.png", 3.4, "Bản đồ RAM trong 00H-7FH", "slide 26")
    add_table(
        doc,
        ["Vùng", "Địa chỉ", "Nội dung"],
        [
            ("Băng thanh ghi", "00H-1FH", "4 băng, mỗi băng R0-R7"),
            ("Định địa chỉ bit", "20H-2FH", "16 byte = 128 bit, truy cập theo byte hoặc bit"),
            ("RAM đa dụng", "30H-7FH", "80 byte lưu dữ liệu/tham số"),
        ],
        widths=[1.5, 1.2, 3.6],
    )
    add_table(
        doc,
        ["RS1", "RS0", "Băng", "Địa chỉ R0-R7"],
        [
            ("0", "0", "0", "00H-07H"),
            ("0", "1", "1", "08H-0FH"),
            ("1", "0", "2", "10H-17H"),
            ("1", "1", "3", "18H-1FH"),
        ],
        widths=[0.8, 0.8, 1.1, 3.6],
    )
    add_callout(
        doc,
        "Stack",
        "Sau reset SP=07H. PUSH/call/ngắt làm SP tăng trước khi ghi; POP/RET/RETI lấy dữ liệu rồi làm SP giảm. "
        "Slide khuyên chuyển SP sang vùng 30H-7FH nếu cần stack lớn và không dùng vùng bit 20H-2FH làm stack.",
        "formula",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 26-28; vxl.docx, câu 12, 62, 87.")

    add_heading(doc, "2.4. Bộ nhớ ngoài và tính dung lượng", 2)
    add_body(
        doc,
        "8051 tách không gian chương trình và dữ liệu; mỗi không gian có thể mở rộng tối đa 64 KB. "
        "P0 mang AD0-AD7, ALE điều khiển latch để tách A0-A7; P2 mang A8-A15."
    )
    add_picture(doc, "image14.png", 4.7, "Không gian bộ nhớ chương trình trong/ngoài", "slide 29")
    add_picture(doc, "image15.png", 4.7, "Truy xuất ROM ngoài bằng P0, P2, ALE và PSEN", "slide 30")
    add_callout(
        doc,
        "Công thức dung lượng",
        "Nếu có n đường địa chỉ biến thiên thì dung lượng = 2^n byte. "
        "Một miền từ địa chỉ đầu đến địa chỉ cuối có số byte = địa chỉ cuối - địa chỉ đầu + 1.",
        "formula",
    )
    add_bullets(
        doc,
        [
            "MOVC đọc bộ nhớ chương trình.",
            "MOVX đọc/ghi bộ nhớ dữ liệu ngoài.",
            "MOV thao tác thanh ghi, RAM trong và SFR theo các dạng hợp lệ.",
            "R0/R1 chỉ mang địa chỉ 8 bit trong dạng @Ri; DPTR mang địa chỉ 16 bit.",
        ],
    )
    add_source(doc, "VXL_CNTT.pptx, slide 29-39, 52; vxl.docx, câu 6, 26, 34, 52, 74, 88, 99.")

    # Registers
    add_page_break(doc)
    add_heading(doc, "3. Thanh ghi và cờ của 89C51", 1)
    add_heading(doc, "3.1. Nhóm thanh ghi cốt lõi", 2)
    add_table(
        doc,
        ["Thanh ghi", "Địa chỉ", "Vai trò"],
        [
            ("A/ACC", "E0H", "Thanh tích lũy, toán hạng/đích chính của số học và logic"),
            ("B", "F0H", "Dùng với A trong MUL AB và DIV AB"),
            ("PSW", "D0H", "Cờ trạng thái và chọn băng thanh ghi"),
            ("SP", "81H", "Con trỏ đỉnh stack"),
            ("DPL/DPH", "82H/83H", "Hai byte thấp/cao của DPTR 16 bit"),
            ("TL0/TH0", "8AH/8CH", "Giá trị Timer 0"),
            ("TL1/TH1", "8BH/8DH", "Giá trị Timer 1"),
            ("TMOD", "89H", "Chọn Timer/Counter, GATE và chế độ"),
            ("TCON", "88H", "Điều khiển Timer và ngắt ngoài"),
            ("SCON", "98H", "Cấu hình/cờ cổng nối tiếp"),
            ("SBUF", "99H", "Đệm dữ liệu truyền hoặc nhận"),
            ("IE", "A8H", "Cho phép/cấm ngắt"),
            ("IP", "B8H", "Mức ưu tiên ngắt"),
            ("PCON", "87H", "Điều khiển nguồn; SMOD có thể nhân đôi baud"),
        ],
        widths=[1.2, 1.0, 4.1],
        font_size=8.9,
    )
    add_source(doc, "VXL_CNTT.pptx, slide 41-46, 72-76, 90-100, 112-116.")

    add_heading(doc, "3.2. PSW và cách tính cờ", 2)
    add_picture(doc, "image23.png", 3.3, "Các SFR có thể định địa chỉ bit", "slide 41")
    add_table(
        doc,
        ["Bit", "Ký hiệu", "Ý nghĩa"],
        [
            ("PSW.7", "CY", "Nhớ/mượn chính"),
            ("PSW.6", "AC", "Nhớ từ bit 3 sang bit 4"),
            ("PSW.5", "-", "Người dùng"),
            ("PSW.4", "RS1", "Chọn băng thanh ghi"),
            ("PSW.3", "RS0", "Chọn băng thanh ghi"),
            ("PSW.2", "OV", "Tràn số có dấu"),
            ("PSW.1", "-", "Người dùng"),
            ("PSW.0", "P", "1 khi A có số bit 1 lẻ; 0 khi số bit 1 chẵn"),
        ],
        widths=[1.1, 1.0, 4.2],
    )
    add_callout(
        doc,
        "Quy trình tính cờ",
        "Viết phép tính nhị phân hoặc hexa; xác định carry ra bit 7 (CY), carry từ nibble thấp (AC), "
        "tràn dấu (OV), rồi đếm số bit 1 của kết quả A để tính P.",
        "tip",
    )
    add_table(
        doc,
        ["Ví dụ từ ngân hàng", "Kết quả"],
        [
            ("A=AFH, (20H)=81H, CY=0; ADDC A,20H", "A=30H, CY=1"),
            ("A=FBH, R4=53H, CY=1; SUBB A,R4", "A=A7H"),
            ("A=D7H, R0=2FH; ADD A,R0", "A=06H; AC=1, OV=0, CY=1, P=0"),
            ("A=B9H, R1=81H, PSW=78H; ADDC A,R1", "CY,AC,OV,P = 1,0,1,0"),
        ],
        widths=[4.1, 2.2],
        font_size=9,
    )
    add_source(doc, "vxl.docx, câu 10, 48, 53, 65.")

    # Addressing and instructions
    add_page_break(doc)
    add_heading(doc, "4. Chế độ địa chỉ và tập lệnh 8051", 1)
    add_heading(doc, "4.1. Bốn chế độ địa chỉ trọng tâm", 2)
    add_table(
        doc,
        ["Chế độ", "Dấu hiệu", "Ví dụ", "Ý nghĩa"],
        [
            ("Tức thời", "#data", "MOV A,#25H", "Nạp hằng 25H"),
            ("Thanh ghi", "Rn/A/B", "MOV A,R0", "Lấy nội dung thanh ghi"),
            ("Trực tiếp", "địa chỉ không #", "MOV A,35H", "Lấy nội dung ô 35H"),
            ("Gián tiếp qua thanh ghi", "@R0 hoặc @R1", "MOV A,@R0", "R0 chứa địa chỉ ô nhớ"),
        ],
        widths=[1.35, 1.15, 1.45, 2.35],
    )
    add_callout(
        doc,
        "Bẫy số 1",
        "MOV A,#73H nạp hằng 73H; MOV A,73H lấy nội dung ô nhớ 73H; MOV A,@R1 lấy nội dung ô nhớ có địa chỉ nằm trong R1.",
        "trap",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 47-48; vxl.docx, câu 22, 89, 97.")

    add_heading(doc, "4.2. MOV, MOVC, MOVX", 2)
    add_table(
        doc,
        ["Lệnh", "Không gian", "Mẫu quan trọng"],
        [
            ("MOV", "Thanh ghi, RAM trong, SFR", "MOV A,30H; MOV @R0,A; MOV DPTR,#2000H"),
            ("MOVC", "Bộ nhớ chương trình", "MOVC A,@A+DPTR hoặc MOVC A,@A+PC"),
            ("MOVX", "RAM ngoài", "MOVX A,@DPTR; MOVX @DPTR,A; dạng @Ri"),
        ],
        widths=[1.0, 2.0, 3.3],
    )
    add_code(
        doc,
        "MOV  A,20H\nMOV  DPTR,#2020H\nMOVX @DPTR,A",
        "Chuyển dữ liệu từ RAM trong 20H sang RAM ngoài 2020H",
    )
    add_callout(
        doc,
        "MOVC với PC",
        "PC dùng trong địa chỉ hiệu dụng là vị trí sau khi CPU đã lấy lệnh MOVC; khi đọc bảng DB đặt ngay sau lệnh, "
        "phải tính đúng độ lệch. Đây là dạng bẫy xuất hiện trong ngân hàng.",
        "trap",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 51-52; vxl.docx, câu 6, 32, 74, 88.")

    add_heading(doc, "4.3. Nhóm lệnh số học", 2)
    add_table(
        doc,
        ["Lệnh", "Phép toán", "Điểm cần nhớ"],
        [
            ("ADD A,src", "A=A+src", "Không cộng CY đầu vào"),
            ("ADDC A,src", "A=A+src+CY", "Dùng cho cộng nhiều byte"),
            ("SUBB A,src", "A=A-src-CY", "Phải kiểm tra/xóa CY nếu bài không muốn trừ mượn"),
            ("DA A", "Hiệu chỉnh BCD sau ADD/ADDC", "Nibble thấp >9 hoặc AC=1: +06H; nibble cao >9 hoặc CY=1: +60H"),
            ("MUL AB", "A*B", "A=byte thấp, B=byte cao"),
            ("DIV AB", "A/B", "A=thương, B=dư"),
            ("INC/DEC", "Tăng/giảm 1", "INC DPTR tăng đủ 16 bit"),
        ],
        widths=[1.3, 1.7, 3.3],
        font_size=9,
    )
    add_callout(
        doc,
        "Cơ số",
        "Trong tài liệu, số có H là hexa, B là nhị phân; số không hậu tố thường được hiểu là thập phân. "
        "Ví dụ MOV A,#60 rồi DEC A cho A=59 thập phân = 3BH.",
        "trap",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 55-59; vxl.docx, câu 9-10, 13, 48, 58, 61, 79, 96.")

    add_heading(doc, "4.4. Logic, bit và rẽ nhánh", 2)
    add_table(
        doc,
        ["Nhóm", "Lệnh tiêu biểu", "Ghi nhớ"],
        [
            ("Logic byte", "ANL, ORL, XRL, CPL A", "Dùng mặt nạ bit"),
            ("Quay/dịch", "RR, RL, RRC, RLC, SWAP", "RRC/RLC đi qua CY"),
            ("Bit", "SETB bit, CLR bit, CPL bit, MOV C,bit", "Chỉ dùng toán hạng bit hợp lệ"),
            ("Nhảy theo A", "JZ, JNZ", "Kiểm tra A bằng 0/khác 0"),
            ("Nhảy theo CY", "JC, JNC", "Kiểm tra cờ nhớ"),
            ("Nhảy theo bit", "JB, JNB, JBC", "JBC còn xóa bit sau khi kiểm tra bằng 1"),
            ("Vòng lặp", "DJNZ Rn,rel; DJNZ direct,rel", "Giảm rồi nhảy nếu khác 0"),
            ("So sánh", "CJNE ... ,rel", "So sánh và nhảy nếu khác"),
            ("Gọi/trả về", "ACALL/LCALL, RET, RETI", "RETI chỉ kết thúc ISR"),
        ],
        widths=[1.25, 2.2, 2.85],
        font_size=8.9,
    )
    add_callout(
        doc,
        "Đặt/xóa bit cao nhất A",
        "Đặt: ORL A,#80H hoặc SETB ACC.7. Xóa: ANL A,#7FH hoặc CLR ACC.7. "
        "Ngân hàng có nhiều đáp án tương đương, cần đọc yêu cầu chọn một hay chọn nhiều.",
        "tip",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 60-64; vxl.docx, câu 2, 11, 21, 37, 42, 54, 57, 66, 77-78, 83.")

    add_heading(doc, "4.5. Chỉ dẫn hợp dịch và khung chương trình", 2)
    add_table(
        doc,
        ["Chỉ dẫn", "Chức năng"],
        [
            ("ORG address", "Đặt địa chỉ bắt đầu đoạn mã/dữ liệu"),
            ("EQU value", "Định nghĩa hằng"),
            ("BIT bit_address", "Định nghĩa tên cho bit"),
            ("DB ...", "Khai báo các byte dữ liệu"),
            ("END", "Kết thúc nguồn hợp dịch"),
        ],
        widths=[1.6, 4.7],
    )
    add_code(
        doc,
        "ORG 0000H\n    LJMP MAIN\n\nORG 000BH\n    LJMP TIMER0_ISR\n\nORG 0030H\nMAIN:\n    ; khoi tao\nHERE: SJMP HERE\n\nTIMER0_ISR:\n    ; xu ly ngat\n    RETI\nEND",
        "Khung chương trình có ngắt",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 65-66, 122.")

    # Trace method
    add_page_break(doc)
    add_heading(doc, "5. Phương pháp truy vết chương trình", 1)
    add_heading(doc, "5.1. Bảng trạng thái bắt buộc", 2)
    add_body(
        doc,
        "Với mỗi lệnh, tạo một dòng gồm: PC/nhãn, lệnh, A, B, R0-R7 liên quan, DPTR, SP, CY/AC/OV/P và các ô nhớ bị thay đổi. "
        "Không tính nhẩm cả đoạn mã trong đầu."
    )
    add_table(
        doc,
        ["Bước", "Lệnh", "A", "CY", "Ô nhớ/thanh ghi khác", "Nhánh?"],
        [
            ("0", "Trạng thái đầu", "...", "...", "...", "-"),
            ("1", "Lệnh 1", "...", "...", "...", "Có/không"),
            ("2", "Lệnh 2", "...", "...", "...", "Có/không"),
        ],
        widths=[0.55, 1.5, 0.8, 0.65, 2.0, 0.8],
    )

    add_heading(doc, "5.2. Ví dụ nhánh theo bit", 2)
    add_code(
        doc,
        "MOV A,#9AH\nMOV 30H,#10H\nJB  ACC.5,SKIP\nADD A,30H\nMOV 30H,A\nSJMP EXIT\nSKIP: SUBB A,30H\n      MOV R0,30H\nEXIT: SJMP $",
    )
    add_body(
        doc,
        "9AH = 10011010B nên ACC.5=0; không nhảy SKIP. A=9AH+10H=AAH, sau đó (30H)=AAH. "
        "Lưu ý SUBB chỉ chạy khi ACC.5=1."
    )
    add_source(doc, "vxl.docx, câu 9.")

    add_heading(doc, "5.3. Chương trình con và stack", 2)
    add_bullets(
        doc,
        [
            "LCALL/ACALL cất địa chỉ trở về lên stack; RET lấy địa chỉ về.",
            "Khi có ngắt, địa chỉ lệnh kế tiếp được cất lên stack; RETI lấy lại PC và kết thúc trạng thái phục vụ ngắt.",
            "Nếu ISR không có RETI mà nhảy về chương trình chính, hai byte địa chỉ vẫn nằm trên stack; các lần ngắt tiếp theo làm SP tiếp tục tăng.",
            "Các lệnh ảnh hưởng SP trong ngân hàng: MOV SP,#data, PUSH, LCALL/ACALL, RET và RETI; ngắt phần cứng cũng làm SP tăng.",
        ],
    )
    add_callout(
        doc,
        "Bẫy",
        "SJMP/LJMP không tự động cất địa chỉ trở về. RET không thay cho RETI trong chương trình phục vụ ngắt.",
        "trap",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 28, 109-110; vxl.docx, câu 12, 62, 95.")

    # Timer
    add_page_break(doc)
    add_heading(doc, "6. Timer/Counter 8051", 1)
    add_heading(doc, "6.1. Nguyên lý", 2)
    add_body(
        doc,
        "Timer/Counter 0 và 1 là chuỗi flip-flop đếm tăng. Khi dùng Timer, nguồn xung là clock nội; "
        "khi dùng Counter, nguồn xung ngoài đi vào T0/P3.4 hoặc T1/P3.5. Khi tràn, TF0 hoặc TF1 được lập."
    )
    add_picture(doc, "image30.png", 5.4, "Thanh ghi TMOD: nửa cao cho Timer 1, nửa thấp cho Timer 0", "slide 73")
    add_table(
        doc,
        ["Bit", "Giá trị", "Ý nghĩa"],
        [
            ("GATE", "0", "Khởi động/dừng bằng TRx trong phần mềm"),
            ("GATE", "1", "Chạy khi TRx=1 và chân INTx ở mức cho phép"),
            ("C/T", "0", "Timer, xung nội = fosc/12"),
            ("C/T", "1", "Counter, xung ngoài tại T0/T1"),
            ("M1 M0", "00", "Chế độ 0, 13 bit"),
            ("M1 M0", "01", "Chế độ 1, 16 bit"),
            ("M1 M0", "10", "Chế độ 2, 8 bit tự nạp"),
            ("M1 M0", "11", "Chế độ 3, chia tách"),
        ],
        widths=[1.0, 1.0, 4.3],
    )
    add_picture(doc, "image32.png", 5.0, "Thanh ghi TCON", "slide 76")
    add_table(
        doc,
        ["Bit", "Chức năng"],
        [
            ("TF1/TF0", "Cờ tràn Timer 1/0"),
            ("TR1/TR0", "Bit chạy/dừng Timer 1/0"),
            ("IE1/IE0", "Cờ ngắt ngoài 1/0"),
            ("IT1/IT0", "Chọn kích hoạt ngắt ngoài theo sườn hay mức"),
        ],
        widths=[1.4, 4.9],
    )
    add_source(doc, "VXL_CNTT.pptx, slide 71-77.")

    add_heading(doc, "6.2. Công thức chế độ 1", 2)
    add_callout(
        doc,
        "Công thức",
        "Tchu_ky_may = 12/fosc. Số đếm N = t/Tchu_ky_may. Giá trị nạp 16 bit = 65536 - N. "
        "THx là byte cao, TLx là byte thấp.",
        "formula",
    )
    add_table(
        doc,
        ["Bài mẫu", "Tính toán", "Kết quả"],
        [
            ("fosc=12 MHz, t=1 ms", "Tmc=1 us; N=1000; 65536-1000=64536", "FC18H -> TH=FCH, TL=18H"),
            ("fosc=6 MHz, t=50 ms", "Tmc=2 us; N=25000; 65536-25000=40536", "9E58H -> TH=9EH, TL=58H"),
            ("fosc=6 MHz, TH0=C5H, TL0=35H", "N=65536-C535H=15051; t=15051*2 us", "30.102 ms"),
        ],
        widths=[1.7, 3.25, 1.35],
        font_size=8.8,
    )
    add_code(
        doc,
        "MOV TMOD,#01H\nMOV TH0,#0FCH\nMOV TL0,#18H\nSETB TR0\nWAIT: JNB TF0,WAIT\nCLR TR0\nCLR TF0\nRET",
        "Tạo trễ 1 ms bằng Timer 0 chế độ 1, fosc=12 MHz",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 78-83; vxl.docx, câu 35-36, 43, 60, 86, 93.")

    add_heading(doc, "6.3. Công thức chế độ 2", 2)
    add_callout(
        doc,
        "Công thức",
        "N = 256 - THx. Khi TLx tràn, TLx tự động nhận lại giá trị từ THx. "
        "Nếu mỗi lần tràn đảo chân bằng CPL thì thời gian tràn là nửa chu kỳ sóng vuông.",
        "formula",
    )
    add_table(
        doc,
        ["Tình huống", "Kết quả"],
        [
            ("TH1=5, fosc=12 MHz", "N=251 -> nửa chu kỳ 251 us -> T=502 us -> f≈1.992 kHz"),
            ("Đếm 10 xung ở chế độ 2", "Giá trị nạp = 256-10 = 246 = F6H; TH1=TL1=F6H khi khởi tạo"),
            ("Timer 1 chế độ 2", "TMOD nửa cao M1M0=10 -> 20H nếu GATE=0, C/T=0"),
            ("Counter 1 chế độ 2", "TMOD=60H nếu GATE=0, C/T=1"),
            ("Counter 0 chế độ 2", "TMOD=06H nếu GATE=0, C/T=1"),
        ],
        widths=[2.25, 4.05],
    )
    add_source(doc, "VXL_CNTT.pptx, slide 84-89; vxl.docx, câu 3, 8, 28, 40, 73, 80, 84-85.")

    add_heading(doc, "6.4. Quy trình giải câu Timer", 2)
    add_numbered(
        doc,
        [
            "Tách đúng nửa TMOD: Timer 1 ở 4 bit cao, Timer 0 ở 4 bit thấp.",
            "Đọc GATE, C/T, M1, M0; kết luận Timer hay Counter và chế độ.",
            "Nếu Timer, tính chu kỳ máy 12/fosc; nếu Counter, số đếm là số xung ngoài.",
            "Xác định giá trị nạp theo mô-đun 65536 hoặc 256.",
            "Nếu chân được CPL sau mỗi tràn, một lần tràn chỉ là nửa chu kỳ.",
            "Kiểm tra chương trình có nạp lại TH/TL và xóa TF hay không.",
        ],
    )
    add_callout(
        doc,
        "Bẫy JNB/JB",
        "Vòng chờ tràn thường là JNB TF0,WAIT: còn TF0=0 thì quay lại. "
        "JB TF0,label lại nhảy khi cờ đã bằng 1; ý nghĩa khác nhau.",
        "trap",
    )

    # UART
    add_page_break(doc)
    add_heading(doc, "7. Truyền thông nối tiếp", 1)
    add_heading(doc, "7.1. Khối và thanh ghi", 2)
    add_bullets(
        doc,
        [
            "TxD=P3.1 phát; RxD=P3.0 nhận.",
            "Cổng nối tiếp full duplex: có thể thu và phát đồng thời.",
            "SBUF chứa byte cần truyền hoặc byte vừa nhận.",
            "SCON cấu hình chế độ và chứa RI/TI.",
            "Timer 1 chế độ 2 thường tạo tốc độ baud thay đổi.",
        ],
    )
    add_picture(doc, "image38.png", 5.2, "Thanh ghi SCON", "slide 95 và 100")
    add_table(
        doc,
        ["Bit", "Ý nghĩa"],
        [
            ("SM0, SM1", "Chọn chế độ nối tiếp"),
            ("SM2", "Hỗ trợ chế độ đa xử lý"),
            ("REN", "1 cho phép thu; 0 khóa bộ thu"),
            ("TB8/RB8", "Bit dữ liệu thứ 9 trong mode 2/3"),
            ("TI", "1 khi truyền xong; phải xóa bằng phần mềm"),
            ("RI", "1 khi nhận xong; đọc SBUF rồi xóa bằng phần mềm"),
        ],
        widths=[1.2, 5.1],
    )
    add_table(
        doc,
        ["SM0", "SM1", "Mode", "Khung/chức năng", "Baud"],
        [
            ("0", "0", "0", "Thanh ghi dịch 8 bit", "Cố định"),
            ("0", "1", "1", "UART 8 bit: start + 8 data + stop", "Thay đổi qua Timer 1"),
            ("1", "0", "2", "UART 9 bit", "Cố định"),
            ("1", "1", "3", "UART 9 bit", "Thay đổi"),
        ],
        widths=[0.55, 0.55, 0.6, 3.5, 1.1],
        font_size=8.7,
    )
    add_source(doc, "VXL_CNTT.pptx, slide 90-100.")

    add_heading(doc, "7.2. Truyền và nhận ở mode 1", 2)
    add_code(
        doc,
        "MOV TMOD,#20H\nMOV TH1,#-6\nMOV SCON,#50H\nSETB TR1\n\nSEND: MOV SBUF,#'A'\nWAIT_T: JNB TI,WAIT_T\n        CLR TI\n        SJMP SEND",
        "Truyền liên tục ký tự A ở 4800 baud theo ví dụ slide",
    )
    add_code(
        doc,
        "MOV TMOD,#20H\nMOV TH1,#-6\nMOV SCON,#50H\nSETB TR1\n\nWAIT_R: JNB RI,WAIT_R\n        MOV A,SBUF\n        MOV P1,A\n        CLR RI\n        SJMP WAIT_R",
        "Nhận byte và đưa ra P1",
    )
    add_callout(
        doc,
        "Trình tự truyền",
        "Khởi tạo -> ghi SBUF -> phần cứng phát qua TxD -> TI=1 -> phần mềm xóa TI.",
        "tip",
    )
    add_callout(
        doc,
        "Trình tự nhận",
        "Khởi tạo REN -> dữ liệu vào RxD -> RI=1 -> đọc SBUF -> phần mềm xóa RI.",
        "tip",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 101-104; vxl.docx, câu 5, 16, 24, 38, 41, 50, 55-56, 72, 76, 83, 91, 98.")

    add_heading(doc, "7.3. Tốc độ baud", 2)
    add_body(
        doc,
        "Trong tài liệu, Timer 1 được đặt mode 2 để tạo baud. Với XTAL=11.0592 MHz, TH1=-6 được dùng cho 4800 baud; "
        "PCON.7/SMOD=1 làm nhân đôi tốc độ baud."
    )
    add_table(
        doc,
        ["Cấu hình", "Ý nghĩa"],
        [
            ("TMOD=20H", "Timer 1 mode 2, nguồn clock nội"),
            ("SCON=50H", "Mode 1, REN=1, RI=TI=0"),
            ("TH1=-6 tại 11.0592 MHz", "4800 baud theo ví dụ/bảng trong slide"),
            ("SMOD=1", "Nhân đôi baud"),
        ],
        widths=[2.2, 4.1],
    )
    add_callout(
        doc,
        "Bẫy mode 2 và mode 3",
        "Slide mô tả mode 2 có baud cố định và mode 3 có baud thay đổi; ngân hàng có câu hỏi dễ nhầm Timer 1 với PCON. "
        "PCON/SMOD chỉ nhân đôi, còn nguồn tạo baud thay đổi vẫn liên quan Timer 1.",
        "trap",
    )

    # Interrupt
    add_page_break(doc)
    add_heading(doc, "8. Ngắt trong 8051", 1)
    add_heading(doc, "8.1. Trình tự đáp ứng ngắt", 2)
    add_numbered(
        doc,
        [
            "CPU hoàn tất lệnh đang thực hiện.",
            "Địa chỉ lệnh kế tiếp được cất lên stack; SP tăng 2.",
            "CPU nạp địa chỉ vector ngắt vào PC.",
            "ISR thực thi.",
            "RETI lấy lại PC từ stack; SP giảm 2; chương trình chính tiếp tục.",
        ],
    )
    add_callout(
        doc,
        "Thuật ngữ",
        "Slide có chỗ ghi IRET nhưng toàn bộ khung chương trình 8051 dùng RETI. Khi làm bài 89C51, dùng RETI.",
        "trap",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 109-110.")

    add_heading(doc, "8.2. Bảng vector", 2)
    add_table(
        doc,
        ["Nguồn", "Cờ", "Vector"],
        [
            ("Reset", "-", "0000H"),
            ("INT0", "IE0", "0003H"),
            ("Timer 0", "TF0", "000BH"),
            ("INT1", "IE1", "0013H"),
            ("Timer 1", "TF1", "001BH"),
            ("Nối tiếp", "RI/TI", "0023H"),
        ],
        widths=[2.4, 1.4, 2.5],
    )
    add_callout(
        doc,
        "Mẹo nhớ",
        "Sau INT0 0003H, các vector cách nhau 8 byte: 000BH, 0013H, 001BH, 0023H.",
        "tip",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 111; vxl.docx, câu 15, 49.")

    add_heading(doc, "8.3. IE và IP", 2)
    add_picture(doc, "image45.png", 5.2, "Thanh ghi IE", "slide 112")
    add_table(
        doc,
        ["Bit IE", "Nguồn"],
        [
            ("EA", "Cho phép tổng; EA=0 cấm mọi ngắt"),
            ("ES", "Ngắt nối tiếp"),
            ("ET1", "Timer 1"),
            ("EX1", "Ngắt ngoài 1"),
            ("ET0", "Timer 0"),
            ("EX0", "Ngắt ngoài 0"),
        ],
        widths=[1.3, 5.0],
    )
    add_picture(doc, "image46.png", 5.0, "Thanh ghi IP", "slide 116")
    add_body(
        doc,
        "Sau reset, thứ tự ưu tiên mặc định từ cao xuống thấp trong tài liệu: INT0, Timer 0, INT1, Timer 1, nối tiếp. "
        "IP đặt nguồn tương ứng lên mức ưu tiên cao; nếu nhiều nguồn cùng mức, vẫn xét theo thứ tự mặc định."
    )
    add_callout(
        doc,
        "Điều kiện ngắt Timer 0",
        "TF0=1, ET0=1 và EA=1. Chỉ SETB ET0 chưa đủ nếu EA đang bằng 0.",
        "trap",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 112-118; vxl.docx, câu 23, 27, 29, 64, 67, 92.")

    add_heading(doc, "8.4. Ngắt ngoài và ngắt nối tiếp", 2)
    add_table(
        doc,
        ["Loại", "Điểm cần nhớ"],
        [
            ("INT0/INT1 theo mức", "Mặc định sau reset; mức thấp kích hoạt; phải về cao trước RETI để tránh ngắt lặp"),
            ("INT0/INT1 theo sườn", "IT0/IT1 chọn sườn xuống; IE0/IE1 được CPU lập và xóa khi xử lý"),
            ("Ngắt nối tiếp", "Chung một vector 0023H cho RI và TI; ISR phải kiểm tra cờ nào gây ngắt"),
            ("Cờ RI/TI", "Không tự xóa như một số cờ Timer trong quy trình ISR; phần mềm phải xóa"),
        ],
        widths=[2.0, 4.3],
    )
    add_callout(
        doc,
        "ISR Timer",
        "Theo slide, khi CPU nhảy vào vector Timer, TFx được xóa nên ISR không bắt buộc CLR TFx trước RETI. "
        "Trong chương trình polling thì phải xóa TFx bằng phần mềm.",
        "tip",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 122, 124-129.")

    # ARM
    add_page_break(doc)
    add_heading(doc, "9. Tổng quan ARM", 1)
    add_heading(doc, "9.1. Tư tưởng RISC và load-store", 2)
    add_bullets(
        doc,
        [
            "ARM là kiến trúc RISC 32 bit và 64 bit trong nội dung slide.",
            "Tập thanh ghi đồng dạng; kiểu địa chỉ đơn giản; lệnh có độ dài cố định và đồng dạng.",
            "Kiến trúc load-store: toán hạng được nạp vào thanh ghi, xử lý trên thanh ghi, sau đó lưu về bộ nhớ.",
            "ALU có thể kết hợp bộ dịch; hỗ trợ tự tăng/tự giảm địa chỉ và load/store nhiều thanh ghi.",
            "Mục tiêu cân bằng hiệu năng, kích thước mã và công suất.",
        ],
    )
    add_picture(doc, "image49.png", 3.7, "Kiến trúc khối ARM", "slide 138")
    add_source(doc, "VXL_CNTT.pptx, slide 137-140, 143.")

    add_heading(doc, "9.2. Hệ thanh ghi", 2)
    add_picture(doc, "image50.png", 5.6, "Các thanh ghi banked của ARM", "slide 141")
    add_table(
        doc,
        ["Thanh ghi", "Vai trò"],
        [
            ("r0-r12", "Đa dụng"),
            ("r13 / SP", "Con trỏ stack"),
            ("r14 / LR", "Địa chỉ quay lại khi gọi hàm"),
            ("r15 / PC", "Bộ đếm chương trình"),
            ("CPSR", "Trạng thái chương trình hiện tại"),
            ("SPSR", "Lưu CPSR khi xảy ra ngoại lệ"),
        ],
        widths=[1.6, 4.7],
    )
    add_body(
        doc,
        "Slide nêu ARM có 37 thanh ghi, 31 thanh ghi đa dụng; tại một thời điểm nhìn thấy 16 thanh ghi đa dụng và 2 thanh ghi trạng thái."
    )
    add_source(doc, "VXL_CNTT.pptx, slide 141-142.")

    add_heading(doc, "9.3. Lệnh và bus", 2)
    add_picture(doc, "image52.png", 5.5, "Các bit trạng thái CPSR được minh họa trong slide", "slide 144")
    add_bullets(
        doc,
        [
            "Lệnh ARM trong phần trình bày có độ dài 32 bit và dạng ba địa chỉ.",
            "Hầu hết lệnh có thể thực thi có điều kiện dựa trên cờ CPSR.",
            "Thumb nén lệnh 32 bit thành 16 bit để tăng mật độ mã.",
            "Mô hình giao tiếp gồm bus hiệu suất cao AHB, cầu AHB-APB và bus ngoại vi.",
        ],
    )
    add_picture(doc, "image53.png", 5.0, "Mô hình AHB/APB và ngoại vi ARM", "slide 146")
    add_table(
        doc,
        ["Dòng lõi theo slide", "Định hướng"],
        [
            ("Cortex-A", "Thiết bị ứng dụng, xử lý song song/SIMD"),
            ("Cortex-R", "Hệ nhúng thời gian thực, tính sẵn sàng cao"),
            ("Cortex-M", "Vi điều khiển nhỏ gọn"),
            ("Ethos-N/U", "Tác vụ AI/máy học"),
            ("Neoverse", "Máy chủ và trung tâm dữ liệu"),
            ("SecurCore", "Thẻ thông minh, bảo mật"),
        ],
        widths=[1.6, 4.7],
    )
    add_callout(
        doc,
        "Liên hệ gợi hình",
        "STM32 là ví dụ quen thuộc để hình dung nhóm vi điều khiển dùng lõi Cortex-M; "
        "điểm cần nhớ trong môn là tư tưởng load-store, hệ thanh ghi, CPSR/SPSR và mô hình bus AHB/APB.",
        "note",
    )
    add_source(doc, "VXL_CNTT.pptx, slide 144-148.")

    # Traps
    add_page_break(doc)
    add_heading(doc, "10. Bẫy và lỗi sai phổ biến", 1)
    traps = [
        ("Dấu #", "Quên # biến địa chỉ thành dữ liệu hoặc ngược lại."),
        ("Dấu @", "@R0/@R1 là gián tiếp; không dùng @R2-@R7 trong các dạng được học."),
        ("MOV/MOVC/MOVX", "Nhầm không gian RAM trong, ROM chương trình và RAM ngoài."),
        ("R0/R1 và DPTR", "R0/R1 mang địa chỉ 8 bit; DPTR mang địa chỉ 16 bit."),
        ("CY trước SUBB/ADDC", "Nếu đề cho CY, phải dùng; nếu không khởi tạo mà chương trình dùng SUBB thì kết quả phụ thuộc CY hiện tại."),
        ("P", "P=1 khi A có số bit 1 lẻ, không phải khi kết quả là số lẻ."),
        ("OV và CY", "OV xét tràn có dấu; CY xét nhớ/mượn không dấu."),
        ("SP", "PUSH tăng SP trước khi ghi; ngắt/call cất 2 byte địa chỉ."),
        ("RET và RETI", "RET cho chương trình con; RETI cho ISR."),
        ("TMOD", "Timer 1 ở nibble cao; Timer 0 ở nibble thấp."),
        ("C/T", "0 là Timer dùng clock nội; 1 là Counter dùng T0/T1."),
        ("GATE", "Muốn điều khiển thuần phần mềm bằng TRx thì GATE=0."),
        ("Mode 1/2 Timer", "Mode 1 phải nạp lại TH/TL; mode 2 tự nạp TL từ TH."),
        ("Chu kỳ sóng", "CPL sau mỗi lần tràn tạo nửa chu kỳ, phải nhân đôi để ra T."),
        ("TI/RI", "TI báo truyền xong, RI báo nhận xong; cả hai phải xóa bằng phần mềm."),
        ("Ngắt nối tiếp", "RI và TI chung vector, phải kiểm tra nguyên nhân trong ISR."),
        ("EA", "EA trong IE là cho phép tổng ngắt; chân EA của chip lại chọn bộ nhớ chương trình."),
        ("Địa chỉ vector", "0003H, 000BH, 0013H, 001BH, 0023H; đừng nhầm Timer 0 với INT0."),
        ("Cơ số", "Không hậu tố thường là thập phân; H là hexa; B là nhị phân."),
        ("Lỗi gõ trong ngân hàng", "Có câu hỏi nhắc P2 nhưng đoạn mã sửa P1, hoặc đáp án/lệnh bị gõ sai. Phải ưu tiên logic của đoạn mã và kiến thức trong slide."),
    ]
    add_table(doc, ["Bẫy", "Cách tránh"], traps, widths=[1.55, 4.75], font_size=8.9)
    add_source(doc, "Đối chiếu VXL_CNTT.pptx với các câu 2, 12, 21, 35-43, 47, 57, 62, 65-69, 80-86, 93 trong vxl.docx.")

    add_heading(doc, "Các điểm chưa nhất quán trong nguồn", 2)
    add_table(
        doc,
        ["Vị trí", "Vấn đề", "Cách xử lý khi ôn"],
        [
            ("Slide 18 vs 111", "Ghi 6 nguồn ngắt nhưng bảng vector có 5 nguồn", "Học bảng vector 5 nguồn cho 89C51 trong chương ngắt"),
            ("Slide 44", "Ví dụ dùng MOV @DPTR,A", "Phần tập lệnh sau đó quy định RAM ngoài phải dùng MOVX @DPTR,A"),
            ("Slide 110", "Ghi IRET", "Khung chương trình và các ví dụ dùng RETI"),
            ("Slide 114", "Chú thích ET0/ET1 có chỗ gõ nhầm", "Dùng bảng IE: IE.1=ET0, IE.3=ET1"),
            ("vxl.docx câu 47", "Hỏi P2 nhưng mã thay đổi P1", "Xem là lỗi đề; không suy ra giá trị P2 từ mã"),
            ("Một số câu cú pháp", "LACLL/DJNE hoặc tham số không hợp lệ do gõ sai", "Đối chiếu cú pháp bảng lệnh trong slide"),
        ],
        widths=[1.4, 2.25, 2.65],
        font_size=8.8,
    )

    # Quick sheet
    add_page_break(doc)
    add_heading(doc, "11. Tờ tra nhanh trước giờ thi", 1)
    add_heading(doc, "Thanh ghi và địa chỉ", 2)
    add_table(
        doc,
        ["Nhóm", "Chuỗi cần thuộc"],
        [
            ("PSW", "CY AC - RS1 RS0 OV - P"),
            ("TMOD", "GATE C/T M1 M0 | GATE C/T M1 M0"),
            ("TCON", "TF1 TR1 TF0 TR0 IE1 IT1 IE0 IT0"),
            ("SCON", "SM0 SM1 SM2 REN TB8 RB8 TI RI"),
            ("IE", "EA - ET2 ES ET1 EX1 ET0 EX0"),
            ("Vector", "0003 INT0; 000B T0; 0013 INT1; 001B T1; 0023 UART"),
        ],
        widths=[1.3, 5.0],
    )
    add_heading(doc, "Công thức", 2)
    add_table(
        doc,
        ["Bài", "Công thức"],
        [
            ("Chu kỳ máy 8051", "Tmc = 12/fosc"),
            ("Mode 1", "Init = 65536 - t/Tmc"),
            ("Mode 2", "Init = 256 - t/Tmc"),
            ("Dung lượng bộ nhớ", "2^n byte với n đường địa chỉ"),
            ("Số byte miền địa chỉ", "End - Start + 1"),
            ("Sóng vuông CPL mỗi tràn", "T = 2 * thời gian tràn; f = 1/T"),
        ],
        widths=[2.1, 4.2],
    )
    add_heading(doc, "Checklist 10 giây", 2)
    add_bullets(
        doc,
        [
            "Có dấu # không?",
            "Có @R0/@R1 hay DPTR không?",
            "MOV, MOVC hay MOVX?",
            "CY ban đầu bằng bao nhiêu?",
            "TMOD đang đọc nibble nào?",
            "C/T=0 hay 1?",
            "Mode 1 hay mode 2?",
            "Một lần tràn là trễ hay nửa chu kỳ?",
            "TI hay RI?",
            "RET hay RETI?",
        ],
    )

    # Practice plan
    add_heading(doc, "Lộ trình ôn 5 buổi", 2)
    add_table(
        doc,
        ["Buổi", "Nội dung", "Mục tiêu"],
        [
            ("1", "Kiến trúc, chân, bộ nhớ, SFR", "Làm nhanh câu nhận biết và bản đồ địa chỉ"),
            ("2", "Địa chỉ + MOV/MOVC/MOVX + số học/cờ", "Truy vết từng lệnh không nhầm không gian nhớ"),
            ("3", "Timer/Counter", "Tự tính TH/TL, TMOD và chu kỳ sóng"),
            ("4", "UART + ngắt", "Thuộc SCON, IE/IP, vector và khung chương trình"),
            ("5", "Làm 40 câu trong 60 phút", "Phân loại lỗi theo mục Bẫy, ôn lại tờ tra nhanh"),
        ],
        widths=[0.65, 2.65, 3.0],
    )

    add_heading(doc, "Nguồn sử dụng", 1)
    add_bullets(
        doc,
        [
            "ma-tran-de-ky-thuat-vi-xu-ly.pdf: đề cương, CLO và ma trận 40 câu.",
            "VXL_CNTT.pptx: 150 slide bài giảng, chương 1-5.",
            "vxl.docx: ngân hàng 100 câu hỏi 89C51 và các hình/đoạn mã đi kèm.",
        ],
    )
    add_callout(
        doc,
        "Giới hạn",
        "Tài liệu không tự khẳng định đáp án cho những câu bị thiếu hình hoặc bị gõ sai trong vxl.docx; "
        "thay vào đó, nó cung cấp phương pháp giải và đánh dấu điểm mơ hồ.",
        "note",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
