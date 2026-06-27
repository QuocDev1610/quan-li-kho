from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"D:\ProjectWeb\Project - Sao chép\Outline-PLDC-19-cau-de-hoc-thuoc.docx"

QUESTIONS = [
    (
        "Câu 1. Nhà nước: khái niệm, bản chất, ý nghĩa",
        "Mã nhớ: ĐẶC BIỆT – GIAI CẤP – XÃ HỘI – TOÀN DIỆN",
        [
            ("Khái niệm", "Tổ chức đặc biệt của quyền lực chính trị; có bộ máy quản lý và cưỡng chế; quản lý dân cư theo lãnh thổ; ban hành pháp luật, thu thuế và đại diện quốc gia."),
            ("Tính giai cấp", "Công cụ quyền lực của giai cấp/lực lượng cầm quyền; bảo vệ chế độ và trật tự phù hợp lợi ích của lực lượng thống trị."),
            ("Tính xã hội", "Giải quyết công việc chung: an ninh, kinh tế, hạ tầng, giáo dục, y tế, môi trường, quyền con người."),
            ("Ý nghĩa", "Phải xem xét đồng thời hai mặt giai cấp và xã hội; đặt Nhà nước trong điều kiện lịch sử – kinh tế – chính trị cụ thể."),
        ],
        "Không tuyệt đối hóa tính giai cấp hoặc tính xã hội; không coi Nhà nước là hiện tượng vĩnh viễn, bất biến.",
    ),
    (
        "Câu 2. Văn bản quy phạm pháp luật",
        "Mã nhớ đặc điểm: ĐÚNG CHỦ THỂ – ĐÚNG HÌNH THỨC – QUY TẮC CHUNG – LẶP LẠI – BẢO ĐẢM",
        [
            ("Khái niệm", "Văn bản chứa quy phạm pháp luật, được ban hành đúng thẩm quyền, hình thức, trình tự, thủ tục."),
            ("Đặc điểm", "Do chủ thể có thẩm quyền ban hành; tên gọi luật định; chứa quy tắc xử sự chung; áp dụng nhiều lần; có hiệu lực xác định; Nhà nước bảo đảm."),
            ("Nguyên tắc 1", "Hợp hiến, hợp pháp, thống nhất trong hệ thống pháp luật."),
            ("Nguyên tắc 2", "Đúng thẩm quyền, nội dung, hình thức, trình tự và thủ tục."),
            ("Nguyên tắc 3", "Dân chủ, công khai, minh bạch; lấy ý kiến đối tượng chịu tác động."),
            ("Nguyên tắc 4", "Khả thi, kịp thời, ổn định, hiệu quả; bảo vệ quyền con người và phù hợp điều ước quốc tế."),
        ],
        "Không phải mọi văn bản của cơ quan nhà nước đều là văn bản quy phạm pháp luật.",
    ),
    (
        "Câu 3. Nhà nước Cộng hòa xã hội chủ nghĩa Việt Nam",
        "Mã nhớ: CỦA DÂN – DO DÂN – VÌ DÂN; PHÁP QUYỀN – THỐNG NHẤT – KIỂM SOÁT – ĐẢNG LÃNH ĐẠO",
        [
            ("Cơ sở", "Điều 2 Hiến pháp 2013: Nhà nước pháp quyền XHCN của Nhân dân, do Nhân dân, vì Nhân dân; quyền lực thuộc về Nhân dân."),
            ("Của Nhân dân", "Nhân dân là chủ thể tối cao của quyền lực nhà nước."),
            ("Do Nhân dân", "Bộ máy hình thành từ sự ủy quyền, tham gia và giám sát của Nhân dân."),
            ("Vì Nhân dân", "Bảo vệ Tổ quốc, quyền con người, quyền công dân; nâng cao đời sống Nhân dân."),
            ("Đặc trưng", "Pháp quyền XHCN; quyền lực thống nhất nhưng có phân công, phối hợp, kiểm soát; Đảng lãnh đạo; tập trung dân chủ; quốc gia thống nhất của các dân tộc; đối ngoại hòa bình."),
        ],
        "“Dân giàu, nước mạnh…” là đặc trưng của xã hội XHCN, không phải toàn bộ đặc trưng pháp lý riêng của Nhà nước.",
    ),
    (
        "Câu 4. Luật hình sự và các nguyên tắc",
        "Mã nhớ: PHÁP CHẾ – BÌNH ĐẲNG – CÓ LỖI – CÁ THỂ HÓA – NHÂN ĐẠO – PHÒNG NGỪA",
        [
            ("Khái niệm", "Ngành luật quy định tội phạm, hình phạt, biện pháp hình sự, căn cứ trách nhiệm hình sự và nguyên tắc xử lý."),
            ("Pháp chế", "Chỉ Bộ luật Hình sự quy định tội phạm, hình phạt; xử lý đúng pháp luật, không suy diễn để buộc tội."),
            ("Bình đẳng", "Không phân biệt giới tính, dân tộc, tôn giáo, địa vị, tài sản."),
            ("Có lỗi", "Chỉ chịu trách nhiệm khi có lỗi cố ý hoặc vô ý."),
            ("Công bằng/cá thể hóa", "Hình phạt tương xứng hành vi, lỗi, nhân thân, vai trò và tình tiết."),
            ("Nhân đạo/phòng ngừa", "Không trả thù, hạ nhục; kết hợp trừng trị, giáo dục, cải tạo và phòng ngừa."),
        ],
        "Pháp nhân thương mại có thể chịu trách nhiệm hình sự đối với những tội luật định; không còn đúng khi nói mọi tổ chức đều bị loại trừ.",
    ),
    (
        "Câu 5. Quan hệ pháp luật",
        "Mã nhớ cấu thành: CHỦ THỂ – NỘI DUNG – KHÁCH THỂ; phát sinh nhờ QUY PHẠM + SỰ KIỆN",
        [
            ("Khái niệm", "Quan hệ xã hội được pháp luật điều chỉnh, các bên có quyền và nghĩa vụ pháp lý được Nhà nước bảo đảm."),
            ("Đặc điểm", "Có tính ý chí; dựa trên quy phạm và sự kiện pháp lý; quyền/nghĩa vụ xác định; được Nhà nước bảo đảm."),
            ("Chủ thể", "Cá nhân, pháp nhân, tổ chức, Nhà nước; phải có năng lực pháp luật và năng lực hành vi phù hợp."),
            ("Nội dung", "Quyền chủ thể và nghĩa vụ pháp lý tương ứng."),
            ("Khách thể", "Lợi ích vật chất, tinh thần hoặc xã hội mà các bên hướng tới."),
            ("Sự kiện pháp lý", "Căn cứ làm phát sinh, thay đổi, chấm dứt quan hệ; gồm hành vi và biến cố."),
        ],
        "Khách thể không phải là người bị thiệt hại; đó là lợi ích hoặc quan hệ xã hội mà chủ thể hướng tới.",
    ),
    (
        "Câu 6. Hình thức pháp luật",
        "Mã nhớ: TẬP QUÁN – ÁN LỆ – VĂN BẢN – ĐIỀU ƯỚC – BỔ TRỢ",
        [
            ("Khái niệm", "Phương thức tồn tại, biểu hiện của pháp luật; theo nghĩa ngoài còn gọi là nguồn pháp luật."),
            ("Tập quán pháp", "Tập quán được Nhà nước thừa nhận, áp dụng khi đủ điều kiện."),
            ("Án lệ", "Lập luận, phán quyết được lựa chọn, công bố để nghiên cứu và áp dụng cho vụ việc tương tự."),
            ("Văn bản QPPL", "Nguồn chủ yếu ở Việt Nam; rõ ràng, phổ biến, dễ công bố và hệ thống hóa."),
            ("Điều ước quốc tế", "Nguồn quan trọng trong quan hệ có yếu tố quốc tế; áp dụng theo Hiến pháp và luật."),
            ("Nguồn bổ trợ", "Nguyên tắc pháp luật, lẽ công bằng, tương tự pháp luật trong trường hợp luật cho phép."),
        ],
        "Không kết luận án lệ luôn tùy tiện; án lệ giúp thống nhất xét xử nếu được lựa chọn và kiểm soát chặt.",
    ),
    (
        "Câu 7. Luật dân sự",
        "Mã nhớ: TÀI SẢN + NHÂN THÂN; BÌNH ĐẲNG – TỰ DO – ĐỘC LẬP – TỰ CHỊU",
        [
            ("Khái niệm", "Ngành luật điều chỉnh quan hệ tài sản và nhân thân trên cơ sở bình đẳng, tự do ý chí, độc lập tài sản và tự chịu trách nhiệm."),
            ("Quan hệ tài sản", "Sở hữu, hợp đồng, nghĩa vụ, bồi thường, thừa kế; tài sản gồm vật, tiền, giấy tờ có giá, quyền tài sản."),
            ("Quan hệ nhân thân", "Tên, hình ảnh, danh dự, nhân phẩm, uy tín, đời tư, quyền tác giả; có loại gắn và không gắn tài sản."),
            ("Phương pháp", "Bình đẳng địa vị; tự do thỏa thuận; độc lập tài sản; tự chịu trách nhiệm; tự bảo vệ hoặc yêu cầu cơ quan bảo vệ."),
        ],
        "Quyền kết hôn, ly hôn chủ yếu thuộc Luật Hôn nhân và gia đình, không phải ví dụ trung tâm của Luật dân sự.",
    ),
    (
        "Câu 8. Quy phạm pháp luật",
        "Mã nhớ cấu trúc: GIẢ ĐỊNH – QUY ĐỊNH – CHẾ TÀI = AI/KHI NÀO – LÀM GÌ – HẬU QUẢ",
        [
            ("Khái niệm", "Quy tắc xử sự chung do Nhà nước ban hành hoặc thừa nhận và bảo đảm thực hiện."),
            ("Đặc điểm", "Bắt buộc chung; thể hiện ý chí nhà nước; áp dụng nhiều lần; hình thức chặt chẽ; xác định quyền/nghĩa vụ."),
            ("Giả định", "Chủ thể, hoàn cảnh, điều kiện áp dụng: ai, khi nào, ở đâu?"),
            ("Quy định", "Cách xử sự: được làm, phải làm, không được làm, làm thế nào?"),
            ("Chế tài", "Hậu quả pháp lý khi vi phạm: hình sự, hành chính, dân sự, kỷ luật."),
        ],
        "Quy phạm pháp luật không đồng nhất với điều luật; ba bộ phận có thể nằm ở nhiều điều hoặc được thể hiện ngầm.",
    ),
    (
        "Câu 9. Sở hữu và quyền sở hữu",
        "Mã nhớ quyền năng: CHIẾM HỮU – SỬ DỤNG – ĐỊNH ĐOẠT",
        [
            ("Sở hữu", "Quan hệ giữa người với người về chiếm hữu, sử dụng, định đoạt tài sản."),
            ("Quyền sở hữu", "Quyền chiếm hữu, sử dụng và định đoạt tài sản theo luật."),
            ("Xác lập", "Lao động/sản xuất; chuyển quyền; hoa lợi/lợi tức; tạo vật mới; thừa kế; chiếm hữu theo luật; thời hiệu."),
            ("Chấm dứt", "Chuyển quyền; từ bỏ; tiêu dùng/tiêu hủy; xử lý nghĩa vụ; trưng mua; tịch thu; người khác được xác lập."),
            ("Chiếm hữu", "Nắm giữ, chi phối tài sản."),
            ("Sử dụng", "Khai thác công dụng, hưởng hoa lợi, lợi tức."),
            ("Định đoạt", "Chuyển quyền, từ bỏ, tiêu dùng hoặc tiêu hủy."),
        ],
        "Ôn theo Bộ luật Dân sự 2015, không dùng Bộ luật Dân sự 2005 đã hết hiệu lực.",
    ),
    (
        "Câu 10. Vi phạm pháp luật",
        "Mã nhớ dấu hiệu: HÀNH VI – TRÁI LUẬT – CÓ LỖI – ĐỦ NĂNG LỰC – XÂM HẠI",
        [
            ("Khái niệm", "Hành vi trái pháp luật, có lỗi, do chủ thể đủ năng lực trách nhiệm thực hiện, xâm hại quan hệ được pháp luật bảo vệ."),
            ("Hành vi", "Biểu hiện bằng hành động hoặc không hành động; ý nghĩ chưa biểu hiện chưa phải vi phạm."),
            ("Trái luật", "Làm điều cấm, không làm điều bắt buộc, vượt giới hạn hoặc lạm dụng quyền."),
            ("Có lỗi", "Cố ý hoặc vô ý."),
            ("Đủ năng lực", "Đủ tuổi, khả năng nhận thức/điều khiển và điều kiện theo ngành luật."),
            ("Phân loại", "Hình sự – hành chính – dân sự – kỷ luật."),
        ],
        "Không phải mọi vi phạm đều cần hậu quả vật chất thực tế; nhiều cấu thành chỉ cần hành vi hoặc nguy cơ.",
    ),
    (
        "Câu 11. Thừa kế theo di chúc",
        "Mã nhớ: NGƯỜI LẬP – HỢP PHÁP – HÌNH THỨC – HIỆU LỰC – SUẤT BẮT BUỘC – THANH TOÁN",
        [
            ("Thừa kế", "Chuyển di sản của người chết cho người còn sống/tổ chức theo di chúc hoặc pháp luật."),
            ("Di chúc", "Ý chí cá nhân chuyển tài sản cho người khác sau khi chết."),
            ("Hợp pháp", "Minh mẫn, tự nguyện; nội dung không trái luật/đạo đức; hình thức đúng luật."),
            ("Hình thức", "Văn bản; miệng chỉ khi tính mạng bị đe dọa và không thể lập văn bản."),
            ("Hiệu lực", "Từ thời điểm mở thừa kế – thời điểm người để lại di sản chết."),
            ("Suất bắt buộc", "Con chưa thành niên; cha mẹ, vợ chồng; con thành niên không có khả năng lao động: 2/3 suất theo pháp luật."),
            ("Chia di sản", "Thanh toán nghĩa vụ, chi phí trước; phần còn lại chia theo di chúc, phần không định đoạt chia theo pháp luật."),
        ],
        "Người 15–dưới 18 tuổi cần cha mẹ/người giám hộ đồng ý việc lập di chúc, không nhất thiết phải ký vào di chúc.",
    ),
    (
        "Câu 12. Luật lao động",
        "Mã nhớ: QUAN HỆ CÁ NHÂN – TẬP THỂ – LIÊN QUAN; THỎA THUẬN – MỆNH LỆNH – BẢO VỆ – ĐẠI DIỆN",
        [
            ("Khái niệm", "Ngành luật điều chỉnh quan hệ lao động có trả công và các quan hệ liên quan trực tiếp."),
            ("Quan hệ cá nhân", "Tuyển dụng, hợp đồng, lương, giờ làm/nghỉ, an toàn, kỷ luật, chấm dứt."),
            ("Quan hệ tập thể", "Đối thoại, thương lượng, thỏa ước, đại diện, tranh chấp, đình công."),
            ("Quan hệ liên quan", "Việc làm, đào tạo, bảo hiểm, thanh tra, quản lý, giải quyết tranh chấp."),
            ("Phương pháp", "Thỏa thuận; quyền quản lý của người sử dụng; bảo vệ người lao động; sự tham gia của tổ chức đại diện."),
        ],
        "Quan hệ công vụ của cán bộ, công chức chủ yếu không thuộc Luật lao động.",
    ),
    (
        "Câu 13. Hợp đồng lao động",
        "Mã nhớ: THỎA THUẬN – CÓ TRẢ CÔNG – CÓ QUẢN LÝ; 2 LOẠI HỢP ĐỒNG",
        [
            ("Khái niệm", "Thỏa thuận về việc làm có trả công, tiền lương, điều kiện lao động và quyền/nghĩa vụ; có sự quản lý, điều hành, giám sát."),
            ("Nguyên tắc", "Tự nguyện, bình đẳng, thiện chí, hợp tác, trung thực; tự do nhưng không trái luật."),
            ("Hình thức", "Văn bản/điện tử; lời nói với hợp đồng dưới 1 tháng, trừ trường hợp bắt buộc văn bản."),
            ("Hai loại", "Không xác định thời hạn; xác định thời hạn không quá 36 tháng."),
            ("Nội dung", "Chủ thể, công việc, địa điểm, thời hạn, lương, giờ làm/nghỉ, bảo hộ, bảo hiểm, đào tạo."),
            ("Chấm dứt", "Theo căn cứ luật định; đơn phương phải đúng căn cứ và báo trước, nếu trái luật phải bồi thường/khắc phục."),
        ],
        "Không còn loại hợp đồng mùa vụ riêng theo Bộ luật Lao động 2019.",
    ),
    (
        "Câu 14. Tranh chấp lao động",
        "Mã nhớ: CÁ NHÂN/TẬP THỂ – QUYỀN/LỢI ÍCH; HÒA GIẢI – TRỌNG TÀI – TÒA ÁN",
        [
            ("Khái niệm", "Tranh chấp về quyền, nghĩa vụ, lợi ích trong xác lập, thực hiện hoặc chấm dứt quan hệ lao động."),
            ("Phân loại", "Cá nhân; tập thể về quyền; tập thể về lợi ích."),
            ("Thẩm quyền", "Hòa giải viên lao động; Hội đồng trọng tài lao động; Tòa án."),
            ("Cá nhân", "Thường phải hòa giải trước; một số tranh chấp được khởi kiện thẳng như sa thải, đơn phương chấm dứt, bảo hiểm…"),
            ("Tập thể về quyền", "Hòa giải → trọng tài hoặc Tòa án."),
            ("Tập thể về lợi ích", "Hòa giải → trọng tài hoặc tiến hành thủ tục đình công khi đủ điều kiện."),
        ],
        "Chỉ tranh chấp tập thể về lợi ích mới có thể dẫn đến đình công; cơ chế Chủ tịch UBND huyện là kiến thức cũ.",
    ),
    (
        "Câu 15. Luật hành chính",
        "Mã nhớ: QUẢN LÝ NHÀ NƯỚC – QUYỀN LỰC/PHỤC TÙNG – QUYẾT ĐỊNH ĐƠN PHƯƠNG",
        [
            ("Khái niệm", "Ngành luật điều chỉnh quan hệ phát sinh trong quản lý hành chính nhà nước và thực hiện quyền hành pháp."),
            ("Nhóm 1", "Quản lý của Chính phủ, bộ, UBND và cơ quan hành chính."),
            ("Nhóm 2", "Quản lý nội bộ cơ quan nhà nước."),
            ("Nhóm 3", "Hoạt động hành chính nội bộ của cơ quan lập pháp, tư pháp."),
            ("Nhóm 4", "Hoạt động quản lý của tổ chức/cá nhân được Nhà nước trao quyền."),
            ("Phương pháp", "Quyền lực – phục tùng: một bên nhân danh Nhà nước ra quyết định, kiểm tra, cưỡng chế; bên kia chấp hành quyết định hợp pháp."),
        ],
        "Quyết định hành chính phải đúng thẩm quyền, căn cứ, thủ tục; cá nhân có quyền khiếu nại hoặc khởi kiện.",
    ),
    (
        "Câu 16. Tội phạm",
        "Mã nhớ đặc điểm: NGUY HIỂM – ĐƯỢC QUY ĐỊNH – CÓ LỖI – ĐỦ NĂNG LỰC – CHỊU HÌNH PHẠT",
        [
            ("Khái niệm", "Hành vi nguy hiểm được Bộ luật Hình sự quy định, do cá nhân đủ năng lực hoặc pháp nhân thương mại thực hiện có lỗi."),
            ("Nguy hiểm", "Gây hoặc đe dọa gây thiệt hại đáng kể cho quan hệ được luật hình sự bảo vệ."),
            ("Được quy định", "Không có tội nếu Bộ luật Hình sự không quy định."),
            ("Có lỗi", "Cố ý trực tiếp/gián tiếp; vô ý quá tự tin/cẩu thả."),
            ("Năng lực", "Đủ tuổi và khả năng nhận thức, điều khiển; pháp nhân chỉ chịu với tội luật định."),
            ("Phân loại", "Ít nghiêm trọng: đến 3 năm; nghiêm trọng: trên 3–7; rất nghiêm trọng: trên 7–15; đặc biệt nghiêm trọng: trên 15, chung thân hoặc tử hình."),
        ],
        "Dùng Bộ luật Hình sự 2015 sửa đổi 2017; không dùng định nghĩa của Bộ luật Hình sự 1999.",
    ),
    (
        "Câu 17. Tố cáo",
        "Mã nhớ trình tự: THỤ LÝ – XÁC MINH – KẾT LUẬN – XỬ LÝ",
        [
            ("Khái niệm", "Cá nhân báo cho chủ thể có thẩm quyền về hành vi vi phạm gây/đe dọa gây thiệt hại."),
            ("Đặc điểm", "Người tố cáo là cá nhân; đối tượng là hành vi vi phạm; phải trung thực; được bảo mật và bảo vệ."),
            ("Thẩm quyền", "Người đứng đầu quản lý người bị tố cáo; cơ quan quản lý lĩnh vực; dấu hiệu tội phạm chuyển cơ quan tố tụng."),
            ("Hình thức", "Bằng đơn hoặc trình bày trực tiếp."),
            ("Trình tự", "Thụ lý → xác minh → kết luận → xử lý kết luận."),
            ("Thời hạn", "30 ngày; phức tạp gia hạn 1 lần; đặc biệt phức tạp gia hạn 2 lần, mỗi lần tối đa 30 ngày."),
        ],
        "Thông tin nặc danh không giải quyết theo thủ tục tố cáo, nhưng có chứng cứ rõ vẫn có thể được dùng để thanh tra/kiểm tra; không gọi tuyệt đối là “bất hợp pháp”.",
    ),
    (
        "Câu 18. Luật Hôn nhân và gia đình",
        "Mã nhớ: TỰ NGUYỆN – MỘT VỢ MỘT CHỒNG – BÌNH ĐẲNG – KHÔNG PHÂN BIỆT – BẢO VỆ – TRUYỀN THỐNG",
        [
            ("Khái niệm", "Ngành luật điều chỉnh quan hệ nhân thân, tài sản từ hôn nhân và giữa các thành viên gia đình."),
            ("Nguyên tắc 1", "Hôn nhân tự nguyện, tiến bộ, một vợ một chồng, vợ chồng bình đẳng."),
            ("Nguyên tắc 2", "Không phân biệt dân tộc, tôn giáo; bảo vệ hôn nhân có yếu tố nước ngoài hợp pháp."),
            ("Nguyên tắc 3", "Gia đình ấm no, tiến bộ, hạnh phúc; thành viên tôn trọng, chăm sóc, giúp đỡ nhau."),
            ("Nguyên tắc 4", "Không phân biệt các con; bảo vệ trẻ em, người cao tuổi, người khuyết tật và người yếu thế."),
            ("Ưu việt", "Kết hợp quyền con người, bình đẳng hiện đại với tình nghĩa, hiếu thảo và tương trợ truyền thống."),
        ],
        "Ôn theo Luật Hôn nhân và gia đình 2014, không theo Luật năm 2000.",
    ),
    (
        "Câu 19. Hiến pháp năm 2013",
        "Mã nhớ 11 chương: CHÍNH TRỊ – QUYỀN – KINH TẾ/XÃ HỘI – BẢO VỆ – QH – CTN – CP – TA/VKS – ĐỊA PHƯƠNG – HĐBC/KTNN – HIỆU LỰC",
        [
            ("Khái quát", "Thông qua 28/11/2013; hiệu lực 01/01/2014; Lời nói đầu, 11 chương, 120 điều; hiệu lực pháp lý cao nhất."),
            ("Chương I–IV", "Chế độ chính trị; quyền con người/công dân; kinh tế–xã hội–văn hóa–giáo dục–khoa học–môi trường; bảo vệ Tổ quốc."),
            ("Chương V–VIII", "Quốc hội; Chủ tịch nước; Chính phủ; Tòa án và Viện kiểm sát."),
            ("Chương IX–XI", "Chính quyền địa phương; Hội đồng bầu cử quốc gia và Kiểm toán nhà nước; hiệu lực và sửa đổi Hiến pháp."),
            ("Điểm mới lớn", "Chủ quyền Nhân dân; quyền con người; hạn chế quyền bằng luật; kiểm soát quyền lực; Chính phủ thực hiện hành pháp; Tòa án thực hiện tư pháp."),
            ("Ý nghĩa", "Nền tảng tổ chức và giới hạn quyền lực nhà nước, bảo vệ quyền và định hướng phát triển đất nước."),
        ],
        "Khi phân tích không chỉ kể tên chương: cần nêu ít nhất 3 giá trị nổi bật về quyền con người, kiểm soát quyền lực và phân công lập pháp–hành pháp–tư pháp.",
    ),
]


def set_run_font(run, size=11, bold=None, italic=None, color="222222"):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=8.5, color="666666")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.72)
sec.bottom_margin = Inches(0.68)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.32)
sec.footer_distance = Inches(0.32)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string("222222")
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.16

h1 = styles["Heading 1"]
h1.font.name = "Arial"
h1._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
h1.font.size = Pt(15)
h1.font.bold = True
h1.font.color.rgb = RGBColor.from_string("1F4E79")
h1.paragraph_format.space_before = Pt(14)
h1.paragraph_format.space_after = Pt(6)
h1.paragraph_format.keep_with_next = True

bullet = styles["List Bullet"]
bullet.font.name = "Arial"
bullet.font.size = Pt(10.3)
bullet.paragraph_format.left_indent = Inches(0.37)
bullet.paragraph_format.first_line_indent = Inches(-0.18)
bullet.paragraph_format.space_after = Pt(3)
bullet.paragraph_format.line_spacing = 1.15

header = sec.header.paragraphs[0]
header.text = "OUTLINE PHÁP LUẬT ĐẠI CƯƠNG • 19 CÂU"
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(header.runs[0], size=8.3, bold=True, color="6B7280")
add_page_number(sec.footer.paragraphs[0])

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(55)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OUTLINE HỌC THUỘC NHANH")
set_run_font(r, size=24, bold=True, color="17365D")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("PHÁP LUẬT ĐẠI CƯƠNG – 19 CÂU")
set_run_font(r, size=17, bold=True, color="2E75B6")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(22)
r = p.add_run("Từ khóa chấm điểm • Mã nhớ • Điểm dễ nhầm")
set_run_font(r, size=11.5, italic=True, color="555555")

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.25)
p.paragraph_format.right_indent = Inches(0.25)
p.paragraph_format.space_after = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run(
    "Cách học: (1) đọc “Mã nhớ”; (2) che phần nội dung và tự nói lại các ý; "
    "(3) mở ra kiểm tra; (4) viết mỗi ý thành 2–3 câu khi làm bài thi. "
    "Outline này dùng để thuộc khung, còn bản 19 câu hoàn chỉnh dùng để học phần phân tích."
)
set_run_font(r, size=10.5)
shade(p, "EAF1F8")

p = doc.add_paragraph()
r = p.add_run("BẢN ĐỒ 5 CỤM KIẾN THỨC")
set_run_font(r, size=13, bold=True, color="1F4E79")

clusters = [
    "Lý luận chung: Câu 1, 2, 3, 5, 6, 8, 10",
    "Dân sự – gia đình: Câu 7, 9, 11, 18",
    "Hình sự: Câu 4, 16",
    "Lao động – hành chính: Câu 12, 13, 14, 15, 17",
    "Hiến pháp: Câu 19",
]
for text in clusters:
    q = doc.add_paragraph(style="List Bullet")
    q.add_run(text)

doc.add_page_break()

for title, mnemonic, points, trap in QUESTIONS:
    doc.add_paragraph(title, style="Heading 1")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(mnemonic)
    set_run_font(r, size=10.2, bold=True, color="1F4E79")
    shade(p, "E8EEF5")

    for label, text in points:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.3, bold=True, color="222222")
        r = p.add_run(text)
        set_run_font(r, size=10.3)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Dễ nhầm: ")
    set_run_font(r, size=9.8, bold=True, color="9B1C1C")
    r = p.add_run(trap)
    set_run_font(r, size=9.8, italic=True, color="7A1F1F")
    shade(p, "FFF2F2")

doc.add_page_break()
p = doc.add_paragraph("CÔNG THỨC LÀM BÀI 5 BƯỚC", style="Heading 1")
steps = [
    "1. Viết khái niệm chuẩn trong 2–3 dòng.",
    "2. Chia đúng các vế mà đề yêu cầu thành đề mục.",
    "3. Mỗi ý: nêu từ khóa → giải thích 1–2 câu → cho ví dụ/phân biệt nếu cần.",
    "4. Ghi đúng văn bản hiện hành; nếu đề dùng luật cũ, nêu lưu ý cập nhật.",
    "5. Kết luận 1 câu về ý nghĩa hoặc vai trò.",
]
for text in steps:
    q = doc.add_paragraph(style="List Bullet")
    q.add_run(text)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Lịch ôn gợi ý trong 4 lượt")
set_run_font(r, size=12, bold=True, color="1F4E79")
for text in [
    "Lượt 1: Đọc toàn bộ mã nhớ và tiêu đề.",
    "Lượt 2: Thuộc khái niệm + số lượng ý chính của từng câu.",
    "Lượt 3: Tự nói lại không nhìn tài liệu; đánh dấu câu còn vấp.",
    "Lượt 4: Viết thử 3–5 câu ngẫu nhiên theo thời gian thi.",
]:
    q = doc.add_paragraph(style="List Bullet")
    q.add_run(text)

doc.core_properties.title = "Outline Pháp luật đại cương – 19 câu dễ học thuộc"
doc.core_properties.subject = "Từ khóa, mã nhớ và điểm dễ nhầm"
doc.core_properties.author = "Codex"
doc.save(OUTPUT)
print(OUTPUT)
